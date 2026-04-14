import re
from pathlib import Path
import sys
from zipfile import ZipFile

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

REPO_ROOT = Path(__file__).resolve().parents[3]
PACKAGE_SCRIPTS_DIR = REPO_ROOT / "packages" / "bensz-paper" / "scripts"
if str(PACKAGE_SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(PACKAGE_SCRIPTS_DIR))

import manuscript_tool


def _read_document_xml(docx_path: Path) -> str:
    with ZipFile(docx_path) as archive:
        return archive.read("word/document.xml").decode("utf-8", errors="replace")


def _build_docx_via_production_pipeline(markdown: str, docx_path: Path, tmp_path: Path) -> None:
    bib_path = tmp_path / "refs.bib"
    bib_path.write_text("", encoding="utf-8")
    csl_path = REPO_ROOT / "projects" / "paper-sci-01" / "artifacts" / "manuscript.csl"
    reference_doc = REPO_ROOT / "projects" / "paper-sci-01" / "artifacts" / "reference.docx"

    manuscript_tool.build_docx_from_markdown(
        manuscript_md=markdown,
        docx_path=docx_path,
        csl_path=csl_path,
        bibliography_path=bib_path,
        reference_doc=reference_doc,
    )
    manuscript_tool.fix_docx_spacing(docx_path)


def test_package_version_matches_cli_version():
    scripts_init = PACKAGE_SCRIPTS_DIR / "__init__.py"
    namespace: dict[str, str] = {}
    exec(scripts_init.read_text(encoding="utf-8"), namespace)

    assert namespace["__version__"] == manuscript_tool.VERSION


def test_collect_extra_tex_inputs_follows_main_tex_order(tmp_path):
    project_dir = tmp_path / "paper-demo"
    (project_dir / "extraTex" / "front").mkdir(parents=True)
    (project_dir / "extraTex" / "body").mkdir(parents=True)
    (project_dir / "extraTex" / "back").mkdir(parents=True)

    (project_dir / "main.tex").write_text(
        "\n".join(
            [
                r"\documentclass{article}",
                r"\begin{document}",
                r"\input{extraTex/front/frontmatter.tex}",
                r"\input{extraTex/body/introduction}",
                r"\input{ignored/not-part-of-docx.tex}",
                r"\input{extraTex/back/figure-legends.tex}",
                r"\end{document}",
                "",
            ]
        ),
        encoding="utf-8",
    )
    (project_dir / "extraTex" / "front" / "frontmatter.tex").write_text("front", encoding="utf-8")
    (project_dir / "extraTex" / "body" / "introduction.tex").write_text("body", encoding="utf-8")
    (project_dir / "extraTex" / "back" / "figure-legends.tex").write_text("back", encoding="utf-8")

    inputs = manuscript_tool.collect_extra_tex_inputs(project_dir)

    assert inputs == [
        Path("extraTex/front/frontmatter.tex"),
        Path("extraTex/body/introduction.tex"),
        Path("extraTex/back/figure-legends.tex"),
    ]


def test_collect_extra_tex_inputs_ignores_commented_inputs(tmp_path):
    project_dir = tmp_path / "paper-demo"
    (project_dir / "extraTex" / "body").mkdir(parents=True)

    (project_dir / "main.tex").write_text(
        "\n".join(
            [
                r"\documentclass{article}",
                r"\begin{document}",
                r"% \input{extraTex/body/commented-out.tex}",
                r"\input{extraTex/body/kept.tex} % keep this one",
                r"\end{document}",
                "",
            ]
        ),
        encoding="utf-8",
    )
    (project_dir / "extraTex" / "body" / "kept.tex").write_text("kept", encoding="utf-8")

    inputs = manuscript_tool.collect_extra_tex_inputs(project_dir)

    assert inputs == [Path("extraTex/body/kept.tex")]


def test_pandoc_latex_to_markdown_preserves_headings_and_citations():
    latex = "\n".join(
        [
            r"\section{Results}",
            r"Text with \supercite{KeyA,KeyB} and \texttt{inline code}.",
            r"\subsection*{Inner}",
            r"More text.",
            "",
        ]
    )

    markdown = manuscript_tool.pandoc_latex_to_markdown(latex)

    assert "# Results" in markdown
    assert "[@KeyA; @KeyB]" in markdown
    assert "## Inner" in markdown
    assert "`inline code`" in markdown


def test_pandoc_latex_to_markdown_preserves_frontmatter_superscripts():
    latex = "\n".join(
        [
            r"Author One\textsuperscript{1†}, Author Two\textsuperscript{2*}",
            r"\noindent\textsuperscript{1}Department A",
            r"\noindent\textsuperscript{2}Department B",
            "",
        ]
    )

    markdown = manuscript_tool.pandoc_latex_to_markdown(latex)
    normalized = " ".join(markdown.split())

    assert "Author One^1†^, Author Two^2\\*^" in normalized
    assert "^1^Department A" in normalized
    assert "^2^Department B" in normalized


def test_convert_sup_tags_to_superscript_preserves_escaped_asterisk():
    markdown = manuscript_tool._convert_sup_tags_to_superscript(
        r"<sup>\*</sup>Correspondence: template.author@example.org."
    )

    assert markdown == r"^\*^Correspondence: template.author@example.org."


def test_count_visible_words_ignores_latex_commands_and_math(tmp_path):
    tex_path = tmp_path / "section.tex"
    tex_path.write_text(
        "\n".join(
            [
                r"\section{Introduction}",
                r"Visible \textbf{bold text} and \emph{emphasis}.",
                r"Citation \cite{Key2026} and inline math $\alpha + \beta$ stay out.",
                r"\label{sec:intro}",
                r"\begin{figure}",
                r"\includegraphics[width=\textwidth]{figure-1.png}",
                r"\caption{Figure caption words.}",
                r"\end{figure}",
                r"% commented words should disappear",
                "",
            ]
        ),
        encoding="utf-8",
    )

    summary = manuscript_tool.count_words_for_tex_sources([tex_path])

    assert summary.total_words == 15
    assert summary.file_counts == [(tex_path, 15)]


def test_count_visible_words_follows_input_chain_and_simple_macros(tmp_path):
    project_dir = tmp_path / "paper-demo"
    (project_dir / "extraTex" / "body").mkdir(parents=True)

    main_tex = project_dir / "main.tex"
    intro_tex = project_dir / "extraTex" / "body" / "introduction.tex"
    results_tex = project_dir / "extraTex" / "body" / "results.tex"

    main_tex.write_text(
        "\n".join(
            [
                r"\documentclass{article}",
                r"\newcommand{\JournalName}{Nature Medicine}",
                r"\begin{document}",
                r"\input{extraTex/body/introduction}",
                r"\input{extraTex/body/results.tex}",
                r"\end{document}",
                "",
            ]
        ),
        encoding="utf-8",
    )
    intro_tex.write_text(
        "\n".join(
            [
                r"\section{Introduction}",
                r"We submitted to \JournalName.",
                "",
            ]
        ),
        encoding="utf-8",
    )
    results_tex.write_text(
        "\n".join(
            [
                r"\section{Results}",
                r"Observed robust treatment benefit.",
                "",
            ]
        ),
        encoding="utf-8",
    )

    summary = manuscript_tool.count_words_for_tex_sources([main_tex])

    assert summary.total_words == 11
    assert summary.file_counts == [(main_tex, 11)]


def test_count_words_cli_prints_per_file_and_total(tmp_path, monkeypatch, capsys):
    first_tex = tmp_path / "abstract.tex"
    second_tex = tmp_path / "discussion.tex"
    first_tex.write_text(r"\section*{Abstract}Short abstract text.", encoding="utf-8")
    second_tex.write_text(r"\section{Discussion}Discussion text only.", encoding="utf-8")

    monkeypatch.setattr(
        sys,
        "argv",
        [
            "paper_project_tool.py",
            "count-words",
            str(first_tex),
            str(second_tex),
        ],
    )

    manuscript_tool.main()

    captured = capsys.readouterr()
    assert f"{first_tex}: 4" in captured.out
    assert f"{second_tex}: 4" in captured.out
    assert "Total visible words: 8" in captured.out


def test_pandoc_latex_to_markdown_converts_simple_tables():
    latex = "\n".join(
        [
            r"\section{Methods}",
            r"\begin{table}[htbp]",
            r"\centering",
            r"\caption{Demo table}",
            r"\begin{tabular}{ll}",
            r"\hline",
            r"Item & Value \\",
            r"\hline",
            r"A & 1 \\",
            r"B & 2 \\",
            r"\hline",
            r"\end{tabular}",
            r"\end{table}",
            "",
        ]
    )

    markdown = manuscript_tool.pandoc_latex_to_markdown(latex)

    assert "# Methods" in markdown
    assert "Item   Value" in markdown
    assert "A      1" in markdown
    assert "Demo table" in markdown


def test_pandoc_latex_to_markdown_keeps_standard_math_syntax():
    latex = r"Inline $\rho$ and $\Delta$. Display: $$\log_{10}(x+1)$$."

    markdown = manuscript_tool.pandoc_latex_to_markdown(latex)

    assert "$`\\rho`$" not in markdown
    assert "$`\\Delta`$" not in markdown
    assert "``` math" not in markdown
    assert r"$\rho$" in markdown
    assert r"$\Delta$" in markdown
    assert r"$$\log_{10}(x+1)$$" in markdown


def test_build_markdown_for_docx_reads_extra_tex_without_source_manifest(tmp_path):
    project_dir = tmp_path / "paper-demo"
    (project_dir / "extraTex" / "front").mkdir(parents=True)
    (project_dir / "extraTex" / "body").mkdir(parents=True)
    (project_dir / "extraTex" / "back").mkdir(parents=True)

    (project_dir / "main.tex").write_text(
        "\n".join(
            [
                r"\documentclass{article}",
                r"\begin{document}",
                r"\input{extraTex/front/frontmatter.tex}",
                r"\input{extraTex/body/introduction.tex}",
                r"\input{extraTex/back/figure-legends.tex}",
                r"\end{document}",
                "",
            ]
        ),
        encoding="utf-8",
    )
    (project_dir / "extraTex" / "front" / "frontmatter.tex").write_text(
        "\n".join(
            [
                r"\begin{center}",
                r"\textbf{Template Example}",
                "",
                r"Author One\textsuperscript{1}",
                r"\end{center}",
                "",
            ]
        ),
        encoding="utf-8",
    )
    (project_dir / "extraTex" / "body" / "introduction.tex").write_text(
        "\n".join(
            [
                r"\section{Introduction}",
                r"Body text with \supercite{Demo2026}.",
                "",
            ]
        ),
        encoding="utf-8",
    )
    (project_dir / "extraTex" / "back" / "figure-legends.tex").write_text(
        "\n".join(
            [
                r"\section*{Figure legends}",
                r"Legend text.",
                "",
            ]
        ),
        encoding="utf-8",
    )

    markdown = manuscript_tool.build_markdown_for_docx(project_dir)

    assert "Template Example" in markdown
    assert "Author One" in markdown
    assert "# Introduction" in markdown
    assert "[@Demo2026]" in markdown
    assert "# Figure legends" in markdown


def test_normalize_frontmatter_markdown_marks_centered_author_block():
    markdown = "\n".join(
        [
            "::: center",
            "**Template Example**",
            "",
            "Author One^1^, Author Two^2^",
            ":::",
            "",
            "^1^Department A",
            "",
        ]
    )

    normalized = manuscript_tool._normalize_frontmatter_markdown(markdown)

    assert normalized.splitlines()[0] == "# Template Example"
    assert manuscript_tool.DOCX_FRONTMATTER_CENTER_START in normalized
    assert "Author One^1^, Author Two^2^" in normalized
    assert manuscript_tool.DOCX_FRONTMATTER_CENTER_END in normalized
    assert "^1^Department A" in normalized


def test_build_markdown_for_docx_frontmatter_superscripts_survive_docx_roundtrip(tmp_path):
    project_dir = tmp_path / "paper-demo"
    (project_dir / "extraTex" / "front").mkdir(parents=True)

    (project_dir / "main.tex").write_text(
        "\n".join(
            [
                r"\documentclass{article}",
                r"\begin{document}",
                r"\input{extraTex/front/frontmatter.tex}",
                r"\end{document}",
                "",
            ]
        ),
        encoding="utf-8",
    )
    (project_dir / "extraTex" / "front" / "frontmatter.tex").write_text(
        "\n".join(
            [
                r"\begin{center}",
                r"\textbf{Template Example}",
                "",
                r"Author One\textsuperscript{1†}, Author Two\textsuperscript{2*}",
                r"\end{center}",
                "",
                r"\noindent\textsuperscript{1}Department A\par",
                r"\noindent\textsuperscript{2}Department B\par",
                "",
            ]
        ),
        encoding="utf-8",
    )

    markdown = manuscript_tool.build_markdown_for_docx(project_dir)
    docx_path = tmp_path / "frontmatter.docx"
    _build_docx_via_production_pipeline(markdown, docx_path, tmp_path)

    doc = Document(docx_path)
    author_para = next(para for para in doc.paragraphs if "Author One" in para.text)
    affiliation_para = next(para for para in doc.paragraphs if "Department A" in para.text)
    second_affiliation_para = next(para for para in doc.paragraphs if "Department B" in para.text)

    assert author_para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert affiliation_para.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.CENTER
    assert second_affiliation_para.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.CENTER
    assert manuscript_tool.DOCX_FRONTMATTER_CENTER_START not in "\n".join(
        para.text for para in doc.paragraphs
    )
    assert manuscript_tool.DOCX_FRONTMATTER_CENTER_END not in "\n".join(
        para.text for para in doc.paragraphs
    )
    assert any(run.text == "1†" and run.font.superscript for run in author_para.runs)
    assert any(run.text == "2*" and run.font.superscript for run in author_para.runs)
    assert any(run.text == "1" and run.font.superscript for run in affiliation_para.runs)
    assert any(run.text == "2" and run.font.superscript for run in second_affiliation_para.runs)


def test_build_markdown_for_docx_correspondence_asterisk_survives_docx_roundtrip(tmp_path):
    project_dir = tmp_path / "paper-demo"
    (project_dir / "extraTex" / "front").mkdir(parents=True)

    (project_dir / "main.tex").write_text(
        "\n".join(
            [
                r"\documentclass{article}",
                r"\begin{document}",
                r"\input{extraTex/front/frontmatter.tex}",
                r"\end{document}",
                "",
            ]
        ),
        encoding="utf-8",
    )
    (project_dir / "extraTex" / "front" / "frontmatter.tex").write_text(
        "\n".join(
            [
                r"\begin{center}",
                r"\textbf{Template Example}",
                r"\end{center}",
                "",
                r"\noindent\textsuperscript{*}Correspondence: template.author@example.org.\par",
                "",
            ]
        ),
        encoding="utf-8",
    )

    markdown = manuscript_tool.build_markdown_for_docx(project_dir)
    docx_path = tmp_path / "correspondence.docx"
    _build_docx_via_production_pipeline(markdown, docx_path, tmp_path)

    correspondence_para = next(
        para for para in Document(docx_path).paragraphs if "Correspondence:" in para.text
    )

    assert "^*^" not in correspondence_para.text
    assert any(run.text == "*" and run.font.superscript for run in correspondence_para.runs)


def test_build_markdown_for_docx_expands_simple_newcommand_metadata_across_fragments(tmp_path):
    project_dir = tmp_path / "paper-coverletter-demo"
    (project_dir / "extraTex" / "front").mkdir(parents=True)
    (project_dir / "extraTex" / "body").mkdir(parents=True)

    (project_dir / "main.tex").write_text(
        "\n".join(
            [
                r"\documentclass{article}",
                r"\begin{document}",
                r"\input{extraTex/front/metadata.tex}",
                r"\input{extraTex/body/letter.tex}",
                r"\end{document}",
                "",
            ]
        ),
        encoding="utf-8",
    )
    (project_dir / "extraTex" / "front" / "metadata.tex").write_text(
        "\n".join(
            [
                r"\newcommand{\CoverLetterDate}{31 March 2026}",
                r"\newcommand{\CoverLetterRecipient}{Editors of [Target Journal]}",
                r"\newcommand{\CoverLetterSalutation}{Dear Editors,}",
                r"\newcommand{\CoverLetterManuscriptTitle}{Template Example: A Multi-Cohort Translational Oncology Study}",
                r"\newcommand{\CoverLetterJournal}{[Target Journal]}",
                r"\newcommand{\CoverLetterClosingName}{Feng BaoBao}",
                r"\newcommand{\CoverLetterClosingRole}{Corresponding Author}",
                r"\newcommand{\CoverLetterClosingAffiliation}{Department of Translational Medicine, Example Research University}",
                r"\newcommand{\CoverLetterClosingEmail}{feng.baobao@example.edu}",
                "",
            ]
        ),
        encoding="utf-8",
    )
    (project_dir / "extraTex" / "body" / "letter.tex").write_text(
        "\n".join(
            [
                r"\noindent \CoverLetterDate\par",
                r"\noindent \CoverLetterRecipient\par",
                r"\noindent \CoverLetterSalutation\par",
                r"\noindent We are pleased to submit our manuscript, entitled ``\CoverLetterManuscriptTitle,'' for consideration at \CoverLetterJournal.",
                r"\noindent Sincerely,\par",
                r"\noindent \CoverLetterClosingName\par",
                r"\noindent \CoverLetterClosingRole\par",
                r"\noindent \CoverLetterClosingAffiliation\par",
                r"\noindent \CoverLetterClosingEmail\par",
                "",
            ]
        ),
        encoding="utf-8",
    )

    markdown = manuscript_tool.build_markdown_for_docx(project_dir)

    assert "31 March 2026" in markdown
    assert "Editors of \\[Target Journal\\]" in markdown
    assert "Dear Editors," in markdown
    assert "Template Example: A" in markdown
    assert "Multi-Cohort Translational Oncology Study" in markdown
    assert "Feng BaoBao" in markdown
    assert "Corresponding Author" in markdown
    assert "feng.baobao@example.edu" in markdown
    assert r"\CoverLetter" not in markdown


def test_build_docx_from_markdown_promotes_math_to_omml(tmp_path):
    bib_path = tmp_path / "refs.bib"
    bib_path.write_text(
        "\n".join(
            [
                "@article{Demo2026,",
                "  author = {Alpha, Alice and Beta, Bob},",
                "  title = {Synthetic Math Regression},",
                "  journal = {Journal of Regression Tests},",
                "  year = {2026},",
                "  doi = {10.1234/demo.2026.002},",
                "}",
                "",
            ]
        ),
        encoding="utf-8",
    )

    docx_path = tmp_path / "math.docx"
    csl_path = REPO_ROOT / "projects" / "paper-sci-01" / "artifacts" / "manuscript.csl"
    reference_doc = REPO_ROOT / "projects" / "paper-sci-01" / "artifacts" / "reference.docx"
    markdown = "\n".join(
        [
            "Body text with a citation [@Demo2026]. Inline $\\rho$ and $\\Delta$.",
            "",
            "Display:",
            "$$\\log_{10}(x+1)$$",
            "",
        ]
    )

    manuscript_tool.build_docx_from_markdown(
        manuscript_md=markdown,
        docx_path=docx_path,
        csl_path=csl_path,
        bibliography_path=bib_path,
        reference_doc=reference_doc,
    )

    document_xml = _read_document_xml(docx_path)
    assert document_xml.count("<m:oMath") >= 3
    assert "\\rho" not in document_xml
    assert "\\Delta" not in document_xml
    assert "\\log_{10}(x+1)" not in document_xml
    assert "`" not in document_xml


def test_build_docx_from_markdown_supports_projects_without_bibliography(tmp_path):
    docx_path = tmp_path / "cover-letter.docx"
    reference_doc = REPO_ROOT / "projects" / "paper-sci-01" / "artifacts" / "reference.docx"
    markdown = "\n".join(
        [
            "31 March 2026",
            "",
            "Dear Editors,",
            "",
            "We are pleased to submit our anonymized manuscript.",
            "",
            "Sincerely,",
            "",
            "Feng BaoBao",
            "",
        ]
    )

    manuscript_tool.build_docx_from_markdown(
        manuscript_md=markdown,
        docx_path=docx_path,
        reference_doc=reference_doc,
    )

    assert docx_path.exists()
    assert "Feng BaoBao" in "\n".join(para.text for para in Document(docx_path).paragraphs)


def test_main_tex_uses_bibliography_detects_biblatex_commands():
    assert manuscript_tool.main_tex_uses_bibliography(
        "\n".join(
            [
                r"\documentclass{article}",
                r"\addbibresource{references/refs.bib}",
                r"\begin{document}",
                r"\printbibliography[heading=none]",
                r"\end{document}",
            ]
        )
    )


def test_main_tex_uses_bibliography_returns_false_for_plain_cover_letter():
    assert not manuscript_tool.main_tex_uses_bibliography(
        "\n".join(
            [
                r"\documentclass{article}",
                r"\begin{document}",
                r"\input{extraTex/front/metadata.tex}",
                r"\input{extraTex/body/letter.tex}",
                r"\end{document}",
            ]
        )
    )


def test_remove_legacy_docx_intermediates_cleans_old_cache(tmp_path):
    cache_dir = tmp_path / ".latex-cache"
    (cache_dir / "extraTex" / "body").mkdir(parents=True)
    (cache_dir / "main.md").write_text("legacy markdown", encoding="utf-8")
    (cache_dir / "extraTex" / "body" / "intro.tex").write_text("legacy tex", encoding="utf-8")

    manuscript_tool.remove_legacy_docx_intermediates(cache_dir)

    assert not (cache_dir / "main.md").exists()
    assert not (cache_dir / "extraTex").exists()


def test_fix_docx_spacing_keeps_title_centered_and_left_aligns_section_headings(tmp_path):
    docx_path = tmp_path / "heading-alignment.docx"
    doc = Document()
    doc.styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Template Example Title", style="Heading 1")
    doc.add_paragraph("Template author list", style="Normal")
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("Body paragraph.", style="Normal")
    doc.save(docx_path)

    manuscript_tool.fix_docx_spacing(docx_path)

    fixed_doc = Document(docx_path)
    title_para = fixed_doc.paragraphs[0]
    section_para = fixed_doc.paragraphs[2]

    assert title_para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert section_para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.LEFT


def test_fix_docx_spacing_normalizes_frontmatter_title_and_author_fonts(tmp_path):
    docx_path = tmp_path / "frontmatter-style.docx"
    doc = Document()

    title_para = doc.add_paragraph("Template Example Title", style="Heading 1")
    title_para.runs[0].font.size = Pt(30)
    title_para.runs[0].font.bold = False

    doc.add_paragraph(manuscript_tool.DOCX_FRONTMATTER_CENTER_START, style="Normal")
    author_para = doc.add_paragraph(
        "Template Author One1†, Template Author Two2*",
        style="Normal",
    )
    author_para.runs[0].font.size = Pt(8)
    author_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph(manuscript_tool.DOCX_FRONTMATTER_CENTER_END, style="Normal")

    doc.add_paragraph("Introduction", style="Heading 1")
    doc.save(docx_path)

    manuscript_tool.fix_docx_spacing(docx_path)

    fixed_doc = Document(docx_path)
    fixed_title = fixed_doc.paragraphs[0]
    fixed_author = next(
        para for para in fixed_doc.paragraphs if "Template Author One" in para.text
    )

    assert fixed_title.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert fixed_title.runs[0].font.bold is True
    assert fixed_title.runs[0].font.size.pt == pytest.approx(15.0, abs=0.01)
    assert fixed_author.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert fixed_author.runs[0].font.size.pt == pytest.approx(12.0, abs=0.01)


def test_fix_docx_spacing_reorders_references_before_figure_titles_and_legends(tmp_path):
    docx_path = tmp_path / "references-order.docx"
    doc = Document()
    doc.styles.add_style("Bibliography", WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("Body paragraph.", style="Normal")
    doc.add_paragraph("Figure titles and legends", style="Heading 1")
    doc.add_paragraph("Figure legend paragraph.", style="Normal")
    doc.add_paragraph("[1]\tReference one.", style="Bibliography")
    doc.add_paragraph("[2]\tReference two.", style="Bibliography")
    doc.save(docx_path)

    manuscript_tool.fix_docx_spacing(docx_path)

    fixed_doc = Document(docx_path)
    paragraphs = [para for para in fixed_doc.paragraphs if para.text.strip()]
    paragraph_texts = [para.text.strip() for para in paragraphs]
    first_bibliography_index = next(
        index for index, para in enumerate(paragraphs) if para.style.name == "Bibliography"
    )
    last_bibliography_index = max(
        index for index, para in enumerate(paragraphs) if para.style.name == "Bibliography"
    )
    references_heading = next(para for para in paragraphs if para.text.strip() == "References")

    assert paragraph_texts.index("References") < paragraph_texts.index("Figure titles and legends")
    assert paragraph_texts.index("References") < first_bibliography_index
    assert last_bibliography_index < paragraph_texts.index("Figure titles and legends")
    assert references_heading.style.name == "Heading 1"


def test_fix_docx_spacing_matches_pdf_like_bibliography_labels_and_hanging_indent(tmp_path):
    docx_path = tmp_path / "references-layout.docx"
    doc = Document()
    doc.styles.add_style("Bibliography", WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("Body paragraph.", style="Normal")
    first_ref = doc.add_paragraph("1. Reference one.", style="Bibliography")
    second_ref = doc.add_paragraph("2. Reference two.", style="Bibliography")
    first_ref.paragraph_format.left_indent = Pt(0)
    first_ref.paragraph_format.first_line_indent = Pt(0)
    second_ref.paragraph_format.left_indent = Pt(0)
    second_ref.paragraph_format.first_line_indent = Pt(0)
    doc.save(docx_path)

    manuscript_tool.fix_docx_spacing(docx_path)

    fixed_doc = Document(docx_path)
    bibliography_paragraphs = [
        para for para in fixed_doc.paragraphs if para.style.name == "Bibliography"
    ]

    assert bibliography_paragraphs[0].text == "[1] Reference one."
    assert bibliography_paragraphs[1].text == "[2] Reference two."
    assert bibliography_paragraphs[0].paragraph_format.left_indent.pt == pytest.approx(
        19.44, abs=0.02
    )
    assert bibliography_paragraphs[0].paragraph_format.first_line_indent.pt == pytest.approx(
        -19.44, abs=0.02
    )


def test_fix_docx_spacing_promotes_existing_references_heading_to_heading_1(tmp_path):
    docx_path = tmp_path / "references-heading-level.docx"
    doc = Document()
    doc.styles.add_style("Bibliography", WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("Body paragraph.", style="Normal")
    doc.add_paragraph("References", style="Heading 2")
    doc.add_paragraph("1. Reference one.", style="Bibliography")
    doc.save(docx_path)

    manuscript_tool.fix_docx_spacing(docx_path)

    fixed_doc = Document(docx_path)
    references_heading = next(para for para in fixed_doc.paragraphs if para.text.strip() == "References")

    assert references_heading.style.name == "Heading 1"


def test_fix_docx_spacing_keeps_alias_tail_sections_flush_left(tmp_path):
    docx_path = tmp_path / "tail-alias-indent.docx"
    doc = Document()
    doc.add_paragraph("Title", style="Heading 1")
    doc.add_paragraph("Front matter", style="Normal")
    doc.add_paragraph("Figure titles and legends", style="Heading 1")
    doc.add_paragraph("Figure legend paragraph.", style="Normal")
    doc.add_paragraph("Supplemental information titles and legends", style="Heading 1")
    doc.add_paragraph("Supplement paragraph.", style="Normal")
    doc.save(docx_path)

    manuscript_tool.fix_docx_spacing(docx_path)

    fixed_doc = Document(docx_path)
    figure_para = next(para for para in fixed_doc.paragraphs if para.text == "Figure legend paragraph.")
    supplement_para = next(para for para in fixed_doc.paragraphs if para.text == "Supplement paragraph.")

    assert figure_para.paragraph_format.first_line_indent.pt == 0
    assert supplement_para.paragraph_format.first_line_indent.pt == 0


def test_paper_sci_01_csl_keeps_three_authors_before_et_al(tmp_path):
    bib_path = tmp_path / "refs.bib"
    bib_path.write_text(
        "\n".join(
            [
                "@article{Demo2026,",
                (
                    "  author = {Alpha, Alice and Beta, Bob and Gamma, Carol and "
                    "Delta, David and Epsilon, Erin and Zeta, Frank},"
                ),
                "  title = {Synthetic Reference for CSL Regression},",
                "  journal = {Journal of Regression Tests},",
                "  year = {2026},",
                "  volume = {12},",
                "  pages = {34--56},",
                "}",
                "",
            ]
        ),
        encoding="utf-8",
    )
    docx_path = tmp_path / "bibliography.docx"
    csl_path = REPO_ROOT / "projects" / "paper-sci-01" / "artifacts" / "manuscript.csl"
    reference_doc = REPO_ROOT / "projects" / "paper-sci-01" / "artifacts" / "reference.docx"

    manuscript_tool.run_cmd(
        [
            manuscript_tool.resolve_executable("pandoc"),
            "-",
            "-f",
            "markdown",
            "--citeproc",
            "--csl",
            str(csl_path),
            "--bibliography",
            str(bib_path),
            "--reference-doc",
            str(reference_doc),
            "-o",
            str(docx_path),
        ],
        input_text="Body text with a citation [@Demo2026].\n",
    )

    doc = Document(docx_path)
    bibliography_para = next(para for para in doc.paragraphs if "Synthetic reference for CSL regression" in para.text)

    assert bibliography_para.text.startswith("1. ")
    assert "Alpha A et al" not in bibliography_para.text
    assert "Alpha A, Beta B, Gamma C, et al." in bibliography_para.text
    assert "et al." in bibliography_para.text
    assert "Delta D" not in bibliography_para.text


def test_paper_sci_01_csl_prints_single_doi_without_duplicate_url(tmp_path):
    bib_path = tmp_path / "refs.bib"
    bib_path.write_text(
        "\n".join(
            [
                "@article{Demo2026,",
                "  author = {Alpha, Alice and Beta, Bob and Gamma, Carol},",
                "  title = {Synthetic DOI Regression},",
                "  journal = {Journal of Regression Tests},",
                "  year = {2026},",
                "  doi = {10.1234/demo.2026.001},",
                "  url = {https://doi.org/10.1234/demo.2026.001},",
                "}",
                "",
            ]
        ),
        encoding="utf-8",
    )
    docx_path = tmp_path / "bibliography.docx"
    csl_path = REPO_ROOT / "projects" / "paper-sci-01" / "artifacts" / "manuscript.csl"
    reference_doc = REPO_ROOT / "projects" / "paper-sci-01" / "artifacts" / "reference.docx"

    manuscript_tool.run_cmd(
        [
            manuscript_tool.resolve_executable("pandoc"),
            "-",
            "-f",
            "markdown",
            "--citeproc",
            "--csl",
            str(csl_path),
            "--bibliography",
            str(bib_path),
            "--reference-doc",
            str(reference_doc),
            "-o",
            str(docx_path),
        ],
        input_text="Body text with a citation [@Demo2026].\n",
    )

    bibliography_para = next(
        para
        for para in Document(docx_path).paragraphs
        if "synthetic doi regression" in para.text.lower()
    )

    assert "doi:10.1234/demo.2026.001" in bibliography_para.text.replace(" ", "").lower()
    assert "https://doi.org/10.1234/demo.2026.001" not in bibliography_para.text.lower()


def test_paper_sci_01_csl_keeps_reference_number_and_text_in_single_paragraph(tmp_path):
    bib_path = tmp_path / "refs.bib"
    bib_path.write_text(
        "\n".join(
            [
                "@article{Demo2026,",
                (
                    "  author = {Alpha, Alice and Beta, Bob and Gamma, Carol and "
                    "Delta, David},"
                ),
                "  title = {Synthetic Reference Layout Regression},",
                "  journal = {Journal of Regression Tests},",
                "  year = {2026},",
                "  volume = {12},",
                "  pages = {34--56},",
                "  doi = {10.1234/demo.2026.123},",
                "}",
                "",
            ]
        ),
        encoding="utf-8",
    )
    docx_path = tmp_path / "bibliography-layout.docx"
    csl_path = REPO_ROOT / "projects" / "paper-sci-01" / "artifacts" / "manuscript.csl"
    reference_doc = REPO_ROOT / "projects" / "paper-sci-01" / "artifacts" / "reference.docx"

    manuscript_tool.build_docx_from_markdown(
        manuscript_md="Body text with a citation [@Demo2026].\n",
        docx_path=docx_path,
        csl_path=csl_path,
        bibliography_path=bib_path,
        reference_doc=reference_doc,
    )

    bibliography_paragraphs = [
        para for para in Document(docx_path).paragraphs if para.style.name == "Bibliography" and para.text.strip()
    ]

    assert len(bibliography_paragraphs) == 1
    assert bibliography_paragraphs[0].text.startswith("1. ")
    assert "Synthetic reference layout regression" in bibliography_paragraphs[0].text


def test_paper_sci_01_final_docx_bibliography_text_tracks_pdf_style(tmp_path):
    bib_path = tmp_path / "refs.bib"
    bib_path.write_text(
        "\n".join(
            [
                "@article{Demo2026,",
                (
                    "  author = {Alpha, Alice and Beta, Bob and Gamma, Carol and "
                    "Delta, David},"
                ),
                "  title = {Synthetic Reference Layout Regression},",
                "  journal = {Journal of Regression Tests},",
                "  year = {2026},",
                "  volume = {12},",
                "  number = {3},",
                "  pages = {34--56},",
                "  doi = {10.1234/demo.2026.123},",
                "}",
                "",
            ]
        ),
        encoding="utf-8",
    )
    docx_path = tmp_path / "bibliography-final.docx"
    csl_path = REPO_ROOT / "projects" / "paper-sci-01" / "artifacts" / "manuscript.csl"
    reference_doc = REPO_ROOT / "projects" / "paper-sci-01" / "artifacts" / "reference.docx"

    manuscript_tool.build_docx_from_markdown(
        manuscript_md="Body text with a citation [@Demo2026].\n",
        docx_path=docx_path,
        csl_path=csl_path,
        bibliography_path=bib_path,
        reference_doc=reference_doc,
    )
    project_dir = tmp_path / "paper-sci-01"
    project_dir.mkdir()
    (project_dir / "main.tex").write_text(
        "\n".join(
            [
                r"\documentclass{article}",
                r"\usepackage{bensz-paper}",
                r"\BenszPaperSetup{template = paper-sci-01, bibliography-style = vancouver}",
                r"\begin{document}",
                r"\end{document}",
                "",
            ]
        ),
        encoding="utf-8",
    )
    manuscript_tool.fix_docx_spacing(docx_path, project_dir=project_dir)

    bibliography_para = next(
        para for para in Document(docx_path).paragraphs if para.style.name == "Bibliography"
    )
    normalized = " ".join(bibliography_para.text.split())

    assert normalized.startswith("1. ")
    assert "[J]." not in normalized
    assert re.search(
        r"Journal of Regression Tests\. 2026(?: [A-Za-z]+)?;12\(3\):34[-–]56\.",
        normalized,
    )
    assert "doi: 10.1234/demo.2026.123" in normalized
    assert "DOI:" not in normalized


def test_paper_sci_01_final_docx_bibliography_drops_doi_hyperlink_and_inline_emphasis(tmp_path):
    bib_path = tmp_path / "refs.bib"
    bib_path.write_text(
        "\n".join(
            [
                "@article{Demo2026,",
                "  author = {Alpha, Alice and Beta, Bob and Gamma, Carol},",
                "  title = {Synthetic DOI Regression},",
                "  journal = {Journal of Regression Tests},",
                "  year = {2026},",
                "  volume = {12},",
                "  number = {3},",
                "  pages = {34--56},",
                "  doi = {10.1234/demo.2026.001},",
                "}",
                "",
            ]
        ),
        encoding="utf-8",
    )
    docx_path = tmp_path / "bibliography-no-hyperlink.docx"
    csl_path = REPO_ROOT / "projects" / "paper-sci-01" / "artifacts" / "manuscript.csl"
    reference_doc = REPO_ROOT / "projects" / "paper-sci-01" / "artifacts" / "reference.docx"

    manuscript_tool.build_docx_from_markdown(
        manuscript_md="Body text with a citation [@Demo2026].\n",
        docx_path=docx_path,
        csl_path=csl_path,
        bibliography_path=bib_path,
        reference_doc=reference_doc,
    )
    project_dir = tmp_path / "paper-sci-01"
    project_dir.mkdir()
    (project_dir / "main.tex").write_text(
        "\n".join(
            [
                r"\documentclass{article}",
                r"\usepackage{bensz-paper}",
                r"\BenszPaperSetup{template = paper-sci-01, bibliography-style = vancouver}",
                r"\begin{document}",
                r"\end{document}",
                "",
            ]
        ),
        encoding="utf-8",
    )
    manuscript_tool.fix_docx_spacing(docx_path, project_dir=project_dir)

    bibliography_para = next(
        para for para in Document(docx_path).paragraphs if para.style.name == "Bibliography"
    )
    with ZipFile(docx_path) as archive:
        document_rels = archive.read("word/_rels/document.xml.rels").decode("utf-8", errors="replace")

    assert all(run.font.italic in (None, False) for run in bibliography_para.runs)
    assert all(run.font.bold in (None, False) for run in bibliography_para.runs)
    assert "https://doi.org/10.1234/demo.2026.001" in document_rels


def test_fix_docx_spacing_vancouver_adds_blue_hyperlinked_in_text_citations(tmp_path):
    bib_path = tmp_path / "refs.bib"
    bib_path.write_text(
        "\n".join(
            [
                "@article{Demo2026,",
                "  author = {Alpha, Alice and Beta, Bob and Gamma, Carol},",
                "  title = {Synthetic Citation Hyperlink Regression},",
                "  journal = {Journal of Regression Tests},",
                "  year = {2026},",
                "  volume = {12},",
                "  number = {3},",
                "  pages = {34--56},",
                "  doi = {10.1234/demo.2026.001},",
                "}",
                "",
            ]
        ),
        encoding="utf-8",
    )
    docx_path = tmp_path / "citation-hyperlinks.docx"
    csl_path = REPO_ROOT / "projects" / "paper-sci-01" / "artifacts" / "manuscript.csl"
    reference_doc = REPO_ROOT / "projects" / "paper-sci-01" / "artifacts" / "reference.docx"

    manuscript_tool.build_docx_from_markdown(
        manuscript_md="Body text with grouped citations [@Demo2026; @Demo2026].\n",
        docx_path=docx_path,
        csl_path=csl_path,
        bibliography_path=bib_path,
        reference_doc=reference_doc,
    )
    project_dir = tmp_path / "paper-sci-01"
    project_dir.mkdir()
    (project_dir / "main.tex").write_text(
        "\n".join(
            [
                r"\documentclass{article}",
                r"\usepackage{bensz-paper}",
                r"\BenszPaperSetup{template = paper-sci-01, bibliography-style = vancouver}",
                r"\begin{document}",
                r"\end{document}",
                "",
            ]
        ),
        encoding="utf-8",
    )
    manuscript_tool.fix_docx_spacing(docx_path, project_dir=project_dir)

    document_xml = _read_document_xml(docx_path)

    assert 'w:bookmarkStart' in document_xml
    assert 'w:name="ref-1"' in document_xml
    assert 'w:hyperlink w:anchor="ref-1"' in document_xml
    assert '<w:color w:val="0000FF"/>' in document_xml
    assert '<w:vertAlign w:val="superscript"/>' in document_xml


def test_fix_docx_spacing_adds_visible_horizontal_borders_to_normal_tables(tmp_path):
    docx_path = tmp_path / "table.docx"
    csl_path = REPO_ROOT / "projects" / "paper-sci-01" / "artifacts" / "manuscript.csl"
    reference_doc = REPO_ROOT / "projects" / "paper-sci-01" / "artifacts" / "reference.docx"

    manuscript_tool.run_cmd(
        [
            manuscript_tool.resolve_executable("pandoc"),
            "-",
            "-f",
            "markdown",
            "--reference-doc",
            str(reference_doc),
            "--csl",
            str(csl_path),
            "-o",
            str(docx_path),
        ],
        input_text="| A | B |\n| --- | --- |\n| 1 | 2 |\n",
    )

    manuscript_tool.fix_docx_spacing(docx_path)

    xml = _read_document_xml(docx_path)
    assert "<w:tblBorders>" in xml
    assert '<w:top w:val="single"' in xml
    assert '<w:bottom w:val="single"' in xml
    assert '<w:insideH w:val="single"' in xml
    assert '<w:left w:val="nil"' in xml
    assert '<w:right w:val="nil"' in xml
    assert '<w:insideV w:val="nil"' in xml


def test_fix_docx_spacing_preserves_existing_table_border_overrides(tmp_path):
    docx_path = tmp_path / "custom-table.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"

    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge, value in (
        ("top", "double"),
        ("bottom", "double"),
        ("left", "single"),
        ("right", "single"),
        ("insideH", "double"),
        ("insideV", "single"),
    ):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), value)
        borders.append(element)
    tbl_pr.append(borders)
    doc.save(docx_path)

    manuscript_tool.fix_docx_spacing(docx_path)

    xml = _read_document_xml(docx_path)
    assert '<w:top w:val="double"' in xml
    assert '<w:bottom w:val="double"' in xml
    assert '<w:insideH w:val="double"' in xml
    assert '<w:insideV w:val="single"' in xml
    assert '<w:left w:val="single"' in xml
    assert '<w:right w:val="single"' in xml


def test_bml_bibliography_prefers_doi_without_printing_urls():
    bibliography_style = (REPO_ROOT / "packages" / "bensz-paper" / "bml-bibliography.sty").read_text(
        encoding="utf-8"
    )

    assert "url=true" in bibliography_style
    assert "doi=true" in bibliography_style
    assert r"\iffieldundef{doi}" in bibliography_style
    assert r"\clearfield{url}\clearfield{urlyear}\clearfield{urlmonth}\clearfield{urlday}" in bibliography_style
    assert r"\ifdefstring{\bml@bibstyle}{vancouver}" in bibliography_style


def test_paper_sci_01_main_tex_uses_vancouver_reference_style():
    main_tex = (REPO_ROOT / "projects" / "paper-sci-01" / "main.tex").read_text(encoding="utf-8")

    assert "bibliography-style = vancouver" in main_tex
    assert r"\ExecuteBibliographyOptions{url=true, doi=true}" in main_tex
    assert r"\DeclareCiteCommand{\supercite}[\mkbibsuperscript]" in main_tex
    assert "colorlinks=true" in main_tex
    assert "citecolor=blue" in main_tex
    assert "urlcolor=blue" in main_tex
    assert r"\iffieldundef{doi}" in main_tex
    assert r"\section{References}" in main_tex
    assert r"\printbibliography[heading=none]" in main_tex


def test_bml_core_loads_profile_from_repo_package_root():
    bml_core = (REPO_ROOT / "packages" / "bensz-paper" / "bml-core.sty").read_text(encoding="utf-8")

    assert r"\RequirePackage{currfile}" in bml_core
    assert r"\edef\bml@packageroot{\currfiledir}" in bml_core
    assert r"\InputIfFileExists{\bml@packageroot" in bml_core
    assert r"profiles/bml-profile-\bml@template.def}{}{%" in bml_core


def test_paper_sci_01_profile_sets_shared_bibliography_size():
    profile_text = (
        REPO_ROOT / "packages" / "bensz-paper" / "profiles" / "bml-profile-paper-sci-01.def"
    ).read_text(encoding="utf-8")

    assert r"\def\bml@bibfontsize{10.5pt}" in profile_text
    assert r"\def\bml@bibledaing{15pt}" in profile_text


def test_paper_sci_01_bib_data_drops_doi_mirror_urls():
    refs_bib = (REPO_ROOT / "projects" / "paper-sci-01" / "references" / "refs.bib").read_text(
        encoding="utf-8"
    )

    assert "https://doi.org/" not in refs_bib
