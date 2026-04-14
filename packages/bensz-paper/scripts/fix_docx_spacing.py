#!/usr/bin/env python3
r"""DOCX 段间距修复工具（bensz-paper 公共包配套脚本）。

修复 DOCX 文件的行间距、段间距和首行缩进，使其尽量接近 LaTeX PDF 输出版式。
在 build_project() 的 DOCX 生成管线末尾自动调用。

LaTeX 参数（权威来源：texmf/.../profiles/bml-profile-bensz-manu-01.def）：
- 行距：\setstretch{1.5}  →  ONE_POINT_FIVE
- 段首缩进：\parindent = 1.5em（12pt × 1.5 = 18pt）
  - heading 后第一段：Pt(0)（titlespacing* 抑制）
  - Abstract 节所有段：Pt(0)（摘要环境无缩进）
  - 其余 body 段落（第二段起）：Pt(18)
- 段间距：\parskip = 0.2\baselineskip
    ≈ 0.2 × (12pt × 1.5) = 0.2 × 18pt = 3.6pt  →  近似取 Pt(4)

注意：DOCX 段间距（4pt）是对 LaTeX parskip（3.6pt）的近似，
两者在视觉上接近但不完全等价。详见 docs/package/BenszManuscriptLaTeX.md
"PDF 与 DOCX 对齐策略"一节。

用法：
  python fix_docx_spacing.py --project-dir <项目路径>
  python fix_docx_spacing.py --docx <DOCX文件路径> --no-backup
"""

from __future__ import annotations

import argparse
import re
import shutil
import sys
from dataclasses import dataclass
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

# 用于判定目录是否为项目根的标记文件
PROJECT_ROOT_MARKERS = ("main.tex",)


# DOCX spacing constants — keep in sync with bml-profile-bensz-manu-01.def
# LaTeX: \setstretch{1.5}
DOCX_LINE_SPACING_RULE = WD_LINE_SPACING.ONE_POINT_FIVE
# LaTeX: \parindent = 1.5em → 12pt × 1.5 = 18pt (body paragraphs, 2nd onwards)
DOCX_BODY_INDENT = Pt(18)
DOCX_NO_INDENT   = Pt(0)    # heading后第一段 & Abstract节所有段
# LaTeX: \parskip = 0.2\baselineskip ≈ 3.6pt → rounded to 4pt for Word
DOCX_SPACE_AFTER = Pt(4)
DOCX_SPACE_BEFORE = Pt(0)
# Abstract section heading text — used to suppress indentation for the whole section
ABSTRACT_HEADING = "Abstract"
# References section heading text — inserted before Bibliography paragraphs if absent
REFERENCES_HEADING = "References"
REFERENCES_HEADING_STYLE = "Heading 1"
FIGURE_LEGENDS_HEADINGS = ("Figure legends", "Figure titles and legends")
SUPPLEMENTARY_HEADINGS = ("Supplementary materials", "Supplemental information titles and legends")
DOCX_FRONTMATTER_CENTER_START = "BENSZ_DOCX_FRONTMATTER_CENTER_START"
DOCX_FRONTMATTER_CENTER_END = "BENSZ_DOCX_FRONTMATTER_CENTER_END"
FRONTMATTER_TITLE_FONT_SIZE_PT = 15.0
FRONTMATTER_TITLE_LEADING_PT = 20.0
FRONTMATTER_TITLE_SPACE_AFTER_PT = 21.6
FRONTMATTER_AUTHOR_SPACE_AFTER_PT = 0.0
DEFAULT_BODY_FONT_SIZE_PT = 12.0
DEFAULT_BODY_LEADING_PT = 18.0
PROFILE_TEMPLATE_PATTERN = re.compile(r"template\s*=\s*([A-Za-z0-9._-]+)")
PROFILE_BIBLIOGRAPHY_STYLE_PATTERN = re.compile(r"bibliography-style\s*=\s*([A-Za-z0-9._-]+)")
PROFILE_DEF_PATTERN = re.compile(r"\\def\\bml@([A-Za-z@]+)\{([^}]*)\}")
BIBLIOGRAPHY_LABEL_PATTERN = re.compile(r"^\s*(\d+)\.\s*")
DOCX_LINK_BLUE = "0000FF"
DOCX_BOOKMARK_PREFIX = "ref-"
DOCX_CITATION_PATTERN = re.compile(
    r"^\[(\d+(?:\s*[-\u2013]\s*\d+)?(?:,\s*\d+(?:\s*[-\u2013]\s*\d+)?)*)\]$"
)


def _normalize_heading_text(value: str) -> str:
    """将标题文本标准化：合并空白、去前后空白、转小写。用于标题匹配比较。"""
    return " ".join(value.split()).strip().lower()


# Section headings whose body paragraphs should all be flush left (no first-line indent)
NO_INDENT_SECTIONS = {
    _normalize_heading_text(heading)
    for heading in (ABSTRACT_HEADING, *FIGURE_LEGENDS_HEADINGS, *SUPPLEMENTARY_HEADINGS)
}
# 使用默认样式的 DOCX 表格（无自定义边框），需要补充水平线
DEFAULT_DOCX_TABLE_STYLES = {"Normal Table"}
# 默认表格边框配置：上下和水平内线为细实线，左右和垂直内线为无边框
DEFAULT_DOCX_TABLE_BORDERS = {
    "top": {"val": "single", "sz": "8", "space": "0", "color": "auto"},
    "bottom": {"val": "single", "sz": "8", "space": "0", "color": "auto"},
    "insideH": {"val": "single", "sz": "8", "space": "0", "color": "auto"},
    "left": {"val": "nil"},
    "right": {"val": "nil"},
    "insideV": {"val": "nil"},
}


@dataclass(frozen=True)
class DocxStyleProfile:
    main_font: str = "Times New Roman"
    bibliography_style: str = "gb7714-2015"
    body_font_size_pt: float = DEFAULT_BODY_FONT_SIZE_PT
    body_leading_pt: float = DEFAULT_BODY_LEADING_PT
    section_font_size_pt: float = 14.0
    section_space_before_pt: float = 0.0
    section_space_after_pt: float = 3.0
    bibliography_font_size_pt: float = 7.875
    bibliography_leading_pt: float = 9.5
    bibliography_hang_pt: float = 19.44
    title_font_size_pt: float = FRONTMATTER_TITLE_FONT_SIZE_PT
    title_leading_pt: float = FRONTMATTER_TITLE_LEADING_PT
    title_space_after_pt: float = FRONTMATTER_TITLE_SPACE_AFTER_PT
    author_font_size_pt: float = DEFAULT_BODY_FONT_SIZE_PT
    author_space_after_pt: float = FRONTMATTER_AUTHOR_SPACE_AFTER_PT


def _latex_dimension_to_points(raw_value: str, *, em_base_pt: float = DEFAULT_BODY_FONT_SIZE_PT) -> float:
    normalized = raw_value.strip()
    if not normalized:
        raise ValueError("empty latex dimension")

    if normalized.endswith("pt"):
        return float(normalized[:-2])
    if normalized.endswith("in"):
        return float(normalized[:-2]) * 72.0
    if normalized.endswith("em"):
        return float(normalized[:-2]) * em_base_pt
    return float(normalized)


def _read_profile_definitions(profile_path: Path) -> dict[str, str]:
    if not profile_path.exists():
        return {}
    text = profile_path.read_text(encoding="utf-8")
    return {
        match.group(1): match.group(2).strip()
        for match in PROFILE_DEF_PATTERN.finditer(text)
    }


def _resolve_template_name(project_dir: Path | None) -> str | None:
    if project_dir is None:
        return None
    main_tex = project_dir / "main.tex"
    if not main_tex.exists():
        return None
    match = PROFILE_TEMPLATE_PATTERN.search(main_tex.read_text(encoding="utf-8"))
    if match is None:
        return None
    return match.group(1)


def _resolve_bibliography_style(project_dir: Path | None) -> str | None:
    if project_dir is None:
        return None
    main_tex = project_dir / "main.tex"
    if not main_tex.exists():
        return None
    match = PROFILE_BIBLIOGRAPHY_STYLE_PATTERN.search(main_tex.read_text(encoding="utf-8"))
    if match is None:
        return None
    return match.group(1)


def _load_docx_style_profile(project_dir: Path | None) -> DocxStyleProfile:
    package_root = Path(__file__).resolve().parents[1]
    profiles_dir = package_root / "profiles"
    template_name = _resolve_template_name(project_dir)

    base_defs = _read_profile_definitions(profiles_dir / "bml-profile-bensz-manu-01.def")
    template_defs = (
        _read_profile_definitions(profiles_dir / f"bml-profile-{template_name}.def")
        if template_name
        else {}
    )

    defs = {**base_defs, **template_defs}
    if not defs:
        return DocxStyleProfile()

    bibliography_style = _resolve_bibliography_style(project_dir) or "gb7714-2015"
    main_font = defs.get("mainfont", DocxStyleProfile.main_font)
    section_font_size_pt = _latex_dimension_to_points(
        defs.get("secfontsize", str(DocxStyleProfile.section_font_size_pt))
    )
    section_space_before_pt = _latex_dimension_to_points(
        defs.get("secspacebefore", f"{DocxStyleProfile.section_space_before_pt}pt")
    )
    section_space_after_pt = _latex_dimension_to_points(
        defs.get("secspaceafter", f"{DocxStyleProfile.section_space_after_pt}pt")
    )
    bibliography_font_size_pt = _latex_dimension_to_points(
        defs.get("bibfontsize", str(DocxStyleProfile.bibliography_font_size_pt))
    )
    bibliography_leading_pt = _latex_dimension_to_points(
        defs.get("bibledaing", str(DocxStyleProfile.bibliography_leading_pt))
    )
    bibliography_hang_pt = _latex_dimension_to_points(
        defs.get("bibhang", f"{DocxStyleProfile.bibliography_hang_pt}pt")
    )

    return DocxStyleProfile(
        main_font=main_font,
        bibliography_style=bibliography_style,
        section_font_size_pt=section_font_size_pt,
        section_space_before_pt=section_space_before_pt,
        section_space_after_pt=section_space_after_pt,
        bibliography_font_size_pt=bibliography_font_size_pt,
        bibliography_leading_pt=bibliography_leading_pt,
        bibliography_hang_pt=bibliography_hang_pt,
    )


def configure_windows_stdio_utf8() -> None:
    if sys.platform != "win32":
        return
    for stream_name in ("stdout", "stderr"):
        stream = getattr(sys, stream_name, None)
        if hasattr(stream, "reconfigure"):
            stream.reconfigure(encoding="utf-8", errors="replace")


def is_project_root(path: Path) -> bool:
    """Return True when the directory looks like a manuscript project root."""
    return any((path / marker).exists() for marker in PROJECT_ROOT_MARKERS)


def find_project_root(start: Path | None = None) -> Path:
    """Walk upward until a manuscript project root is found."""
    origin = (start or Path.cwd()).expanduser().resolve()
    for candidate in [origin, *origin.parents]:
        if is_project_root(candidate):
            return candidate
    raise FileNotFoundError(
        "Cannot find project root. Run inside the manuscript project tree or "
        "pass --project-dir explicitly."
    )


def resolve_project_dir(project_dir: Path | None) -> Path:
    """Resolve the manuscript project directory from CLI input or the current working tree."""
    if project_dir is None:
        return find_project_root()

    candidate = project_dir.expanduser().resolve()
    if not candidate.exists():
        raise FileNotFoundError(f"Project directory not found: {candidate}")
    if candidate.is_file():
        candidate = candidate.parent
    if is_project_root(candidate):
        return candidate
    return find_project_root(candidate)


def _resolve_profile_project_dir(docx_path: Path, project_dir: Path | None) -> Path | None:
    if project_dir is not None:
        return project_dir
    try:
        return find_project_root(docx_path.parent)
    except FileNotFoundError:
        return None


def _add_references_heading_if_missing(
    doc: "Document", heading_text: str = REFERENCES_HEADING
) -> None:
    """确保 References 在 DOCX 中始终是 Heading 1，并位于第一个 Bibliography 段落前。"""
    first_bib = next(
        (p for p in doc.paragraphs if p.style.name == "Bibliography"), None
    )
    if first_bib is None:
        return

    heading = _find_heading_paragraph(doc, (heading_text,))
    if heading is None:
        heading = doc.add_heading(heading_text, level=1)
        heading_elem = heading._element
        heading_elem.getparent().remove(heading_elem)
        first_bib._element.addprevious(heading_elem)
        print(f"✓ 已插入 '{heading_text}' 标题")
    else:
        body_elem = first_bib._element.getparent()
        children = list(body_elem)
        if children.index(heading._element) > children.index(first_bib._element):
            body_elem.remove(heading._element)
            first_bib._element.addprevious(heading._element)
            print(f"✓ 已将现有 '{heading_text}' 标题移动到参考文献前")

    if heading.style.name != REFERENCES_HEADING_STYLE:
        heading.style = doc.styles[REFERENCES_HEADING_STYLE]
        print(f"✓ 已将 '{heading_text}' 标题样式统一为 {REFERENCES_HEADING_STYLE}")


def _find_heading_paragraph(doc: "Document", headings: tuple[str, ...] | list[str] | set[str]):
    """在文档中查找第一个匹配指定标题文本的 Heading 段落，未找到返回 None。"""
    normalized_headings = {_normalize_heading_text(heading) for heading in headings}
    return next(
        (
            para
            for para in doc.paragraphs
            if para.style.name.startswith("Heading")
            and _normalize_heading_text(para.text) in normalized_headings
        ),
        None,
    )


def _reorder_references_before_figure_legends(
    doc: "Document",
    references_heading: str = REFERENCES_HEADING,
    figure_legends_headings: tuple[str, ...] = FIGURE_LEGENDS_HEADINGS,
) -> None:
    """将 References 节（标题 + 所有 Bibliography 段落）移动到 Figure legends 前。

    Pandoc citeproc 默认将参考文献追加至文档末尾，但 SCI 手稿惯例是
    References 位于正文之后、Figure legends 之前。
    """
    paras = doc.paragraphs

    # Locate the References heading
    ref_heading_para = _find_heading_paragraph(doc, (references_heading,))
    if ref_heading_para is None:
        return  # Nothing to move

    # Locate the Figure legends heading
    fig_legends_para = _find_heading_paragraph(doc, figure_legends_headings)
    if fig_legends_para is None:
        return  # No Figure legends section found

    # Collect ALL Bibliography paragraph elements via python-docx style
    # resolution (raw XML pStyle may use an opaque ID like "af4" instead of
    # the human-readable name "Bibliography").
    bib_elements = [p._element for p in paras if p.style.name == "Bibliography"]
    if not bib_elements:
        return

    body_elem = ref_heading_para._element.getparent()
    children = list(body_elem)
    fig_idx = children.index(fig_legends_para._element)

    # Only skip if heading AND every Bibliography entry are already before Figure legends
    ref_idx = children.index(ref_heading_para._element)
    all_before = ref_idx < fig_idx and all(
        children.index(e) < fig_idx for e in bib_elements
    )
    if all_before:
        return

    # Build the block: References heading + all Bibliography paragraphs
    ref_block = [ref_heading_para._element] + bib_elements

    for elem in ref_block:
        body_elem.remove(elem)

    anchor = fig_legends_para._element
    for elem in ref_block:
        anchor.addprevious(elem)

    print(f"✓ 已将 References 节移动到 '{fig_legends_para.text.strip()}' 前")


def _find_table_borders_element(table) -> OxmlElement | None:
    """查找表格的 w:tblBorders XML 元素，用于判断是否已有自定义边框。"""
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        return None

    for child in tbl_pr.iterchildren():
        if child.tag == qn("w:tblBorders"):
            return child
    return None


def _ensure_default_horizontal_table_borders(table) -> None:
    """为使用 Normal Table 样式且无边框的表格添加默认水平线边框。

    SCI 论文表格通常只需上下边框和水平分隔线，不需要垂直线和左右边框。
    仅当表格使用 Normal Table 样式且尚无自定义边框时才补充。
    """
    style_name = table.style.name if table.style else ""
    if style_name not in DEFAULT_DOCX_TABLE_STYLES:
        return
    if _find_table_borders_element(table) is not None:
        return

    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        return

    borders = OxmlElement("w:tblBorders")
    for edge, attrs in DEFAULT_DOCX_TABLE_BORDERS.items():
        element = OxmlElement(f"w:{edge}")
        for attr_name, value in attrs.items():
            element.set(qn(f"w:{attr_name}"), value)
        borders.append(element)
    tbl_pr.append(borders)


def _remove_paragraph(para) -> None:
    """从文档中移除指定段落。"""
    para._element.getparent().remove(para._element)


def _remove_xml_child(parent, child_tag: str) -> None:
    child = parent.find(qn(child_tag))
    if child is not None:
        parent.remove(child)


def _set_run_font(run, *, font_name: str, size_pt: float, bold: bool | None = None) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold


def _apply_title_paragraph_style(para, profile: DocxStyleProfile) -> None:
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = DOCX_NO_INDENT
    pf.space_before = DOCX_SPACE_BEFORE
    pf.space_after = Pt(profile.title_space_after_pt)
    pf.line_spacing = Pt(profile.title_leading_pt)
    for run in para.runs:
        _set_run_font(
            run,
            font_name=profile.main_font,
            size_pt=profile.title_font_size_pt,
            bold=True,
        )


def _apply_author_paragraph_style(para, profile: DocxStyleProfile) -> None:
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = DOCX_NO_INDENT
    pf.space_before = DOCX_SPACE_BEFORE
    pf.space_after = Pt(profile.author_space_after_pt)
    pf.line_spacing = Pt(profile.body_leading_pt)
    for run in para.runs:
        _set_run_font(
            run,
            font_name=profile.main_font,
            size_pt=profile.author_font_size_pt,
        )


def _apply_section_heading_style(para, profile: DocxStyleProfile) -> None:
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = DOCX_NO_INDENT
    pf.space_before = Pt(profile.section_space_before_pt)
    pf.space_after = Pt(profile.section_space_after_pt)
    pf.line_spacing = Pt(profile.body_leading_pt)
    for run in para.runs:
        _set_run_font(
            run,
            font_name=profile.main_font,
            size_pt=profile.section_font_size_pt,
            bold=True,
        )


def _normalize_bibliography_label(para) -> None:
    if not para.runs:
        return
    for run in para.runs:
        if "\t" in run.text:
            run.text = run.text.replace("\t", "")
        if not run.text.strip():
            continue
        match = BIBLIOGRAPHY_LABEL_PATTERN.match(run.text)
        if match is not None:
            run.text = BIBLIOGRAPHY_LABEL_PATTERN.sub(
                f"[{match.group(1)}] ",
                run.text,
                count=1,
            )
        break


def _unwrap_hyperlinks(para) -> None:
    paragraph_element = para._element
    for child in list(paragraph_element):
        if child.tag != qn("w:hyperlink"):
            continue
        for nested in list(child):
            child.addprevious(nested)
        paragraph_element.remove(child)


def _clear_bibliography_run_formatting(run) -> None:
    run.font.bold = False
    run.font.italic = False
    run.font.underline = False
    run.font.color.rgb = None
    run._element.attrib.pop(qn("w:rsidRPr"), None)

    run_properties = run._element.rPr
    if run_properties is None:
        return

    for tag in ("w:rStyle", "w:b", "w:bCs", "w:i", "w:iCs", "w:u", "w:color"):
        _remove_xml_child(run_properties, tag)


def _apply_bibliography_style(para, profile: DocxStyleProfile) -> None:
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.left_indent = Pt(profile.bibliography_hang_pt)
    pf.first_line_indent = Pt(-profile.bibliography_hang_pt)
    pf.space_before = DOCX_SPACE_BEFORE
    pf.space_after = DOCX_SPACE_AFTER
    pf.line_spacing = Pt(profile.bibliography_leading_pt)
    if profile.bibliography_style == "vancouver":
        for run in para.runs:
            _set_run_font(
                run,
                font_name=profile.main_font,
                size_pt=profile.bibliography_font_size_pt,
            )
        return

    _unwrap_hyperlinks(para)
    _normalize_bibliography_label(para)
    for run in para.runs:
        _clear_bibliography_run_formatting(run)
        _set_run_font(
            run,
            font_name=profile.main_font,
            size_pt=profile.bibliography_font_size_pt,
        )


def _next_bookmark_id(doc: "Document") -> int:
    bookmark_ids = [
        int(node.get(qn("w:id")))
        for node in doc.element.xpath("//w:bookmarkStart")
        if node.get(qn("w:id")) and node.get(qn("w:id")).isdigit()
    ]
    return max(bookmark_ids, default=0) + 1


def _remove_existing_bookmarks(paragraph) -> None:
    for node in list(paragraph._p.xpath("./w:bookmarkStart | ./w:bookmarkEnd")):
        paragraph._p.remove(node)


def _insert_paragraph_bookmark(paragraph, bookmark_id: int, bookmark_name: str) -> None:
    _remove_existing_bookmarks(paragraph)

    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), str(bookmark_id))
    bookmark_start.set(qn("w:name"), bookmark_name)

    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), str(bookmark_id))

    insert_at = 1 if len(paragraph._p) > 0 and paragraph._p[0].tag == qn("w:pPr") else 0
    paragraph._p.insert(insert_at, bookmark_start)
    paragraph._p.insert(insert_at + 1, bookmark_end)


def _ensure_bibliography_bookmarks(doc: "Document") -> dict[str, str]:
    reference_targets: dict[str, str] = {}
    bookmark_id = _next_bookmark_id(doc)

    for para in doc.paragraphs:
        if para.style.name != "Bibliography":
            continue

        match = BIBLIOGRAPHY_LABEL_PATTERN.match(para.text.strip())
        if match is None:
            continue

        reference_number = match.group(1)
        bookmark_name = f"{DOCX_BOOKMARK_PREFIX}{reference_number}"
        _insert_paragraph_bookmark(para, bookmark_id, bookmark_name)
        reference_targets[reference_number] = bookmark_name
        bookmark_id += 1

    return reference_targets


def _make_docx_run(
    text: str,
    *,
    superscript: bool,
    color: str,
    font_name: str,
    size_pt: float,
) -> OxmlElement:
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), font_name)
    fonts.set(qn("w:hAnsi"), font_name)
    run_properties.append(fonts)

    size = OxmlElement("w:sz")
    size.set(qn("w:val"), str(int(round(size_pt * 2))))
    run_properties.append(size)

    if superscript:
        vertical_align = OxmlElement("w:vertAlign")
        vertical_align.set(qn("w:val"), "superscript")
        run_properties.append(vertical_align)

    color_element = OxmlElement("w:color")
    color_element.set(qn("w:val"), color)
    run_properties.append(color_element)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    run_properties.append(underline)

    run.append(run_properties)

    text_element = OxmlElement("w:t")
    text_element.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_element.text = text
    run.append(text_element)
    return run


def _make_docx_anchor_hyperlink(
    text: str,
    anchor: str,
    *,
    profile: DocxStyleProfile,
) -> OxmlElement:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.append(
        _make_docx_run(
            text,
            superscript=True,
            color=DOCX_LINK_BLUE,
            font_name=profile.main_font,
            size_pt=profile.body_font_size_pt,
        )
    )
    return hyperlink


def _citation_anchor_for_token(token: str, reference_targets: dict[str, str]) -> str | None:
    token = token.strip()
    if not token:
        return None

    first_number = re.split(r"\s*[-\u2013]\s*", token, maxsplit=1)[0].strip()
    return reference_targets.get(first_number)


def _replace_citation_run(
    run_element,
    reference_targets: dict[str, str],
    profile: DocxStyleProfile,
) -> bool:
    run_text = "".join(run_element.xpath(".//w:t/text()"))
    match = DOCX_CITATION_PATTERN.fullmatch(run_text.strip())
    if match is None:
        return False

    cluster = match.group(1)
    parent = run_element.getparent()
    insert_at = parent.index(run_element)
    new_elements: list[OxmlElement] = []
    tokens = [token.strip() for token in cluster.split(",") if token.strip()]

    if len(tokens) == 1:
        anchor = _citation_anchor_for_token(tokens[0], reference_targets)
        if anchor is not None:
            new_elements.append(
                _make_docx_anchor_hyperlink(f"[{tokens[0]}]", anchor, profile=profile)
            )
        else:
            new_elements.append(
                _make_docx_run(
                    f"[{tokens[0]}]",
                    superscript=True,
                    color=DOCX_LINK_BLUE,
                    font_name=profile.main_font,
                    size_pt=profile.body_font_size_pt,
                )
            )
    else:
        new_elements.append(
            _make_docx_run(
                "[",
                superscript=True,
                color=DOCX_LINK_BLUE,
                font_name=profile.main_font,
                size_pt=profile.body_font_size_pt,
            )
        )
        for index, token in enumerate(tokens):
            anchor = _citation_anchor_for_token(token, reference_targets)
            if anchor is not None:
                new_elements.append(_make_docx_anchor_hyperlink(token, anchor, profile=profile))
            else:
                new_elements.append(
                    _make_docx_run(
                        token,
                        superscript=True,
                        color=DOCX_LINK_BLUE,
                        font_name=profile.main_font,
                        size_pt=profile.body_font_size_pt,
                    )
                )

            if index < len(tokens) - 1:
                new_elements.append(
                    _make_docx_run(
                        ",",
                        superscript=True,
                        color=DOCX_LINK_BLUE,
                        font_name=profile.main_font,
                        size_pt=profile.body_font_size_pt,
                    )
                )

        new_elements.append(
            _make_docx_run(
                "]",
                superscript=True,
                color=DOCX_LINK_BLUE,
                font_name=profile.main_font,
                size_pt=profile.body_font_size_pt,
            )
        )

    for offset, element in enumerate(new_elements):
        parent.insert(insert_at + offset, element)
    parent.remove(run_element)
    return True


def _style_docx_in_text_citations(
    doc: "Document",
    reference_targets: dict[str, str],
    profile: DocxStyleProfile,
) -> bool:
    changed = False

    for para in doc.paragraphs:
        if para.style.name == "Bibliography":
            continue

        paragraph = para._p
        for child in list(paragraph):
            if child.tag != qn("w:r"):
                continue
            changed = _replace_citation_run(child, reference_targets, profile) or changed

    return changed


def _center_frontmatter_author_blocks(doc: "Document", profile: DocxStyleProfile) -> None:
    """将 frontmatter 中显式标记的作者块设为居中，并删除中间标记段落。"""
    search_from = 0

    while True:
        paragraphs = list(doc.paragraphs)
        start_index = next(
            (
                index
                for index in range(search_from, len(paragraphs))
                if paragraphs[index].text.strip() == DOCX_FRONTMATTER_CENTER_START
            ),
            None,
        )
        if start_index is None:
            return

        end_index = next(
            (
                index
                for index in range(start_index + 1, len(paragraphs))
                if paragraphs[index].text.strip() == DOCX_FRONTMATTER_CENTER_END
            ),
            None,
        )
        if end_index is None:
            return

        for para in paragraphs[start_index + 1 : end_index]:
            if not para.text.strip():
                continue
            _apply_author_paragraph_style(para, profile)

        _remove_paragraph(paragraphs[end_index])
        _remove_paragraph(paragraphs[start_index])
        search_from = start_index


def fix_docx_spacing(docx_path: Path, project_dir: Path | None = None) -> None:
    """修复 DOCX 文件的行间距、段间距和首行缩进。

    主要修复逻辑：
    1. 确保参考文献前有 References 标题（Heading 1）。
    2. 将 References 节移动到 Figure legends 之前（SCI 手稿惯例）。
    3. 统一所有段落的行距（1.5 倍）和段后间距（4pt）。
    4. 按规则设置首行缩进：
       - Heading 段落和其后首段不缩进
       - Abstract/Figure legends/Supplementary 节内所有段落不缩进
       - 正文其余段落首行缩进 18pt
    5. Bibliography 段落改为与 PDF 一致的悬挂缩进、方括号编号和独立字号。
    6. 为 Normal Table 样式的表格补充水平边框。
    """
    print(f"正在修复: {docx_path}")

    doc = Document(docx_path)
    profile = _load_docx_style_profile(_resolve_profile_project_dir(docx_path, project_dir))

    # Insert "References" heading before bibliography if absent
    _add_references_heading_if_missing(doc)

    # Move References section before Figure legends (matches PDF layout)
    _reorder_references_before_figure_legends(doc)
    _center_frontmatter_author_blocks(doc, profile)

    total_paragraphs = 0
    fixed_paragraphs = 0

    in_no_indent_section = False  # True for Abstract, Figure legends, Supplementary materials
    prev_was_heading     = True   # 文档起始视为 heading 后，首段不缩进
    seen_section_heading = False  # 见到第一个 H2+ 之前（frontmatter）不缩进
    seen_title_heading   = False  # 首个 Heading 1 是论文标题，保持居中

    for para in doc.paragraphs:
        total_paragraphs += 1
        style_name = para.style.name if para.style else ""
        is_heading = style_name.startswith("Heading")
        is_bibliography = style_name == "Bibliography"

        pf = para.paragraph_format
        pf.line_spacing_rule = DOCX_LINE_SPACING_RULE
        pf.space_after  = DOCX_SPACE_AFTER
        pf.space_before = DOCX_SPACE_BEFORE

        if is_heading:
            if style_name == "Heading 1" and not seen_title_heading:
                _apply_title_paragraph_style(para, profile)
                seen_title_heading = True
            else:
                _apply_section_heading_style(para, profile)
            if style_name != "Heading 1":
                seen_section_heading = True
            in_no_indent_section = _normalize_heading_text(para.text) in NO_INDENT_SECTIONS
            prev_was_heading = True
            pf.first_line_indent = DOCX_NO_INDENT
        elif is_bibliography:
            _apply_bibliography_style(para, profile)
            prev_was_heading = False
        else:
            if not seen_section_heading or in_no_indent_section or prev_was_heading:
                pf.first_line_indent = DOCX_NO_INDENT
            else:
                pf.first_line_indent = DOCX_BODY_INDENT
            for run in para.runs:
                _set_run_font(
                    run,
                    font_name=profile.main_font,
                    size_pt=profile.body_font_size_pt,
                )
            prev_was_heading = False

        fixed_paragraphs += 1

    for table in doc.tables:
        _ensure_default_horizontal_table_borders(table)

    if profile.bibliography_style == "vancouver":
        reference_targets = _ensure_bibliography_bookmarks(doc)
        _style_docx_in_text_citations(doc, reference_targets, profile)

    doc.save(docx_path)

    print(f"✓ 已处理 {fixed_paragraphs}/{total_paragraphs} 个段落")
    print(f"✓ 文件已保存: {docx_path}")


def parse_args() -> argparse.Namespace:
    """Parse CLI arguments."""
    parser = argparse.ArgumentParser(description="Fix paragraph spacing in manuscript DOCX output.")
    parser.add_argument(
        "--project-dir",
        type=Path,
        default=None,
        help="Project directory. Defaults to the nearest parent containing main.tex.",
    )
    parser.add_argument(
        "--docx",
        type=Path,
        default=None,
        help="Explicit DOCX file path. Overrides --project-dir/main.docx when provided.",
    )
    parser.add_argument(
        "--no-backup",
        action="store_true",
        help="Skip writing a .backup copy before modifying the DOCX file.",
    )
    return parser.parse_args()


def main() -> None:
    """命令行入口：解析参数后对目标 DOCX 执行间距修复。默认修改前先备份原文件。"""
    configure_windows_stdio_utf8()
    args = parse_args()
    if args.docx is not None:
        docx_path = args.docx.expanduser().resolve()
        project_dir = docx_path.parent
    else:
        project_dir = resolve_project_dir(args.project_dir)
        docx_path = project_dir / "main.docx"

    if not docx_path.exists():
        print(f"错误: 找不到文件 {docx_path}")
        print("请先运行: bml build")
        return

    if not args.no_backup:
        backup_path = docx_path.with_suffix(".docx.backup")
        shutil.copy2(docx_path, backup_path)
        print(f"✓ 已备份原文件: {backup_path}\n")

    # 修复文档
    fix_docx_spacing(docx_path)

    if not args.no_backup:
        print("\n提示: 如需恢复原文件，请运行:")
        print(f"  cp \"{backup_path}\" \"{docx_path}\"")


if __name__ == '__main__':
    main()
