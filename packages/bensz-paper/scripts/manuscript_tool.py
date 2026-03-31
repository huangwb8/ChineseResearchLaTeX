#!/usr/bin/env python3
"""
SCI manuscript build tool bundled with bensz-paper.
"""

from __future__ import annotations

import argparse
import os
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt

VERSION = "1.3.0"
PROJECT_ROOT_MARKERS = ("main.tex",)
EXTRA_TEX_INPUT_PATTERN = re.compile(r"\\input\{(extraTex/[^}]+)\}")
CITATION_PATTERN = re.compile(
    r"\\(supercite|cite|autocite|parencite|textcite|citep|citet)\{([^}]*)\}"
)
TOOL_CANDIDATES = {
    "xelatex": (
        "/Library/TeX/texbin/xelatex",
        "/usr/local/texlive/2024/bin/universal-darwin/xelatex",
        "/usr/local/texlive/2024/bin/x86_64-linux/xelatex",
    ),
    "biber": (
        "/Library/TeX/texbin/biber",
        "/usr/local/texlive/2024/bin/universal-darwin/biber",
        "/usr/local/texlive/2024/bin/x86_64-linux/biber",
    ),
    "pandoc": ("/opt/homebrew/bin/pandoc", "/usr/local/bin/pandoc"),
    "soffice": (
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "/usr/bin/soffice",
        "/usr/local/bin/soffice",
    ),
}


def run_cmd(args: list[str], cwd: Path | None = None, input_text: str | None = None) -> str:
    result = subprocess.run(
        args,
        cwd=cwd,
        input=input_text,
        text=True,
        capture_output=True,
        check=True,
    )
    return result.stdout


def configure_windows_stdio_utf8() -> None:
    if sys.platform != "win32":
        return
    for stream_name in ("stdout", "stderr"):
        stream = getattr(sys, stream_name, None)
        if hasattr(stream, "reconfigure"):
            stream.reconfigure(encoding="utf-8", errors="replace")


def run_best_effort(
    args: list[str],
    cwd: Path | None = None,
    env: dict[str, str] | None = None,
) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        args,
        cwd=cwd,
        env=env,
        text=True,
        capture_output=True,
        check=False,
    )


def collect_extra_tex_inputs(project_dir: Path) -> list[Path]:
    main_tex = project_dir / "main.tex"
    if not main_tex.exists():
        raise FileNotFoundError(f"Missing main.tex: {main_tex}")

    ordered_inputs: list[Path] = []
    main_tex_text = []
    for line in main_tex.read_text(encoding="utf-8").splitlines():
        comment_index = None
        for index, char in enumerate(line):
            if char == "%" and (index == 0 or line[index - 1] != "\\"):
                comment_index = index
                break
        main_tex_text.append(line if comment_index is None else line[:comment_index])

    for matched_path in EXTRA_TEX_INPUT_PATTERN.findall("\n".join(main_tex_text)):
        rel_path = Path(matched_path)
        if rel_path.suffix != ".tex":
            rel_path = rel_path.with_suffix(".tex")
        source_path = project_dir / rel_path
        if not source_path.exists():
            raise FileNotFoundError(f"Missing extraTex source referenced by main.tex: {source_path}")
        ordered_inputs.append(rel_path)
    return ordered_inputs


def _replace_latex_citations_with_tokens(latex_text: str) -> tuple[str, dict[str, str]]:
    placeholder_map: dict[str, str] = {}

    def repl(match: re.Match[str]) -> str:
        token = f"CITETOKEN{len(placeholder_map) + 1:04d}"
        keys = [part.strip() for part in match.group(2).split(",") if part.strip()]
        placeholder_map[token] = "[" + "; ".join(f"@{key}" for key in keys) + "]"
        return token

    return CITATION_PATTERN.sub(repl, latex_text), placeholder_map


_SUP_TAG_RE = re.compile(r"<sup>(.*?)</sup>", re.DOTALL)


def _convert_sup_tags_to_superscript(md_text: str) -> str:
    """Convert HTML <sup> tags to pandoc native ^superscript^ syntax."""

    def _replace_sup(match: re.Match[str]) -> str:
        content = match.group(1).replace("\\*", "*")
        return f"^{content}^"

    return _SUP_TAG_RE.sub(_replace_sup, md_text)


def _normalize_frontmatter_markdown(md_text: str) -> str:
    lines = [line for line in md_text.splitlines() if line.strip() not in {"<div class=\"center\">", "</div>"}]
    normalized: list[str] = []
    title_promoted = False

    for line in lines:
        stripped = line.strip()
        if not title_promoted and stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
            normalized.append("# " + stripped[2:-2].strip())
            title_promoted = True
            continue
        normalized.append(line)

    return "\n".join(normalized).strip()


def pandoc_latex_to_markdown(latex_text: str) -> str:
    prepared_text, placeholder_map = _replace_latex_citations_with_tokens(latex_text)
    markdown = run_cmd(
        [resolve_executable("pandoc"), "-f", "latex", "-t", "gfm+raw_html"],
        input_text=prepared_text,
    ).strip()
    for token, replacement in placeholder_map.items():
        markdown = markdown.replace(token, replacement)
    markdown = _convert_sup_tags_to_superscript(markdown)
    return markdown.strip() + "\n" if markdown.strip() else ""


def _add_references_heading_if_missing(doc: "Document", heading_text: str = "References") -> None:
    first_bib = next((p for p in doc.paragraphs if p.style.name == "Bibliography"), None)
    if first_bib is None:
        return

    already_exists = any(
        p.style.name.startswith("Heading") and p.text.strip().lower() == heading_text.lower()
        for p in doc.paragraphs
    )
    if already_exists:
        return

    heading = doc.add_heading(heading_text, level=2)
    heading_elem = heading._element
    heading_elem.getparent().remove(heading_elem)
    first_bib._element.addprevious(heading_elem)


def _reorder_references_before_figure_legends(
    doc: "Document",
    references_heading: str = "References",
    figure_legends_heading: str = "Figure legends",
) -> None:
    paras = doc.paragraphs
    ref_heading_para = next(
        (
            p
            for p in paras
            if p.style.name.startswith("Heading")
            and p.text.strip().lower() == references_heading.lower()
        ),
        None,
    )
    if ref_heading_para is None:
        return

    fig_legends_para = next(
        (
            p
            for p in paras
            if p.style.name.startswith("Heading")
            and p.text.strip().lower() == figure_legends_heading.lower()
        ),
        None,
    )
    if fig_legends_para is None:
        return

    bib_elements = [p._element for p in paras if p.style.name == "Bibliography"]
    if not bib_elements:
        return

    body_elem = ref_heading_para._element.getparent()
    children = list(body_elem)
    fig_idx = children.index(fig_legends_para._element)
    ref_idx = children.index(ref_heading_para._element)
    all_before = ref_idx < fig_idx and all(children.index(e) < fig_idx for e in bib_elements)
    if all_before:
        return

    ref_block = [ref_heading_para._element] + bib_elements
    for elem in ref_block:
        body_elem.remove(elem)

    anchor = fig_legends_para._element
    for elem in ref_block:
        anchor.addprevious(elem)


def fix_docx_spacing(docx_path: Path) -> None:
    body_indent = Pt(18)
    no_indent = Pt(0)
    no_indent_sections = {"Abstract", "Figure legends", "Supplementary materials"}

    doc = Document(docx_path)
    _add_references_heading_if_missing(doc)
    _reorder_references_before_figure_legends(doc)

    in_no_indent_section = False
    prev_was_heading = True
    seen_section_heading = False

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        is_heading = style_name.startswith("Heading")
        is_bibliography = style_name == "Bibliography"

        pf = para.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.space_after = Pt(4)
        pf.space_before = Pt(0)

        if is_heading:
            if style_name != "Heading 1":
                seen_section_heading = True
            in_no_indent_section = para.text.strip() in no_indent_sections
            prev_was_heading = True
            pf.first_line_indent = no_indent
        elif is_bibliography:
            pf.first_line_indent = no_indent
            pf.left_indent = no_indent
            for run in para.runs:
                if "\t" in run.text:
                    run.text = run.text.replace("\t", "")
            prev_was_heading = False
        else:
            if not seen_section_heading or in_no_indent_section or prev_was_heading:
                pf.first_line_indent = no_indent
            else:
                pf.first_line_indent = body_indent
            prev_was_heading = False

    doc.save(docx_path)


def is_project_root(path: Path) -> bool:
    return any((path / marker).exists() for marker in PROJECT_ROOT_MARKERS)


def find_project_root(start: Path | None = None) -> Path:
    origin = (start or Path.cwd()).expanduser().resolve()
    for candidate in [origin, *origin.parents]:
        if is_project_root(candidate):
            return candidate
    raise FileNotFoundError(
        "Cannot find project root. Run inside the manuscript project tree or "
        "pass --project-dir explicitly."
    )


def resolve_project_dir(project_dir: Path | None) -> Path:
    if project_dir is None:
        return find_project_root()

    candidate = project_dir.expanduser().resolve()
    if not candidate.exists():
        raise FileNotFoundError(f"Project directory not found: {candidate}")
    if candidate.is_file():
        candidate = candidate.parent
    if is_project_root(candidate):
        return candidate
    return find_project_root(candidate)


def resolve_executable(name: str) -> str:
    resolved = shutil.which(name)
    if resolved:
        return resolved
    for candidate in TOOL_CANDIDATES.get(name, ()):
        if Path(candidate).exists():
            return candidate
    raise FileNotFoundError(f"Required executable not found: {name}")


def build_texinputs(prefixes: list[Path], existing: str) -> str:
    normalized = [f"{path.resolve()}//" for path in prefixes]
    normalized.append("")
    if existing:
        normalized.append(existing)
    return os.pathsep.join(normalized)


def resolve_tex_search_roots(project_dir: Path) -> list[Path]:
    roots: list[Path] = []

    project_tex_root = project_dir / "texmf" / "tex" / "latex"
    if project_tex_root.exists():
        roots.append(project_tex_root)

    package_root = Path(__file__).resolve().parents[1]
    if (package_root / "bml-core.sty").exists() and package_root not in roots:
        roots.append(package_root)

    fonts_package_root = package_root.parent / "bensz-fonts"
    if (fonts_package_root / "bensz-fonts.sty").exists() and fonts_package_root not in roots:
        roots.append(fonts_package_root)

    return roots


def summarize_process_output(label: str, result: subprocess.CompletedProcess[str]) -> str:
    lines = [line for line in (result.stdout + "\n" + result.stderr).splitlines() if line.strip()]
    tail = "\n".join(lines[-20:]) if lines else "(no output)"
    return f"[{label}] exit={result.returncode}\n{tail}"


def remove_legacy_docx_intermediates(cache_dir: Path) -> None:
    legacy_markdown = cache_dir / "main.md"
    legacy_extra_tex = cache_dir / "extraTex"

    if legacy_markdown.exists():
        legacy_markdown.unlink()
    if legacy_extra_tex.exists():
        shutil.rmtree(legacy_extra_tex)


def build_markdown_for_docx(project_dir: Path) -> str:
    parts: list[str] = []

    for rel_path in collect_extra_tex_inputs(project_dir):
        source_text = (project_dir / rel_path).read_text(encoding="utf-8").strip()
        if not source_text:
            continue
        markdown = pandoc_latex_to_markdown(source_text).strip()
        if not markdown:
            continue
        if rel_path.name == "frontmatter.tex":
            markdown = _normalize_frontmatter_markdown(markdown)
        parts.append(markdown)

    if not parts:
        raise RuntimeError(f"No DOCX source fragments found under {project_dir / 'extraTex'}")

    return "\n\n".join(parts).rstrip() + "\n"


def build_project(project_dir: Path) -> None:
    print(f"Building project: {project_dir}")

    main_tex = project_dir / "main.tex"
    if not main_tex.exists():
        raise FileNotFoundError(f"Missing main.tex: {main_tex}")

    print("Building PDF...")
    cache_dir = project_dir / ".latex-cache"
    cache_dir.mkdir(parents=True, exist_ok=True)
    remove_legacy_docx_intermediates(cache_dir)

    tex_env = os.environ.copy()
    tex_roots = resolve_tex_search_roots(project_dir)
    if tex_roots:
        tex_env["TEXINPUTS"] = build_texinputs(tex_roots, tex_env.get("TEXINPUTS", ""))
        print("TeX search roots:")
        for root in tex_roots:
            print(f"  - {root}")

    xelatex_cmd = [
        resolve_executable("xelatex"),
        "-interaction=nonstopmode",
        "-file-line-error",
        "-synctex=1",
        f"-output-directory={cache_dir}",
        "main.tex",
    ]

    xelatex_run_1 = run_best_effort(xelatex_cmd, cwd=project_dir, env=tex_env)
    biber_run = run_best_effort(
        [
            resolve_executable("biber"),
            "--input-directory",
            str(cache_dir),
            "--output-directory",
            str(cache_dir),
            "main",
        ],
        cwd=project_dir,
        env=tex_env,
    )
    xelatex_run_2 = run_best_effort(xelatex_cmd, cwd=project_dir, env=tex_env)
    xelatex_run_3 = run_best_effort(xelatex_cmd, cwd=project_dir, env=tex_env)

    pdf_source = cache_dir / "main.pdf"
    if not pdf_source.exists():
        compiler_logs = "\n\n".join(
            [
                summarize_process_output("xelatex pass 1", xelatex_run_1),
                summarize_process_output("biber", biber_run),
                summarize_process_output("xelatex pass 2", xelatex_run_2),
                summarize_process_output("xelatex pass 3", xelatex_run_3),
            ]
        )
        raise RuntimeError(
            f"PDF compilation failed. Expected output not found: {pdf_source}\n\n{compiler_logs}"
        )

    for label, result in (
        ("xelatex pass 1", xelatex_run_1),
        ("biber", biber_run),
        ("xelatex pass 2", xelatex_run_2),
        ("xelatex pass 3", xelatex_run_3),
    ):
        if result.returncode != 0:
            print(f"Warning: {label} exited with code {result.returncode}; output PDF was still generated.")

    shutil.copy2(pdf_source, project_dir / "main.pdf")
    print(f"✓ PDF generated: {project_dir / 'main.pdf'}")

    print("Building DOCX...")
    manuscript_md = build_markdown_for_docx(project_dir)

    reference_doc = project_dir / "artifacts" / "reference.docx"
    reference_doc_arg = ["--reference-doc", str(reference_doc)] if reference_doc.exists() else []

    csl_path = project_dir / "artifacts" / "manuscript.csl"
    if not csl_path.exists():
        raise FileNotFoundError(f"Missing CSL file: {csl_path}")

    docx_path = project_dir / "main.docx"
    pandoc_cmd = [
        resolve_executable("pandoc"),
        "-",
        "-f",
        "markdown+raw_html+superscript",
        "--citeproc",
        *reference_doc_arg,
        "--csl",
        str(csl_path),
        "--bibliography",
        str(project_dir / "references" / "refs.bib"),
        "-o",
        str(docx_path),
    ]
    run_cmd(pandoc_cmd, cwd=project_dir, input_text=manuscript_md)
    print(f"✓ DOCX generated: {docx_path}")

    print("Fixing DOCX spacing...")
    fix_docx_spacing(docx_path)
    print("✓ DOCX spacing fixed")

    soffice = shutil.which("soffice") or next(
        (candidate for candidate in TOOL_CANDIDATES["soffice"] if Path(candidate).exists()),
        None,
    )
    if soffice:
        try:
            with tempfile.TemporaryDirectory(prefix="paper-word-pdf-") as tmp_dir:
                word_pdf_dir = Path(tmp_dir)
                run_cmd(
                    [
                        soffice,
                        "--headless",
                        "--convert-to",
                        "pdf",
                        "--outdir",
                        str(word_pdf_dir),
                        str(docx_path),
                    ],
                    cwd=project_dir,
                )
                generated_word_pdf = word_pdf_dir / "main.pdf"
                if generated_word_pdf.exists():
                    shutil.copy2(generated_word_pdf, cache_dir / "main.word.pdf")
                    print(f"✓ Word-based PDF generated: {cache_dir / 'main.word.pdf'}")
        except Exception as exc:
            print(f"Note: Could not generate Word-based PDF: {exc}")

    print("\n✓ Build complete!")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build manuscript PDF and DOCX from local sources.")
    parser.add_argument("--version", action="version", version=f"%(prog)s {VERSION}")
    parser.add_argument("command", choices=["build"], help="Command to execute")
    parser.add_argument(
        "--project-dir",
        type=Path,
        default=None,
        help="Project directory. Defaults to the nearest parent containing main.tex.",
    )
    return parser.parse_args()


def main() -> None:
    configure_windows_stdio_utf8()
    args = parse_args()
    if args.command == "build":
        project_dir = resolve_project_dir(args.project_dir)
        build_project(project_dir)
        return
    raise ValueError(f"Unsupported command: {args.command}")


if __name__ == "__main__":
    main()
