#!/usr/bin/env python3
"""将 thesis-ucas-doctor 的 LaTeX 源尽量转换为可编辑的 Word（docx）。

设计目标：
1. 优先用源码而不是 PDF 提取，以提升结构和公式可保真度。
2. 对复杂 LaTeX 环境做可控降级，避免 pandoc 直接解析失败。
3. 通过 --reference-doc 套用学校 Word 模板样式。
4. 转换后自动做样式归一化与质量报告（标题层级/段落样式/版面参数）。
"""

from __future__ import annotations

import argparse
import os
import re
import shutil
import subprocess
import sys
import uuid
import xml.etree.ElementTree as ET
import zipfile
from collections import Counter
from datetime import datetime
from pathlib import Path, PurePosixPath
from tempfile import NamedTemporaryFile, TemporaryDirectory
from typing import Dict, Iterable, List, Optional, Set, Tuple

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from word_export.format_utils import (
    INLINE_FIGURE_MACRO_USE_RE,
    INLINE_TABLE_MACRO_USE_RE,
    MINIPAGE_ENV_RE,
    NOTE_MACRO_USE_RE,
    _build_caption_marker,
    _caption_latex_to_text,
    _extract_graphicspaths,
    _format_caption_text,
    _format_caption_marker_sequence,
    _format_chapter_sequence,
    _format_export_sequence,
    _italicize_plain_stat_tokens,
    _normalize_cjk_number_classifier_spacing,
    _normalize_simple_stat_inline_math,
    _parse_includegraphics_path,
    _parse_optional_and_braced_args,
    _resolve_graphics_path,
    _strip_caption_terminal_punctuation,
    _strip_inline_figure_macro_definitions,
    _strip_macro_calls,
    latex_inline_to_text,
)

try:
    from docx import Document as DocxDocument  # type: ignore
except Exception:
    DocxDocument = None


DEFAULT_PROJECT_DIR = Path(__file__).resolve().parents[1]
PREPARE_TEX_SCRIPT_REL = Path("scripts") / "prepare_tex_for_word_export.py"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PKG_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
W = {"w": W_NS, "r": R_NS}
DEFAULT_REFERENCE_DOC_NAMES = [
    "中国科学院大学资环学科群研究生学位论文word模板.docx",
    "附件2 中国科学院大学学位论文word模板2024.10.23——建议使用office 2021及以上版本.docx",
]
DEFAULT_REFERENCE_DOC_DIRS = [
    ("docs", "official"),
    (),
    ("docs",),
]
DEFAULT_PANDOC_CSL = Path(__file__).resolve().parent / "word_export" / "china-national-standard-gb-t-7714-2015-author-date.csl"
CAPTION_MARKER_RE = re.compile(r"^\[\[BENSZ(FIG|TBL)CAP:(zh|en):([^\]]+)\]\]\s*(.*)$")
STRICT_PYTHON_ENV = "BENSZ_DOCX_PYTHON"
BIB_PLACEHOLDER_TEXT = "文献列表请在 Word 中按模板粘贴补全。"
BIB_PLACEHOLDER_LINE = f"> {BIB_PLACEHOLDER_TEXT}"
BIB_REFS_DIV = "::: {#refs}\n:::"
SECTION_MARKERS_WITH_TEMPLATE_INDEX: List[Tuple[str, int]] = [
    ("摘要", 2),
    ("abstract", 3),
    ("目录", 4),
    ("图表目录", 5),
    ("绪论", 6),
    ("附录", 8),
    ("参考文献", 7),
    ("致谢", 11),
    ("作者简历及攻读学位期间发表的学术论文与其他相关学术成果", 12),
]
DOCX_HEADER_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/header"
EMU_PER_TWIP = 635

OOXML_NAMESPACE_PREFIXES = {
    "wpc": "http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas",
    "mc": MC_NS,
    "o": "urn:schemas-microsoft-com:office:office",
    "r": R_NS,
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    "v": "urn:schemas-microsoft-com:vml",
    "wp14": "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "w10": "urn:schemas-microsoft-com:office:word",
    "w": W_NS,
    "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
    "w15": "http://schemas.microsoft.com/office/word/2012/wordml",
    "w16cex": "http://schemas.microsoft.com/office/word/2018/wordml/cex",
    "w16cid": "http://schemas.microsoft.com/office/word/2016/wordml/cid",
    "w16": "http://schemas.microsoft.com/office/word/2018/wordml",
    "w16sdtdh": "http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash",
    "w16se": "http://schemas.microsoft.com/office/word/2015/wordml/symex",
    "wpg": "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup",
    "wpi": "http://schemas.microsoft.com/office/word/2010/wordprocessingInk",
    "wne": "http://schemas.microsoft.com/office/word/2006/wordml",
    "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
}

IGNORABLE_NAMESPACE_URIS = {
    "w14": OOXML_NAMESPACE_PREFIXES["w14"],
    "w15": OOXML_NAMESPACE_PREFIXES["w15"],
    "w16se": OOXML_NAMESPACE_PREFIXES["w16se"],
    "w16cid": OOXML_NAMESPACE_PREFIXES["w16cid"],
    "w16": OOXML_NAMESPACE_PREFIXES["w16"],
    "w16cex": OOXML_NAMESPACE_PREFIXES["w16cex"],
    "w16sdtdh": OOXML_NAMESPACE_PREFIXES["w16sdtdh"],
    "wp14": OOXML_NAMESPACE_PREFIXES["wp14"],
}


for prefix, uri in OOXML_NAMESPACE_PREFIXES.items():
    ET.register_namespace(prefix, uri)


def _serialize_xml(root: ET.Element, *, default_namespace: Optional[str] = None) -> bytes:
    if default_namespace:
        ET.register_namespace("", default_namespace)
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _w_attr(key: str) -> str:
    return f"{{{W_NS}}}{key}"


def _r_attr(key: str) -> str:
    return f"{{{R_NS}}}{key}"


def _m_attr(key: str) -> str:
    return f"{{{OOXML_NAMESPACE_PREFIXES['m']}}}{key}"


def _v_attr(key: str) -> str:
    return f"{{{OOXML_NAMESPACE_PREFIXES['v']}}}{key}"


def discover_reference_doc(project_dir: Path, explicit: Optional[Path]) -> Path:
    if explicit is not None:
        expanded = explicit.expanduser()
        if expanded.is_absolute() or expanded.exists():
            return expanded.resolve()
        return (project_dir / expanded).resolve()

    candidates: List[Path] = []
    for name in DEFAULT_REFERENCE_DOC_NAMES:
        for rel_dir in DEFAULT_REFERENCE_DOC_DIRS:
            candidates.append(project_dir.joinpath(*rel_dir, name))

    for candidate in candidates:
        if candidate.exists():
            return candidate.resolve()

    available = sorted(p.name for p in project_dir.glob("*.docx"))
    available_official = sorted(p.name for p in (project_dir / "docs" / "official").glob("*.docx"))
    available_docs = sorted(p.name for p in (project_dir / "docs").glob("*.docx"))
    hint = f" 当前项目根目录可见 docx: {available}" if available else ""
    hint_official = f" docs/official 可见 docx: {available_official}" if available_official else ""
    hint_docs = f" docs/ 可见 docx: {available_docs}" if available_docs else ""
    raise FileNotFoundError(
        "未找到参考 Word 模板。请将资环 Word 模板放在项目 docs/official/，"
        "或显式传入 --reference-doc。"
        + hint
        + hint_official
        + hint_docs
    )


def parse_args(argv: Optional[List[str]] = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Export thesis LaTeX source to docx with reference template.")
    parser.add_argument(
        "--mode",
        choices=["portable", "strict"],
        default="strict",
        help="导出模式：portable=跨平台可用，strict=最高对齐优先（含严格前置检查）",
    )
    parser.add_argument("--project-dir", type=Path, default=DEFAULT_PROJECT_DIR, help="论文项目目录")
    parser.add_argument("--tex-file", type=str, default="main.tex", help="主 tex 文件名")
    parser.add_argument("--reference-doc", type=Path, default=None, help="参考 Word 模板 .docx")
    parser.add_argument("--output", type=Path, default=None, help="输出 docx 路径")
    parser.add_argument("--markdown-output", type=Path, default=None, help="中间 markdown 路径")
    parser.add_argument("--quality-report", type=Path, default=None, help="质量报告路径")
    parser.add_argument(
        "--skip-style-normalization",
        action="store_true",
        help="仅导出 docx，不做样式归一化与质量报告",
    )
    parser.add_argument(
        "--postprocess-only-docx",
        type=Path,
        default=None,
        help="仅对已有 docx 做样式归一化与质量报告，不重新跑 pandoc",
    )
    parser.add_argument(
        "--prepare-tex",
        action="store_true",
        help="导出前先执行 TeX 前处理（空格重建 + 时间单位规范化）。",
    )
    parser.add_argument(
        "--word-update-fields",
        action="store_true",
        help="导出后调用 Windows Word 自动更新目录/图表目录等域并保存（仅 Windows + Office Word）",
    )
    parser.add_argument(
        "--ensure-page-cache",
        action="store_true",
        help="导出前先构建 PDF 以生成 .latex-cache/main.lof/main.lot，供 Word 图表目录写入页码。",
    )
    parser.add_argument(
        "--word-visible",
        action="store_true",
        help="与 --word-update-fields 配合，调试时显示 Word 窗口（默认隐藏）",
    )
    parser.add_argument(
        "--word-update-timeout",
        type=int,
        default=300,
        help="Word 自动更新域超时秒数（默认 300）",
    )
    parser.add_argument(
        "--omml-gate-max",
        type=int,
        default=None,
        help="导出后允许的最大 OMML 数学节点数（按 m:oMath 计数）；超出则报错。",
    )
    return parser.parse_args(argv)


def build_prepare_tex_command(python_cmd: str, project_dir: Path) -> List[str]:
    prepare_script = (project_dir / PREPARE_TEX_SCRIPT_REL).resolve()
    return [
        python_cmd,
        str(prepare_script),
        "--project-dir",
        str(project_dir),
        "--apply",
    ]


def _find_thesis_project_tool_for_page_cache() -> Path:
    repo_candidate = Path(__file__).resolve().parents[3] / "packages" / "bensz-thesis" / "scripts" / "thesis_project_tool.py"
    if repo_candidate.exists():
        return repo_candidate

    kpsewhich = subprocess.run(
        ["kpsewhich", "bensz-thesis.sty"],
        capture_output=True,
        check=False,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    if kpsewhich.returncode == 0 and kpsewhich.stdout.strip():
        package_root = Path(kpsewhich.stdout.strip()).resolve().parent
        installed_candidate = package_root / "scripts" / "thesis_project_tool.py"
        if installed_candidate.exists():
            return installed_candidate

    raise FileNotFoundError("未找到 thesis_project_tool.py，无法为 Word 图表目录生成页码缓存。")


def build_page_cache_command(python_cmd: str, project_dir: Path, thesis_tool: Optional[Path] = None) -> List[str]:
    tool = thesis_tool or _find_thesis_project_tool_for_page_cache()
    return [
        python_cmd,
        str(tool),
        "build",
        "--project-dir",
        str(project_dir),
    ]


def run_page_cache_build_for_word_export(project_dir: Path, python_cmd: Optional[str] = None) -> None:
    command = build_page_cache_command(python_cmd or sys.executable, project_dir)
    print(f"[RUN] page cache build: {' '.join(command)}")
    result = subprocess.run(command, cwd=str(project_dir), check=False)
    if result.returncode != 0:
        raise RuntimeError(f"Word 图表目录页码缓存构建失败，退出码：{result.returncode}")

    cache_dir = project_dir / ".latex-cache"
    if not (cache_dir / "main.lof").exists() and not (cache_dir / "main.lot").exists():
        raise RuntimeError(f"Word 图表目录页码缓存缺失：{cache_dir / 'main.lof'} / {cache_dir / 'main.lot'}")
    print("[OK] page cache build: completed")


def run_prepare_tex_for_export(project_dir: Path, python_cmd: Optional[str] = None) -> None:
    prepare_script = (project_dir / PREPARE_TEX_SCRIPT_REL).resolve()
    if not prepare_script.exists():
        raise FileNotFoundError(f"未找到 Word 前处理脚本：{prepare_script}")

    command = build_prepare_tex_command(python_cmd or sys.executable, project_dir)
    print(f"[RUN] prepare tex: {' '.join(command)}")
    result = subprocess.run(command, cwd=str(project_dir), check=False)
    if result.returncode != 0:
        raise RuntimeError(f"Word 前处理失败，退出码：{result.returncode}")
    print("[OK] prepare tex: completed")


def _replace_docx_with_lock_hint(tmp_path: Path, docx_path: Path) -> None:
    try:
        tmp_path.replace(docx_path)
    except PermissionError as exc:
        raise _build_docx_lock_permission_error(exc, docx_path) from exc


def _build_docx_lock_permission_error(exc: PermissionError, docx_path: Path) -> PermissionError:
    return PermissionError(
        exc.errno,
        (
            f"{exc.strerror}: 目标 DOCX 可能正被 Word、WPS、资源管理器预览窗格或同步程序占用。"
            f" 请先关闭 `{docx_path.name}` 的打开窗口/预览后重试。"
        ),
        str(docx_path),
    )


def parse_braced(text: str, start: int) -> Tuple[str, int]:
    if start >= len(text) or text[start] != "{":
        raise ValueError("parse_braced expects '{' at start")

    depth = 0
    i = start
    out: List[str] = []

    while i < len(text):
        ch = text[i]
        if ch == "\\":
            if i + 1 < len(text):
                out.append(text[i : i + 2])
                i += 2
                continue
            out.append(ch)
            i += 1
            continue

        if ch == "{":
            depth += 1
            if depth > 1:
                out.append(ch)
            i += 1
            continue

        if ch == "}":
            depth -= 1
            if depth == 0:
                return "".join(out), i + 1
            out.append(ch)
            i += 1
            continue

        out.append(ch)
        i += 1

    raise ValueError("unbalanced braces")


def parse_macro_args(text: str, macro: str, n_args: int) -> Optional[List[str]]:
    pat = re.compile(r"\\" + re.escape(macro) + r"\s*")
    m = pat.search(text)
    if not m:
        return None

    i = m.end()
    args: List[str] = []
    for _ in range(n_args):
        while i < len(text) and text[i].isspace():
            i += 1
        if i >= len(text) or text[i] != "{":
            return None
        value, i = parse_braced(text, i)
        args.append(value)
    return args


def parse_bracketed(text: str, start: int) -> Tuple[str, int]:
    if start >= len(text) or text[start] != "[":
        raise ValueError("parse_bracketed expects '[' at start")

    depth = 0
    i = start
    out: List[str] = []

    while i < len(text):
        ch = text[i]
        if ch == "\\":
            if i + 1 < len(text):
                out.append(text[i : i + 2])
                i += 2
                continue
            out.append(ch)
            i += 1
            continue

        if ch == "[":
            depth += 1
            if depth > 1:
                out.append(ch)
            i += 1
            continue

        if ch == "]":
            depth -= 1
            if depth == 0:
                return "".join(out), i + 1
            out.append(ch)
            i += 1
            continue

        out.append(ch)
        i += 1

    raise ValueError("unbalanced brackets")


def strip_comments(text: str) -> str:
    lines = []
    for line in text.splitlines():
        line = re.sub(r"(?<!\\)%.*$", "", line)
        lines.append(line)
    return "\n".join(lines)


def flatten_texorpdfstring(text: str) -> str:
    pattern = re.compile(r"\\texorpdfstring\s*\{")

    while True:
        m = pattern.search(text)
        if not m:
            break
        try:
            _, p1 = parse_braced(text, m.end() - 1)
            while p1 < len(text) and text[p1].isspace():
                p1 += 1
            if p1 >= len(text) or text[p1] != "{":
                break
            arg2, p2 = parse_braced(text, p1)
        except Exception:
            break

        text = text[: m.start()] + arg2 + text[p2:]

    return text


def _is_appendix_tex(text: str) -> bool:
    text = strip_comments(flatten_texorpdfstring(text))
    match = re.search(r"\\chapter\*?\{([^{}]*)\}", text)
    if not match:
        return False
    title = latex_inline_to_text(match.group(1)).strip()
    return title.startswith("附录")


def _format_chinese_ordinal(value: int) -> str:
    numerals = "零一二三四五六七八九"
    if value <= 0:
        return str(value)
    if value < 10:
        return numerals[value]
    if value == 10:
        return "十"
    if value < 20:
        return "十" + numerals[value - 10]
    if value < 100:
        tens, ones = divmod(value, 10)
        return numerals[tens] + "十" + (numerals[ones] if ones else "")
    return str(value)


def _format_appendix_chapter_heading(title: str, appendix_index: int) -> str:
    plain_title = latex_inline_to_text(title).strip()
    if plain_title.startswith("附录"):
        return plain_title
    return f"附录{_format_chinese_ordinal(appendix_index)} {plain_title}".strip()


def _extract_label(text: str) -> Optional[str]:
    match = re.search(r"\\label\{([^{}]+)\}", text)
    if not match:
        return None
    label = match.group(1).strip()
    return label or None


def _extract_caption_arg(text: str) -> Optional[str]:
    pat = re.compile(r"\\caption\s*")
    m = pat.search(text)
    if not m:
        return None

    i = m.end()
    while i < len(text) and text[i].isspace():
        i += 1
    if i < len(text) and text[i] == "[":
        try:
            _, i = parse_bracketed(text, i)
        except Exception:
            return None
    while i < len(text) and text[i].isspace():
        i += 1
    if i >= len(text) or text[i] != "{":
        return None
    try:
        value, _ = parse_braced(text, i)
    except Exception:
        return None
    return value


def _replace_cross_references(text: str, reference_map: Dict[str, str]) -> str:
    if not reference_map:
        return text

    def _explicit_repl(m: re.Match[str]) -> str:
        prefix = m.group(1)
        label = m.group(2).strip()
        return f"{prefix}{reference_map.get(label, label)}"

    text = re.sub(r"(附图|附表|图|表)\s*~?\s*\\ref\{([^{}]+)\}", _explicit_repl, text)

    def _generic_repl(m: re.Match[str]) -> str:
        label = m.group(1).strip()
        return reference_map.get(label, label)

    for macro in ["autoref", "ref", "cref", "Cref"]:
        text = re.sub(r"\\" + macro + r"\{([^{}]+)\}", _generic_repl, text)
    text = re.sub(
        r"\\eqref\{([^{}]+)\}",
        lambda m: f"({reference_map.get(m.group(1).strip(), m.group(1).strip())})",
        text,
    )
    text = re.sub(
        r"\\pageref\{([^{}]+)\}",
        lambda m: reference_map.get(m.group(1).strip(), m.group(1).strip()),
        text,
    )
    return text


def _extract_trailing_label(text: str, start: int) -> Tuple[Optional[str], int]:
    i = start
    while i < len(text) and text[i].isspace():
        i += 1
    if not text.startswith("\\label", i):
        return None, start
    i += len("\\label")
    while i < len(text) and text[i].isspace():
        i += 1
    if i >= len(text) or text[i] != "{":
        return None, start
    try:
        label, end = parse_braced(text, i)
    except Exception:
        return None, start
    label = label.strip()
    return (label or None), end


def _collect_structural_reference_map(text: str, chapter_no: int) -> Dict[str, str]:
    reference_map: Dict[str, str] = {}
    section_no = 0
    subsection_no = 0
    subsubsection_no = 0
    heading_re = re.compile(r"\\(subsubsection|subsection|section|chapter)(\*)?\s*")

    cursor = 0
    while True:
        match = heading_re.search(text, cursor)
        if not match:
            break

        command = match.group(1)
        numbered = match.group(2) != "*"
        i = match.end()
        while i < len(text) and text[i].isspace():
            i += 1
        if i >= len(text) or text[i] != "{":
            cursor = match.end()
            continue

        try:
            _, i = parse_braced(text, i)
        except Exception:
            cursor = match.end()
            continue

        label, _ = _extract_trailing_label(text, i)
        ref_value: Optional[str] = None
        if numbered:
            if command == "chapter":
                section_no = 0
                subsection_no = 0
                subsubsection_no = 0
                ref_value = str(chapter_no) if chapter_no > 0 else None
            elif chapter_no > 0 and command == "section":
                section_no += 1
                subsection_no = 0
                subsubsection_no = 0
                ref_value = f"{chapter_no}.{section_no}"
            elif chapter_no > 0 and command == "subsection":
                subsection_no += 1
                subsubsection_no = 0
                ref_value = f"{chapter_no}.{section_no}.{subsection_no}" if section_no > 0 else None
            elif chapter_no > 0 and command == "subsubsection":
                subsubsection_no += 1
                ref_value = (
                    f"{chapter_no}.{section_no}.{subsection_no}.{subsubsection_no}"
                    if section_no > 0 and subsection_no > 0
                    else None
                )

        if label and ref_value:
            reference_map[label] = ref_value

        cursor = i

    equation_no = 0
    for match in re.finditer(r"\\begin\{equation\}(?:\[[^\]]*\])?(.*?)\\end\{equation\}", text, re.S):
        label = _extract_label(match.group(1))
        if not label:
            continue
        equation_no += 1
        reference_map[label] = _format_chapter_sequence(chapter_no, equation_no)

    return reference_map


def _extract_tabular_body(block: str) -> str:
    for env_name, arg_count in (("tabular", 1), ("tabularx", 2), ("longtable", 1)):
        match = re.search(r"\\begin\{" + re.escape(env_name) + r"\}", block)
        if not match:
            continue

        i = match.end()
        while i < len(block) and block[i].isspace():
            i += 1

        if i < len(block) and block[i] == "[":
            try:
                _, i = parse_bracketed(block, i)
            except Exception:
                continue

        try:
            for _ in range(arg_count):
                while i < len(block) and block[i].isspace():
                    i += 1
                if i >= len(block) or block[i] != "{":
                    raise ValueError("missing tabular argument")
                _, i = parse_braced(block, i)
        except Exception:
            continue

        end_match = re.search(r"\\end\{" + re.escape(env_name) + r"\}", block[i:])
        if end_match:
            return block[i : i + end_match.start()]

    return ""


def _split_tex_top_level(text: str, delimiter: str) -> List[str]:
    parts: List[str] = []
    buf: List[str] = []
    brace_depth = 0
    bracket_depth = 0
    i = 0

    while i < len(text):
        ch = text[i]

        if ch == "\\" and i + 1 < len(text):
            if (
                delimiter == "\\\\"
                and text[i + 1] == "\\"
                and brace_depth == 0
                and bracket_depth == 0
            ):
                parts.append("".join(buf))
                buf = []
                i += 2
                continue

            buf.append(ch)
            buf.append(text[i + 1])
            i += 2
            continue

        if ch == "{":
            brace_depth += 1
        elif ch == "}" and brace_depth > 0:
            brace_depth -= 1
        elif ch == "[":
            bracket_depth += 1
        elif ch == "]" and bracket_depth > 0:
            bracket_depth -= 1

        if delimiter == "&" and ch == "&" and brace_depth == 0 and bracket_depth == 0:
            parts.append("".join(buf))
            buf = []
            i += 1
            continue

        buf.append(ch)
        i += 1

    parts.append("".join(buf))
    return parts


def _normalize_longtable_body(body: str) -> str:
    if "\\endfirsthead" not in body:
        return body

    first_head, remainder = body.split("\\endfirsthead", 1)
    data_part = remainder
    if "\\endlastfoot" in data_part:
        data_part = data_part.split("\\endlastfoot", 1)[1]
    elif "\\endfoot" in data_part:
        data_part = data_part.split("\\endfoot", 1)[1]
    elif "\\endhead" in data_part:
        data_part = data_part.split("\\endhead", 1)[1]

    return f"{first_head}\n{data_part}"


def _normalize_shortstack_cell(content: str) -> str:
    parts = [_caption_latex_to_text(part).strip() for part in _split_tex_top_level(content, "\\\\")]
    parts = [part for part in parts if part]
    if not parts:
        return ""
    if len(parts) == 1:
        return parts[0]
    if len(parts) == 2 and parts[1].startswith("±"):
        return f"{parts[0]}{parts[1]}"
    if len(parts) == 2 and re.fullmatch(r"\(?n\s*=\s*\d+\)?", parts[1], re.IGNORECASE):
        return f"{parts[0]} {parts[1]}"
    if all(_CJK_CHAR_RE.search(part) for part in parts):
        return "".join(parts)
    if any(re.fullmatch(r"[A-Za-z]{1,6}", part) for part in parts):
        return " ".join(parts)
    return "".join(parts)


def _extract_table_rows(block: str) -> List[List[str]]:
    body = _extract_tabular_body(block)
    if not body:
        return []
    body = _normalize_longtable_body(body)
    body = re.sub(r"\\caption(?:\[[^\]]*\])?\{[^{}]*\}", "", body)
    body = re.sub(r"\\label\{[^{}]*\}", "", body)
    body = re.sub(r"\\(?:hline|toprule|midrule|bottomrule)\b", "", body)
    body = re.sub(r"\\cline\{[^{}]*\}", "", body)
    body = re.sub(r"\\(?:endfirsthead|endhead|endfoot|endlastfoot)\b", "", body)
    raw_rows = _split_tex_top_level(body, "\\\\")
    rows: List[List[str]] = []
    for raw_row in raw_rows:
        row = raw_row.strip()
        if not row:
            continue
        cells = [cell.strip() for cell in _split_tex_top_level(row, "&")]
        cleaned_cells = []
        for cell in cells:
            if not cell:
                cleaned_cells.append("")
                continue
            cell = re.sub(
                r"\\multicolumn\{[^{}]*\}\{[^{}]*\}\{([^{}]*)\}",
                lambda m: m.group(1),
                cell,
            )
            cell = re.sub(
                r"\\shortstack\{([^{}]*)\}",
                lambda m: _normalize_shortstack_cell(m.group(1)),
                cell,
            )
            cell = cell.replace(r"\%", "%")
            cell = cell.replace(r"\_", "_")
            cell = _caption_latex_to_text(cell)
            cell = cell.replace("|", r"\|").strip()
            cleaned_cells.append(cell)
        if any(cleaned_cells) and not any("续下页" in c for c in cleaned_cells):
            rows.append(cleaned_cells)
    return rows


def _markdown_table_from_rows(rows: List[List[str]]) -> str:
    if not rows:
        return ""

    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    separator = ["---"] * width
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(separator) + " |",
    ]
    for row in normalized[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def _render_table_markdown(
    block: str,
    *,
    state: Dict[str, object],
) -> str:
    bi_args = parse_macro_args(block, "bicaption", 2)
    cap_zh = _caption_latex_to_text(bi_args[0]) if bi_args else ""
    cap_en = _caption_latex_to_text(bi_args[1]) if bi_args else ""
    if not cap_zh and not cap_en:
        cap_arg = _extract_caption_arg(block)
        if cap_arg:
            cap_zh = _caption_latex_to_text(cap_arg)

    rows = _extract_table_rows(block)
    if rows:
        state["table_no"] = int(state.get("table_no", 0)) + 1
        if state.get("appendix_mode"):
            state["appendix_table_no"] = int(state.get("appendix_table_no", 0)) + 1
            appendix_chapter_no = int(state.get("appendix_chapter_no", 0) or 0)
            sequence = _format_export_sequence(
                appendix_chapter_no,
                int(state["appendix_table_no"]),
                appendix_mode=True,
            )
            marker_sequence = _format_caption_marker_sequence(
                appendix_chapter_no,
                int(state["appendix_table_no"]),
                appendix_mode=True,
            )
        else:
            sequence = _format_export_sequence(
                int(state.get("chapter_no", 0)),
                int(state["table_no"]),
                appendix_mode=False,
            )
            marker_sequence = sequence
        notes = _extract_bensz_notes(block)
        lines: List[str] = []
        if cap_zh:
            lines.append(_build_caption_marker("TBL", "zh", marker_sequence, cap_zh))
        if cap_en:
            lines.append(_build_caption_marker("TBL", "en", marker_sequence, cap_en))
        if lines:
            lines.append("")
        lines.append(_markdown_table_from_rows(rows))
        for note in notes:
            lines.append("")
            lines.append(f"**注：** {note}")
        return "\n\n" + "\n\n".join(lines) + "\n\n"

    cap = cap_zh if not cap_en else f"{cap_zh} / {cap_en}"
    if cap:
        return f"\n\n> 表：{cap}（需在 Word 中人工整理）\n\n"
    return "\n\n> 表：内容略（需在 Word 中人工整理）\n\n"


def _replace_inline_table_macros(
    text: str,
    *,
    state: Dict[str, object],
) -> str:
    cursor = 0
    chunks: List[str] = []

    while True:
        match = INLINE_TABLE_MACRO_USE_RE.search(text, cursor)
        if not match:
            chunks.append(text[cursor:])
            break

        chunks.append(text[cursor : match.start()])
        parsed = _parse_optional_and_braced_args(text, match.end(), required_args=4)
        if not parsed:
            chunks.append(text[match.start() : match.end()])
            cursor = match.end()
            continue

        end, _, args = parsed
        synthetic_block = (
            "\\captionsetup{type=table}\n"
            f"\\bicaption{{{args[0]}}}{{{args[1]}}}\n"
            f"\\label{{{args[2]}}}\n"
            f"{args[3]}"
        )
        chunks.append(_render_table_markdown(synthetic_block, state=state))
        cursor = end

    return "".join(chunks)


def _collect_reference_map_with_counters(
    text: str,
    chapter_no: int,
    *,
    appendix_mode: bool = False,
    appendix_chapter_no: int = 0,
    appendix_figure_start: int = 0,
    appendix_table_start: int = 0,
) -> Tuple[Dict[str, str], int, int]:
    text = strip_comments(text)
    text = _strip_inline_figure_macro_definitions(text)
    effective_chapter_no = appendix_chapter_no if appendix_mode else chapter_no
    reference_map: Dict[str, str] = _collect_structural_reference_map(text, effective_chapter_no)
    figure_items: List[Tuple[int, str]] = []
    table_items: List[Tuple[int, str]] = []

    cursor = 0
    while True:
        match = INLINE_FIGURE_MACRO_USE_RE.search(text, cursor)
        if not match:
            break
        parsed = _parse_optional_and_braced_args(text, match.end(), required_args=4)
        if not parsed:
            cursor = match.end()
            continue
        end, _, args = parsed
        label = args[3].strip()
        if label:
            figure_items.append((match.start(), label))
        cursor = end

    for match in re.finditer(r"\\begin\{figure\}(?:\[[^\]]*\])?(.*?)\\end\{figure\}", text, re.S):
        label = _extract_label(match.group(1))
        if label:
            figure_items.append((match.start(), label))

    for match in re.finditer(r"\\begin\{landscape\}(.*?)\\end\{landscape\}", text, re.S):
        block = match.group(1)
        if not _parse_includegraphics_path(block):
            continue
        has_figure_hint = "\\captionsetup{type=figure" in block
        if not has_figure_hint:
            has_figure_hint = bool(parse_macro_args(block, "bicaption", 2) or _extract_caption_arg(block))
        if not has_figure_hint:
            continue
        label = _extract_label(block)
        if label:
            figure_items.append((match.start(), label))

    figure_count = appendix_figure_start
    for index, (_, label) in enumerate(sorted(figure_items), start=1):
        if appendix_mode:
            figure_count += 1
            reference_map[label] = _format_export_sequence(effective_chapter_no, figure_count, appendix_mode=True)
        else:
            reference_map[label] = _format_export_sequence(chapter_no, index, appendix_mode=False)

    cursor = 0
    while True:
        match = INLINE_TABLE_MACRO_USE_RE.search(text, cursor)
        if not match:
            break
        parsed = _parse_optional_and_braced_args(text, match.end(), required_args=4)
        if not parsed:
            cursor = match.end()
            continue
        end, _, args = parsed
        label = args[2].strip()
        if label:
            table_items.append((match.start(), label))
        cursor = end

    for match in re.finditer(r"\\begin\{table\}(?:\[[^\]]*\])?(.*?)\\end\{table\}", text, re.S):
        label = _extract_label(match.group(1))
        if label:
            table_items.append((match.start(), label))

    for match in re.finditer(r"\\begin\{longtable\}(?:\[[^\]]*\])?(.*?)\\end\{longtable\}", text, re.S):
        label = _extract_label(match.group(1))
        if label:
            table_items.append((match.start(), label))

    for match in MINIPAGE_ENV_RE.finditer(text):
        block = match.group(1)
        if "\\captionsetup{type=table}" not in block:
            continue
        if not re.search(r"\\begin\{tabularx?\}", block):
            continue
        label = _extract_label(block)
        if label:
            table_items.append((match.start(), label))

    table_count = appendix_table_start
    for index, (_, label) in enumerate(sorted(table_items), start=1):
        if appendix_mode:
            table_count += 1
            reference_map[label] = _format_export_sequence(effective_chapter_no, table_count, appendix_mode=True)
        else:
            reference_map[label] = _format_export_sequence(chapter_no, index, appendix_mode=False)

    return reference_map, figure_count, table_count


def _collect_reference_map(
    text: str,
    chapter_no: int,
    *,
    appendix_mode: bool = False,
    appendix_chapter_no: int = 0,
    appendix_figure_start: int = 0,
    appendix_table_start: int = 0,
) -> Dict[str, str]:
    reference_map, _, _ = _collect_reference_map_with_counters(
        text,
        chapter_no,
        appendix_mode=appendix_mode,
        appendix_chapter_no=appendix_chapter_no,
        appendix_figure_start=appendix_figure_start,
        appendix_table_start=appendix_table_start,
    )
    return reference_map


def _render_figure_markdown(
    image_path: Optional[str],
    cap_zh: str,
    cap_en: str,
    *,
    notes: Optional[List[str]] = None,
    project_dir: Path,
    tex_path: Path,
    graphicspaths: List[str],
    state: Dict[str, int],
) -> str:
    if image_path:
        state["figure_no"] = int(state.get("figure_no", 0)) + 1
        if state.get("appendix_mode"):
            state["appendix_figure_no"] = int(state.get("appendix_figure_no", 0)) + 1
            appendix_chapter_no = int(state.get("appendix_chapter_no", 0) or 0)
            sequence = _format_export_sequence(
                appendix_chapter_no,
                int(state["appendix_figure_no"]),
                appendix_mode=True,
            )
            marker_sequence = _format_caption_marker_sequence(
                appendix_chapter_no,
                int(state["appendix_figure_no"]),
                appendix_mode=True,
            )
        else:
            sequence = _format_export_sequence(
                int(state.get("chapter_no", 0)),
                int(state["figure_no"]),
                appendix_mode=False,
            )
            marker_sequence = sequence
        resolved_img = _resolve_graphics_path(image_path, project_dir, tex_path, graphicspaths)
        lines = [f"![](<{resolved_img}>)"]
        if cap_zh:
            lines.append(_build_caption_marker("FIG", "zh", marker_sequence, cap_zh))
        if cap_en:
            lines.append(_build_caption_marker("FIG", "en", marker_sequence, cap_en))
        for note in notes or []:
            lines.append(f"**注：** {note}")
        return "\n\n" + "\n\n".join(lines) + "\n\n"
    if cap_zh:
        cap = cap_zh if not cap_en else f"{cap_zh} / {cap_en}"
        return f"\n\n> 图：{cap}\n\n"
    return "\n\n> 图：内容略\n\n"


def _replace_inline_figure_macros(
    text: str,
    *,
    project_dir: Path,
    tex_path: Path,
    graphicspaths: List[str],
    state: Dict[str, int],
) -> str:
    cursor = 0
    chunks: List[str] = []

    while True:
        match = INLINE_FIGURE_MACRO_USE_RE.search(text, cursor)
        if not match:
            chunks.append(text[cursor:])
            break

        chunks.append(text[cursor : match.start()])
        parsed = _parse_optional_and_braced_args(text, match.end(), required_args=4)
        if not parsed:
            chunks.append(text[match.start() : match.end()])
            cursor = match.end()
            continue

        end, _, args = parsed
        cap_zh = _caption_latex_to_text(args[1])
        cap_en = _caption_latex_to_text(args[2])
        chunks.append(
            _render_figure_markdown(
                args[0].strip(),
                cap_zh,
                cap_en,
                project_dir=project_dir,
                tex_path=tex_path,
                graphicspaths=graphicspaths,
                state=state,
            )
        )
        cursor = end

    return "".join(chunks)


def _replace_bensz_note_macros(text: str) -> str:
    cursor = 0
    chunks: List[str] = []

    while True:
        match = NOTE_MACRO_USE_RE.search(text, cursor)
        if not match:
            chunks.append(text[cursor:])
            break

        chunks.append(text[cursor : match.start()])
        parsed = _parse_optional_and_braced_args(text, match.end(), required_args=1)
        if not parsed:
            chunks.append(text[match.start() : match.end()])
            cursor = match.end()
            continue

        end, _, args = parsed
        note = _caption_latex_to_text(args[0])
        if note:
            chunks.append(f"\n\n**注：** {note}\n\n")
        cursor = end

    return "".join(chunks)


def _extract_bensz_notes(text: str) -> List[str]:
    notes: List[str] = []
    cursor = 0

    while True:
        match = NOTE_MACRO_USE_RE.search(text, cursor)
        if not match:
            break

        parsed = _parse_optional_and_braced_args(text, match.end(), required_args=1)
        if not parsed:
            cursor = match.end()
            continue

        end, _, args = parsed
        note = _caption_latex_to_text(args[0])
        if note:
            notes.append(note)
        cursor = end

    return notes


def _tex_block_to_markdown_text(text: str) -> str:
    text = strip_comments(text)
    text = flatten_texorpdfstring(text)
    text = re.sub(r"\\begin\{(?:itemize|enumerate)\}[^\n]*", "", text)
    text = re.sub(r"\\end\{(?:itemize|enumerate)\}", "", text)
    text = re.sub(r"^[ \t]*\\item[ \t]*", "- ", text, flags=re.M)
    text = text.replace("\\\\", "\n")

    lines: List[str] = []
    for raw in text.splitlines():
        stripped = raw.strip()
        if not stripped:
            lines.append("")
            continue
        if stripped in {"itemize", "enumerate", "\\itemize", "\\enumerate"}:
            continue
        if stripped.startswith("- "):
            payload = _italicize_plain_stat_tokens(latex_inline_to_text(stripped[2:]))
            if payload:
                lines.append(f"- {payload}")
            continue
        payload = _italicize_plain_stat_tokens(latex_inline_to_text(raw))
        if payload in {"itemize", "enumerate"}:
            continue
        if payload:
            lines.append(payload)

    out = "\n".join(lines)
    out = re.sub(r"\n{3,}", "\n\n", out)
    return out.strip()


def _cleanup_markdown_artifacts(text: str) -> str:
    lines: List[str] = []
    placeholder_seen = False
    in_code = False
    in_math = False
    pending_marker_line: Optional[str] = None
    dropping_pt_cluster = False
    marker_only_re = re.compile(r"^\[\d+\]$")
    pt_only_re = re.compile(r"^\d+(?:\.\d+)?\s*pt$", re.IGNORECASE)
    pt_heading_re = re.compile(r"^#\s*\d+(?:\.\d+)?\s*pt$", re.IGNORECASE)

    def _normalize_cjk_latin_spacing(line_text: str) -> str:
        # 中文正文中，字母缩写与中文之间不留空格。
        out = re.sub(r"([\u4e00-\u9fff])\s+([A-Za-z][A-Za-z0-9+\-]*)", r"\1\2", line_text)
        out = re.sub(r"([A-Za-z][A-Za-z0-9+\-]*)\s+([\u4e00-\u9fff])", r"\1\2", out)
        return out

    for raw in text.splitlines():
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code = not in_code
            lines.append(line)
            continue

        if not in_code and stripped == "$$":
            in_math = not in_math
            lines.append(line)
            continue

        if not in_code and not in_math:
            if dropping_pt_cluster:
                if pt_only_re.fullmatch(stripped) or pt_heading_re.fullmatch(stripped):
                    continue
                dropping_pt_cluster = False

            if pending_marker_line is not None:
                if pt_only_re.fullmatch(stripped) or pt_heading_re.fullmatch(stripped):
                    # Drop the whole marker + pt cluster produced by malformed spacing macros.
                    pending_marker_line = None
                    dropping_pt_cluster = True
                    continue
                lines.append(pending_marker_line)
                pending_marker_line = None

            if stripped in {"itemize", "enumerate", "\\itemize", "\\enumerate"}:
                continue
            if marker_only_re.fullmatch(stripped):
                pending_marker_line = line
                continue
            if "$$" in line and stripped != "$$":
                line = re.sub(r"\s*\$\$\s*", " → ", line)
                line = re.sub(r"[ \t]{2,}", " ", line).strip()
                stripped = line.strip()
            if stripped == BIB_PLACEHOLDER_LINE:
                if placeholder_seen:
                    continue
                placeholder_seen = True
            line = _normalize_cjk_latin_spacing(line)

        lines.append(line)

    if pending_marker_line is not None:
        lines.append(pending_marker_line)

    out = "\n".join(lines)
    out = re.sub(r"\n{3,}", "\n\n", out)
    return out.strip() + "\n"


def _scan_markdown_artifacts(markdown_text: str) -> Dict[str, int]:
    bare_list_env_lines = 0
    abnormal_double_dollar_lines = 0
    placeholder_count = 0
    in_code = False
    in_math = False

    for raw in markdown_text.splitlines():
        stripped = raw.strip()

        if stripped.startswith("```"):
            in_code = not in_code
            continue

        if not in_code and stripped == "$$":
            in_math = not in_math
            continue

        if in_code or in_math:
            continue

        if stripped in {"itemize", "enumerate", "\\itemize", "\\enumerate"}:
            bare_list_env_lines += 1
        if "$$" in raw and stripped != "$$":
            abnormal_double_dollar_lines += 1
        if stripped == BIB_PLACEHOLDER_LINE:
            placeholder_count += 1

    return {
        "bare_list_env_lines": bare_list_env_lines,
        "abnormal_double_dollar_lines": abnormal_double_dollar_lines,
        "duplicate_bibliography_placeholder": max(0, placeholder_count - 1),
        "bibliography_placeholder_count": placeholder_count,
    }


def _ensure_target_writable(path: Path) -> None:
    parent = path.parent
    parent.mkdir(parents=True, exist_ok=True)
    with NamedTemporaryFile(prefix=".write-test-", suffix=".tmp", dir=str(parent), delete=True):
        pass
    if path.exists() and path.suffix.lower() == ".docx":
        probe_path = path.with_name(f".lock-probe-{uuid.uuid4().hex}-{path.name}")
        try:
            path.replace(probe_path)
            probe_path.replace(path)
        except PermissionError as exc:
            raise _build_docx_lock_permission_error(exc, path) from exc
        finally:
            if probe_path.exists() and not path.exists():
                probe_path.replace(path)


def _strict_python_hint() -> Optional[str]:
    configured = os.environ.get(STRICT_PYTHON_ENV, "").strip()
    if not configured:
        return None
    candidate = Path(configured).expanduser()
    status = "exists" if candidate.exists() else "not-found"
    return (
        f"可尝试使用环境变量 {STRICT_PYTHON_ENV} 指定的解释器执行：\n"
        f'  "{configured}" projects/thesis-ucas-doctor/scripts/export_docx.py '
        f"--project-dir projects/thesis-ucas-doctor  [{status}]"
    )


def _find_powershell_executable() -> Optional[str]:
    for candidate in ("powershell", "pwsh"):
        path = shutil.which(candidate)
        if path:
            return path
    return None


def _to_word_automation_path(path: Path) -> str:
    raw = str(path.resolve())
    match = re.match(r"^/mnt/([a-zA-Z])/(.+)$", raw)
    if match:
        drive = match.group(1).upper()
        tail = match.group(2).replace("/", "\\")
        return f"{drive}:\\{tail}"
    return raw


def _escape_powershell_single_quoted(text: str) -> str:
    return text.replace("'", "''")


def _run_export_preflight(
    mode: str,
    reference_doc: Path,
    markdown_path: Path,
    output_docx: Path,
    report_path: Optional[Path],
    word_update_fields: bool = False,
) -> None:
    failures: List[str] = []
    pandoc_path = shutil.which("pandoc")
    if not pandoc_path:
        failures.append("未检测到 pandoc（strict/portable 都需要）")
    else:
        proc = subprocess.run([pandoc_path, "--version"], capture_output=True, text=True)
        if proc.returncode != 0:
            failures.append("pandoc --version 执行失败")

    if not reference_doc.exists():
        failures.append(f"参考 Word 模板不存在：{reference_doc}")

    if mode == "strict" and DocxDocument is None:
        msg = "strict 模式需要 python-docx（当前解释器不可导入 docx）"
        hint = _strict_python_hint()
        if hint:
            msg = f"{msg}\n{hint}"
        failures.append(msg)
    if word_update_fields:
        if sys.platform != "win32":
            failures.append("启用 --word-update-fields 需要在 Windows Python 环境执行（以便调用 Word COM）")
        if _find_powershell_executable() is None:
            failures.append("启用 --word-update-fields 需要 powershell 或 pwsh")

    for target in [markdown_path, output_docx]:
        try:
            _ensure_target_writable(target)
        except Exception as exc:
            failures.append(f"目标路径不可写：{target} ({exc})")

    if report_path is not None:
        try:
            _ensure_target_writable(report_path)
        except Exception as exc:
            failures.append(f"质量报告路径不可写：{report_path} ({exc})")

    if failures:
        detail = "\n".join(f"- {item}" for item in failures)
        raise RuntimeError(f"前置检查未通过：\n{detail}")


def _run_postprocess_preflight(mode: str, docx_path: Path, report_path: Path, word_update_fields: bool = False) -> None:
    failures: List[str] = []
    if not docx_path.exists():
        failures.append(f"待处理 docx 不存在：{docx_path}")
    if mode == "strict" and DocxDocument is None:
        msg = "strict 模式需要 python-docx（当前解释器不可导入 docx）"
        hint = _strict_python_hint()
        if hint:
            msg = f"{msg}\n{hint}"
        failures.append(msg)
    if word_update_fields:
        if sys.platform != "win32":
            failures.append("启用 --word-update-fields 需要在 Windows Python 环境执行（以便调用 Word COM）")
        if _find_powershell_executable() is None:
            failures.append("启用 --word-update-fields 需要 powershell 或 pwsh")
    try:
        _ensure_target_writable(report_path)
    except Exception as exc:
        failures.append(f"质量报告路径不可写：{report_path} ({exc})")
    if failures:
        detail = "\n".join(f"- {item}" for item in failures)
        raise RuntimeError(f"前置检查未通过：\n{detail}")


def update_docx_fields_with_word(
    docx_path: Path,
    *,
    visible: bool = False,
    timeout_seconds: int = 300,
) -> Dict[str, object]:
    result: Dict[str, object] = {
        "applied": False,
        "docx_path": str(docx_path),
        "platform": sys.platform,
        "visible": bool(visible),
    }
    if sys.platform != "win32":
        result["reason"] = f"当前平台 {sys.platform} 不支持 Word COM 自动化"
        return result
    powershell = _find_powershell_executable()
    if not powershell:
        result["reason"] = "未检测到 powershell 或 pwsh"
        return result

    with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
        saved_docx_path = Path(tmp.name)
    saved_docx_path.unlink(missing_ok=True)

    docx_word_path = _to_word_automation_path(docx_path)
    saved_docx_word_path = _to_word_automation_path(saved_docx_path)
    ps_docx_path = _escape_powershell_single_quoted(docx_word_path)
    ps_save_path = _escape_powershell_single_quoted(saved_docx_word_path)
    ps_visible = "$true" if visible else "$false"
    script_lines = [
        "$ErrorActionPreference = 'Stop'",
        "$word = $null",
        "$doc = $null",
        "$wdDoNotSaveChanges = 0",
        "$wdFormatXMLDocument = 12",
        f"$docPath = '{ps_docx_path}'",
        f"$savePath = '{ps_save_path}'",
        "if (-not (Test-Path -LiteralPath $docPath)) { throw \"DOCX not found: $docPath\" }",
        "if (Test-Path -LiteralPath $savePath) { Remove-Item -LiteralPath $savePath -Force }",
        "try {",
        "  $word = New-Object -ComObject Word.Application",
        f"  $word.Visible = {ps_visible}",
        "  $word.DisplayAlerts = 0",
        "  try { $word.Options.UpdateLinksAtOpen = $false } catch {}",
        "  $doc = $word.Documents.Open($docPath, $false, $false)",
        "  foreach ($storyRange in $doc.StoryRanges) {",
        "    $cursor = $storyRange",
        "    while ($null -ne $cursor) {",
        "      try { $null = $cursor.Fields.Update() } catch {}",
        "      $cursor = $cursor.NextStoryRange",
        "    }",
        "  }",
        "  try { $null = $doc.Fields.Update() } catch {}",
        "  foreach ($toc in $doc.TablesOfContents) { try { $null = $toc.Update() } catch {} }",
        "  foreach ($tof in $doc.TablesOfFigures) { try { $null = $tof.Update() } catch {} }",
        "  foreach ($idx in $doc.Indexes) { try { $null = $idx.Update() } catch {} }",
        "  $doc.SaveAs2([ref]$savePath, [ref]$wdFormatXMLDocument)",
        "  $doc.Close([ref]$wdDoNotSaveChanges)",
        "  $doc = $null",
        "  $word.Quit([ref]$wdDoNotSaveChanges)",
        "  $word = $null",
        "  if (-not (Test-Path -LiteralPath $savePath)) { throw \"Updated DOCX not saved: $savePath\" }",
        "  Write-Output 'WORD_FIELD_UPDATE_OK'",
        "} catch {",
        "  if ($doc -ne $null) { try { $doc.Close([ref]$wdDoNotSaveChanges) } catch {} }",
        "  if ($word -ne $null) { try { $word.Quit([ref]$wdDoNotSaveChanges) } catch {} }",
        "  throw",
        "}",
    ]
    script = "\n".join(script_lines)
    try:
        proc = subprocess.run(
            [
                powershell,
                "-NoLogo",
                "-NoProfile",
                "-NonInteractive",
                "-ExecutionPolicy",
                "Bypass",
                "-Command",
                script,
            ],
            capture_output=True,
            text=True,
            timeout=max(30, int(timeout_seconds)),
        )
    except subprocess.TimeoutExpired:
        result["reason"] = f"Word 自动更新域超时（>{timeout_seconds}s）"
        saved_docx_path.unlink(missing_ok=True)
        return result

    stdout = (proc.stdout or "").strip()
    stderr = (proc.stderr or "").strip()
    if proc.returncode != 0:
        result["reason"] = stderr or stdout or f"PowerShell 退出码 {proc.returncode}"
        result["returncode"] = proc.returncode
        saved_docx_path.unlink(missing_ok=True)
        return result
    if "WORD_FIELD_UPDATE_OK" not in stdout:
        result["reason"] = "Word 自动化未返回成功标记"
        result["stdout"] = stdout
        saved_docx_path.unlink(missing_ok=True)
        return result
    if not saved_docx_path.exists():
        result["reason"] = f"Word 自动化未生成更新后的 DOCX：{saved_docx_path}"
        return result

    package_check = _validate_docx_package_integrity(saved_docx_path)
    if not package_check["ok"]:
        details: List[str] = []
        for key in [
            "missing_entries",
            "xml_parse_failures",
            "missing_rel_targets",
            "missing_override_parts",
            "missing_rid_refs",
            "unmatched_bookmarks",
            "invalid_mc_ignorable",
        ]:
            items = package_check.get(key, [])
            if items:
                details.append(f"{key}={len(items)}")
        detail_text = ", ".join(details) if details else "unknown package issue"
        result["reason"] = f"Word 自动更新后 DOCX 结构完整性校验失败：{detail_text}"
        saved_docx_path.unlink(missing_ok=True)
        return result

    _replace_docx_with_lock_hint(saved_docx_path, docx_path)

    result["applied"] = True
    result["powershell"] = powershell
    result["docx_word_path"] = docx_word_path
    result["saved_via_temp_docx"] = True
    return result


def repair_docx_with_word_open_and_repair(
    docx_path: Path,
    *,
    visible: bool = False,
    timeout_seconds: int = 300,
) -> Dict[str, object]:
    result: Dict[str, object] = {
        "applied": False,
        "docx_path": str(docx_path),
        "platform": sys.platform,
        "visible": bool(visible),
    }
    if sys.platform != "win32":
        result["reason"] = f"当前平台 {sys.platform} 不支持 Word COM 自动化"
        return result
    powershell = _find_powershell_executable()
    if not powershell:
        result["reason"] = "未检测到 powershell 或 pwsh"
        return result

    with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
        saved_docx_path = Path(tmp.name)
    saved_docx_path.unlink(missing_ok=True)

    docx_word_path = _to_word_automation_path(docx_path)
    saved_docx_word_path = _to_word_automation_path(saved_docx_path)
    ps_docx_path = _escape_powershell_single_quoted(docx_word_path)
    ps_save_path = _escape_powershell_single_quoted(saved_docx_word_path)
    ps_visible = "$true" if visible else "$false"
    script_lines = [
        "$ErrorActionPreference = 'Stop'",
        "$word = $null",
        "$doc = $null",
        "$wdDoNotSaveChanges = 0",
        "$wdFormatXMLDocument = 12",
        "$OpenAndRepair = $true",
        f"$docPath = '{ps_docx_path}'",
        f"$savePath = '{ps_save_path}'",
        "if (-not (Test-Path -LiteralPath $docPath)) { throw \"DOCX not found: $docPath\" }",
        "if (Test-Path -LiteralPath $savePath) { Remove-Item -LiteralPath $savePath -Force }",
        "try {",
        "  $word = New-Object -ComObject Word.Application",
        f"  $word.Visible = {ps_visible}",
        "  $word.DisplayAlerts = 0",
        "  try { $word.Options.UpdateLinksAtOpen = $false } catch {}",
        "  $doc = $word.Documents.Open($docPath, $false, $false, $false, '', '', $false, '', '', 0, 65001, $true, $OpenAndRepair, 0, $true)",
        "  $doc.SaveAs2([ref]$savePath, [ref]$wdFormatXMLDocument)",
        "  $doc.Close([ref]$wdDoNotSaveChanges)",
        "  $doc = $null",
        "  $word.Quit([ref]$wdDoNotSaveChanges)",
        "  $word = $null",
        "  if (-not (Test-Path -LiteralPath $savePath)) { throw \"Repaired DOCX not saved: $savePath\" }",
        "  Write-Output 'WORD_OPEN_AND_REPAIR_OK'",
        "} catch {",
        "  if ($doc -ne $null) { try { $doc.Close([ref]$wdDoNotSaveChanges) } catch {} }",
        "  if ($word -ne $null) { try { $word.Quit([ref]$wdDoNotSaveChanges) } catch {} }",
        "  throw",
        "}",
    ]
    script = "\n".join(script_lines)
    try:
        proc = subprocess.run(
            [
                powershell,
                "-NoLogo",
                "-NoProfile",
                "-NonInteractive",
                "-ExecutionPolicy",
                "Bypass",
                "-Command",
                script,
            ],
            capture_output=True,
            text=True,
            timeout=max(30, int(timeout_seconds)),
        )
    except subprocess.TimeoutExpired:
        result["reason"] = f"Word OpenAndRepair 超时（>{timeout_seconds}s）"
        saved_docx_path.unlink(missing_ok=True)
        return result

    stdout = (proc.stdout or "").strip()
    stderr = (proc.stderr or "").strip()
    if proc.returncode != 0:
        result["reason"] = stderr or stdout or f"PowerShell 退出码 {proc.returncode}"
        result["returncode"] = proc.returncode
        saved_docx_path.unlink(missing_ok=True)
        return result
    if "WORD_OPEN_AND_REPAIR_OK" not in stdout:
        result["reason"] = "Word OpenAndRepair 未返回成功标记"
        result["stdout"] = stdout
        saved_docx_path.unlink(missing_ok=True)
        return result
    if not saved_docx_path.exists():
        result["reason"] = f"Word OpenAndRepair 未生成修复后的 DOCX：{saved_docx_path}"
        return result

    package_check = _validate_docx_package_integrity(saved_docx_path)
    if not package_check["ok"]:
        details: List[str] = []
        for key in [
            "missing_entries",
            "xml_parse_failures",
            "missing_rel_targets",
            "missing_override_parts",
            "missing_content_type_overrides",
            "missing_rid_refs",
            "unmatched_bookmarks",
            "invalid_mc_ignorable",
        ]:
            items = package_check.get(key, [])
            if items:
                details.append(f"{key}={len(items)}")
        detail_text = ", ".join(details) if details else "unknown package issue"
        result["reason"] = f"Word OpenAndRepair 后 DOCX 结构完整性校验失败：{detail_text}"
        saved_docx_path.unlink(missing_ok=True)
        return result

    _replace_docx_with_lock_hint(saved_docx_path, docx_path)

    result["applied"] = True
    result["powershell"] = powershell
    result["docx_word_path"] = docx_word_path
    result["saved_via_temp_docx"] = True
    return result


def convert_abstract_file(text: str) -> str:
    text = strip_comments(text)
    abs_args = parse_macro_args(text, "abstract", 2)
    kw_args = parse_macro_args(text, "keywords", 2)

    lines: List[str] = []
    if abs_args:
        lines.append("# 摘  要")
        lines.append("")
        lines.append(_tex_block_to_markdown_text(_normalize_simple_stat_inline_math(abs_args[0])))
        lines.append("")

    if kw_args:
        lines.append(f"关键词：{latex_inline_to_text(kw_args[0])}")
        lines.append("")

    if abs_args:
        lines.append("# Abstract")
        lines.append("")
        lines.append(_tex_block_to_markdown_text(_normalize_simple_stat_inline_math(abs_args[1])))
        lines.append("")

    if kw_args:
        lines.append(f"Key Words: {latex_inline_to_text(kw_args[1])}")
        lines.append("")

    return "\n".join(lines).strip() + "\n"


def convert_info_file(text: str) -> str:
    text = strip_comments(text)
    title = parse_macro_args(text, "title", 2) or ["", ""]
    degree_level = parse_macro_args(text, "degreeLevel", 1) or [""]
    author = parse_macro_args(text, "author", 2) or ["", ""]
    supervisor = parse_macro_args(text, "supervisor", 3) or ["", "", ""]
    degree_type = parse_macro_args(text, "degreeType", 2) or ["", ""]
    subject = parse_macro_args(text, "subject", 2) or ["", ""]
    institute = parse_macro_args(text, "institute", 2) or ["", ""]
    grad_year = parse_macro_args(text, "gradYear", 1) or [""]
    grad_month = parse_macro_args(text, "gradMonth", 2) or ["", ""]

    lines = [
        "# 论文信息",
        "",
        f"- 中文题目：{latex_inline_to_text(title[0])}",
        f"- 英文题目：{latex_inline_to_text(title[1])}",
        f"- 学位层次：{latex_inline_to_text(degree_level[0])}",
        f"- 作者：{latex_inline_to_text(author[0])} / {latex_inline_to_text(author[1])}",
        f"- 导师：{latex_inline_to_text(supervisor[0])} / {latex_inline_to_text(supervisor[1])}",
        f"- 导师单位：{latex_inline_to_text(supervisor[2])}",
        f"- 学位类别：{latex_inline_to_text(degree_type[0])} / {latex_inline_to_text(degree_type[1])}",
        f"- 学科专业：{latex_inline_to_text(subject[0])} / {latex_inline_to_text(subject[1])}",
        f"- 培养单位：{latex_inline_to_text(institute[0])} / {latex_inline_to_text(institute[1])}",
        f"- 毕业时间：{latex_inline_to_text(grad_year[0])}-{latex_inline_to_text(grad_month[0])}",
        "",
    ]
    return "\n".join(lines)


def _parse_info_fields(text: str) -> Dict[str, List[str]]:
    text = strip_comments(text)
    return {
        "title": parse_macro_args(text, "title", 2) or ["", ""],
        "degree_level": parse_macro_args(text, "degreeLevel", 1) or [""],
        "author": parse_macro_args(text, "author", 2) or ["", ""],
        "supervisor": parse_macro_args(text, "supervisor", 3) or ["", "", ""],
        "degree_type": parse_macro_args(text, "degreeType", 2) or ["", ""],
        "subject": parse_macro_args(text, "subject", 2) or ["", ""],
        "institute": parse_macro_args(text, "institute", 2) or ["", ""],
        "grad_year": parse_macro_args(text, "gradYear", 1) or [""],
        "grad_month": parse_macro_args(text, "gradMonth", 2) or ["", ""],
    }


def _clean_info_value(value: str) -> str:
    text = latex_inline_to_text(value or "")
    text = text.replace("\\\\", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _strip_markdown_emphasis_markers(text: str) -> str:
    """移除 markdown 风格强调标记，避免其泄漏到页眉/页脚纯文本区域。"""
    if not text:
        return ""
    out = text
    out = re.sub(r"\*([^*\n]+)\*", r"\1", out)
    out = re.sub(r"_([^_\n]+)_", r"\1", out)
    return out


def replace_env_blocks(text: str, env_name: str, replacer) -> str:
    pattern = re.compile(
        r"\\begin\{" + re.escape(env_name) + r"\}(?:\[[^\]]*\])?(.*?)\\end\{" + re.escape(env_name) + r"\}",
        re.S,
    )
    return pattern.sub(lambda m: replacer(m.group(1)), text)


def _render_landscape_block(
    block: str,
    *,
    project_dir: Path,
    tex_path: Path,
    graphicspaths: List[str],
    state: Dict[str, int],
) -> str:
    """将 landscape 环境中可识别的图块转换为标准 Markdown 图像。

    说明：
    - 技术路线图等内容常写在 landscape 环境中而非 figure 环境。
    - 若检测到 includegraphics + 图题语义，则按图处理；
      否则原样返回块体内容，交给后续 table/longtable 处理。
    """
    cap_zh = ""
    cap_en = ""
    bi_args = parse_macro_args(block, "bicaption", 2)
    if bi_args:
        cap_zh = _caption_latex_to_text(bi_args[0])
        cap_en = _caption_latex_to_text(bi_args[1])
    else:
        cap_arg = _extract_caption_arg(block)
        if cap_arg:
            cap_zh = _caption_latex_to_text(cap_arg)

    image_path = _parse_includegraphics_path(block)
    has_figure_hint = "\\captionsetup{type=figure" in block or bool(cap_zh) or bool(cap_en)
    if image_path and has_figure_hint:
        notes = _extract_bensz_notes(block)
        return _render_figure_markdown(
            image_path,
            cap_zh,
            cap_en,
            notes=notes,
            project_dir=project_dir,
            tex_path=tex_path,
            graphicspaths=graphicspaths,
            state=state,
        )

    return block


def convert_body_tex(
    text: str,
    tex_path: Path,
    project_dir: Path,
    state: Dict[str, object],
    global_reference_map: Optional[Dict[str, str]] = None,
) -> str:
    text = strip_comments(text)
    text = flatten_texorpdfstring(text)
    text = re.sub(r"^[ \t]*\\setlength[^\n]*$", "", text, flags=re.M)
    text = re.sub(r"^[ \t]*\\renewcommand[^\n]*$", "", text, flags=re.M)
    text = re.sub(r"^[ \t]*\\newcolumntype[^\n]*$", "", text, flags=re.M)
    text = re.sub(r"\\setlength\{[^{}]*\}\{[^{}]*\}", "", text)
    text = re.sub(r"\\renewcommand\{[^{}]*\}\{[^{}]*\}", "", text)
    text = re.sub(r"\\newcolumntype\{[^{}]*\}(?:\[[^\]]*\])?\{[^{}]*\}", "", text)
    text = re.sub(r"\\(?:begingroup|endgroup)\b", "", text)
    text = re.sub(r"\\(?:tiny|scriptsize|footnotesize|small|normalsize|large|Large|LARGE|huge|Huge)\b", "", text)
    text = re.sub(r"\\thispagestyle\{[^{}]*\}", "", text)
    text = re.sub(r"\\special\{[^{}]*\}", "", text)
    text = _strip_inline_figure_macro_definitions(text)
    graphicspaths = _extract_graphicspaths(text)
    text = _strip_macro_calls(text, "graphicspath")
    reference_map = global_reference_map or _collect_reference_map(
        text,
        int(state.get("chapter_no", 0)),
        appendix_mode=bool(state.get("appendix_mode")),
        appendix_figure_start=int(state.get("appendix_figure_no", 0)),
        appendix_table_start=int(state.get("appendix_table_no", 0)),
    )
    text = _replace_cross_references(text, reference_map)
    text = re.sub(r"\\label\{[^{}]*\}", "", text)
    text = _normalize_simple_stat_inline_math(text)
    text = _replace_inline_figure_macros(
        text,
        project_dir=project_dir,
        tex_path=tex_path,
        graphicspaths=graphicspaths,
        state=state,
    )
    text = replace_env_blocks(
        text,
        "landscape",
        lambda block: _render_landscape_block(
            block,
            project_dir=project_dir,
            tex_path=tex_path,
            graphicspaths=graphicspaths,
            state=state,
        ),
    )

    text = re.sub(r"\\mintinline\{[^{}]+\}\{([^{}]*)\}", lambda m: f"`{m.group(1)}`", text)
    text = re.sub(
        r"\\mint\{([^{}]+)\}\{([^{}]*)\}",
        lambda m: f"\n\n> 代码（{m.group(1)}）：`{m.group(2)}`\n\n",
        text,
    )
    text = re.sub(
        r"\\inputminted(?:\[[^\]]*\])?\{([^{}]+)\}\{([^{}]+)\}",
        lambda m: f"\n\n```{m.group(1)}\n[代码来自 {m.group(2)}]\n```\n\n",
        text,
    )

    def _equation_repl(block: str) -> str:
        block = re.sub(r"\\label\{[^{}]*\}", "", block)
        block = block.replace("\\begin{aligned}", "")
        block = block.replace("\\end{aligned}", "")
        block = block.strip()
        if not block:
            return ""
        return f"\n\n$$\n{block}\n$$\n\n"

    def _display_math_repl(m: re.Match[str]) -> str:
        block = m.group(1).strip()
        if not block:
            return ""
        return f"\n\n$$\n{block}\n$$\n\n"

    text = replace_env_blocks(text, "equation", _equation_repl)
    text = re.sub(r"\\\[(.*?)\\\]", _display_math_repl, text, flags=re.S)

    def _figure_repl(block: str) -> str:
        cap_zh = ""
        cap_en = ""
        bi_args = parse_macro_args(block, "bicaption", 2)
        if bi_args:
            cap_zh = _caption_latex_to_text(bi_args[0])
            cap_en = _caption_latex_to_text(bi_args[1])
        else:
            cap_arg = _extract_caption_arg(block)
            if cap_arg:
                cap_zh = _caption_latex_to_text(cap_arg)

        img = _parse_includegraphics_path(block)
        notes = _extract_bensz_notes(block)
        return _render_figure_markdown(
            img,
            cap_zh,
            cap_en,
            notes=notes,
            project_dir=project_dir,
            tex_path=tex_path,
            graphicspaths=graphicspaths,
            state=state,
        )

    text = replace_env_blocks(text, "figure", _figure_repl)

    def _caption_only(prefix: str):
        def _inner(block: str) -> str:
            bi_args = parse_macro_args(block, "bicaption", 2)
            if bi_args:
                cap = _caption_latex_to_text(bi_args[0])
            else:
                cap_arg = _extract_caption_arg(block)
                cap = _caption_latex_to_text(cap_arg) if cap_arg else ""
            if cap:
                return f"\n\n> {prefix}：{cap}（需在 Word 中人工整理）\n\n"
            return f"\n\n> {prefix}：内容略（需在 Word 中人工整理）\n\n"

        return _inner

    text = replace_env_blocks(text, "table", lambda block: _render_table_markdown(block, state=state))
    text = replace_env_blocks(
        text,
        "longtable",
        lambda block: _render_table_markdown(f"\\begin{{longtable}}{block}\\end{{longtable}}", state=state),
    )
    text = replace_env_blocks(text, "algorithm", _caption_only("算法"))
    text = replace_env_blocks(text, "listing", _caption_only("代码"))
    text = _replace_inline_table_macros(text, state=state)
    text = MINIPAGE_ENV_RE.sub(
        lambda m: _render_table_markdown(m.group(1), state=state)
        if "\\captionsetup{type=table}" in m.group(1) and re.search(r"\\begin\{tabularx?\}", m.group(1))
        else m.group(1),
        text,
    )
    text = _replace_bensz_note_macros(text)
    # Keep PDF-side TOC helper commands from leaking into Word正文（如 toccapter/tocsection）。
    text = re.sub(r"\\phantomsection\b", "", text)
    text = re.sub(
        r"\\addcontentsline\s*\{[^{}]*\}\s*\{[^{}]*\}\s*\{[^{}]*\}",
        "",
        text,
    )

    text = re.sub(r"\\begin\{(?:tabularx?|longtable|refsection|enumerate|itemize|aligned|algorithmic)\}[^\n]*", "", text)
    text = re.sub(r"\\end\{(?:tabularx?|longtable|refsection|enumerate|itemize|aligned|algorithmic)\}", "", text)
    text = re.sub(r"\\nocite\{[^{}]*\}", "", text)
    text = re.sub(r"\\printbibliography(?:\[[^\]]*\])?", f"\n\n{BIB_REFS_DIV}\n\n", text)
    text = text.replace("\\acknowledgementsDate", "")

    heading_map = {
        "chapter": "#",
        "chapter*": "#",
        "section": "##",
        "section*": "##",
        "subsection": "###",
        "subsection*": "###",
        "subsubsection": "###",
    }

    for cmd, h in heading_map.items():
        pattern = re.compile(r"\\" + re.escape(cmd) + r"\{([^{}]*)\}")

        def _heading_repl(match: re.Match[str], command: str = cmd, heading: str = h) -> str:
            title = match.group(1)
            if command == "chapter" and bool(state.get("appendix_mode")):
                appendix_index = int(state.get("appendix_chapter_no", 0) or 0)
                title_text = _format_appendix_chapter_heading(title, appendix_index)
            else:
                title_text = latex_inline_to_text(title)
            return f"\n\n{heading} {title_text}\n\n"

        text = pattern.sub(_heading_repl, text)

    text = re.sub(r"^[ \t]*\\item[ \t]*", "- ", text, flags=re.M)
    text = text.replace("\\\\", "\n")

    new_lines = []
    in_math = False
    in_code = False
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code = not in_code
            new_lines.append(line)
            continue

        if stripped == "$$":
            in_math = not in_math
            new_lines.append(line)
            continue

        if in_code or in_math:
            new_lines.append(line.rstrip())
            continue

        new_lines.append(latex_inline_to_text(line))

    text = "\n".join(new_lines)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return _cleanup_markdown_artifacts(text)


def parse_main_includes(main_tex: str) -> List[str]:
    includes: List[str] = []
    for m in re.finditer(r"\\(?:input|include)\{([^{}]+)\}", main_tex):
        name = m.group(1).strip()
        if not name.endswith(".tex"):
            name += ".tex"
        includes.append(name)
    return includes


def parse_main_appendix_includes(main_tex: str) -> Set[str]:
    appendix_includes: Set[str] = set()
    in_appendix = False
    pattern = re.compile(r"\\appendix\b|\\(?:input|include)\{([^{}]+)\}")
    for match in pattern.finditer(main_tex):
        if match.group(0).startswith(r"\appendix"):
            in_appendix = True
            continue
        if not in_appendix:
            continue
        name = (match.group(1) or "").strip()
        if not name.endswith(".tex"):
            name += ".tex"
        if Path(name).stem.lower().startswith("appendix"):
            appendix_includes.add(name)
    return appendix_includes


def render_markdown(project_dir: Path, tex_file: Path) -> str:
    main_text = tex_file.read_text(encoding="utf-8")
    include_files = parse_main_includes(main_text)
    appendix_include_files = parse_main_appendix_includes(main_text)
    global_reference_map: Dict[str, str] = {}
    scan_chapter_no = 0
    scan_appendix_chapter_no = 0
    scan_appendix_figure_no = 0
    scan_appendix_table_no = 0

    skip_names = {"config-pre.tex"}
    for rel in include_files:
        if Path(rel).name in skip_names:
            continue
        p = (project_dir / rel).resolve()
        if not p.exists():
            continue
        raw = p.read_text(encoding="utf-8")
        if p.name in {"info.tex", "abstract.tex"}:
            continue
        appendix_mode = rel in appendix_include_files or _is_appendix_tex(raw)
        if re.search(r"\\chapter\*?\{", raw):
            scan_chapter_no += 1
            if appendix_mode:
                scan_appendix_chapter_no += 1
        local_reference_map, scan_appendix_figure_no, scan_appendix_table_no = _collect_reference_map_with_counters(
            strip_comments(flatten_texorpdfstring(raw)),
            scan_chapter_no,
            appendix_mode=appendix_mode,
            appendix_chapter_no=scan_appendix_chapter_no,
            appendix_figure_start=scan_appendix_figure_no,
            appendix_table_start=scan_appendix_table_no,
        )
        global_reference_map.update(local_reference_map)

    state: Dict[str, object] = {
        "chapter_no": 0,
        "figure_no": 0,
        "table_no": 0,
        "appendix_mode": False,
        "appendix_chapter_no": 0,
        "appendix_figure_no": 0,
        "appendix_table_no": 0,
    }

    chunks: List[str] = [
        "# 学位论文（LaTeX 源转换稿）",
        "",
        "说明：本文件由 LaTeX 源自动转换，复杂表格/算法/代码块可能需要人工微调。",
        "",
    ]

    inserted_reference_heading = False
    for rel in include_files:
        if Path(rel).name in skip_names:
            continue
        p = (project_dir / rel).resolve()
        if not p.exists():
            chunks.append(f"> 警告：未找到文件 {rel}")
            chunks.append("")
            continue

        raw = p.read_text(encoding="utf-8")
        if p.name == "info.tex":
            chunks.append(convert_info_file(raw))
        elif p.name == "abstract.tex":
            chunks.append(convert_abstract_file(raw))
            chunks.extend(
                [
                    "# 目录",
                    "",
                    "> 目录域请在 Word 中右键更新。",
                    "",
                    "# 图表目录",
                    "",
                    "## 图目录",
                    "",
                    "> 图目录请在 Word 中右键更新。",
                    "",
                    "## 表目录",
                    "",
                    "> 表目录请在 Word 中右键更新。",
                    "",
                ]
            )
        else:
            is_current_appendix = rel in appendix_include_files or _is_appendix_tex(raw)
            if not inserted_reference_heading and (
                is_current_appendix or p.name in {"acknowledgements.tex", "cv.tex"}
            ):
                chunks.extend(
                    [
                        "# 参考文献",
                        "",
                        BIB_REFS_DIV,
                        "",
                    ]
                )
                inserted_reference_heading = True
            state["appendix_mode"] = is_current_appendix
            if re.search(r"\\chapter\*?\{", raw):
                state["chapter_no"] += 1
                state["figure_no"] = 0
                state["table_no"] = 0
                if bool(state["appendix_mode"]):
                    state["appendix_chapter_no"] = int(state.get("appendix_chapter_no", 0) or 0) + 1
            chunks.append(convert_body_tex(raw, p, project_dir, state, global_reference_map=global_reference_map))
        chunks.append("")

    if not inserted_reference_heading:
        chunks.extend(
            [
                "# 参考文献",
                "",
                BIB_REFS_DIV,
                "",
            ]
        )

    markdown = "\n".join(chunks).strip() + "\n"
    return _cleanup_markdown_artifacts(markdown)


_BIBTEX_REFERENCE_LOCATOR_FIELDS = {
    "pages",
    "eid",
    "article-number",
    "article_number",
    "articleno",
    "elocation-id",
    "elocationid",
}


def _is_doi_url(value: str) -> bool:
    normalized = (value or "").strip().casefold()
    return normalized.startswith(("https://doi.org/", "http://doi.org/", "https://dx.doi.org/", "http://dx.doi.org/"))


def _has_bibtex_reference_locator(fields: Dict[str, str]) -> bool:
    for name in _BIBTEX_REFERENCE_LOCATOR_FIELDS:
        if (fields.get(name) or "").strip():
            return True
    return False


def _remove_bibtex_field(body: str, field_name: str) -> str:
    field_re = re.compile(
        rf"(?ms)^[ \t]*{re.escape(field_name)}\s*=\s*"
        r"(?:\{(?:[^{}]|\{[^{}]*\})*\}|\"(?:[^\"\\]|\\.)*\"|[^,\n\r]*)"
        r"\s*,?[ \t]*(?:\r?\n|$)"
    )
    return field_re.sub("", body)


def _append_bibtex_field(body: str, field_name: str, value: str) -> str:
    suffix = "\n" if body.endswith("\n") else ""
    stripped = body.rstrip()
    if stripped and not stripped.endswith(","):
        stripped += ","
    return f"{stripped}\n  {field_name} = {{{value}}},{suffix}"


def _is_official_document_entry(entry_type: str, fields: Dict[str, str]) -> bool:
    if entry_type not in {"misc", "online", "report"}:
        return False
    marker_text = " ".join(
        fields.get(name, "")
        for name in ("type", "howpublished", "note", "title", "publisher", "institution")
    )
    official_markers = (
        "国务院文件",
        "部门通知",
        "官方公报",
        "国家标准",
        "行业标准",
        "环境保护行业标准",
        "生态环境行业标准",
        "农业行业标准",
        "林业行业标准",
        "供销合作行业标准",
        "中华人民共和国国家标准",
    )
    return any(marker in marker_text for marker in official_markers)


def _sanitize_bibtex_for_docx(text: str) -> str:
    """Remove noisy locators from temporary bibliography used by DOCX export.

    The source ``references.bib`` keeps DOI and official URLs for traceability.
    Word output uses this sanitized copy to keep the bibliography compact.
    """
    source = strip_comments(text)
    chunks: List[str] = []
    idx = 0

    while True:
        at = source.find("@", idx)
        if at < 0:
            chunks.append(source[idx:])
            break
        brace = source.find("{", at)
        if brace < 0:
            chunks.append(source[idx:])
            break
        comma = source.find(",", brace + 1)
        if comma < 0:
            chunks.append(source[idx:])
            break

        depth = 1
        end = comma + 1
        while end < len(source) and depth:
            if source[end] == "{":
                depth += 1
            elif source[end] == "}":
                depth -= 1
            end += 1
        if depth != 0:
            chunks.append(source[idx:])
            break

        chunks.append(source[idx:at])
        entry_type = source[at + 1 : brace].strip().casefold()
        body = source[comma + 1 : end - 1]
        fields = _parse_bibtex_fields(body)

        body = _remove_bibtex_field(body, "doi")
        body = _remove_bibtex_field(body, "url")
        body = _remove_bibtex_field(body, "urldate")
        fields = _parse_bibtex_fields(body)

        if "langid" not in fields and _CJK_CHAR_RE.search(
            " ".join(fields.get(name, "") for name in ("author", "title", "journal", "booktitle", "publisher"))
        ):
            body = _append_bibtex_field(body, "langid", "chinese")

        chunks.append(source[at : comma + 1])
        chunks.append(body)
        chunks.append(source[end - 1 : end])
        idx = end

    return "".join(chunks)


def run_pandoc(
    markdown_path: Path,
    output_docx: Path,
    reference_doc: Path,
    project_dir: Path,
    bibliography: Iterable[Path],
) -> None:
    with TemporaryDirectory(prefix="thesis-docx-bib-") as tmp_dir:
        tmp_root = Path(tmp_dir)
        sanitized_bibs: List[Path] = []
        for bib in bibliography:
            if not bib.exists():
                continue
            cleaned = _sanitize_bibtex_for_docx(bib.read_text(encoding="utf-8"))
            sanitized_path = tmp_root / bib.name
            sanitized_path.write_text(cleaned, encoding="utf-8")
            sanitized_bibs.append(sanitized_path)

        with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(output_docx.parent)) as tmp:
            pandoc_output = Path(tmp.name)
        try:
            cmd = [
                "pandoc",
                str(markdown_path),
                "--from=markdown+tex_math_dollars+raw_tex+fenced_divs-superscript-subscript",
                "--to=docx",
                "--standalone",
                f"--reference-doc={reference_doc}",
                f"--resource-path={project_dir}",
                "--citeproc",
                f"--csl={DEFAULT_PANDOC_CSL}",
                "-o",
                str(pandoc_output),
            ]

            for bib in sanitized_bibs:
                cmd.append(f"--bibliography={bib}")

            proc = subprocess.run(cmd, capture_output=True, text=True, cwd=project_dir)
            if proc.returncode != 0:
                raise RuntimeError(
                    "pandoc 转换失败\n"
                    f"cmd: {' '.join(cmd)}\n"
                    f"stdout:\n{proc.stdout}\n"
                    f"stderr:\n{proc.stderr}\n"
                )
            _replace_docx_with_lock_hint(pandoc_output, output_docx)
        finally:
            pandoc_output.unlink(missing_ok=True)


_BIB_FIELD_NAME_RE = re.compile(r"[A-Za-z][A-Za-z0-9_-]*")
_CJK_LANGIDS = {"chinese", "zh", "zh-cn", "zh_cn", "chinese-journal"}


def _strip_outer_braces(text: str) -> str:
    value = (text or "").strip()
    while value.startswith("{") and value.endswith("}"):
        depth = 0
        balanced_outer = True
        for idx, char in enumerate(value):
            if char == "{":
                depth += 1
            elif char == "}":
                depth -= 1
                if depth == 0 and idx != len(value) - 1:
                    balanced_outer = False
                    break
        if not balanced_outer or depth != 0:
            break
        value = value[1:-1].strip()
    return value


def _parse_bibtex_value(body: str, start: int) -> Tuple[str, int]:
    idx = start
    while idx < len(body) and body[idx].isspace():
        idx += 1
    if idx >= len(body):
        return "", idx

    if body[idx] == "{":
        depth = 0
        chunks: List[str] = []
        while idx < len(body):
            char = body[idx]
            if char == "{":
                depth += 1
                if depth > 1:
                    chunks.append(char)
            elif char == "}":
                depth -= 1
                if depth == 0:
                    idx += 1
                    break
                chunks.append(char)
            else:
                chunks.append(char)
            idx += 1
        return "".join(chunks).strip(), idx

    if body[idx] == '"':
        idx += 1
        chunks = []
        depth = 0
        while idx < len(body):
            char = body[idx]
            if char == "{":
                depth += 1
            elif char == "}":
                depth = max(0, depth - 1)
            elif char == '"' and depth == 0:
                idx += 1
                break
            chunks.append(char)
            idx += 1
        return "".join(chunks).strip(), idx

    end = idx
    while end < len(body) and body[end] not in ",\n\r":
        end += 1
    return body[idx:end].strip(), end


def _parse_bibtex_fields(body: str) -> Dict[str, str]:
    fields: Dict[str, str] = {}
    idx = 0
    while idx < len(body):
        match = _BIB_FIELD_NAME_RE.search(body, idx)
        if not match:
            break
        name = match.group(0).casefold()
        idx = match.end()
        while idx < len(body) and body[idx].isspace():
            idx += 1
        if idx >= len(body) or body[idx] != "=":
            continue
        value, idx = _parse_bibtex_value(body, idx + 1)
        fields[name] = value
        while idx < len(body) and body[idx] != ",":
            if not body[idx].isspace():
                break
            idx += 1
        if idx < len(body) and body[idx] == ",":
            idx += 1
    return fields


def _iter_bibtex_entries(text: str) -> Iterable[Tuple[str, Dict[str, str]]]:
    source = strip_comments(text)
    idx = 0
    while True:
        at = source.find("@", idx)
        if at < 0:
            break
        brace = source.find("{", at)
        if brace < 0:
            break
        comma = source.find(",", brace + 1)
        if comma < 0:
            break
        key = source[brace + 1 : comma].strip()
        depth = 1
        end = comma + 1
        while end < len(source) and depth:
            if source[end] == "{":
                depth += 1
            elif source[end] == "}":
                depth -= 1
            end += 1
        if key and depth == 0:
            yield key, _parse_bibtex_fields(source[comma + 1 : end - 1])
        idx = max(end, comma + 1)


def _split_bibtex_names(author_field: str) -> List[str]:
    names: List[str] = []
    depth = 0
    start = 0
    idx = 0
    source = author_field or ""
    while idx < len(source):
        char = source[idx]
        if char == "{":
            depth += 1
        elif char == "}":
            depth = max(0, depth - 1)
        elif depth == 0 and source.startswith(" and ", idx):
            name = source[start:idx].strip()
            if name:
                names.append(name)
            idx += 5
            start = idx
            continue
        idx += 1
    tail = source[start:].strip()
    if tail:
        names.append(tail)
    return names


def _clean_bibtex_person_name(name: str) -> str:
    value = _strip_outer_braces(name)
    previous = None
    while previous != value:
        previous = value
        value = re.sub(r"\\[A-Za-z@]+\*?\{([^{}]*)\}", r"\1", value)
    value = latex_inline_to_text(value)
    value = value.replace("\\", "")
    value = value.replace("{", "").replace("}", "")
    value = re.sub(r"\s+", " ", value).strip()
    if "," in value and not _CJK_CHAR_RE.search(value):
        parts = [part.strip() for part in value.split(",") if part.strip()]
        if parts:
            return parts[0]
    return value


def _is_chinese_bib_entry(fields: Dict[str, str]) -> bool:
    lang = (fields.get("langid") or fields.get("language") or "").strip().casefold()
    if lang in _CJK_LANGIDS:
        return True
    probe = " ".join([fields.get("author", ""), fields.get("title", "")])
    return bool(_CJK_CHAR_RE.search(probe))


def _collect_chinese_citation_replacements(bibliography: Iterable[Path]) -> List[Tuple[re.Pattern[str], str]]:
    replacements: List[Tuple[re.Pattern[str], str]] = []
    seen: Set[Tuple[str, str, str]] = set()

    for bib in bibliography:
        if not bib.exists():
            continue
        for _key, fields in _iter_bibtex_entries(bib.read_text(encoding="utf-8")):
            if not _is_chinese_bib_entry(fields):
                continue
            names = _split_bibtex_names(fields.get("author", ""))
            if len(names) < 2:
                continue
            first_author = _clean_bibtex_person_name(names[0])
            second_author = _clean_bibtex_person_name(names[1])
            year = (fields.get("year") or fields.get("date") or "").strip()
            year_match = re.search(r"\d{4}", year)
            if not first_author or not year_match:
                continue
            year_value = year_match.group(0)
            if len(names) == 2 and second_author:
                cache_key = (first_author, second_author, year_value)
                if cache_key not in seen:
                    seen.add(cache_key)
                    pattern = re.compile(
                        rf"(?<![A-Za-z0-9_]){re.escape(first_author)}\s+and\s+"
                        rf"{re.escape(second_author)}\s+"
                        rf"({re.escape(year_value)}[a-z]?)"
                        rf"(?![A-Za-z0-9_])"
                    )
                    replacements.append((pattern, rf"{first_author}和{second_author}，\1"))
                    pattern = re.compile(
                        rf"(?<![A-Za-z0-9_]){re.escape(first_author)}\s+等[，,]\s*"
                        rf"({re.escape(year_value)}[a-z]?)"
                        rf"(?![A-Za-z0-9_])"
                    )
                    replacements.append((pattern, rf"{first_author}和{second_author}，\1"))
                continue

            bad_author_variants = {first_author}
            cjk_match = _CJK_CHAR_RE.search(first_author)
            if cjk_match:
                bad_author_variants.add(cjk_match.group(0))

            for bad_author in bad_author_variants:
                cache_key = (bad_author, first_author, year_value)
                if cache_key in seen:
                    continue
                seen.add(cache_key)
                pattern = re.compile(
                    rf"(?<![A-Za-z0-9_]){re.escape(bad_author)}\s+et\s+al\.\s+"
                    rf"({re.escape(year_value)}[a-z]?)"
                    rf"(?![A-Za-z0-9_])"
                )
                replacements.append((pattern, rf"{first_author}等，\1"))
                pattern = re.compile(
                    rf"(?<![A-Za-z0-9_]){re.escape(bad_author)}\s+等[，,]\s*"
                    rf"({re.escape(year_value)}[a-z]?)"
                    rf"(?![A-Za-z0-9_])"
                )
                replacements.append((pattern, rf"{first_author}等，\1"))

    # Longer author strings must run first, otherwise the single-character
    # fallback for rare-glyph names could consume a full author name.
    replacements.sort(key=lambda item: len(item[0].pattern), reverse=True)
    return replacements


def _apply_text_replacements(text: str, replacements: Iterable[Tuple[re.Pattern[str], str]]) -> Tuple[str, int]:
    updated = text
    changed = 0
    for pattern, replacement in replacements:
        updated, count = pattern.subn(replacement, updated)
        changed += count
    return updated, changed


def _collect_text_replacement_matches(
    text: str,
    replacements: Iterable[Tuple[re.Pattern[str], str]],
) -> List[Tuple[int, int, str]]:
    matches: List[Tuple[int, int, str]] = []
    occupied: List[Tuple[int, int]] = []
    for pattern, replacement in replacements:
        for match in pattern.finditer(text):
            start, end = match.span()
            if any(not (end <= used_start or start >= used_end) for used_start, used_end in occupied):
                continue
            matches.append((start, end, match.expand(replacement)))
            occupied.append((start, end))
    matches.sort(key=lambda item: item[0])
    return matches


def _find_text_node_at_offset(
    spans: List[Tuple[int, int, ET.Element]],
    offset: int,
) -> Optional[Tuple[int, int, int, ET.Element]]:
    for idx, (start, end, node) in enumerate(spans):
        if start <= offset < end:
            return idx, start, end, node
    if spans and offset == spans[-1][1]:
        idx = len(spans) - 1
        start, end, node = spans[idx]
        return idx, start, end, node
    return None


def _replace_text_across_paragraph_nodes(
    para: ET.Element,
    replacements: Iterable[Tuple[re.Pattern[str], str]],
) -> int:
    text_nodes = [node for node in para.findall(".//w:t", W) if node.text]
    if not text_nodes:
        return 0

    spans: List[Tuple[int, int, ET.Element]] = []
    cursor = 0
    chunks: List[str] = []
    for node in text_nodes:
        node_text = node.text or ""
        chunks.append(node_text)
        spans.append((cursor, cursor + len(node_text), node))
        cursor += len(node_text)

    full_text = "".join(chunks)
    matches = _collect_text_replacement_matches(full_text, replacements)
    if not matches:
        return 0

    changed = 0
    for start, end, replacement in reversed(matches):
        start_info = _find_text_node_at_offset(spans, start)
        end_info = _find_text_node_at_offset(spans, end - 1)
        if start_info is None or end_info is None:
            continue
        start_idx, start_node_start, _start_node_end, start_node = start_info
        end_idx, end_node_start, _end_node_end, end_node = end_info
        start_text = start_node.text or ""
        end_text = end_node.text or ""
        prefix = start_text[: start - start_node_start]
        suffix = end_text[end - end_node_start :]

        if start_idx == end_idx:
            start_node.text = prefix + replacement + suffix
        else:
            start_node.text = prefix + replacement
            for idx in range(start_idx + 1, end_idx):
                spans[idx][2].text = ""
            end_node.text = suffix

        for node in {start_node, end_node}:
            node_text = node.text or ""
            if node_text[:1].isspace() or node_text[-1:].isspace():
                node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        changed += 1

    return changed


def _localize_chinese_intext_citations(docx_path: Path, bibliography: Iterable[Path]) -> Dict[str, object]:
    replacements = _collect_chinese_citation_replacements(bibliography)
    if not replacements:
        return {"applied": False, "updated_text_nodes": 0, "replacements": 0}

    updated_paragraphs = 0
    replacement_count = 0
    samples: List[str] = []

    with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
        tmp_path = Path(tmp.name)

    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
            tmp_path, "w", compression=zipfile.ZIP_DEFLATED
        ) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    root = ET.fromstring(data)
                    body = root.find("w:body", W)
                    if body is not None:
                        in_bibliography = False
                        for para in body.findall("w:p", W):
                            para_text = _extract_para_text(para).strip()
                            para_norm = _norm_plain_text(para_text)
                            if para_norm == _norm_plain_text("参考文献"):
                                in_bibliography = True
                                continue
                            if in_bibliography and para_norm in _BIBLIOGRAPHY_STOP_MARKERS:
                                in_bibliography = False
                            if in_bibliography:
                                continue

                            count = _replace_text_across_paragraph_nodes(para, replacements)
                            if count <= 0:
                                continue
                            updated_paragraphs += 1
                            replacement_count += count
                            if len(samples) < 8:
                                samples.append(_extract_para_text(para).strip()[:80])
                    if replacement_count:
                        data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

        _replace_docx_with_lock_hint(tmp_path, docx_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)

    return {
        "applied": replacement_count > 0,
        "updated_text_nodes": updated_paragraphs,
        "updated_paragraphs": updated_paragraphs,
        "replacements": replacement_count,
        "samples": samples,
    }


_ENGLISH_INTEXT_CITATION_TERM_REPLACEMENTS: List[Tuple[re.Pattern[str], str]] = [
    (
        re.compile(
            r"\b([A-Z][A-Za-z]*(?:\s+[A-Z]){0,3}(?:,\s+[A-Z][A-Za-z]*(?:\s+[A-Z]){0,3})*,?)\s+等[，,]\s*(\d{4}[a-z]?)"
        ),
        r"\1 et al., \2",
    ),
]


def _localize_english_intext_citations(docx_path: Path) -> Dict[str, object]:
    updated_paragraphs = 0
    replacements = 0
    samples: List[str] = []

    with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
        tmp_path = Path(tmp.name)

    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
            tmp_path, "w", compression=zipfile.ZIP_DEFLATED
        ) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    root = ET.fromstring(data)
                    body = root.find("w:body", W)
                    if body is not None:
                        in_bibliography = False
                        for para in body.findall("w:p", W):
                            para_text = _extract_para_text(para).strip()
                            para_norm = _norm_plain_text(para_text)
                            if para_norm == _norm_plain_text("参考文献"):
                                in_bibliography = True
                                continue
                            if in_bibliography:
                                continue

                            count = _replace_text_across_paragraph_nodes(
                                para,
                                _ENGLISH_INTEXT_CITATION_TERM_REPLACEMENTS,
                            )
                            if count <= 0:
                                continue
                            updated_paragraphs += 1
                            replacements += count
                            if len(samples) < 8:
                                samples.append(_extract_para_text(para).strip()[:120])
                    if replacements:
                        data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

        _replace_docx_with_lock_hint(tmp_path, docx_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)

    return {
        "applied": replacements > 0,
        "updated_paragraphs": updated_paragraphs,
        "replacements": replacements,
        "samples": samples,
    }


_CHINESE_BIBLIOGRAPHY_AUTHOR_TERM_REPLACEMENTS: List[Tuple[re.Pattern[str], str]] = [
    (re.compile(r",\s+et\s+al\.", re.IGNORECASE), ", 等."),
    (re.compile(r"\bet\s+al\.", re.IGNORECASE), "等."),
    (re.compile(r",\s+and\s+", re.IGNORECASE), "和"),
]

_CHINESE_BIBLIOGRAPHY_ENTRY_REPLACEMENTS: List[Tuple[re.Pattern[str], str]] = [
    (
        re.compile(
            r"\s*https?://(?:oversea\.)?cnki\.net/\S+|\s*https?://kns\.cnki\.net/\S+|\s*https?://kns\.cnki\.net/kcms2/\S+",
            re.IGNORECASE,
        ),
        "",
    ),
    (
        re.compile(r",\s+no\.\s+([0-9A-Za-z\-–—]+)(?=[:.,;])", re.IGNORECASE),
        r", (\1)",
    ),
]


def _is_chinese_bibliography_paragraph(text: str) -> bool:
    probe = re.sub(r"[等和卷版]", "", text or "")
    return bool(_CJK_CHAR_RE.search(probe))


_ENGLISH_BIBLIOGRAPHY_TERM_REPLACEMENTS: List[Tuple[re.Pattern[str], str]] = [
    (re.compile(r",\s*等\s*,"), ", et al.,"),
    (re.compile(r",\s*等\."), ", et al."),
    (re.compile(r"(?<=[:\s])卷\s+([0-9A-Za-z][0-9A-Za-z\-–—]*)"), r"vol. \1"),
    (re.compile(r"\b([0-9]+(?:st|nd|rd|th))\s+版(?=[.。])", re.IGNORECASE), r"\1 ed"),
]


def _is_english_bibliography_paragraph(text: str) -> bool:
    if not re.search(r"[A-Za-z]", text or ""):
        return False
    probe = re.sub(r"[等和卷版第册]", "", text or "")
    return not _CJK_CHAR_RE.search(probe)


def _localize_english_bibliography_terms(docx_path: Path) -> Dict[str, object]:
    updated_paragraphs = 0
    replacements = 0
    samples: List[str] = []

    with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
        tmp_path = Path(tmp.name)

    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
            tmp_path, "w", compression=zipfile.ZIP_DEFLATED
        ) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    root = ET.fromstring(data)
                    body = root.find("w:body", W)
                    if body is not None:
                        in_bibliography = False
                        for para in body.findall("w:p", W):
                            para_text = _extract_para_text(para).strip()
                            para_norm = _norm_plain_text(para_text)
                            if para_norm == _norm_plain_text("参考文献"):
                                in_bibliography = True
                                continue
                            if in_bibliography and para_norm in _BIBLIOGRAPHY_STOP_MARKERS:
                                in_bibliography = False
                            if not in_bibliography or not _is_english_bibliography_paragraph(para_text):
                                continue

                            count = _replace_text_across_paragraph_nodes(
                                para,
                                _ENGLISH_BIBLIOGRAPHY_TERM_REPLACEMENTS,
                            )
                            if count <= 0:
                                continue
                            updated_paragraphs += 1
                            replacements += count
                            if len(samples) < 8:
                                samples.append(_extract_para_text(para).strip()[:120])
                    if replacements:
                        data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

        _replace_docx_with_lock_hint(tmp_path, docx_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)

    return {
        "applied": replacements > 0,
        "updated_paragraphs": updated_paragraphs,
        "replacements": replacements,
        "samples": samples,
    }


def _localize_chinese_bibliography_author_terms(docx_path: Path) -> Dict[str, object]:
    updated_paragraphs = 0
    replacements = 0
    samples: List[str] = []

    with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
        tmp_path = Path(tmp.name)

    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
            tmp_path, "w", compression=zipfile.ZIP_DEFLATED
        ) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    root = ET.fromstring(data)
                    body = root.find("w:body", W)
                    if body is not None:
                        in_bibliography = False
                        for para in body.findall("w:p", W):
                            para_text = _extract_para_text(para).strip()
                            para_norm = _norm_plain_text(para_text)
                            if para_norm == _norm_plain_text("参考文献"):
                                in_bibliography = True
                                continue
                            if in_bibliography and para_norm in _BIBLIOGRAPHY_STOP_MARKERS:
                                in_bibliography = False
                            if not in_bibliography or not _is_chinese_bibliography_paragraph(para_text):
                                continue

                            count = _replace_text_across_paragraph_nodes(
                                para,
                                _CHINESE_BIBLIOGRAPHY_AUTHOR_TERM_REPLACEMENTS,
                            )
                            if count <= 0:
                                continue
                            updated_paragraphs += 1
                            replacements += count
                            if len(samples) < 8:
                                samples.append(_extract_para_text(para).strip()[:100])
                    if replacements:
                        data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

        _replace_docx_with_lock_hint(tmp_path, docx_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)

    return {
        "applied": replacements > 0,
        "updated_paragraphs": updated_paragraphs,
        "replacements": replacements,
        "samples": samples,
    }


def _normalize_chinese_bibliography_entries(docx_path: Path) -> Dict[str, object]:
    updated_paragraphs = 0
    replacements = 0
    samples: List[str] = []

    with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
        tmp_path = Path(tmp.name)

    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
            tmp_path, "w", compression=zipfile.ZIP_DEFLATED
        ) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    root = ET.fromstring(data)
                    body = root.find("w:body", W)
                    if body is not None:
                        in_bibliography = False
                        for para in body.findall("w:p", W):
                            para_text = _extract_para_text(para).strip()
                            para_norm = _norm_plain_text(para_text)
                            if para_norm == _norm_plain_text("参考文献"):
                                in_bibliography = True
                                continue
                            if in_bibliography and para_norm in _BIBLIOGRAPHY_STOP_MARKERS:
                                in_bibliography = False
                            if not in_bibliography or not _is_chinese_bibliography_paragraph(para_text):
                                continue

                            count = _replace_text_across_paragraph_nodes(
                                para,
                                _CHINESE_BIBLIOGRAPHY_ENTRY_REPLACEMENTS,
                            )
                            if count <= 0:
                                continue

                            count += _replace_text_across_paragraph_nodes(
                                para,
                                [
                                    (re.compile(r"\s+\."), "."),
                                    (re.compile(r"\.\."), "."),
                                    (re.compile(r"\s{2,}"), " "),
                                ],
                            )
                            updated_paragraphs += 1
                            replacements += count
                            if len(samples) < 8:
                                samples.append(_extract_para_text(para).strip()[:100])
                    if replacements:
                        data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

        _replace_docx_with_lock_hint(tmp_path, docx_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)

    return {
        "applied": replacements > 0,
        "updated_paragraphs": updated_paragraphs,
        "replacements": replacements,
        "samples": samples,
    }


def _clone_xml(elem: ET.Element) -> ET.Element:
    return ET.fromstring(ET.tostring(elem, encoding="utf-8"))


def _norm_heading_key(text: str) -> str:
    return re.sub(r"[\s\u3000]+", "", (text or "")).casefold()


def _extract_para_text(para: ET.Element) -> str:
    return "".join((t.text or "") for t in para.findall(".//w:t", W)).strip()


def _extract_para_style_id(para: ET.Element) -> str:
    ppr = para.find("w:pPr", W)
    if ppr is None:
        return ""
    pstyle = ppr.find("w:pStyle", W)
    if pstyle is None:
        return ""
    return pstyle.get(_w_attr("val"), "")


def _is_heading1_paragraph(para: ET.Element) -> bool:
    sid = _extract_para_style_id(para)
    sid_norm = re.sub(r"[\s_-]+", "", sid.casefold())
    return sid_norm in {"1", "heading1"}


def _is_toc_entry_paragraph(para: ET.Element) -> bool:
    sid = _extract_para_style_id(para)
    sid_norm = re.sub(r"[\s_-]+", "", sid.casefold())
    return bool(re.fullmatch(r"toc\d*", sid_norm))


def _match_marker(text_norm: str, marker_norm: str) -> bool:
    if not text_norm or not marker_norm:
        return False
    strict_exact = {"摘要", "abstract", "目录", "图表目录", "绪论", "参考文献", "致谢"}
    if marker_norm in strict_exact:
        return text_norm == marker_norm
    return text_norm == marker_norm or text_norm.startswith(marker_norm)


def _collect_template_sections(reference_doc_root: ET.Element) -> List[ET.Element]:
    body = reference_doc_root.find("w:body", W)
    if body is None:
        return []

    sections: List[ET.Element] = []
    for child in list(body):
        if child.tag == _w_attr("p"):
            ppr = child.find("w:pPr", W)
            if ppr is None:
                continue
            sect = ppr.find("w:sectPr", W)
            if sect is not None:
                sections.append(_clone_xml(sect))
        elif child.tag == _w_attr("sectPr"):
            sections.append(_clone_xml(child))
    return sections


def _build_section_plan(output_doc_root: ET.Element) -> List[Tuple[int, int, str]]:
    body = output_doc_root.find("w:body", W)
    if body is None:
        return []

    paragraphs = [p for p in list(body) if p.tag == _w_attr("p")]
    if not paragraphs:
        return []

    plan: List[Tuple[int, int, str]] = [(0, 1, "doc-start")]
    last_idx = 0
    for marker, template_idx in SECTION_MARKERS_WITH_TEMPLATE_INDEX:
        marker_norm = _norm_heading_key(marker)
        found_idx: Optional[int] = None
        for i in range(last_idx + 1, len(paragraphs)):
            para = paragraphs[i]
            if _is_toc_entry_paragraph(para):
                continue
            text_norm = _norm_heading_key(_extract_para_text(para))
            if not _match_marker(text_norm, marker_norm):
                continue
            is_h1 = _is_heading1_paragraph(para)
            allow_non_h1 = marker_norm in {
                _norm_heading_key("摘要"),
                _norm_heading_key("abstract"),
                _norm_heading_key("目录"),
                _norm_heading_key("图表目录"),
                _norm_heading_key("参考文献"),
                _norm_heading_key("附录"),
                _norm_heading_key("致谢"),
                _norm_heading_key("作者简历及攻读学位期间发表的学术论文与其他相关学术成果"),
            }
            if is_h1 or allow_non_h1:
                found_idx = i
                break
        if found_idx is None:
            continue
        plan.append((found_idx, template_idx, marker))
        last_idx = found_idx
    return plan


def _set_paragraph_section(para: ET.Element, sect_pr: ET.Element) -> None:
    ppr = para.find("w:pPr", W)
    if ppr is None:
        ppr = ET.Element(_w_attr("pPr"))
        para.insert(0, ppr)
    old = ppr.find("w:sectPr", W)
    if old is not None:
        ppr.remove(old)
    ppr.append(_clone_xml(sect_pr))


def _collect_sect_rel_ids(sections: List[ET.Element]) -> List[str]:
    rel_ids: List[str] = []
    for sect in sections:
        for ref in sect.findall("w:headerReference", W) + sect.findall("w:footerReference", W):
            rid = ref.get(_r_attr("id"), "")
            if rid:
                rel_ids.append(rid)
    return rel_ids


_SECTION_PROPERTY_ORDER = {
    _w_attr(name): idx
    for idx, name in enumerate(
        [
            "headerReference",
            "footerReference",
            "footnotePr",
            "endnotePr",
            "type",
            "pgSz",
            "pgMar",
            "paperSrc",
            "pgBorders",
            "lnNumType",
            "pgNumType",
            "cols",
            "formProt",
            "vAlign",
            "noEndnote",
            "titlePg",
            "textDirection",
            "bidi",
            "rtlGutter",
            "docGrid",
            "printerSettings",
            "sectPrChange",
        ]
    )
}


_GENERATED_MAIN_HEADER_START = 37
_LEGACY_GENERATED_MAIN_HEADER_START = 200


def _normalize_section_properties_order(sect_pr: ET.Element) -> None:
    children = list(sect_pr)
    ordered = sorted(
        enumerate(children),
        key=lambda item: (_SECTION_PROPERTY_ORDER.get(item[1].tag, 10_000), item[0]),
    )
    if [child for _, child in ordered] == children:
        return
    for child in children:
        sect_pr.remove(child)
    for _, child in ordered:
        sect_pr.append(child)


def _sync_docx_section_layout(docx_path: Path, reference_doc: Path, mode: str) -> Dict[str, object]:
    with zipfile.ZipFile(reference_doc, "r") as zref:
        ref_doc_root = ET.fromstring(zref.read("word/document.xml"))
        ref_rels_root = ET.fromstring(zref.read("word/_rels/document.xml.rels"))
        ref_rel_by_id = {
            rel.get("Id", ""): rel
            for rel in ref_rels_root.findall("{http://schemas.openxmlformats.org/package/2006/relationships}Relationship")
        }
        template_sections = _collect_template_sections(ref_doc_root)

        if not template_sections:
            raise RuntimeError("参考模板未找到 section 定义（word/document.xml）")

        tmp_path: Optional[Path] = None
        with zipfile.ZipFile(docx_path, "r") as zin:
            out_doc_root = ET.fromstring(zin.read("word/document.xml"))
            out_rels_root = ET.fromstring(zin.read("word/_rels/document.xml.rels"))
            plan = _build_section_plan(out_doc_root)
            if len(plan) <= 1:
                if mode == "strict":
                    raise RuntimeError("strict 模式未识别到可用于分节映射的标题锚点")
                return {
                    "applied": False,
                    "section_count": len(plan),
                    "template_section_count": len(template_sections),
                    "anchors": [],
                }

            body = out_doc_root.find("w:body", W)
            if body is None:
                raise RuntimeError("导出 docx 缺少 word/document.xml 的 w:body")

            paragraphs = [p for p in list(body) if p.tag == _w_attr("p")]
            if not paragraphs:
                raise RuntimeError("导出 docx 正文为空（无段落）")

            for para in paragraphs:
                ppr = para.find("w:pPr", W)
                if ppr is None:
                    continue
                old = ppr.find("w:sectPr", W)
                if old is not None:
                    ppr.remove(old)

            old_body_sect = body.find("w:sectPr", W)
            if old_body_sect is not None:
                body.remove(old_body_sect)

            assignments: List[Tuple[int, ET.Element]] = []
            used_sections: List[ET.Element] = []
            for i in range(len(plan) - 1):
                start_idx, template_idx, _ = plan[i]
                next_start_idx, _, _ = plan[i + 1]
                end_idx = next_start_idx - 1
                if end_idx < start_idx:
                    continue
                if template_idx < 1 or template_idx > len(template_sections):
                    raise RuntimeError(f"模板 section 下标越界: {template_idx}")
                sect = _clone_xml(template_sections[template_idx - 1])
                used_sections.append(sect)
                assignments.append((end_idx, sect))

            last_template_idx = plan[-1][1]
            if last_template_idx < 1 or last_template_idx > len(template_sections):
                raise RuntimeError(f"模板 section 下标越界: {last_template_idx}")
            final_section = _clone_xml(template_sections[last_template_idx - 1])
            used_sections.append(final_section)

            rel_ids = _collect_sect_rel_ids(used_sections)
            rid_map, required_targets = _merge_required_relationships(rel_ids, ref_rel_by_id, out_rels_root)
            _remap_relationship_ids(used_sections, rid_map)
            _restore_markup_compatibility(out_doc_root, ref_doc_root)

            for end_idx, sect in assignments:
                _set_paragraph_section(paragraphs[end_idx], sect)
            body.append(final_section)

            required_part_files = _collect_related_parts_from_targets(
                zref,
                source_part="word/document.xml",
                targets=required_targets,
            )
            with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
                tmp_path = Path(tmp.name)

            with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    if item.filename in required_part_files:
                        continue
                    if item.filename == "word/document.xml":
                        data = _serialize_xml(out_doc_root)
                        zout.writestr(item, data)
                    elif item.filename == "word/_rels/document.xml.rels":
                        data = _serialize_xml(out_rels_root, default_namespace=PKG_REL_NS)
                        zout.writestr(item, data)
                    else:
                        zout.writestr(item, zin.read(item.filename))

                for part in sorted(required_part_files):
                    if part in zref.namelist():
                        zout.writestr(part, zref.read(part))

        if tmp_path is not None:
            try:
                _replace_docx_with_lock_hint(tmp_path, docx_path)
            finally:
                if tmp_path.exists():
                    tmp_path.unlink(missing_ok=True)

    return {
        "applied": True,
        "section_count": len(plan),
        "template_section_count": len(template_sections),
        "anchors": [{"index": idx, "template_section": sec, "marker": marker} for idx, sec, marker in plan],
    }


def _build_center_header_xml(text: str, style_id: str = "a3") -> bytes:
    root = ET.Element(_w_attr("hdr"))
    para = ET.SubElement(root, _w_attr("p"))
    ppr = ET.SubElement(para, _w_attr("pPr"))
    pstyle = ET.SubElement(ppr, _w_attr("pStyle"))
    pstyle.set(_w_attr("val"), style_id)
    ind = ET.SubElement(ppr, _w_attr("ind"))
    ind.set(_w_attr("firstLine"), "0")
    ind.set(_w_attr("firstLineChars"), "0")
    ind.set(_w_attr("left"), "0")
    ind.set(_w_attr("leftChars"), "0")
    ind.set(_w_attr("right"), "0")
    ind.set(_w_attr("rightChars"), "0")
    ind.set(_w_attr("hanging"), "0")
    ind.set(_w_attr("hangingChars"), "0")
    jc = ET.SubElement(ppr, _w_attr("jc"))
    jc.set(_w_attr("val"), "center")

    run = ET.SubElement(para, _w_attr("r"))
    text_node = ET.SubElement(run, _w_attr("t"))
    if text[:1].isspace() or text[-1:].isspace():
        text_node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_node.text = text
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _build_center_styleref_header_xml(
    fallback_text: str,
    instr_text: str,
    style_id: str = "a3",
) -> bytes:
    instr_text = (instr_text or "").strip()
    if not instr_text:
        instr_text = 'STYLEREF "heading 1" \\* MERGEFORMAT'
    if "MERGEFORMAT" not in instr_text or not instr_text.startswith("STYLEREF"):
        instr_text = 'STYLEREF "heading 1" \\* MERGEFORMAT'

    root = ET.Element(_w_attr("hdr"))
    para = ET.SubElement(root, _w_attr("p"))
    ppr = ET.SubElement(para, _w_attr("pPr"))
    pstyle = ET.SubElement(ppr, _w_attr("pStyle"))
    pstyle.set(_w_attr("val"), style_id)
    ind = ET.SubElement(ppr, _w_attr("ind"))
    ind.set(_w_attr("firstLine"), "0")
    ind.set(_w_attr("firstLineChars"), "0")
    ind.set(_w_attr("left"), "0")
    ind.set(_w_attr("leftChars"), "0")
    ind.set(_w_attr("right"), "0")
    ind.set(_w_attr("rightChars"), "0")
    ind.set(_w_attr("hanging"), "0")
    ind.set(_w_attr("hangingChars"), "0")
    jc = ET.SubElement(ppr, _w_attr("jc"))
    jc.set(_w_attr("val"), "center")

    run_begin = ET.SubElement(para, _w_attr("r"))
    fld_begin = ET.SubElement(run_begin, _w_attr("fldChar"))
    fld_begin.set(_w_attr("fldCharType"), "begin")

    run_instr = ET.SubElement(para, _w_attr("r"))
    instr = ET.SubElement(run_instr, _w_attr("instrText"))
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = f" {instr_text} "

    run_sep = ET.SubElement(para, _w_attr("r"))
    fld_sep = ET.SubElement(run_sep, _w_attr("fldChar"))
    fld_sep.set(_w_attr("fldCharType"), "separate")

    run_text = ET.SubElement(para, _w_attr("r"))
    text_node = ET.SubElement(run_text, _w_attr("t"))
    text_node.text = fallback_text

    run_end = ET.SubElement(para, _w_attr("r"))
    fld_end = ET.SubElement(run_end, _w_attr("fldChar"))
    fld_end.set(_w_attr("fldCharType"), "end")

    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _paragraphs_inside_toc_field(paragraphs: List[ET.Element]) -> List[bool]:
    mask: List[bool] = []
    in_toc_field = False
    pending_field_begin = False
    for para in paragraphs:
        para_in_toc = in_toc_field
        for node in para.iter():
            if node.tag == _w_attr("fldChar"):
                fld_type = node.get(_w_attr("fldCharType"), "")
                if fld_type == "begin":
                    pending_field_begin = True
                elif fld_type == "end":
                    if in_toc_field:
                        para_in_toc = True
                    in_toc_field = False
                    pending_field_begin = False
                continue
            if node.tag == _w_attr("instrText"):
                instr = re.sub(r"\s+", " ", node.text or "").strip()
                if "TOC" in instr and (pending_field_begin or in_toc_field):
                    in_toc_field = True
                    para_in_toc = True
        mask.append(para_in_toc)
    return mask


def _first_heading1_text(paragraphs: List[ET.Element], start_idx: int, end_idx: int) -> str:
    if start_idx < 0:
        start_idx = 0
    end_idx = min(end_idx, len(paragraphs) - 1)
    fallback = ""
    for idx in range(start_idx, end_idx + 1):
        para = paragraphs[idx]
        if _is_toc_entry_paragraph(para):
            continue
        text = _extract_para_text(para).strip()
        if not text:
            continue
        if _is_heading1_paragraph(para):
            return text
        if fallback:
            continue
        compact = _norm_plain_text(text)
        if compact in {
            _norm_plain_text("摘要"),
            _norm_plain_text("abstract"),
            _norm_plain_text("目录"),
            _norm_plain_text("图表目录"),
            _norm_plain_text("参考文献"),
            _norm_plain_text("致谢"),
            _norm_plain_text("作者简历及攻读学位期间发表的学术论文与其他相关学术成果"),
        }:
            fallback = text
            continue
        if compact.startswith(_norm_plain_text("附录")):
            fallback = text
    return fallback


def _classify_section_header_role(heading_text: str) -> str:
    compact = _norm_plain_text(heading_text)
    if not compact:
        return ""
    if compact == _norm_plain_text("摘要"):
        return "abstract_cn"
    if compact == _norm_plain_text("abstract"):
        return "abstract_en"
    if compact == _norm_plain_text("目录"):
        return "toc"
    if compact in {
        _norm_plain_text("图表目录"),
        _norm_plain_text("图目录"),
        _norm_plain_text("表目录"),
    }:
        return "catalog"
    if compact == _norm_plain_text("参考文献"):
        return "bibliography"
    if compact == _norm_plain_text("致谢"):
        return "ack"
    if compact.startswith(_norm_plain_text("作者简历及攻读学位期间发表的学术论文与其他相关学术成果")):
        return "cv"
    if compact.startswith(_norm_plain_text("附录")) or _is_appendix_heading_text(heading_text):
        return "appendix"
    return "main"


def _normalize_docx_headers_by_rules(docx_path: Path, project_dir: Path) -> Dict[str, object]:
    info_path = project_dir / "extraTex" / "info.tex"
    if not info_path.exists():
        return {"applied": False, "updated_sections": 0, "reason": f"missing info.tex: {info_path}"}

    info = _parse_info_fields(info_path.read_text(encoding="utf-8"))
    zh_title = _strip_markdown_emphasis_markers(_clean_info_value(info["title"][0])) or "论文题目"
    en_title = _strip_markdown_emphasis_markers(_clean_info_value(info["title"][1])) or zh_title

    tmp_path: Optional[Path] = None
    with zipfile.ZipFile(docx_path, "r") as zin:
        out_doc_root = ET.fromstring(zin.read("word/document.xml"))
        out_rels_root = ET.fromstring(zin.read("word/_rels/document.xml.rels"))
        settings_root = ET.fromstring(zin.read("word/settings.xml"))
        content_types_root = ET.fromstring(zin.read("[Content_Types].xml"))

        body = out_doc_root.find("w:body", W)
        if body is None:
            return {"applied": False, "updated_sections": 0, "reason": "missing w:body"}
        paragraphs = [p for p in list(body) if p.tag == _w_attr("p")]
        if not paragraphs:
            return {"applied": False, "updated_sections": 0, "reason": "empty body"}
        toc_field_mask = _paragraphs_inside_toc_field(paragraphs)

        def _first_text_anchor_index(marker: str) -> Optional[int]:
            marker_norm = _norm_plain_text(marker)
            for idx, para in enumerate(paragraphs):
                if idx < len(toc_field_mask) and toc_field_mask[idx]:
                    continue
                if _is_toc_entry_paragraph(para):
                    continue
                style_norm = re.sub(r"[\s_-]+", "", _paragraph_style_id(para).casefold())
                if style_norm.startswith("toc"):
                    continue
                text = _extract_para_text(para).strip()
                if not text:
                    continue
                text_norm = _norm_plain_text(text)
                if marker == "图表目录":
                    if text_norm in {_norm_plain_text("图表目录"), _norm_plain_text("图目录")}:
                        return idx
                    continue
                if marker == "附录":
                    if text_norm.startswith(_norm_plain_text("附录")) or _is_appendix_heading_text(text):
                        return idx
                    continue
                if text_norm == marker_norm:
                    return idx
            return None

        def _paragraph_style_id(para: ET.Element) -> str:
            ppr = para.find("w:pPr", W)
            if ppr is None:
                return ""
            pstyle = ppr.find("w:pStyle", W)
            if pstyle is None:
                return ""
            return pstyle.get(_w_attr("val"), "")

        def _insert_break_before_paragraph(anchor_idx: int) -> bool:
            if anchor_idx <= 0 or anchor_idx >= len(paragraphs):
                return False
            break_idx = anchor_idx - 1
            break_para = paragraphs[break_idx]
            break_ppr = break_para.find("w:pPr", W)
            if break_ppr is not None and break_ppr.find("w:sectPr", W) is not None:
                return False

            template_sect: Optional[ET.Element] = None
            for idx in range(anchor_idx, len(paragraphs)):
                ppr = paragraphs[idx].find("w:pPr", W)
                if ppr is None:
                    continue
                sect = ppr.find("w:sectPr", W)
                if sect is not None:
                    template_sect = sect
                    break
            if template_sect is None:
                template_sect = body.find("w:sectPr", W)
            if template_sect is None:
                return False

            if _is_toc_entry_paragraph(break_para):
                empty_break_para = ET.Element(_w_attr("p"))
                empty_break_ppr = ET.SubElement(empty_break_para, _w_attr("pPr"))
                empty_break_ppr.append(_clone_xml(template_sect))
                body.insert(list(body).index(paragraphs[anchor_idx]), empty_break_para)
                paragraphs.insert(anchor_idx, empty_break_para)
                return True

            if break_ppr is None:
                break_ppr = ET.Element(_w_attr("pPr"))
                break_para.insert(0, break_ppr)
            break_ppr.append(_clone_xml(template_sect))
            return True

        inserted_breaks = 0
        split_anchors = ["目录", "图表目录", "绪论", "附录"]
        for anchor in split_anchors:
            anchor_idx = _first_text_anchor_index(anchor)
            if anchor_idx is not None and _insert_break_before_paragraph(anchor_idx):
                inserted_breaks += 1

        # Give each chapter-level heading its own section so odd-page headers can be
        # written as stable plain text instead of relying on Word STYLEREF fields.
        heading1_indices = [
            idx
            for idx, para in enumerate(paragraphs)
            if not (idx < len(toc_field_mask) and toc_field_mask[idx])
            and re.sub(r"[\s_-]+", "", _paragraph_style_id(para).casefold()) in {"1", "heading1"}
            and _extract_para_text(para).strip()
        ]
        for anchor_idx in heading1_indices:
            if _insert_break_before_paragraph(anchor_idx):
                inserted_breaks += 1

        sections: List[Dict[str, object]] = []
        start_idx = 0
        for idx, para in enumerate(paragraphs):
            ppr = para.find("w:pPr", W)
            if ppr is None:
                continue
            sect = ppr.find("w:sectPr", W)
            if sect is None:
                continue
            sections.append({"sect": sect, "start": start_idx, "end": idx})
            start_idx = idx + 1
        body_sect = body.find("w:sectPr", W)
        if body_sect is not None:
            sections.append({"sect": body_sect, "start": start_idx, "end": len(paragraphs) - 1})
        if not sections:
            return {"applied": False, "updated_sections": 0, "reason": "no sections"}

        main_heading_fallback = "第1章"
        section_roles: List[str] = []
        heading_snapshots: List[str] = []
        seen_abstract = False
        last_role = ""
        appendix_mode = False
        for sec in sections:
            start_idx = int(sec.get("start", 0))
            end_idx = int(sec.get("end", len(paragraphs) - 1))
            visible_paragraphs = [
                para
                for idx, para in enumerate(paragraphs)
                if start_idx <= idx <= end_idx
                and not (idx < len(toc_field_mask) and toc_field_mask[idx])
                and not re.sub(r"[\s_-]+", "", _paragraph_style_id(para).casefold()).startswith("toc")
            ]
            heading = _first_heading1_text(
                visible_paragraphs,
                0,
                len(visible_paragraphs) - 1,
            ) if visible_paragraphs else ""
            heading_snapshots.append(heading)
            role = _classify_section_header_role(heading)
            if role == "main" and heading and main_heading_fallback == "第1章":
                main_heading_fallback = heading
            if role in {"abstract_cn", "abstract_en"}:
                seen_abstract = True
            if not seen_abstract:
                section_roles.append("")
                continue
            if not role:
                role = last_role
            if appendix_mode and role in {"", "main"}:
                role = "appendix"
            if not role:
                role = "main"
            if role == "appendix":
                appendix_mode = True
            elif role in {"bibliography", "ack", "cv"}:
                appendix_mode = False
            section_roles.append(role)
            last_role = role

        rel_nodes = out_rels_root.findall(f"{{{PKG_REL_NS}}}Relationship")
        existing_ids: Set[str] = {rel.get("Id", "") for rel in rel_nodes}
        package_part_names = set(zin.namelist())
        stale_generated_header_parts = set()
        for name in package_part_names:
            if re.fullmatch(r"word/header_bensz_main_\d+\.xml", name):
                stale_generated_header_parts.add(name)
                continue
            match = re.fullmatch(r"word/header(\d+)\.xml", name)
            if match and int(match.group(1)) >= _LEGACY_GENERATED_MAIN_HEADER_START:
                stale_generated_header_parts.add(name)

        def _resolve_target_name(target: str) -> str:
            return _resolve_relationship_target("word/document.xml", target)

        if stale_generated_header_parts:
            for rel in list(rel_nodes):
                if rel.get("Type", "") != DOCX_HEADER_REL_TYPE:
                    continue
                if _resolve_target_name(rel.get("Target", "")) not in stale_generated_header_parts:
                    continue
                out_rels_root.remove(rel)
                rel_nodes.remove(rel)
            existing_ids = {rel.get("Id", "") for rel in rel_nodes}
            for override in list(content_types_root.findall(f"{{{CT_NS}}}Override")):
                part_name = _normalize_part_name(override.get("PartName", ""))
                if part_name in stale_generated_header_parts:
                    content_types_root.remove(override)

        def _ensure_header_rel_id(target_name: str) -> str:
            target_resolved = _resolve_target_name(target_name)
            for rel in rel_nodes:
                if rel.get("Type", "") != DOCX_HEADER_REL_TYPE:
                    continue
                rel_target = _resolve_target_name(rel.get("Target", ""))
                if rel_target == target_resolved:
                    rid = rel.get("Id", "")
                    if rid:
                        return rid
            rid = _next_available_rid(existing_ids)
            rel = ET.Element(f"{{{PKG_REL_NS}}}Relationship")
            rel.set("Id", rid)
            rel.set("Type", DOCX_HEADER_REL_TYPE)
            rel.set("Target", target_name)
            out_rels_root.append(rel)
            rel_nodes.append(rel)
            existing_ids.add(rid)
            return rid

        def _ensure_even_default_footer_refs(sect: ET.Element) -> None:
            footer_refs = sect.findall("w:footerReference", W)
            rid_by_type = {
                ref.get(_w_attr("type"), ""): ref.get(_r_attr("id"), "")
                for ref in footer_refs
                if ref.get(_r_attr("id"), "")
            }
            children = list(sect)
            if "default" in rid_by_type and "even" not in rid_by_type and footer_refs:
                even_ref = ET.Element(_w_attr("footerReference"))
                even_ref.set(_w_attr("type"), "even")
                even_ref.set(_r_attr("id"), rid_by_type["default"])
                sect.insert(children.index(footer_refs[0]), even_ref)
            elif "even" in rid_by_type and "default" not in rid_by_type and footer_refs:
                default_ref = ET.Element(_w_attr("footerReference"))
                default_ref.set(_w_attr("type"), "default")
                default_ref.set(_r_attr("id"), rid_by_type["even"])
                insert_after = max(children.index(ref) for ref in footer_refs)
                sect.insert(insert_after + 1, default_ref)

        main_instr_text = 'STYLEREF "heading 1" \\* MERGEFORMAT'
        if "word/header11.xml" in zin.namelist():
            try:
                main_hdr_root = ET.fromstring(zin.read("word/header11.xml"))
                for node in main_hdr_root.findall(".//w:instrText", W):
                    text = (node.text or "").strip()
                    if (
                        text
                        and "MERGEFORMAT" in text
                        and text.startswith("STYLEREF")
                        and "|" not in text
                    ):
                        main_instr_text = text
                        break
            except Exception:
                main_instr_text = 'STYLEREF "heading 1" \\* MERGEFORMAT'

        header_payload_by_target: Dict[str, bytes] = {
            "header4.xml": _build_center_header_xml(zh_title),
            "header5.xml": _build_center_header_xml("摘要"),
            "header6.xml": _build_center_header_xml(en_title),
            "header7.xml": _build_center_header_xml("Abstract"),
            "header8.xml": _build_center_header_xml("附录"),
            "header9.xml": _build_center_header_xml("目录"),
            "header10.xml": _build_center_header_xml("图表目录"),
            "header12.xml": _build_center_header_xml("参考文献"),
            "header13.xml": _build_center_header_xml("致谢"),
            "header14.xml": _build_center_header_xml("作者简历及攻读学位期间发表的学术论文与其他相关学术成果"),
        }

        role_to_default_target = {
            "abstract_cn": "header5.xml",
            "abstract_en": "header7.xml",
            "toc": "header9.xml",
            "catalog": "header10.xml",
            "bibliography": "header12.xml",
            "appendix": "header8.xml",
            "ack": "header13.xml",
            "cv": "header14.xml",
        }

        updated_sections = 0
        page_numbering_updates = 0
        front_pagination_started = False
        main_pagination_started = False
        front_roles = {"abstract_cn", "abstract_en", "toc", "catalog"}
        continuous_decimal_roles = {"bibliography", "appendix", "ack", "cv"}
        main_header_counter = 0
        for sec, role, heading in zip(sections, section_roles, heading_snapshots):
            if not role:
                continue
            sect = sec["sect"]
            assert isinstance(sect, ET.Element)

            if role == "main":
                main_header_counter += 1
                default_target = f"header{_GENERATED_MAIN_HEADER_START + main_header_counter}.xml"
                header_payload_by_target[default_target] = _build_center_header_xml(heading or main_heading_fallback)
            else:
                default_target = role_to_default_target[role]
            even_target = "header6.xml" if role == "abstract_en" else "header4.xml"
            default_rid = _ensure_header_rel_id(default_target)
            even_rid = _ensure_header_rel_id(even_target)

            for node in list(sect.findall("w:headerReference", W)):
                sect.remove(node)
            title_pg = sect.find("w:titlePg", W)
            if title_pg is not None:
                sect.remove(title_pg)

            insert_idx = 0
            for idx, child in enumerate(list(sect)):
                if child.tag == _w_attr("footerReference"):
                    insert_idx = idx
                    break
            else:
                insert_idx = len(list(sect))

            even_ref = ET.Element(_w_attr("headerReference"))
            even_ref.set(_w_attr("type"), "even")
            even_ref.set(_r_attr("id"), even_rid)
            sect.insert(insert_idx, even_ref)
            insert_idx += 1

            default_ref = ET.Element(_w_attr("headerReference"))
            default_ref.set(_w_attr("type"), "default")
            default_ref.set(_r_attr("id"), default_rid)
            sect.insert(insert_idx, default_ref)
            _ensure_even_default_footer_refs(sect)
            updated_sections += 1

            # Page numbering policy:
            # 1) front matter (摘要/目录/图表目录) keeps continuous Roman numerals;
            # 2) every main section starts on an odd page, while only the first
            #    main section restarts with decimal 1;
            # 3) later main sections keep decimal numbering continuously.
            pg_num = sect.find("w:pgNumType", W)
            if role in front_roles and not main_pagination_started:
                if pg_num is None:
                    pg_num = ET.Element(_w_attr("pgNumType"))
                    sect.append(pg_num)
                before = (
                    pg_num.get(_w_attr("start"), ""),
                    pg_num.get(_w_attr("fmt"), ""),
                )
                pg_num.set(_w_attr("fmt"), "upperRoman")
                if not front_pagination_started:
                    pg_num.set(_w_attr("start"), "1")
                    front_pagination_started = True
                elif _w_attr("start") in pg_num.attrib:
                    pg_num.attrib.pop(_w_attr("start"), None)
                after = (
                    pg_num.get(_w_attr("start"), ""),
                    pg_num.get(_w_attr("fmt"), ""),
                )
                if before != after:
                    page_numbering_updates += 1
            elif role == "main" and not main_pagination_started:
                if pg_num is None:
                    pg_num = ET.Element(_w_attr("pgNumType"))
                    sect.append(pg_num)
                before = (
                    pg_num.get(_w_attr("start"), ""),
                    pg_num.get(_w_attr("fmt"), ""),
                )
                pg_num.set(_w_attr("start"), "1")
                pg_num.set(_w_attr("fmt"), "decimal")
                after = (
                    pg_num.get(_w_attr("start"), ""),
                    pg_num.get(_w_attr("fmt"), ""),
                )
                if before != after:
                    page_numbering_updates += 1
                _set_section_break_type(sect, "oddPage")
                main_pagination_started = True
            elif role == "main" and main_pagination_started:
                if pg_num is None:
                    pg_num = ET.Element(_w_attr("pgNumType"))
                    sect.append(pg_num)
                before = (
                    pg_num.get(_w_attr("start"), ""),
                    pg_num.get(_w_attr("fmt"), ""),
                )
                if _w_attr("start") in pg_num.attrib:
                    pg_num.attrib.pop(_w_attr("start"), None)
                if pg_num.get(_w_attr("fmt"), "") != "decimal":
                    pg_num.set(_w_attr("fmt"), "decimal")
                after = (
                    pg_num.get(_w_attr("start"), ""),
                    pg_num.get(_w_attr("fmt"), ""),
                )
                if before != after:
                    page_numbering_updates += 1
                _set_section_break_type(sect, "oddPage")
            elif role in continuous_decimal_roles and main_pagination_started:
                if pg_num is None:
                    pg_num = ET.Element(_w_attr("pgNumType"))
                    sect.append(pg_num)
                before = (
                    pg_num.get(_w_attr("start"), ""),
                    pg_num.get(_w_attr("fmt"), ""),
                )
                pg_num.attrib.pop(_w_attr("start"), None)
                if pg_num.get(_w_attr("fmt"), "") != "decimal":
                    pg_num.set(_w_attr("fmt"), "decimal")
                after = (
                    pg_num.get(_w_attr("start"), ""),
                    pg_num.get(_w_attr("fmt"), ""),
                )
                if before != after:
                    page_numbering_updates += 1
                _set_section_break_type(sect, "oddPage")

        for sec in sections:
            sect = sec["sect"]
            if isinstance(sect, ET.Element):
                _normalize_section_properties_order(sect)

        even_odd = settings_root.find("w:evenAndOddHeaders", W)
        if even_odd is None:
            settings_root.append(ET.Element(_w_attr("evenAndOddHeaders")))

        written_parts = package_part_names - stale_generated_header_parts
        existing_overrides = {
            node.get("PartName", "")
            for node in content_types_root.findall(f"{{{CT_NS}}}Override")
        }
        header_content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml"
        for target_name in header_payload_by_target:
            part_name = f"/word/{target_name}"
            if part_name in existing_overrides:
                continue
            override = ET.Element(f"{{{CT_NS}}}Override")
            override.set("PartName", part_name)
            override.set("ContentType", header_content_type)
            content_types_root.append(override)
            existing_overrides.add(part_name)

        with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
            tmp_path = Path(tmp.name)
        with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename in stale_generated_header_parts:
                    continue
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    data = _serialize_xml(out_doc_root)
                elif item.filename == "word/_rels/document.xml.rels":
                    data = _serialize_xml(out_rels_root, default_namespace=PKG_REL_NS)
                elif item.filename == "word/settings.xml":
                    data = _serialize_xml(settings_root)
                elif item.filename == "[Content_Types].xml":
                    data = _serialize_xml(content_types_root, default_namespace=CT_NS)
                elif item.filename.startswith("word/header") and item.filename.endswith(".xml"):
                    target_name = item.filename.split("/", 1)[1]
                    if target_name in header_payload_by_target:
                        data = header_payload_by_target[target_name]
                zout.writestr(item, data)

            for target_name, data in header_payload_by_target.items():
                part = f"word/{target_name}"
                if part in written_parts:
                    continue
                zout.writestr(part, data)

    if tmp_path is not None:
        try:
            _replace_docx_with_lock_hint(tmp_path, docx_path)
        finally:
            if tmp_path.exists():
                tmp_path.unlink(missing_ok=True)

    used_roles = [role for role in section_roles if role]
    return {
        "applied": updated_sections > 0,
        "updated_sections": updated_sections,
        "page_numbering_updates": page_numbering_updates,
        "inserted_breaks": inserted_breaks,
        "roles": used_roles,
        "headings": [h for h in heading_snapshots if h],
    }


def _enforce_first_main_section_odd_page(docx_path: Path) -> Dict[str, object]:
    tmp_path: Optional[Path] = None
    with zipfile.ZipFile(docx_path, "r") as zin:
        out_doc_root = ET.fromstring(zin.read("word/document.xml"))
        body = out_doc_root.find("w:body", W)
        if body is None:
            return {"applied": False, "reason": "missing_body"}

        paragraphs = [p for p in body.findall("w:p", W)]
        main_heading_indices: List[int] = []
        for idx, para in enumerate(paragraphs):
            text = _extract_para_text(para).strip()
            if not text or not _is_heading1_paragraph(para):
                continue
            if _classify_section_header_role(text) == "main":
                main_heading_indices.append(idx)

        if not main_heading_indices:
            return {"applied": False, "reason": "main_heading_not_found"}

        inserted = 0
        updated = 0
        previous_type = ""
        for heading_idx in main_heading_indices:
            if heading_idx <= 0:
                continue
            target_para = paragraphs[heading_idx - 1]
            ppr = target_para.find("w:pPr", W)
            if ppr is None:
                ppr = ET.Element(_w_attr("pPr"))
                target_para.insert(0, ppr)
            sect = ppr.find("w:sectPr", W)
            if sect is None:
                template_sect: Optional[ET.Element] = None
                for para in paragraphs[heading_idx:]:
                    para_ppr = para.find("w:pPr", W)
                    if para_ppr is None:
                        continue
                    para_sect = para_ppr.find("w:sectPr", W)
                    if para_sect is not None:
                        template_sect = para_sect
                        break
                if template_sect is None:
                    template_sect = body.find("w:sectPr", W)
                sect = _clone_xml(template_sect) if template_sect is not None else ET.Element(_w_attr("sectPr"))
                pg_num = sect.find("w:pgNumType", W)
                if pg_num is None:
                    pg_num = ET.Element(_w_attr("pgNumType"))
                    sect.append(pg_num)
                if pg_num.get(_w_attr("fmt"), "") != "decimal":
                    pg_num.set(_w_attr("fmt"), "decimal")
                if pg_num.get(_w_attr("start"), "") != "1":
                    pg_num.set(_w_attr("start"), "1")
                ppr.append(sect)
                inserted += 1

            before = _get_section_break_type(sect)
            if before != "oddPage":
                _set_section_break_type(sect, "oddPage")
                updated += 1
            if not previous_type:
                previous_type = before or "implicit"
        _restore_markup_compatibility(out_doc_root, None)

        if updated == 0 and inserted == 0:
            return {"applied": False, "reason": "already_odd_page"}

        with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
            tmp_path = Path(tmp.name)
        with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    data = _serialize_xml(out_doc_root)
                zout.writestr(item, data)

    if tmp_path is not None:
        try:
            _replace_docx_with_lock_hint(tmp_path, docx_path)
        finally:
            if tmp_path.exists():
                tmp_path.unlink(missing_ok=True)

    return {
        "applied": True,
        "previous_type": previous_type or "implicit",
        "new_type": "oddPage",
        "inserted_section_break": bool(inserted),
        "updated_headings": len(main_heading_indices),
        "updated_odd_page_breaks": updated,
    }


def _docx_has_toc_field_results(docx_path: Path) -> bool:
    try:
        with zipfile.ZipFile(docx_path, "r") as zf:
            root = ET.fromstring(zf.read("word/document.xml"))
    except Exception:
        return False

    for para in root.findall(".//w:body/w:p", W):
        instr_text = " ".join((node.text or "") for node in para.findall(".//w:instrText", W))
        if " TOC " not in f" {re.sub(r'\\s+', ' ', instr_text)} ":
            continue
        visible_text = re.sub(r"\s+", " ", _extract_para_text(para)).strip()
        if visible_text and "右键更新域" not in visible_text and "TOC" not in visible_text:
            return True
    return False


def _should_skip_header_rule_sync_after_word_update(docx_path: Path, *, word_update_fields: bool) -> bool:
    if not word_update_fields:
        return False
    return _docx_has_toc_field_results(docx_path)


def _normalize_header_paragraph_layout(docx_path: Path) -> Dict[str, object]:
    """Force header paragraphs to be center-aligned with explicit no-indent layout."""
    tmp_path: Optional[Path] = None
    with zipfile.ZipFile(docx_path, "r") as zin:
        header_parts = sorted(
            name for name in zin.namelist() if name.startswith("word/header") and name.endswith(".xml")
        )
        if not header_parts:
            return {"applied": False, "updated_parts": 0, "updated_paragraphs": 0, "cleared_indent_attrs": 0}

        updated_parts = 0
        updated_paragraphs = 0
        cleared_indent_attrs = 0
        rewritten_parts: Dict[str, bytes] = {}

        for part_name in header_parts:
            root = ET.fromstring(zin.read(part_name))
            part_changed = False
            for para in root.findall(".//w:p", W):
                para_changed = False
                ppr = para.find("w:pPr", W)
                if ppr is None:
                    ppr = ET.Element(_w_attr("pPr"))
                    para.insert(0, ppr)
                    para_changed = True

                ind = ppr.find("w:ind", W)
                if ind is None:
                    ind = ET.Element(_w_attr("ind"))
                    ppr.append(ind)
                    para_changed = True

                # 显式写入 0，避免从样式继承“首行缩进两字符”。
                for key in [
                    "firstLine",
                    "firstLineChars",
                    "left",
                    "leftChars",
                    "right",
                    "rightChars",
                    "hanging",
                    "hangingChars",
                    "start",
                    "startChars",
                    "end",
                    "endChars",
                ]:
                    attr = _w_attr(key)
                    if ind.get(attr, "") != "0":
                        ind.set(attr, "0")
                        para_changed = True
                        if key.endswith("Chars"):
                            cleared_indent_attrs += 1
                if _ensure_paragraph_center(para):
                    para_changed = True
                if para_changed:
                    updated_paragraphs += 1
                    part_changed = True
            if part_changed:
                _restore_markup_compatibility(root, None)
                rewritten_parts[part_name] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                updated_parts += 1

        if not rewritten_parts:
            return {
                "applied": False,
                "updated_parts": 0,
                "updated_paragraphs": 0,
                "cleared_indent_attrs": 0,
            }

        with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
            tmp_path = Path(tmp.name)

        with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename in rewritten_parts:
                    data = rewritten_parts[item.filename]
                zout.writestr(item, data)

    if tmp_path is not None:
        try:
            _replace_docx_with_lock_hint(tmp_path, docx_path)
        finally:
            if tmp_path.exists():
                tmp_path.unlink(missing_ok=True)

    return {
        "applied": True,
        "updated_parts": updated_parts,
        "updated_paragraphs": updated_paragraphs,
        "cleared_indent_attrs": cleared_indent_attrs,
    }


def _norm_plain_text(text: str) -> str:
    return re.sub(r"[\s\u3000]+", "", (text or "")).casefold()


def _find_body_child_split_by_heading(doc_root: ET.Element, heading_markers: List[str]) -> Optional[int]:
    body = doc_root.find("w:body", W)
    if body is None:
        return None
    marker_set = {_norm_plain_text(x) for x in heading_markers}
    for idx, child in enumerate(list(body)):
        if child.tag != _w_attr("p"):
            continue
        text = _norm_plain_text(_extract_para_text(child))
        if text in marker_set:
            return idx
    return None


def _collect_relationship_ids_from_elements(elements: List[ET.Element]) -> List[str]:
    rel_ids: List[str] = []
    for elem in elements:
        for node in elem.iter():
            for k, v in node.attrib.items():
                if not v:
                    continue
                if not k.startswith(f"{{{R_NS}}}"):
                    continue
                if v.startswith("rId"):
                    rel_ids.append(v)
    return rel_ids


def _collect_used_namespace_uris(root: ET.Element) -> Set[str]:
    used: Set[str] = set()
    for elem in root.iter():
        if elem.tag.startswith("{"):
            used.add(elem.tag.split("}")[0][1:])
        for key in elem.attrib:
            if key.startswith("{"):
                used.add(key.split("}")[0][1:])
    return used


def _restore_markup_compatibility(output_root: ET.Element, reference_root: Optional[ET.Element] = None) -> None:
    used_uris = _collect_used_namespace_uris(output_root)
    ignorable_prefixes = [
        prefix for prefix, uri in IGNORABLE_NAMESPACE_URIS.items() if uri in used_uris
    ]

    if reference_root is not None:
        ref_ignorable = reference_root.get(f"{{{MC_NS}}}Ignorable", "")
        if ref_ignorable:
            preferred = [token for token in ref_ignorable.split() if token in ignorable_prefixes]
            if preferred:
                ignorable_prefixes = preferred

    if ignorable_prefixes:
        output_root.set(f"{{{MC_NS}}}Ignorable", " ".join(ignorable_prefixes))
    else:
        output_root.attrib.pop(f"{{{MC_NS}}}Ignorable", None)


def _ensure_docx_math_font(docx_path: Path, font_name: str = "Times New Roman") -> Dict[str, object]:
    tmp_path: Optional[Path] = None
    math_ns = {"m": OOXML_NAMESPACE_PREFIXES["m"]}

    with zipfile.ZipFile(docx_path, "r") as zin:
        try:
            settings_root = ET.fromstring(zin.read("word/settings.xml"))
        except Exception:
            return {"applied": False, "reason": "missing_settings"}

        math_pr = settings_root.find("m:mathPr", math_ns)
        created_math_pr = False
        if math_pr is None:
            math_pr = ET.Element(_m_attr("mathPr"))
            settings_root.append(math_pr)
            created_math_pr = True

        math_font = math_pr.find("m:mathFont", math_ns)
        created_math_font = False
        if math_font is None:
            math_font = ET.Element(_m_attr("mathFont"))
            math_pr.insert(0, math_font)
            created_math_font = True

        changed = False
        if math_font.get(_m_attr("val"), "") != font_name:
            math_font.set(_m_attr("val"), font_name)
            changed = True

        if not (changed or created_math_pr or created_math_font):
            return {"applied": False, "font": font_name, "reason": "already_set"}

        with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
            tmp_path = Path(tmp.name)
        with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/settings.xml":
                    data = ET.tostring(settings_root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

    if tmp_path is not None:
        try:
            _replace_docx_with_lock_hint(tmp_path, docx_path)
        finally:
            if tmp_path.exists():
                tmp_path.unlink(missing_ok=True)

    return {
        "applied": True,
        "font": font_name,
        "created_math_pr": created_math_pr,
        "created_math_font": created_math_font,
    }


def _normalize_part_name(path: str) -> str:
    path = (path or "").replace("\\", "/")
    parts: List[str] = []
    for token in path.split("/"):
        if token in {"", "."}:
            continue
        if token == "..":
            if parts:
                parts.pop()
            continue
        parts.append(token)
    return "/".join(parts)


def _resolve_relationship_target(source_part: str, target: str) -> str:
    target = (target or "").strip()
    if not target:
        return ""
    if target.startswith("/"):
        return _normalize_part_name(target)
    source_dir = ""
    if source_part:
        parent = str(PurePosixPath(source_part).parent)
        source_dir = "" if parent == "." else parent
    if source_dir:
        return _normalize_part_name(f"{source_dir}/{target}")
    return _normalize_part_name(target)


def _rels_part_for_source(source_part: str) -> str:
    source_part = _normalize_part_name(source_part)
    if not source_part:
        return "_rels/.rels"
    parent = str(PurePosixPath(source_part).parent)
    parent = "" if parent == "." else parent
    name = PurePosixPath(source_part).name
    if parent:
        return f"{parent}/_rels/{name}.rels"
    return f"_rels/{name}.rels"


def _relationship_signature(rel: ET.Element) -> Tuple[str, str, str]:
    return (rel.get("Type", ""), rel.get("Target", ""), rel.get("TargetMode", ""))


def _next_available_rid(existing_ids: Set[str]) -> str:
    max_idx = 0
    for rid in existing_ids:
        m = re.fullmatch(r"rId(\d+)", rid or "")
        if not m:
            continue
        max_idx = max(max_idx, int(m.group(1)))
    nxt = max_idx + 1
    while f"rId{nxt}" in existing_ids:
        nxt += 1
    return f"rId{nxt}"


def _merge_required_relationships(
    required_rel_ids: List[str],
    ref_rel_by_id: Dict[str, ET.Element],
    out_rels_root: ET.Element,
) -> Tuple[Dict[str, str], List[str]]:
    out_rel_by_id = {
        rel.get("Id", ""): rel for rel in out_rels_root.findall(f"{{{PKG_REL_NS}}}Relationship")
    }
    existing_ids: Set[str] = set(out_rel_by_id.keys())
    rid_map: Dict[str, str] = {}
    required_targets: List[str] = []

    for rid in required_rel_ids:
        if rid in rid_map:
            continue
        ref_rel = ref_rel_by_id.get(rid)
        if ref_rel is None:
            continue

        assigned_rid = rid
        out_rel = out_rel_by_id.get(rid)
        if out_rel is None:
            new_rel = _clone_xml(ref_rel)
            out_rels_root.append(new_rel)
            out_rel_by_id[assigned_rid] = new_rel
            existing_ids.add(assigned_rid)
        elif _relationship_signature(out_rel) != _relationship_signature(ref_rel):
            assigned_rid = _next_available_rid(existing_ids)
            new_rel = _clone_xml(ref_rel)
            new_rel.set("Id", assigned_rid)
            out_rels_root.append(new_rel)
            out_rel_by_id[assigned_rid] = new_rel
            existing_ids.add(assigned_rid)

        rid_map[rid] = assigned_rid
        target = ref_rel.get("Target", "")
        if target:
            required_targets.append(target)

    return rid_map, required_targets


def _remap_relationship_ids(elements: List[ET.Element], rid_map: Dict[str, str]) -> None:
    if not rid_map:
        return
    for elem in elements:
        for node in elem.iter():
            for key, value in list(node.attrib.items()):
                if not key.startswith(f"{{{R_NS}}}"):
                    continue
                new_value = rid_map.get(value)
                if new_value and new_value != value:
                    node.set(key, new_value)


def _collect_related_parts_from_targets(
    zref: zipfile.ZipFile,
    source_part: str,
    targets: List[str],
) -> Set[str]:
    names = set(zref.namelist())
    queue: List[str] = []
    for target in targets:
        resolved = _resolve_relationship_target(source_part, target)
        if resolved:
            queue.append(resolved)

    required_parts: Set[str] = set()
    visited_parts: Set[str] = set()
    while queue:
        part = queue.pop(0)
        part = _normalize_part_name(part)
        if not part or part in visited_parts:
            continue
        visited_parts.add(part)
        if part not in names:
            continue

        required_parts.add(part)
        rels_part = _rels_part_for_source(part)
        if rels_part not in names:
            continue

        required_parts.add(rels_part)
        try:
            rels_root = ET.fromstring(zref.read(rels_part))
        except Exception:
            continue

        for rel in rels_root.findall(f"{{{PKG_REL_NS}}}Relationship"):
            mode = (rel.get("TargetMode", "") or "").casefold()
            target = rel.get("Target", "")
            if mode == "external" or not target:
                continue
            resolved_target = _resolve_relationship_target(part, target)
            if resolved_target and resolved_target not in visited_parts:
                queue.append(resolved_target)

    return required_parts


def _validate_docx_package_integrity(docx_path: Path) -> Dict[str, object]:
    required_entries = {
        "[Content_Types].xml",
        "_rels/.rels",
        "word/document.xml",
        "word/_rels/document.xml.rels",
    }
    missing_entries: List[str] = []
    xml_parse_failures: List[str] = []
    missing_rel_targets: List[str] = []
    missing_override_parts: List[str] = []
    missing_rid_refs: List[str] = []
    unmatched_bookmarks: List[str] = []
    invalid_mc_ignorable: List[str] = []

    with zipfile.ZipFile(docx_path, "r") as zf:
        names = set(zf.namelist())
        missing_entries = sorted(entry for entry in required_entries if entry not in names)

        for name in names:
            if not (name.endswith(".xml") or name.endswith(".rels")):
                continue
            try:
                ET.fromstring(zf.read(name))
            except Exception:
                xml_parse_failures.append(name)

        for rels_name in sorted(n for n in names if n.endswith(".rels")):
            if rels_name == "_rels/.rels":
                source = ""
            elif "/_rels/" in rels_name and rels_name.endswith(".rels"):
                left, right = rels_name.split("/_rels/", 1)
                source = _normalize_part_name(f"{left}/{right[:-5]}")
                if source not in names:
                    missing_rel_targets.append(f"{rels_name} -> [source-missing] {source}")
                    continue
            else:
                continue

            try:
                rels_root = ET.fromstring(zf.read(rels_name))
            except Exception:
                continue

            rel_ids = set()
            for rel in rels_root.findall(f"{{{PKG_REL_NS}}}Relationship"):
                rid = rel.get("Id", "")
                if rid:
                    rel_ids.add(rid)
                mode = (rel.get("TargetMode", "") or "").casefold()
                target = rel.get("Target", "")
                if mode == "external" or not target:
                    continue
                resolved = _resolve_relationship_target(source, target)
                if resolved not in names:
                    missing_rel_targets.append(f"{rels_name}::{rid} -> {resolved}")

            if not source:
                continue
            try:
                source_root = ET.fromstring(zf.read(source))
            except Exception:
                continue
            for elem in source_root.iter():
                for key, value in elem.attrib.items():
                    if key.startswith(f"{{{R_NS}}}") and value.startswith("rId") and value not in rel_ids:
                        missing_rid_refs.append(f"{source}::{value}")

        try:
            ct_root = ET.fromstring(zf.read("[Content_Types].xml"))
            for ov in ct_root.findall("{http://schemas.openxmlformats.org/package/2006/content-types}Override"):
                part = ov.get("PartName", "")
                normalized = _normalize_part_name(part)
                if normalized and normalized not in names:
                    missing_override_parts.append(normalized)
        except Exception:
            xml_parse_failures.append("[Content_Types].xml")

        parts_for_mc_check = ["word/document.xml", "word/numbering.xml", "word/styles.xml"]
        parts_for_mc_check.extend(
            sorted(
                name
                for name in names
                if (name.startswith("word/header") or name.startswith("word/footer")) and name.endswith(".xml")
            )
        )
        for part_name in parts_for_mc_check:
            if part_name not in names:
                continue
            raw = zf.read(part_name).decode("utf-8", "ignore")
            start = raw.find("<", raw.find("?>") + 2)
            end = raw.find(">", start)
            if start < 0 or end < 0:
                continue
            root_open = raw[start : end + 1]
            m = re.search(r'mc:Ignorable="([^"]+)"', root_open)
            if not m:
                continue
            declared = set(re.findall(r'xmlns:([A-Za-z0-9_]+)=', root_open))
            for prefix in m.group(1).split():
                if prefix not in declared:
                    invalid_mc_ignorable.append(f"{part_name}::{prefix}")

        try:
            doc_root = ET.fromstring(zf.read("word/document.xml"))
            bookmark_starts = Counter(
                node.get(_w_attr("id"), "") for node in doc_root.findall(".//w:bookmarkStart", W)
            )
            bookmark_ends = Counter(
                node.get(_w_attr("id"), "") for node in doc_root.findall(".//w:bookmarkEnd", W)
            )
            for bid in sorted(set(bookmark_starts) | set(bookmark_ends), key=lambda x: int(x or 0)):
                if bookmark_starts[bid] != bookmark_ends[bid]:
                    unmatched_bookmarks.append(
                        f"id={bid}: start={bookmark_starts[bid]} end={bookmark_ends[bid]}"
                    )
        except Exception:
            xml_parse_failures.append("word/document.xml")

    return {
        "ok": not (
            missing_entries
            or xml_parse_failures
            or missing_rel_targets
            or missing_override_parts
            or missing_rid_refs
            or unmatched_bookmarks
            or invalid_mc_ignorable
        ),
        "missing_entries": sorted(set(missing_entries)),
        "xml_parse_failures": sorted(set(xml_parse_failures)),
        "missing_rel_targets": sorted(set(missing_rel_targets)),
        "missing_override_parts": sorted(set(missing_override_parts)),
        "missing_rid_refs": sorted(set(missing_rid_refs)),
        "unmatched_bookmarks": unmatched_bookmarks,
        "invalid_mc_ignorable": sorted(set(invalid_mc_ignorable)),
    }


def _scan_docx_omml_nodes(docx_path: Path) -> Dict[str, object]:
    parts: List[Tuple[str, int, int]] = []
    o_math_total = 0
    o_math_para_total = 0
    with zipfile.ZipFile(docx_path, "r") as zf:
        for name in sorted(zf.namelist()):
            if not (name.startswith("word/") and name.endswith(".xml")):
                continue
            try:
                root = ET.fromstring(zf.read(name))
            except Exception:
                continue
            part_o_math = len(root.findall(".//m:oMath", OOXML_NAMESPACE_PREFIXES))
            part_o_math_para = len(root.findall(".//m:oMathPara", OOXML_NAMESPACE_PREFIXES))
            if part_o_math or part_o_math_para:
                parts.append((name, part_o_math, part_o_math_para))
                o_math_total += part_o_math
                o_math_para_total += part_o_math_para
    return {
        "oMath": o_math_total,
        "oMathPara": o_math_para_total,
        "parts": parts,
    }


def _enforce_omml_gate(docx_path: Path, max_omml_nodes: Optional[int]) -> Dict[str, object]:
    scan = _scan_docx_omml_nodes(docx_path)
    o_math = int(scan.get("oMath", 0))
    o_math_para = int(scan.get("oMathPara", 0))
    gate = int(max_omml_nodes) if max_omml_nodes is not None else None
    if gate is not None and o_math > gate:
        raise RuntimeError(f"OMML 门禁失败：m:oMath={o_math}，阈值={gate}。")
    return {
        "applied": gate is not None,
        "gate_max": gate,
        "oMath": o_math,
        "oMathPara": o_math_para,
        "parts": scan.get("parts", []),
    }


def _cleanup_unmatched_bookmarks(docx_path: Path) -> Dict[str, object]:
    with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
        tmp_path = Path(tmp.name)

    removed_starts = 0
    removed_ends = 0
    removed_ids: List[str] = []

    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
            tmp_path, "w", compression=zipfile.ZIP_DEFLATED
        ) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    root = ET.fromstring(data)
                    before_ignorable = root.get(f"{{{MC_NS}}}Ignorable", "")
                    parents = {child: parent for parent in root.iter() for child in parent}
                    open_counts: Counter[str] = Counter()
                    trailing_starts: List[ET.Element] = []
                    orphan_ends: List[ET.Element] = []

                    for node in root.iter():
                        if node.tag == _w_attr("bookmarkStart"):
                            bid = node.get(_w_attr("id"), "")
                            open_counts[bid] += 1
                            trailing_starts.append(node)
                        elif node.tag == _w_attr("bookmarkEnd"):
                            bid = node.get(_w_attr("id"), "")
                            if open_counts[bid] > 0:
                                open_counts[bid] -= 1
                            else:
                                orphan_ends.append(node)
                                removed_ids.append(bid)

                    extra_start_counts = Counter({bid: cnt for bid, cnt in open_counts.items() if cnt > 0})
                    orphan_starts: List[ET.Element] = []
                    if extra_start_counts:
                        for node in reversed(trailing_starts):
                            bid = node.get(_w_attr("id"), "")
                            if extra_start_counts[bid] > 0:
                                orphan_starts.append(node)
                                extra_start_counts[bid] -= 1
                                removed_ids.append(bid)

                    for node in orphan_ends:
                        parent = parents.get(node)
                        if parent is not None:
                            parent.remove(node)
                            removed_ends += 1
                    for node in orphan_starts:
                        parent = parents.get(node)
                        if parent is not None:
                            parent.remove(node)
                            removed_starts += 1

                    _restore_markup_compatibility(root, None)
                    after_ignorable = root.get(f"{{{MC_NS}}}Ignorable", "")
                    if removed_starts or removed_ends or before_ignorable != after_ignorable:
                        data = ET.tostring(root, encoding="utf-8", xml_declaration=True)

                zout.writestr(item, data)

        _replace_docx_with_lock_hint(tmp_path, docx_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)

    return {
        "removed_starts": removed_starts,
        "removed_ends": removed_ends,
        "removed_ids": sorted(removed_ids, key=lambda x: int(x or 0)),
        "applied": bool(removed_starts or removed_ends),
    }


def _replace_paragraph_content_with_template(dst_para: ET.Element, src_para: ET.Element, text: str) -> None:
    new_children: List[ET.Element] = []
    src_ppr = src_para.find("w:pPr", W)
    if src_ppr is not None:
        new_children.append(_clone_xml(src_ppr))

    src_run = src_para.find("w:r", W)
    if src_run is not None:
        new_run = _clone_xml(src_run)
        for child in list(new_run):
            if child.tag != _w_attr("rPr"):
                new_run.remove(child)
        t = ET.Element(_w_attr("t"))
        if text[:1].isspace() or text[-1:].isspace():
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = text
        new_run.append(t)
    else:
        new_run = ET.Element(_w_attr("r"))
        t = ET.Element(_w_attr("t"))
        if text[:1].isspace() or text[-1:].isspace():
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = text
        new_run.append(t)
    new_children.append(new_run)

    for child in list(dst_para):
        dst_para.remove(child)
    for child in new_children:
        dst_para.append(child)


def _fix_preface_heading_numbering(docx_path: Path, reference_doc: Path) -> Dict[str, object]:
    alias_to_template_key = {
        "摘要": "摘要",
        "abstract": "abstract",
        "目录": "图目录",
        "图表目录": "图目录",
        "图目录": "图目录",
        "表目录": "表目录",
    }

    with zipfile.ZipFile(reference_doc, "r") as zref:
        ref_doc_root = ET.fromstring(zref.read("word/document.xml"))
        ref_body = ref_doc_root.find("w:body", W)
        if ref_body is None:
            raise RuntimeError("参考模板缺少 w:body，无法修正前置标题样式")

        template_para_by_key: Dict[str, ET.Element] = {}
        for para in ref_body.findall("w:p", W):
            text = _extract_para_text(para)
            norm = _norm_plain_text(text)
            if norm in {"摘要", "abstract", "图目录", "表目录"} and norm not in template_para_by_key:
                template_para_by_key[norm] = _clone_xml(para)

        tmp_path: Optional[Path] = None
        with zipfile.ZipFile(docx_path, "r") as zin:
            out_doc_root = ET.fromstring(zin.read("word/document.xml"))
            out_body = out_doc_root.find("w:body", W)
            if out_body is None:
                raise RuntimeError("导出 docx 缺少 w:body，无法修正前置标题样式")

            first_formal_idx = _find_body_child_split_by_heading(out_doc_root, ["绪论"])
            if first_formal_idx is None:
                return {"applied": False, "fixed": []}

            fixed: List[str] = []
            body_children = list(out_body)
            for idx, child in enumerate(body_children[:first_formal_idx]):
                if child.tag != _w_attr("p"):
                    continue
                text = _extract_para_text(child)
                norm = _norm_plain_text(text)
                template_key = alias_to_template_key.get(norm)
                if not template_key:
                    continue
                template_para = template_para_by_key.get(template_key)
                if template_para is None:
                    continue
                _replace_paragraph_content_with_template(child, template_para, text)
                fixed.append(text)

            _restore_markup_compatibility(out_doc_root, ref_doc_root)

            with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
                tmp_path = Path(tmp.name)

            with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename == "word/document.xml":
                        data = ET.tostring(out_doc_root, encoding="utf-8", xml_declaration=True)
                    zout.writestr(item, data)

        if tmp_path is not None:
            try:
                _replace_docx_with_lock_hint(tmp_path, docx_path)
            finally:
                if tmp_path.exists():
                    tmp_path.unlink(missing_ok=True)

    return {"applied": bool(fixed), "fixed": fixed}


def _strip_preface_numbering_controls(docx_path: Path) -> Dict[str, object]:
    tmp_path: Optional[Path] = None
    with zipfile.ZipFile(docx_path, "r") as zin:
        out_doc_root = ET.fromstring(zin.read("word/document.xml"))
        first_formal_idx = _find_body_child_split_by_heading(out_doc_root, ["绪论"])
        if first_formal_idx is None:
            return {"applied": False, "removed_numpr": 0, "removed_heading_styles": 0}

        out_body = out_doc_root.find("w:body", W)
        if out_body is None:
            raise RuntimeError("导出 docx 缺少 w:body，无法清理前置编号")

        removed_numpr = 0
        removed_heading_styles = 0
        for child in list(out_body)[:first_formal_idx]:
            if child.tag != _w_attr("p"):
                continue
            ppr = child.find("w:pPr", W)
            if ppr is None:
                continue

            numpr = ppr.find("w:numPr", W)
            if numpr is not None:
                ppr.remove(numpr)
                removed_numpr += 1

            pstyle = ppr.find("w:pStyle", W)
            if pstyle is not None:
                sid = pstyle.get(_w_attr("val"), "")
                if sid in {"1", "2", "3", "4", "5", "6", "7", "8", "9"}:
                    ppr.remove(pstyle)
                    removed_heading_styles += 1

            outline = ppr.find("w:outlineLvl", W)
            if outline is not None:
                ppr.remove(outline)

        with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
            tmp_path = Path(tmp.name)

        with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    data = ET.tostring(out_doc_root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

    if tmp_path is not None:
        try:
            _replace_docx_with_lock_hint(tmp_path, docx_path)
        finally:
            if tmp_path.exists():
                tmp_path.unlink(missing_ok=True)

    return {
        "applied": bool(removed_numpr or removed_heading_styles),
        "removed_numpr": removed_numpr,
        "removed_heading_styles": removed_heading_styles,
    }


def _normalize_heading_multilevel_numbering(docx_path: Path) -> Dict[str, object]:
    tmp_path: Optional[Path] = None
    updated_levels = 0
    removed_restart_rules = 0

    with zipfile.ZipFile(docx_path, "r") as zin:
        numbering_root = ET.fromstring(zin.read("word/numbering.xml"))
        before_ignorable = numbering_root.get(f"{{{MC_NS}}}Ignorable", "")

        for abstract_num in numbering_root.findall("w:abstractNum", W):
            style_map: Dict[str, ET.Element] = {}
            for lvl in abstract_num.findall("w:lvl", W):
                pstyle = lvl.find("w:pStyle", W)
                if pstyle is not None:
                    style_map[pstyle.get(_w_attr("val"), "")] = lvl

            # UCAS 正文章节使用 heading 1/2/3 链路；二级标题必须从 1.1 开始
            if {"1", "2", "3"}.issubset(style_map.keys()):
                lvl1 = style_map["2"]
                start = lvl1.find("w:start", W)
                if start is None:
                    start = ET.Element(_w_attr("start"))
                    lvl1.insert(0, start)
                if start.get(_w_attr("val"), "") != "1":
                    start.set(_w_attr("val"), "1")
                    updated_levels += 1

                # 让二级标题在每个新的一级标题后按 Word 默认行为从 .1 重启。
                lvl_restart = lvl1.find("w:lvlRestart", W)
                if lvl_restart is not None:
                    lvl1.remove(lvl_restart)
                    removed_restart_rules += 1

        _restore_markup_compatibility(numbering_root, None)
        after_ignorable = numbering_root.get(f"{{{MC_NS}}}Ignorable", "")

        if not updated_levels and not removed_restart_rules and before_ignorable == after_ignorable:
            return {"applied": False, "updated_levels": 0, "removed_restart_rules": 0}

        with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
            tmp_path = Path(tmp.name)

        with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/numbering.xml":
                    data = ET.tostring(numbering_root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

    if tmp_path is not None:
        try:
            _replace_docx_with_lock_hint(tmp_path, docx_path)
        finally:
            if tmp_path.exists():
                tmp_path.unlink(missing_ok=True)

    return {
        "applied": True,
        "updated_levels": updated_levels,
        "removed_restart_rules": removed_restart_rules,
    }


def _enforce_heading_number_spacing(docx_path: Path) -> Dict[str, object]:
    """Ensure chapter/section numbering renders with two spaces before heading text."""
    tmp_path: Optional[Path] = None
    updated_formats = 0

    desired_lvl_text = {
        "1": "第%1章  ",
        "2": "%1.%2  ",
        "3": "%1.%2.%3  ",
    }

    with zipfile.ZipFile(docx_path, "r") as zin:
        numbering_root = ET.fromstring(zin.read("word/numbering.xml"))
        before_ignorable = numbering_root.get(f"{{{MC_NS}}}Ignorable", "")

        for abstract_num in numbering_root.findall("w:abstractNum", W):
            style_map: Dict[str, ET.Element] = {}
            for lvl in abstract_num.findall("w:lvl", W):
                pstyle = lvl.find("w:pStyle", W)
                if pstyle is None:
                    continue
                sid_raw = pstyle.get(_w_attr("val"), "")
                sid_norm = re.sub(r"[\s_-]+", "", sid_raw.casefold())
                if sid_norm in {"1", "heading1"}:
                    style_map["1"] = lvl
                elif sid_norm in {"2", "heading2"}:
                    style_map["2"] = lvl
                elif sid_norm in {"3", "heading3"}:
                    style_map["3"] = lvl

            if {"1", "2", "3"}.issubset(style_map.keys()):
                for sid, lvl in style_map.items():
                    lvl_text = lvl.find("w:lvlText", W)
                    if lvl_text is None:
                        lvl_text = ET.Element(_w_attr("lvlText"))
                        lvl.insert(0, lvl_text)
                    if lvl_text.get(_w_attr("val"), "") != desired_lvl_text[sid]:
                        lvl_text.set(_w_attr("val"), desired_lvl_text[sid])
                        updated_formats += 1

                    suff = lvl.find("w:suff", W)
                    if suff is None:
                        suff = ET.Element(_w_attr("suff"))
                        lvl.append(suff)
                    if suff.get(_w_attr("val"), "") != "nothing":
                        suff.set(_w_attr("val"), "nothing")
                        updated_formats += 1

        _restore_markup_compatibility(numbering_root, None)
        after_ignorable = numbering_root.get(f"{{{MC_NS}}}Ignorable", "")

        if not updated_formats and before_ignorable == after_ignorable:
            return {"applied": False, "updated_formats": 0}

        with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
            tmp_path = Path(tmp.name)

        with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/numbering.xml":
                    data = ET.tostring(numbering_root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

    if tmp_path is not None:
        try:
            _replace_docx_with_lock_hint(tmp_path, docx_path)
        finally:
            if tmp_path.exists():
                tmp_path.unlink(missing_ok=True)

    return {
        "applied": True,
        "updated_formats": updated_formats,
    }


def _strip_italic_nodes_from_rpr(rpr: ET.Element) -> ET.Element:
    cleaned = _clone_xml(rpr)
    for node in list(cleaned):
        if node.tag in {_w_attr("i"), _w_attr("iCs")}:
            cleaned.remove(node)
    return cleaned


def _set_rpr_italic(rpr: ET.Element, italic: bool) -> ET.Element:
    target = _clone_xml(rpr)
    i_node = target.find("w:i", W)
    i_cs_node = target.find("w:iCs", W)
    if italic:
        if i_node is None:
            i_node = ET.Element(_w_attr("i"))
            target.append(i_node)
        i_node.set(_w_attr("val"), "1")
        if i_cs_node is None:
            i_cs_node = ET.Element(_w_attr("iCs"))
            target.append(i_cs_node)
        i_cs_node.set(_w_attr("val"), "1")
    else:
        if i_node is not None:
            target.remove(i_node)
        if i_cs_node is not None:
            target.remove(i_cs_node)
    return target


def _set_rpr_vert_align(rpr: ET.Element, vert_align: Optional[str]) -> ET.Element:
    target = _clone_xml(rpr)
    node = target.find("w:vertAlign", W)
    if vert_align:
        if node is None:
            node = ET.Element(_w_attr("vertAlign"))
            target.append(node)
        node.set(_w_attr("val"), vert_align)
    elif node is not None:
        target.remove(node)
    return target


def _run_is_italic(run: ET.Element) -> bool:
    rpr = run.find("w:rPr", W)
    if rpr is None:
        return False
    for key in ("i", "iCs"):
        node = rpr.find(f"w:{key}", W)
        if node is None:
            continue
        if node.get(_w_attr("val"), "1") != "0":
            return True
    return False


def _extract_run_text(run: ET.Element) -> str:
    return "".join((node.text or "") for node in run.findall(".//w:t", W))


def _build_run_like(
    source_run: ET.Element,
    text: str,
    *,
    italic: bool,
    vert_align: Optional[str] = None,
) -> ET.Element:
    new_run = ET.Element(_w_attr("r"))
    source_rpr = source_run.find("w:rPr", W)
    if source_rpr is not None:
        new_rpr = _set_rpr_italic(source_rpr, italic)
        new_rpr = _set_rpr_vert_align(new_rpr, vert_align)
        new_run.append(new_rpr)
    elif italic:
        new_rpr = ET.Element(_w_attr("rPr"))
        i_node = ET.Element(_w_attr("i"))
        i_node.set(_w_attr("val"), "1")
        i_cs_node = ET.Element(_w_attr("iCs"))
        i_cs_node.set(_w_attr("val"), "1")
        new_rpr.extend([i_node, i_cs_node])
        new_rpr = _set_rpr_vert_align(new_rpr, vert_align)
        new_run.append(new_rpr)
    elif vert_align:
        new_rpr = ET.Element(_w_attr("rPr"))
        new_rpr = _set_rpr_vert_align(new_rpr, vert_align)
        new_run.append(new_rpr)

    t = ET.Element(_w_attr("t"))
    if text[:1].isspace() or text[-1:].isspace():
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    new_run.append(t)
    return new_run


_PLAIN_SCRIPT_TOKEN_RE = re.compile(r"[⁰¹²³⁴⁵⁶⁷⁸⁹]+[A-Z][a-z]?|10\^-?\d+|η\^-?\d+|[A-Za-z][A-Za-z0-9()^+\-]*")
_PLAIN_UNIT_EXPONENT_TOKEN_RE = re.compile(r"^(?P<base>[A-Za-z]{1,5})\^(?P<exp>-?\d+)$")
_PLAIN_SCIENTIFIC_EXPONENT_TOKEN_RE = re.compile(r"^(?P<base>10)\^(?P<exp>-?\d+)$")
_PLAIN_GREEK_EXPONENT_TOKEN_RE = re.compile(r"^(?P<base>η)\^(?P<exp>-?\d+)$")
_ION_SINGLE_GROUP_ALLOWLIST = {"H+", "OH-", "Cl-", "Br-", "I-"}
_UNICODE_SUPERSCRIPT_DIGIT_RE = re.compile(r"^[⁰¹²³⁴⁵⁶⁷⁸⁹]+$")
_LEADING_ISOTOPE_TOKEN_RE = re.compile(r"([⁰¹²³⁴⁵⁶⁷⁸⁹]+)([A-Z][a-z]?)")
_UNICODE_SUPERSCRIPT_DIGIT_TO_PLAIN = str.maketrans(
    {
        "⁰": "0",
        "¹": "1",
        "²": "2",
        "³": "3",
        "⁴": "4",
        "⁵": "5",
        "⁶": "6",
        "⁷": "7",
        "⁸": "8",
        "⁹": "9",
    }
)
_PLAIN_SCRIPT_TRAILING_PUNCTUATION = ".,;:!?)]}”\"'"
_CHEMICAL_ELEMENT_SYMBOLS = {
    "H", "He", "Li", "Be", "B", "C", "N", "O", "F", "Ne",
    "Na", "Mg", "Al", "Si", "P", "S", "Cl", "Ar", "K", "Ca",
    "Sc", "Ti", "V", "Cr", "Mn", "Fe", "Co", "Ni", "Cu", "Zn",
    "Ga", "Ge", "As", "Se", "Br", "Kr", "Rb", "Sr", "Y", "Zr",
    "Nb", "Mo", "Tc", "Ru", "Rh", "Pd", "Ag", "Cd", "In", "Sn",
    "Sb", "Te", "I", "Xe", "Cs", "Ba", "La", "Ce", "Pr", "Nd",
    "Pm", "Sm", "Eu", "Gd", "Tb", "Dy", "Ho", "Er", "Tm", "Yb",
    "Lu", "Hf", "Ta", "W", "Re", "Os", "Ir", "Pt", "Au", "Hg",
    "Tl", "Pb", "Bi", "Po", "At", "Rn", "Fr", "Ra", "Ac", "Th",
    "Pa", "U", "Np", "Pu", "Am", "Cm", "Bk", "Cf", "Es", "Fm",
    "Md", "No", "Lr", "Rf", "Db", "Sg", "Bh", "Hs", "Mt", "Ds",
    "Rg", "Cn", "Nh", "Fl", "Mc", "Lv", "Ts", "Og",
}


def _consume_formula_core(token: str, start: int = 0, stop_char: Optional[str] = None) -> Optional[Tuple[int, int]]:
    i = start
    groups = 0
    while i < len(token):
        ch = token[i]
        if stop_char and ch == stop_char:
            break
        if ch == "(":
            inner = _consume_formula_core(token, i + 1, ")")
            if inner is None:
                return None
            end_idx, inner_groups = inner
            if end_idx >= len(token) or token[end_idx] != ")" or inner_groups == 0:
                return None
            i = end_idx + 1
            while i < len(token) and token[i].isdigit():
                i += 1
            groups += inner_groups
            continue
        if not ch.isupper():
            return None
        j = i + 1
        if j < len(token) and token[j].islower():
            j += 1
        symbol = token[i:j]
        if symbol not in _CHEMICAL_ELEMENT_SYMBOLS:
            return None
        while j < len(token) and token[j].isdigit():
            j += 1
        groups += 1
        i = j

    if stop_char:
        if i >= len(token) or token[i] != stop_char:
            return None
        return i, groups
    return i, groups


def _analyze_formula_core(token: str) -> Optional[int]:
    if not token:
        return None
    parsed = _consume_formula_core(token)
    if parsed is None:
        return None
    end_idx, groups = parsed
    if end_idx != len(token) or groups == 0:
        return None
    return groups


def _split_formula_charge(token: str) -> Tuple[str, Optional[str]]:
    if "^" in token:
        base, marker, charge = token.partition("^")
        if marker and charge and _analyze_formula_core(base) is not None:
            return base, charge
        return token, None
    if not token.endswith(("+", "-")):
        return token, None

    sign_only_base = token[:-1]
    sign_only_groups = _analyze_formula_core(sign_only_base)
    trailing_match = re.fullmatch(r"(.+?)(\d+)([+-])", token)
    trailing_base = trailing_match.group(1) if trailing_match else ""
    trailing_groups = _analyze_formula_core(trailing_base) if trailing_match else None

    if sign_only_groups is not None:
        if (
            trailing_match is not None
            and trailing_groups is not None
            and sign_only_groups == 1
            and sign_only_base[-1:].isdigit()
        ):
            return trailing_base, trailing_match.group(2) + trailing_match.group(3)
        return sign_only_base, token[-1]
    if trailing_match is not None and trailing_groups is not None:
        return trailing_base, trailing_match.group(2) + trailing_match.group(3)
    return token, None


def _style_formula_core(token: str) -> List[Tuple[str, Optional[str]]]:
    pieces: List[Tuple[str, Optional[str]]] = []
    i = 0
    while i < len(token):
        ch = token[i]
        if ch.isdigit() and i > 0 and (token[i - 1].isalpha() or token[i - 1] == ")"):
            j = i
            while j < len(token) and token[j].isdigit():
                j += 1
            pieces.append((token[i:j], "subscript"))
            i = j
            continue
        pieces.append((ch, None))
        i += 1
    return pieces


def _should_style_plain_script_token(token: str) -> bool:
    if not token:
        return False
    if token in {"pH"}:
        return False
    if _PLAIN_UNIT_EXPONENT_TOKEN_RE.fullmatch(token):
        return True
    if _PLAIN_SCIENTIFIC_EXPONENT_TOKEN_RE.fullmatch(token):
        return True
    if _PLAIN_GREEK_EXPONENT_TOKEN_RE.fullmatch(token):
        return True
    isotope_match = _LEADING_ISOTOPE_TOKEN_RE.fullmatch(token)
    if isotope_match and isotope_match.group(2) in _CHEMICAL_ELEMENT_SYMBOLS:
        return True
    if token in _ION_SINGLE_GROUP_ALLOWLIST:
        return True
    formula_base, charge = _split_formula_charge(token)
    formula_groups = _analyze_formula_core(formula_base)
    if charge is not None and formula_groups is not None:
        return True
    if formula_groups is not None and formula_groups >= 2:
        return True
    return False


def _style_plain_script_token(token: str) -> List[Tuple[str, Optional[str]]]:
    unit_match = _PLAIN_UNIT_EXPONENT_TOKEN_RE.fullmatch(token)
    if unit_match:
        return [
            (unit_match.group("base"), None),
            (unit_match.group("exp"), "superscript"),
        ]
    scientific_match = _PLAIN_SCIENTIFIC_EXPONENT_TOKEN_RE.fullmatch(token)
    if scientific_match:
        return [
            (scientific_match.group("base"), None),
            (scientific_match.group("exp"), "superscript"),
        ]
    greek_match = _PLAIN_GREEK_EXPONENT_TOKEN_RE.fullmatch(token)
    if greek_match:
        return [
            (greek_match.group("base"), None),
            (greek_match.group("exp"), "superscript"),
        ]

    isotope_match = _LEADING_ISOTOPE_TOKEN_RE.fullmatch(token)
    if isotope_match and isotope_match.group(2) in _CHEMICAL_ELEMENT_SYMBOLS:
        return [
            (isotope_match.group(1).translate(_UNICODE_SUPERSCRIPT_DIGIT_TO_PLAIN), "superscript"),
            (isotope_match.group(2), None),
        ]

    formula_base, charge = _split_formula_charge(token)
    if charge is not None and _analyze_formula_core(formula_base) is not None:
        pieces = _style_formula_core(formula_base)
        pieces.append((charge, "superscript"))
    elif _analyze_formula_core(token) is not None:
        pieces = _style_formula_core(token)
    else:
        return [(token, None)]

    merged: List[Tuple[str, Optional[str]]] = []
    for text, align in pieces:
        if not text:
            continue
        if merged and merged[-1][1] == align:
            merged[-1] = (merged[-1][0] + text, align)
        else:
            merged.append((text, align))
    return merged


def _split_plain_script_token_trailing_punctuation(token: str) -> Tuple[str, str]:
    core = token
    suffix = ""
    while len(core) > 1 and core[-1] in _PLAIN_SCRIPT_TRAILING_PUNCTUATION:
        suffix = core[-1] + suffix
        core = core[:-1]
    return core, suffix


def _split_plain_script_run_text(text: str) -> List[Tuple[str, Optional[str]]]:
    pieces: List[Tuple[str, Optional[str]]] = []
    cursor = 0
    for match in _PLAIN_SCRIPT_TOKEN_RE.finditer(text):
        if match.start() > cursor:
            pieces.append((text[cursor : match.start()], None))
        token = match.group(0)
        if _should_style_plain_script_token(token):
            pieces.extend(_style_plain_script_token(token))
        else:
            core, suffix = _split_plain_script_token_trailing_punctuation(token)
            if core != token and _should_style_plain_script_token(core):
                pieces.extend(_style_plain_script_token(core))
                if suffix:
                    pieces.append((suffix, None))
            else:
                pieces.append((token, None))
        cursor = match.end()
    if cursor < len(text):
        pieces.append((text[cursor:], None))
    filtered = [(segment, align) for segment, align in pieces if segment]
    merged: List[Tuple[str, Optional[str]]] = []
    for segment, align in filtered:
        if merged and merged[-1][1] == align:
            merged[-1] = (merged[-1][0] + segment, align)
        else:
            merged.append((segment, align))
    return merged


def _repair_plain_script_runs_in_paragraph(para: ET.Element) -> int:
    changed = 0
    for run in list(para.findall("w:r", W)):
        if run.find("w:fldChar", W) is not None or run.find("w:instrText", W) is not None:
            continue
        text = _extract_run_text(run)
        if not text:
            continue
        pieces = _split_plain_script_run_text(text)
        if len(pieces) == 1 and pieces[0][0] == text and pieces[0][1] is None:
            continue
        italic = _run_is_italic(run)
        new_runs = [
            _build_run_like(run, segment, italic=italic, vert_align=align) for segment, align in pieces
        ]
        if not new_runs:
            continue
        children = list(para)
        try:
            run_idx = children.index(run)
        except ValueError:
            continue
        for offset, new_run in enumerate(new_runs):
            para.insert(run_idx + offset, new_run)
        para.remove(run)
        changed += 1
    return changed


def _repair_plain_script_runs_in_docx(docx_path: Path) -> Dict[str, object]:
    updated = 0
    with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
        tmp_path = Path(tmp.name)

    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
            tmp_path, "w", compression=zipfile.ZIP_DEFLATED
        ) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    root = ET.fromstring(data)
                    body = root.find("w:body", W)
                    if body is not None:
                        for para in body.findall(".//w:p", W):
                            updated += _repair_plain_script_runs_in_paragraph(para)
                    data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

        _replace_docx_with_lock_hint(tmp_path, docx_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)

    return {"applied": updated > 0, "updated": updated}


def _repair_species_token_runs_in_paragraph(para: ET.Element) -> int:
    changed = 0
    for run in list(para.findall("w:r", W)):
        text = _extract_run_text(run)
        if not text:
            continue

        has_species = _SPECIES_TOKEN_RE.search(text) is not None
        if not has_species:
            continue

        if "*" not in text and _run_is_italic(run):
            continue

        pieces: List[Tuple[str, bool]] = []
        last = 0
        for match in _SPECIES_MARKED_TOKEN_RE.finditer(text):
            if match.start() > last:
                pieces.append((text[last : match.start()], _run_is_italic(run)))
            pieces.append((match.group(1), True))
            last = match.end()
        if last < len(text):
            pieces.append((text[last:], _run_is_italic(run)))

        if not pieces:
            continue

        new_runs = [_build_run_like(run, segment, italic=italic) for segment, italic in pieces if segment]
        if not new_runs:
            continue

        children = list(para)
        try:
            run_idx = children.index(run)
        except ValueError:
            continue
        for offset, new_run in enumerate(new_runs):
            para.insert(run_idx + offset, new_run)
        para.remove(run)
        changed += 1
    return changed


def _repair_species_token_runs_in_docx(docx_path: Path) -> Dict[str, object]:
    updated = 0
    with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
        tmp_path = Path(tmp.name)

    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
            tmp_path, "w", compression=zipfile.ZIP_DEFLATED
        ) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    root = ET.fromstring(data)
                    body = root.find("w:body", W)
                    if body is not None:
                        for para in body.findall("w:p", W):
                            updated += _repair_species_token_runs_in_paragraph(para)
                    data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

        _replace_docx_with_lock_hint(tmp_path, docx_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)

    return {"applied": updated > 0, "updated": updated}


def _replace_paragraph_text_keep_format(
    para: ET.Element,
    text: str,
    *,
    species_italics: bool = False,
) -> None:
    ppr = para.find("w:pPr", W)
    run_rpr: Optional[ET.Element] = None
    for run in para.findall("w:r", W):
        rpr = run.find("w:rPr", W)
        if rpr is not None:
            run_rpr = _clone_xml(rpr)
            break
    if run_rpr is not None:
        run_rpr = _strip_italic_nodes_from_rpr(run_rpr)

    for child in list(para):
        if child.tag != _w_attr("pPr"):
            para.remove(child)

    if species_italics:
        species_text = text or ""
        species_text = re.sub(
            r"\*(Artemisia\s+selengensis|A\.\s*selengensis)\*",
            r"\1",
            species_text,
            flags=re.IGNORECASE,
        )
        pieces: List[Tuple[str, bool]] = []
        last = 0
        for m in _SPECIES_TOKEN_RE.finditer(species_text):
            if m.start() > last:
                pieces.append((species_text[last : m.start()], False))
            pieces.append((m.group(0), True))
            last = m.end()
        if last < len(species_text):
            pieces.append((species_text[last:], False))
        if not pieces:
            pieces = [(species_text, False)]

        for segment, italic in pieces:
            if not segment:
                continue
            new_run = ET.Element(_w_attr("r"))
            if run_rpr is not None:
                new_run.append(_set_rpr_italic(run_rpr, italic))
            elif italic:
                new_rpr = ET.Element(_w_attr("rPr"))
                i_node = ET.Element(_w_attr("i"))
                i_node.set(_w_attr("val"), "1")
                i_cs_node = ET.Element(_w_attr("iCs"))
                i_cs_node.set(_w_attr("val"), "1")
                new_rpr.extend([i_node, i_cs_node])
                new_run.append(new_rpr)
            t = ET.Element(_w_attr("t"))
            if segment[:1].isspace() or segment[-1:].isspace():
                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t.text = segment
            new_run.append(t)
            para.append(new_run)
    else:
        new_run = ET.Element(_w_attr("r"))
        if run_rpr is not None:
            new_run.append(run_rpr)

        t = ET.Element(_w_attr("t"))
        if text[:1].isspace() or text[-1:].isspace():
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = text
        new_run.append(t)
        para.append(new_run)

    if ppr is None:
        ppr_new = ET.Element(_w_attr("pPr"))
        para.insert(0, ppr_new)


def _set_paragraph_style_id(para: ET.Element, style_id: str) -> None:
    if not style_id:
        return
    ppr = para.find("w:pPr", W)
    if ppr is None:
        ppr = ET.Element(_w_attr("pPr"))
        para.insert(0, ppr)
    pstyle = ppr.find("w:pStyle", W)
    if pstyle is None:
        pstyle = ET.Element(_w_attr("pStyle"))
        ppr.insert(0, pstyle)
    pstyle.set(_w_attr("val"), style_id)


def _ensure_paragraph_spacing_zero(para: ET.Element) -> bool:
    ppr = para.find("w:pPr", W)
    if ppr is None:
        ppr = ET.Element(_w_attr("pPr"))
        para.insert(0, ppr)
    spacing = ppr.find("w:spacing", W)
    if spacing is None:
        spacing = ET.Element(_w_attr("spacing"))
        ppr.append(spacing)
    changed = False
    if spacing.get(_w_attr("before"), "") != "0":
        spacing.set(_w_attr("before"), "0")
        changed = True
    if spacing.get(_w_attr("after"), "") != "0":
        spacing.set(_w_attr("after"), "0")
        changed = True
    return changed


def _ensure_paragraph_keep_next(para: ET.Element) -> bool:
    ppr = para.find("w:pPr", W)
    if ppr is None:
        ppr = ET.Element(_w_attr("pPr"))
        para.insert(0, ppr)
    keep_next = ppr.find("w:keepNext", W)
    if keep_next is None:
        keep_next = ET.Element(_w_attr("keepNext"))
        ppr.append(keep_next)
        return True
    return False


def _ensure_paragraph_keep_lines(para: ET.Element) -> bool:
    ppr = para.find("w:pPr", W)
    if ppr is None:
        ppr = ET.Element(_w_attr("pPr"))
        para.insert(0, ppr)
    keep_lines = ppr.find("w:keepLines", W)
    if keep_lines is None:
        keep_lines = ET.Element(_w_attr("keepLines"))
        ppr.append(keep_lines)
        return True
    return False


def _ensure_paragraph_spacing(
    para: ET.Element,
    *,
    before: Optional[str] = None,
    after: Optional[str] = None,
    line: Optional[str] = None,
    line_rule: Optional[str] = None,
) -> bool:
    ppr = para.find("w:pPr", W)
    if ppr is None:
        ppr = ET.Element(_w_attr("pPr"))
        para.insert(0, ppr)
    spacing = ppr.find("w:spacing", W)
    if spacing is None:
        spacing = ET.Element(_w_attr("spacing"))
        ppr.append(spacing)

    changed = False
    targets = {
        "before": before,
        "after": after,
        "line": line,
        "lineRule": line_rule,
    }
    for key, target in targets.items():
        if target is None:
            continue
        if spacing.get(_w_attr(key), "") != target:
            spacing.set(_w_attr(key), target)
            changed = True
    return changed


def _ensure_paragraph_indent_zero(para: ET.Element) -> bool:
    ppr = para.find("w:pPr", W)
    if ppr is None:
        ppr = ET.Element(_w_attr("pPr"))
        para.insert(0, ppr)
    ind = ppr.find("w:ind", W)
    if ind is None:
        ind = ET.Element(_w_attr("ind"))
        ppr.append(ind)
    changed = False
    for key in ["firstLine", "firstLineChars", "left", "right"]:
        if ind.get(_w_attr(key), "") != "0":
            ind.set(_w_attr(key), "0")
            changed = True
    return changed


def _ensure_paragraph_center(para: ET.Element) -> bool:
    ppr = para.find("w:pPr", W)
    if ppr is None:
        ppr = ET.Element(_w_attr("pPr"))
        para.insert(0, ppr)
    jc = ppr.find("w:jc", W)
    if jc is None:
        jc = ET.Element(_w_attr("jc"))
        ppr.append(jc)
    if jc.get(_w_attr("val"), "") != "center":
        jc.set(_w_attr("val"), "center")
        return True
    return False


def _ensure_run_font_size(run: ET.Element, half_points: str) -> bool:
    rpr = run.find("w:rPr", W)
    if rpr is None:
        rpr = ET.Element(_w_attr("rPr"))
        run.insert(0, rpr)
    changed = False
    for key in ["sz", "szCs"]:
        node = rpr.find(f"w:{key}", W)
        if node is None:
            node = ET.Element(_w_attr(key))
            rpr.append(node)
        if node.get(_w_attr("val"), "") != half_points:
            node.set(_w_attr("val"), half_points)
            changed = True
    return changed


def _ensure_run_eastasia_font(run: ET.Element, font_name: str = "宋体") -> bool:
    rpr = run.find("w:rPr", W)
    if rpr is None:
        rpr = ET.Element(_w_attr("rPr"))
        run.insert(0, rpr)
    rfonts = _ensure_child(rpr, "w:rFonts")
    changed = False
    eastasia_attr = _w_attr("eastAsia")
    if rfonts.get(eastasia_attr, "") != font_name:
        rfonts.set(eastasia_attr, font_name)
        changed = True
    hint_attr = _w_attr("hint")
    if rfonts.get(hint_attr, "") != "eastAsia":
        rfonts.set(hint_attr, "eastAsia")
        changed = True
    return changed


def _ensure_run_fonts(
    run: ET.Element,
    *,
    eastasia: Optional[str] = None,
    ascii_font: Optional[str] = None,
    hansi: Optional[str] = None,
    cs: Optional[str] = None,
    hint: Optional[str] = None,
) -> bool:
    rpr = run.find("w:rPr", W)
    if rpr is None:
        rpr = ET.Element(_w_attr("rPr"))
        run.insert(0, rpr)
    rfonts = _ensure_child(rpr, "w:rFonts")
    changed = False
    targets = {
        "eastAsia": eastasia,
        "ascii": ascii_font,
        "hAnsi": hansi,
        "cs": cs,
        "hint": hint,
    }
    for key, target in targets.items():
        if target is None:
            continue
        attr = _w_attr(key)
        if rfonts.get(attr, "") != target:
            rfonts.set(attr, target)
            changed = True
    return changed


def _ensure_run_bold(run: ET.Element) -> bool:
    rpr = run.find("w:rPr", W)
    if rpr is None:
        rpr = ET.Element(_w_attr("rPr"))
        run.insert(0, rpr)
    changed = False
    for key in ["b", "bCs"]:
        node = rpr.find(f"w:{key}", W)
        if node is None:
            node = ET.Element(_w_attr(key))
            rpr.append(node)
        if node.get(_w_attr("val"), "") != "1":
            node.set(_w_attr("val"), "1")
            changed = True
    return changed


def _clear_run_bold(run: ET.Element) -> bool:
    rpr = run.find("w:rPr", W)
    if rpr is None:
        return False
    changed = False
    for key in ["b", "bCs"]:
        node = rpr.find(f"w:{key}", W)
        if node is not None:
            rpr.remove(node)
            changed = True
    return changed


def _replace_paragraph_with_label_value(
    para: ET.Element,
    *,
    label_text: str,
    value_text: str,
) -> bool:
    ppr = para.find("w:pPr", W)
    run_rpr: Optional[ET.Element] = None
    for run in para.findall("w:r", W):
        rpr = run.find("w:rPr", W)
        if rpr is not None:
            run_rpr = _clone_xml(rpr)
            break

    for child in list(para):
        if child.tag != _w_attr("pPr"):
            para.remove(child)

    label_run = ET.Element(_w_attr("r"))
    if run_rpr is not None:
        label_run.append(_clone_xml(run_rpr))
    _ensure_run_bold(label_run)
    label_t = ET.Element(_w_attr("t"))
    label_t.text = label_text
    label_run.append(label_t)
    para.append(label_run)

    value = (value_text or "").strip()
    if value:
        value_run = ET.Element(_w_attr("r"))
        if run_rpr is not None:
            value_run.append(_clone_xml(run_rpr))
        _clear_run_bold(value_run)
        value_t = ET.Element(_w_attr("t"))
        value_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        value_t.text = f" {value}"
        value_run.append(value_t)
        para.append(value_run)

    if ppr is None:
        ppr_new = ET.Element(_w_attr("pPr"))
        para.insert(0, ppr_new)
    return True


def _sync_project_typography_overrides(docx_path: Path) -> Dict[str, object]:
    caption_style_id = _resolve_docx_caption_style_id(docx_path)
    caption_updated = 0
    note_updated = 0
    keyword_updated = 0
    appendix_note_updated = 0

    appendix_note_norm = _norm_plain_text("附录使用说明")
    appendix_note_tail_re = re.compile(rf"^[0-9０-９A-Za-z一二三四五六七八九十\-\.\(\)（）]*{re.escape(appendix_note_norm)}$")

    with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
        tmp_path = Path(tmp.name)

    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
            tmp_path, "w", compression=zipfile.ZIP_DEFLATED
        ) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    root = ET.fromstring(data)
                    body = root.find("w:body", W)
                    if body is not None:
                        for para in body.findall("w:p", W):
                            text = _extract_para_text(para).strip()
                            if not text:
                                continue

                            para_changed = False
                            ppr = para.find("w:pPr", W)
                            pstyle = ppr.find("w:pStyle", W) if ppr is not None else None
                            style_id = pstyle.get(_w_attr("val"), "") if pstyle is not None else ""

                            if style_id == caption_style_id:
                                caption_before = "120"
                                if _is_english_table_caption_text(text):
                                    caption_before = "0"
                                elif _is_figure_caption_paragraph_text(text):
                                    caption_before = "120"
                                if _ensure_paragraph_spacing(para, before=caption_before, after="0"):
                                    para_changed = True
                                for run in para.findall("w:r", W):
                                    if _ensure_run_font_size(run, "21"):
                                        para_changed = True
                                if para_changed:
                                    caption_updated += 1

                            text_norm = _norm_plain_text(text)
                            if text.startswith("注：") or text.startswith("注:"):
                                local_changed = False
                                for run in para.findall("w:r", W):
                                    if _ensure_run_font_size(run, "21"):
                                        local_changed = True
                                if local_changed:
                                    note_updated += 1

                            m_kw_en = re.match(r"(?i)^key\s*words?\s*[:：]\s*(.*)$", text)
                            m_kw_cn = re.match(r"^关键词\s*[:：]\s*(.*)$", text)
                            if m_kw_en:
                                _replace_paragraph_with_label_value(
                                    para, label_text="Key Words:", value_text=m_kw_en.group(1)
                                )
                                if _ensure_paragraph_indent_zero(para):
                                    pass
                                keyword_updated += 1
                            elif m_kw_cn:
                                _replace_paragraph_with_label_value(
                                    para, label_text="关键词：", value_text=m_kw_cn.group(1)
                                )
                                if _ensure_paragraph_indent_zero(para):
                                    pass
                                keyword_updated += 1

                            if (
                                text_norm == appendix_note_norm
                                or text_norm.endswith(appendix_note_norm)
                                or appendix_note_tail_re.fullmatch(text_norm or "")
                            ):
                                local_changed = False
                                for run in para.findall("w:r", W):
                                    if _ensure_run_font_size(run, "28"):
                                        local_changed = True
                                if local_changed:
                                    appendix_note_updated += 1

                    data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

        _replace_docx_with_lock_hint(tmp_path, docx_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)

    return {
        "applied": any([caption_updated, note_updated, keyword_updated, appendix_note_updated]),
        "caption_updated": caption_updated,
        "note_updated": note_updated,
        "keyword_updated": keyword_updated,
        "appendix_note_updated": appendix_note_updated,
    }


def _resolve_docx_caption_style_id(docx_path: Path) -> str:
    try:
        with zipfile.ZipFile(docx_path, "r") as zf:
            styles_root = ET.fromstring(zf.read("word/styles.xml"))
    except Exception:
        return "ae"

    style_ids: Set[str] = set()
    style_name_to_id: Dict[str, str] = {}
    for style in styles_root.findall("w:style", W):
        if style.get(_w_attr("type"), "") != "paragraph":
            continue
        sid = style.get(_w_attr("styleId"), "")
        if not sid:
            continue
        style_ids.add(sid)
        name = style.find("w:name", W)
        style_name = name.get(_w_attr("val"), "") if name is not None else ""
        if style_name:
            style_name_to_id[style_name.casefold()] = sid

    return _resolve_style_id(style_ids, style_name_to_id, ["ae", "Caption", "caption"], ["caption"], "ae")


def _is_table_caption_paragraph_text(text: str) -> bool:
    stripped = (text or "").strip()
    return bool(
        re.match(
            r"(?i)^(附?表(?:\s*[-－—–]?\s*\d+(?:[-\.]\d+)?)|(?:appendix\s+)?table\s+\d+(?:[-\.]\d+)?)",
            stripped,
        )
    )


def _is_figure_caption_paragraph_text(text: str) -> bool:
    stripped = (text or "").strip()
    return bool(
        re.match(
            r"(?i)^(附?图(?:\s*[-－—–]?\s*\d+(?:[-\.]\d+)?)|(?:appendix\s+)?figure\s+[A-Za-z0-9]+(?:[-\.][A-Za-z0-9]+)?)",
            stripped,
        )
    )


def _is_english_table_caption_text(text: str) -> bool:
    stripped = re.sub(r"\s+", " ", (text or "").strip())
    return bool(re.match(r"(?i)^(?:appendix\s+)?table\s+[A-Za-z0-9]+(?:[-\.][A-Za-z0-9]+)*\b", stripped))


def _ensure_child(parent: ET.Element, tag: str) -> ET.Element:
    child = parent.find(tag, W)
    if child is None:
        child = ET.Element(_w_attr(tag.split(":")[-1]))
        parent.append(child)
    return child


def _set_border_attrs(border: ET.Element, val: str, *, sz: Optional[str] = None) -> None:
    border.set(_w_attr("val"), val)
    if sz is not None:
        border.set(_w_attr("sz"), sz)
        border.set(_w_attr("space"), "0")
        border.set(_w_attr("color"), "auto")
    else:
        for key in ["sz", "space", "color", "themeColor", "themeTint", "themeShade"]:
            border.attrib.pop(_w_attr(key), None)


def _set_table_border(tbl_borders: ET.Element, edge: str, val: str, *, sz: Optional[str] = None) -> None:
    border = tbl_borders.find(f"w:{edge}", W)
    if border is None:
        border = ET.Element(_w_attr(edge))
        tbl_borders.append(border)
    _set_border_attrs(border, val, sz=sz)


def _apply_three_line_table_borders(tbl: ET.Element) -> bool:
    tbl_pr = tbl.find("w:tblPr", W)
    if tbl_pr is None:
        tbl_pr = ET.Element(_w_attr("tblPr"))
        tbl.insert(0, tbl_pr)

    tbl_borders = tbl_pr.find("w:tblBorders", W)
    if tbl_borders is None:
        tbl_borders = ET.Element(_w_attr("tblBorders"))
        tbl_pr.append(tbl_borders)

    _set_table_border(tbl_borders, "top", "single", sz="12")
    _set_table_border(tbl_borders, "bottom", "single", sz="12")
    for edge in ["left", "right", "insideH", "insideV"]:
        _set_table_border(tbl_borders, edge, "nil")

    rows = tbl.findall("w:tr", W)
    if not rows:
        return False

    first_row = rows[0]
    for tc in first_row.findall("w:tc", W):
        tc_pr = tc.find("w:tcPr", W)
        if tc_pr is None:
            tc_pr = ET.Element(_w_attr("tcPr"))
            tc.insert(0, tc_pr)
        tc_borders = tc_pr.find("w:tcBorders", W)
        if tc_borders is None:
            tc_borders = ET.Element(_w_attr("tcBorders"))
            tc_pr.append(tc_borders)
        bottom = tc_borders.find("w:bottom", W)
        if bottom is None:
            bottom = ET.Element(_w_attr("bottom"))
            tc_borders.append(bottom)
        _set_border_attrs(bottom, "single", sz="8")
        for edge in ["left", "right"]:
            side = tc_borders.find(f"w:{edge}", W)
            if side is None:
                side = ET.Element(_w_attr(edge))
                tc_borders.append(side)
            _set_border_attrs(side, "nil")
    return True


def _set_tc_vertical_align(tc: ET.Element, align: str = "center") -> None:
    tc_pr = tc.find("w:tcPr", W)
    if tc_pr is None:
        tc_pr = ET.Element(_w_attr("tcPr"))
        tc.insert(0, tc_pr)
    v_align = tc_pr.find("w:vAlign", W)
    if v_align is None:
        v_align = ET.Element(_w_attr("vAlign"))
        tc_pr.append(v_align)
    v_align.set(_w_attr("val"), align)


def _set_table_alignment(tbl: ET.Element, align: str = "center") -> bool:
    tbl_pr = tbl.find("w:tblPr", W)
    if tbl_pr is None:
        tbl_pr = ET.Element(_w_attr("tblPr"))
        tbl.insert(0, tbl_pr)

    changed = False
    jc = tbl_pr.find("w:jc", W)
    if jc is None:
        jc = ET.Element(_w_attr("jc"))
        tbl_pr.append(jc)
        changed = True
    if jc.get(_w_attr("val"), "") != align:
        jc.set(_w_attr("val"), align)
        changed = True

    tbl_ind = tbl_pr.find("w:tblInd", W)
    if tbl_ind is not None:
        tbl_pr.remove(tbl_ind)
        changed = True
    return changed


_WIDE_TABLE_MIN_COLUMNS = 12
_DEFAULT_TEXT_WIDTH_TWIPS = 8312
_DEFAULT_A4_PORTRAIT_W = 11906
_DEFAULT_A4_PORTRAIT_H = 16838
_WIDE_TABLE_TARGET_FILL_RATIO = 0.94
_WIDE_TABLE_CONTEXT_LOOKBACK = 2
_LANDSCAPE_IMAGE_TOKEN = os.environ.get("BENSZ_WORD_LANDSCAPE_IMAGE_TOKEN", "").strip()
_TECHNICAL_ROUTE_TARGET_FILL_RATIO = 0.94


def _extract_table_cell_text(tc: ET.Element) -> str:
    texts: List[str] = []
    for para in tc.findall("w:p", W):
        text = _extract_para_text(para).strip()
        if text:
            texts.append(text)
    return " ".join(texts)


def _parse_twip(value: Optional[str], fallback: int) -> int:
    try:
        parsed = int(value or "")
    except (TypeError, ValueError):
        return fallback
    return parsed if parsed > 0 else fallback


def _set_section_landscape(sect_pr: ET.Element) -> int:
    pg_sz = _ensure_child(sect_pr, "w:pgSz")
    pg_w = _parse_twip(pg_sz.get(_w_attr("w")), _DEFAULT_A4_PORTRAIT_W)
    pg_h = _parse_twip(pg_sz.get(_w_attr("h")), _DEFAULT_A4_PORTRAIT_H)
    land_w = max(pg_w, pg_h)
    land_h = min(pg_w, pg_h)
    pg_sz.set(_w_attr("w"), str(land_w))
    pg_sz.set(_w_attr("h"), str(land_h))
    pg_sz.set(_w_attr("orient"), "landscape")

    pg_mar = _ensure_child(sect_pr, "w:pgMar")
    margin_left = _parse_twip(pg_mar.get(_w_attr("left")), 1797)
    margin_right = _parse_twip(pg_mar.get(_w_attr("right")), 1797)
    return max(3600, land_w - margin_left - margin_right)


def _set_section_portrait(sect_pr: ET.Element) -> None:
    pg_sz = _ensure_child(sect_pr, "w:pgSz")
    pg_w = _parse_twip(pg_sz.get(_w_attr("w")), _DEFAULT_A4_PORTRAIT_W)
    pg_h = _parse_twip(pg_sz.get(_w_attr("h")), _DEFAULT_A4_PORTRAIT_H)
    portrait_w = min(pg_w, pg_h)
    portrait_h = max(pg_w, pg_h)
    pg_sz.set(_w_attr("w"), str(portrait_w))
    pg_sz.set(_w_attr("h"), str(portrait_h))
    if pg_sz.get(_w_attr("orient"), ""):
        pg_sz.attrib.pop(_w_attr("orient"), None)


def _set_section_break_type(sect_pr: ET.Element, break_type: str = "continuous") -> None:
    tp = sect_pr.find("w:type", W)
    if tp is None:
        tp = ET.Element(_w_attr("type"))
        sect_pr.insert(0, tp)
    tp.set(_w_attr("val"), break_type)
    _normalize_section_properties_order(sect_pr)


def _get_section_break_type(sect_pr: Optional[ET.Element]) -> str:
    if sect_pr is None:
        return ""
    tp = sect_pr.find("w:type", W)
    if tp is None:
        return ""
    return tp.get(_w_attr("val"), "") or ""


def _paragraph_has_page_break_before(para: ET.Element) -> bool:
    ppr = para.find("w:pPr", W)
    if ppr is None:
        return False
    page_break_before = ppr.find("w:pageBreakBefore", W)
    if page_break_before is None:
        return False
    current_val = (page_break_before.get(_w_attr("val"), "") or "").strip().casefold()
    return current_val in {"", "1", "true", "on"}


def _expand_wide_table_start_index(children: List[ET.Element], start_idx: int) -> int:
    idx = start_idx
    looked_back = 0
    while idx > 0 and looked_back < _WIDE_TABLE_CONTEXT_LOOKBACK:
        prev = children[idx - 1]
        if prev.tag != _w_attr("p"):
            break
        if not _extract_para_text(prev).strip():
            break
        idx -= 1
        looked_back += 1
        if _paragraph_has_page_break_before(prev):
            break
    return idx


def _make_section_break_paragraph(sect_pr: ET.Element) -> ET.Element:
    para = ET.Element(_w_attr("p"))
    para_pr = ET.Element(_w_attr("pPr"))
    para_pr.append(_clone_xml(sect_pr))
    para.append(para_pr)
    return para


def _paragraph_embeds_image_token(para: ET.Element, token: str) -> bool:
    needle = token.casefold()
    for node in para.findall(".//pic:cNvPr", OOXML_NAMESPACE_PREFIXES):
        for attr_name in ("descr", "name"):
            value = (node.get(attr_name, "") or "").casefold()
            if needle in value:
                return True
    return False


def _set_inline_drawing_size(para: ET.Element, target_cx: int) -> bool:
    drawing = para.find(".//w:drawing", W)
    if drawing is None:
        return False
    inline = drawing.find(".//wp:inline", OOXML_NAMESPACE_PREFIXES)
    extent = inline.find("wp:extent", OOXML_NAMESPACE_PREFIXES) if inline is not None else None
    xfrm_ext = drawing.find(".//a:xfrm/a:ext", OOXML_NAMESPACE_PREFIXES)
    if extent is None or xfrm_ext is None:
        return False
    try:
        old_cx = int(extent.get("cx", "0"))
        old_cy = int(extent.get("cy", "0"))
    except ValueError:
        return False
    if old_cx <= 0 or old_cy <= 0:
        return False
    new_cx = max(1, target_cx)
    new_cy = max(1, int(round(old_cy * (new_cx / old_cx))))
    extent.set("cx", str(new_cx))
    extent.set("cy", str(new_cy))
    xfrm_ext.set("cx", str(new_cx))
    xfrm_ext.set("cy", str(new_cy))
    return new_cx != old_cx or new_cy != old_cy


def _paragraph_section_orientation(para: ET.Element) -> str:
    ppr = para.find("w:pPr", W)
    sect = ppr.find("w:sectPr", W) if ppr is not None else None
    pg_sz = sect.find("w:pgSz", W) if sect is not None else None
    return pg_sz.get(_w_attr("orient"), "") if pg_sz is not None else ""


def _paragraph_has_section_properties(para: ET.Element) -> bool:
    ppr = para.find("w:pPr", W)
    return ppr.find("w:sectPr", W) is not None if ppr is not None else False


def _remove_paragraph_section_properties(para: ET.Element) -> bool:
    ppr = para.find("w:pPr", W)
    if ppr is None:
        return False
    sect = ppr.find("w:sectPr", W)
    if sect is None:
        return False
    ppr.remove(sect)
    return True


def _find_technical_route_caption_end(children: List[ET.Element], image_idx: int) -> int:
    end_idx = image_idx
    scan_idx = image_idx + 1
    while scan_idx < len(children):
        candidate = children[scan_idx]
        if candidate.tag != _w_attr("p"):
            break
        text = _extract_para_text(candidate).strip()
        if not text or not _is_figure_caption_paragraph_text(text):
            break
        end_idx = scan_idx
        scan_idx += 1
    return end_idx


def _technical_route_figure_already_landscape_wrapped(children: List[ET.Element], idx: int) -> bool:
    if idx <= 0 or idx + 1 >= len(children):
        return False
    caption_end_idx = _find_technical_route_caption_end(children, idx)
    prev_para = children[idx - 1]
    if prev_para.tag != _w_attr("p") or not _paragraph_has_section_properties(prev_para):
        return False
    for scan_idx in range(idx, caption_end_idx + 1):
        candidate = children[scan_idx]
        if candidate.tag != _w_attr("p"):
            return False
        if _paragraph_has_section_properties(candidate):
            return False
    post_idx = caption_end_idx + 1
    if post_idx >= len(children):
        return False
    candidate = children[post_idx]
    if candidate.tag != _w_attr("p"):
        return False
    return (
        _paragraph_has_section_properties(candidate)
        and _paragraph_section_orientation(candidate) == "landscape"
    )


def _apply_landscape_section_to_technical_route_figure(docx_path: Path) -> Dict[str, object]:
    if not _LANDSCAPE_IMAGE_TOKEN:
        return {"applied": False, "updated": 0, "reason": "no_landscape_image_token"}

    updated = 0

    with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
        tmp_path = Path(tmp.name)

    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
            tmp_path, "w", compression=zipfile.ZIP_DEFLATED
        ) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    root = ET.fromstring(data)
                    body = root.find("w:body", W)
                    body_sect = body.find("w:sectPr", W) if body is not None else None
                    if body is not None and body_sect is not None:
                        children = list(body)
                        for para in children:
                            if para.tag != _w_attr("p"):
                                continue
                            if not _paragraph_embeds_image_token(para, _LANDSCAPE_IMAGE_TOKEN):
                                continue

                            current_children = list(body)
                            idx = current_children.index(para)
                            caption_end_idx = _find_technical_route_caption_end(current_children, idx)
                            repaired_here = False
                            for block_idx in range(idx, min(caption_end_idx + 1, len(current_children))):
                                candidate = current_children[block_idx]
                                if candidate.tag != _w_attr("p"):
                                    continue
                                if _remove_paragraph_section_properties(candidate):
                                    repaired_here = True
                            current_children = list(body)
                            idx = current_children.index(para)
                            caption_end_idx = _find_technical_route_caption_end(current_children, idx)
                            if _technical_route_figure_already_landscape_wrapped(current_children, idx):
                                if repaired_here:
                                    updated += 1
                                continue
                            portrait_sect = _clone_xml(body_sect)
                            _set_section_portrait(portrait_sect)
                            landscape_sect = _clone_xml(body_sect)
                            landscape_text_width = _set_section_landscape(landscape_sect)
                            target_width_twips = max(
                                3600,
                                int(round(landscape_text_width * _TECHNICAL_ROUTE_TARGET_FILL_RATIO)),
                            )
                            target_cx = target_width_twips * EMU_PER_TWIP
                            _set_section_break_type(portrait_sect, "nextPage")
                            _set_section_break_type(landscape_sect, "nextPage")
                            resized = _set_inline_drawing_size(para, target_cx)

                            children = list(body)
                            idx = children.index(para)
                            if idx <= 0 or children[idx - 1].tag != _w_attr("p"):
                                body.insert(idx, _make_section_break_paragraph(portrait_sect))
                                repaired_here = True
                            else:
                                prev_para = children[idx - 1]
                                prev_ppr = prev_para.find("w:pPr", W)
                                prev_sect = prev_ppr.find("w:sectPr", W) if prev_ppr is not None else None
                                if prev_sect is None:
                                    body.insert(idx, _make_section_break_paragraph(portrait_sect))
                                    repaired_here = True
                                else:
                                    _set_section_portrait(prev_sect)
                                    _set_section_break_type(prev_sect, "nextPage")

                            children = list(body)
                            idx = children.index(para)
                            caption_end_idx = _find_technical_route_caption_end(children, idx)
                            insert_after = caption_end_idx + 1
                            if (
                                insert_after < len(children)
                                and children[insert_after].tag == _w_attr("p")
                                and _paragraph_has_section_properties(children[insert_after])
                                and _paragraph_section_orientation(children[insert_after]) == "landscape"
                            ):
                                post_sect = children[insert_after].find("w:pPr/w:sectPr", W)
                                if post_sect is not None:
                                    _set_section_landscape(post_sect)
                                    _set_section_break_type(post_sect, "nextPage")
                            else:
                                body.insert(insert_after, _make_section_break_paragraph(landscape_sect))
                                repaired_here = True
                            if repaired_here or resized:
                                updated += 1
                            break
                    data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

        _replace_docx_with_lock_hint(tmp_path, docx_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)

    return {"applied": updated > 0, "updated": updated}


def _table_column_count(tbl: ET.Element) -> int:
    first_row = tbl.find("w:tr", W)
    if first_row is None:
        return 0
    return len(first_row.findall("w:tc", W))


def _apply_landscape_section_around_table(
    body: ET.Element,
    tbl: ET.Element,
    caption_para: Optional[ET.Element],
) -> Optional[int]:
    if _table_column_count(tbl) < _WIDE_TABLE_MIN_COLUMNS:
        return None

    body_sect = body.find("w:sectPr", W)
    if body_sect is None:
        return None

    children = list(body)
    if tbl not in children:
        return None
    idx_tbl = children.index(tbl)

    start_idx = idx_tbl
    if caption_para is not None and caption_para in children:
        idx_cap = children.index(caption_para)
        if idx_cap < idx_tbl:
            start_idx = _expand_wide_table_start_index(children, idx_cap)

    portrait_sect = _clone_xml(body_sect)
    _set_section_portrait(portrait_sect)
    landscape_sect = _clone_xml(body_sect)
    landscape_text_width = _set_section_landscape(landscape_sect)
    landscape_target_width = max(3600, int(round(landscape_text_width * _WIDE_TABLE_TARGET_FILL_RATIO)))
    inherited_break_type = _get_section_break_type(body_sect) or "continuous"
    _set_section_break_type(portrait_sect, inherited_break_type)
    _set_section_break_type(landscape_sect, inherited_break_type)

    pre_break = _make_section_break_paragraph(portrait_sect)
    post_break = _make_section_break_paragraph(landscape_sect)

    body.insert(start_idx, pre_break)
    children = list(body)
    idx_tbl = children.index(tbl)
    body.insert(idx_tbl + 1, post_break)

    return landscape_target_width


def _resolve_table_width_twips(tbl: ET.Element) -> int:
    tbl_grid = tbl.find("w:tblGrid", W)
    if tbl_grid is not None:
        widths = []
        for grid_col in tbl_grid.findall("w:gridCol", W):
            try:
                widths.append(int(grid_col.get(_w_attr("w"), "0")))
            except ValueError:
                continue
        if widths and sum(widths) > 0:
            return sum(widths)

    first_row = tbl.find("w:tr", W)
    if first_row is not None:
        widths = []
        for tc in first_row.findall("w:tc", W):
            tc_pr = tc.find("w:tcPr", W)
            tc_w = tc_pr.find("w:tcW", W) if tc_pr is not None else None
            if tc_w is None:
                continue
            try:
                widths.append(int(tc_w.get(_w_attr("w"), "0")))
            except ValueError:
                continue
        if widths and sum(widths) > 0:
            return sum(widths)

    return _DEFAULT_TEXT_WIDTH_TWIPS


def _calculate_wide_table_weights(header_cells: List[str]) -> List[float]:
    weights = [1.0] * len(header_cells)
    if not weights:
        return weights

    for idx, raw_text in enumerate(header_cells):
        text = re.sub(r"\s+", " ", (raw_text or "")).strip()
        lower = text.casefold()
        if "化合物" in text or "compound" in lower:
            weights[idx] = 3.2
            continue
        if "均值" in text or "mean" in lower or "sd" in lower or "±" in text:
            weights[idx] = 2.4
            continue
        if re.search(r"(?:^|\s)n$", lower):
            weights[idx] = 0.75
            continue
        if text in {"F 值", "p 值", "η^2", "η²"}:
            weights[idx] = 0.85
            continue
        if "检出数" in text or "总数" in text:
            weights[idx] = 0.85
            continue

    if len(weights) >= 1:
        weights[0] = max(weights[0], 3.2)
    return weights


def _allocate_column_widths(weights: List[float], total_width: int) -> List[int]:
    if not weights:
        return []
    total_width = max(total_width, len(weights) * 120)
    min_width = 180
    if total_width < min_width * len(weights):
        min_width = max(120, total_width // len(weights))

    base_width = total_width - min_width * len(weights)
    weight_sum = sum(weights) or float(len(weights))
    widths = [min_width + int(round(base_width * weight / weight_sum)) for weight in weights]
    diff = total_width - sum(widths)
    widths[-1] += diff
    if widths[-1] < min_width:
        widths[-1] = min_width
    return widths


def _set_table_column_widths(tbl: ET.Element, widths: List[int]) -> None:
    if not widths:
        return
    total_width = sum(widths)
    tbl_pr = tbl.find("w:tblPr", W)
    if tbl_pr is None:
        tbl_pr = ET.Element(_w_attr("tblPr"))
        tbl.insert(0, tbl_pr)

    tbl_w = tbl_pr.find("w:tblW", W)
    if tbl_w is None:
        tbl_w = ET.Element(_w_attr("tblW"))
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(_w_attr("w"), str(total_width))
    tbl_w.set(_w_attr("type"), "dxa")

    tbl_grid = tbl.find("w:tblGrid", W)
    if tbl_grid is None:
        insert_index = 1 if tbl_pr is not None else 0
        tbl_grid = ET.Element(_w_attr("tblGrid"))
        tbl.insert(insert_index, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = ET.Element(_w_attr("gridCol"))
        grid_col.set(_w_attr("w"), str(width))
        tbl_grid.append(grid_col)

    for row in tbl.findall("w:tr", W):
        cells = row.findall("w:tc", W)
        for idx, tc in enumerate(cells):
            if idx >= len(widths):
                break
            tc_pr = tc.find("w:tcPr", W)
            if tc_pr is None:
                tc_pr = ET.Element(_w_attr("tcPr"))
                tc.insert(0, tc_pr)
            tc_w = tc_pr.find("w:tcW", W)
            if tc_w is None:
                tc_w = ET.Element(_w_attr("tcW"))
                tc_pr.insert(0, tc_w)
            tc_w.set(_w_attr("w"), str(widths[idx]))
            tc_w.set(_w_attr("type"), "dxa")


def _rebalance_wide_table_columns(tbl: ET.Element, preferred_total_width: Optional[int] = None) -> bool:
    rows = tbl.findall("w:tr", W)
    if not rows:
        return False
    header_cells = rows[0].findall("w:tc", W)
    column_count = len(header_cells)
    if column_count < _WIDE_TABLE_MIN_COLUMNS:
        return False

    header_texts = [_extract_table_cell_text(tc) for tc in header_cells]
    weights = _calculate_wide_table_weights(header_texts)
    if len(weights) != column_count:
        return False

    total_width = _resolve_table_width_twips(tbl)
    if preferred_total_width is not None:
        total_width = max(total_width, preferred_total_width)
    widths = _allocate_column_widths(weights, total_width)
    _set_table_column_widths(tbl, widths)
    return True


def _set_table_row_pagination(tbl: ET.Element) -> bool:
    rows = tbl.findall("w:tr", W)
    if not rows:
        return False

    for idx, row in enumerate(rows):
        tr_pr = row.find("w:trPr", W)
        if tr_pr is None:
            tr_pr = ET.Element(_w_attr("trPr"))
            row.insert(0, tr_pr)

        cant_split = tr_pr.find("w:cantSplit", W)
        if cant_split is None:
            cant_split = ET.Element(_w_attr("cantSplit"))
            tr_pr.append(cant_split)

        if idx == 0:
            tbl_header = tr_pr.find("w:tblHeader", W)
            if tbl_header is None:
                tbl_header = ET.Element(_w_attr("tblHeader"))
                tr_pr.append(tbl_header)

    return True


def _normalize_table_paragraph_layout(tbl: ET.Element) -> bool:
    updated = _set_table_alignment(tbl, "center")
    for tc in tbl.findall(".//w:tc", W):
        _set_tc_vertical_align(tc, "center")
        for para in tc.findall("w:p", W):
            if _ensure_paragraph_indent_zero(para):
                updated = True

            ppr = para.find("w:pPr", W)
            if ppr is None:
                ppr = ET.Element(_w_attr("pPr"))
                para.insert(0, ppr)
            spacing = ppr.find("w:spacing", W)
            if spacing is None:
                spacing = ET.Element(_w_attr("spacing"))
                ppr.append(spacing)
            if spacing.get(_w_attr("before"), "") != "0":
                spacing.set(_w_attr("before"), "0")
                updated = True
            if spacing.get(_w_attr("after"), "") != "0":
                spacing.set(_w_attr("after"), "0")
                updated = True
            if spacing.get(_w_attr("line"), "") != "240":
                spacing.set(_w_attr("line"), "240")
                updated = True
            if spacing.get(_w_attr("lineRule"), "") != "auto":
                spacing.set(_w_attr("lineRule"), "auto")
                updated = True

            jc = ppr.find("w:jc", W)
            if jc is None:
                jc = ET.Element(_w_attr("jc"))
                ppr.append(jc)
            if jc.get(_w_attr("val"), "") != "center":
                jc.set(_w_attr("val"), "center")
                updated = True

            for run in para.findall("w:r", W):
                if _ensure_run_font_size(run, "21"):
                    updated = True
    return updated


_CJK_CHAR_RE = re.compile(r"[\u3400-\u4dbf\u4e00-\u9fff\uf900-\ufaff]")
_CJK_QUOTE_RE = re.compile(r"[“”‘’]")
_CJK_PUNCT_RE = re.compile(r"[，。；：？！、（）《》〈〉【】「」『』〔〕…—～·￥％]")
_GREEK_CHAR_RE = re.compile(r"[\u0370-\u03ff\u1f00-\u1fff]")
_BIBLIOGRAPHY_STOP_MARKERS = {
    _norm_plain_text("附录"),
    _norm_plain_text("致谢"),
    _norm_plain_text("作者简历及攻读学位期间发表的学术论文与其他相关学术成果"),
}


def _resolve_style_eastasia_font(
    style_id: str,
    style_fonts: Dict[str, str],
    style_based_on: Dict[str, str],
    default_style_id: str,
    cache: Dict[str, str],
) -> str:
    style_id = style_id or default_style_id
    if style_id in cache:
        return cache[style_id]

    visited: Set[str] = set()
    current = style_id
    while current and current not in visited:
        visited.add(current)
        font = style_fonts.get(current, "")
        if font:
            cache[style_id] = font
            return font
        current = style_based_on.get(current, "")

    fallback = style_fonts.get(default_style_id, "") or "宋体"
    cache[style_id] = fallback
    return fallback


def _run_contains_non_cjk_text(text: str) -> bool:
    for ch in text:
        if ch.isspace():
            continue
        if _CJK_CHAR_RE.fullmatch(ch):
            continue
        if _CJK_PUNCT_RE.fullmatch(ch):
            continue
        return True
    return False


def _run_contains_cjk_char(text: str) -> bool:
    return bool(_CJK_CHAR_RE.search(text or ""))


def _force_cjk_quote_runs_to_eastasia_font(docx_path: Path) -> Dict[str, object]:
    updated_runs = 0
    samples: List[str] = []

    with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
        tmp_path = Path(tmp.name)

    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
            tmp_path, "w", compression=zipfile.ZIP_DEFLATED
        ) as zout:
            styles_root = ET.fromstring(zin.read("word/styles.xml"))
            style_fonts: Dict[str, str] = {}
            style_based_on: Dict[str, str] = {}
            default_style_id = "a"

            for style in styles_root.findall("w:style", W):
                if style.get(_w_attr("type"), "") != "paragraph":
                    continue
                sid = style.get(_w_attr("styleId"), "")
                if not sid:
                    continue
                if style.get(_w_attr("default"), "") == "1":
                    default_style_id = sid
                based_on = style.find("w:basedOn", W)
                if based_on is not None:
                    style_based_on[sid] = based_on.get(_w_attr("val"), "")
                rpr = style.find("w:rPr", W)
                rfonts = rpr.find("w:rFonts", W) if rpr is not None else None
                if rfonts is not None:
                    eastasia = rfonts.get(_w_attr("eastAsia"), "")
                    if eastasia:
                        style_fonts[sid] = eastasia

            font_cache: Dict[str, str] = {}

            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    root = ET.fromstring(data)
                    body = root.find("w:body", W)
                    if body is not None:
                        for para in body.findall(".//w:p", W):
                            para_text = _extract_para_text(para)
                            if not para_text or not _CJK_CHAR_RE.search(para_text) or not _CJK_QUOTE_RE.search(para_text):
                                continue

                            ppr = para.find("w:pPr", W)
                            pstyle = ppr.find("w:pStyle", W) if ppr is not None else None
                            style_id = pstyle.get(_w_attr("val"), "") if pstyle is not None else default_style_id
                            target_font = _resolve_style_eastasia_font(
                                style_id, style_fonts, style_based_on, default_style_id, font_cache
                            )

                            for run in list(para.findall("w:r", W)):
                                run_text = "".join((node.text or "") for node in run.findall("w:t", W))
                                if not run_text or not _CJK_QUOTE_RE.search(run_text):
                                    continue

                                non_rpr_children = [child for child in list(run) if child.tag != _w_attr("rPr")]
                                if any(child.tag != _w_attr("t") for child in non_rpr_children):
                                    continue

                                pieces = re.split(r"([“”‘’])", run_text)
                                if len(pieces) <= 1:
                                    continue

                                insert_at = list(para).index(run)
                                run_rpr = run.find("w:rPr", W)

                                for piece in pieces:
                                    if not piece:
                                        continue
                                    new_run = ET.Element(_w_attr("r"))
                                    if run_rpr is not None:
                                        new_run.append(_clone_xml(run_rpr))
                                    if _CJK_QUOTE_RE.fullmatch(piece):
                                        new_rpr = new_run.find("w:rPr", W)
                                        if new_rpr is None:
                                            new_rpr = ET.Element(_w_attr("rPr"))
                                            new_run.insert(0, new_rpr)
                                        rfonts = _ensure_child(new_rpr, "w:rFonts")
                                        for key in ["ascii", "hAnsi", "eastAsia"]:
                                            rfonts.set(_w_attr(key), target_font)
                                        rfonts.set(_w_attr("hint"), "eastAsia")
                                    t = ET.Element(_w_attr("t"))
                                    if piece[:1].isspace() or piece[-1:].isspace():
                                        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                                    t.text = piece
                                    new_run.append(t)
                                    para.insert(insert_at, new_run)
                                    insert_at += 1
                                    if _CJK_QUOTE_RE.fullmatch(piece):
                                        updated_runs += 1
                                        if len(samples) < 8:
                                            samples.append(piece)
                                para.remove(run)

                    data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

        _replace_docx_with_lock_hint(tmp_path, docx_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)

    return {
        "applied": updated_runs > 0,
        "updated_runs": updated_runs,
        "samples": samples,
    }


def _force_non_cjk_runs_to_times_new_roman(
    docx_path: Path, font_name: str = "Times New Roman"
) -> Dict[str, object]:
    updated_runs = 0
    updated_parts = 0
    samples: List[str] = []

    with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
        tmp_path = Path(tmp.name)

    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
            tmp_path, "w", compression=zipfile.ZIP_DEFLATED
        ) as zout:
            # 仅处理正文分部，避免重写 styles/theme/numbering 等 XML 触发 Word 可读性修复提示。
            target_parts = {"word/document.xml"}
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename not in target_parts:
                    zout.writestr(item, data)
                    continue

                try:
                    root = ET.fromstring(data)
                except Exception:
                    zout.writestr(item, data)
                    continue

                part_updated = False
                for run in root.findall(".//w:r", W):
                    run_text_parts: List[str] = []
                    for tag in ["w:t", "w:delText"]:
                        run_text_parts.extend((node.text or "") for node in run.findall(tag, W))
                    run_text = "".join(run_text_parts)
                    if not run_text or not _run_contains_non_cjk_text(run_text):
                        continue
                    contains_cjk_char = _run_contains_cjk_char(run_text)
                    contains_greek_char = bool(_GREEK_CHAR_RE.search(run_text))

                    rpr = run.find("w:rPr", W)
                    if rpr is None:
                        rpr = ET.Element(_w_attr("rPr"))
                        run.insert(0, rpr)
                    rfonts = _ensure_child(rpr, "w:rFonts")

                    changed = False
                    for key in ["ascii", "hAnsi", "cs"]:
                        attr = _w_attr(key)
                        if rfonts.get(attr, "") != font_name:
                            rfonts.set(attr, font_name)
                            changed = True

                    # Greek symbols (e.g. rho/eta) may fallback to East Asia font in Word.
                    # For runs without CJK characters, pin eastAsia to Times New Roman too.
                    if not contains_cjk_char:
                        eastasia_attr = _w_attr("eastAsia")
                        if rfonts.get(eastasia_attr, "") != font_name:
                            rfonts.set(eastasia_attr, font_name)
                            changed = True
                        hint_attr = _w_attr("hint")
                        if rfonts.get(hint_attr, "") == "eastAsia":
                            rfonts.set(hint_attr, "default")
                            changed = True
                    elif contains_greek_char:
                        # For mixed CJK + Greek runs, avoid forcing eastAsia font (to keep Chinese
                        # typography), but still stop East Asia hint from hijacking Greek glyphs.
                        hint_attr = _w_attr("hint")
                        if rfonts.get(hint_attr, "") == "eastAsia":
                            rfonts.set(hint_attr, "default")
                            changed = True

                    if changed:
                        part_updated = True
                        updated_runs += 1
                        if len(samples) < 8:
                            samples.append(run_text.strip()[:60])

                if part_updated:
                    updated_parts += 1
                    data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

        _replace_docx_with_lock_hint(tmp_path, docx_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)

    return {
        "applied": updated_runs > 0,
        "font": font_name,
        "updated_runs": updated_runs,
        "updated_parts": updated_parts,
        "samples": samples,
    }


def _normalize_bibliography_paragraph_layout(docx_path: Path) -> Dict[str, object]:
    updated = 0
    italic_cleared_runs = 0
    title_updated = 0

    with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
        tmp_path = Path(tmp.name)

    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
            tmp_path, "w", compression=zipfile.ZIP_DEFLATED
        ) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    root = ET.fromstring(data)
                    body = root.find("w:body", W)
                    if body is not None:
                        in_bibliography = False
                        for child in list(body):
                            if child.tag != _w_attr("p"):
                                continue

                            text = _extract_para_text(child).strip()
                            text_norm = _norm_plain_text(text)
                            if text_norm == _norm_plain_text("参考文献"):
                                changed = False
                                if _ensure_paragraph_center(child):
                                    changed = True
                                if _ensure_paragraph_spacing(child, before="0", after="0"):
                                    changed = True
                                if _ensure_paragraph_indent_zero(child):
                                    changed = True
                                for run in child.findall("w:r", W):
                                    if _ensure_run_fonts(
                                        run,
                                        eastasia="黑体",
                                        ascii_font="Times New Roman",
                                        hansi="Times New Roman",
                                        cs="Times New Roman",
                                        hint="eastAsia",
                                    ):
                                        changed = True
                                    if _ensure_run_font_size(run, "28"):
                                        changed = True
                                    if _ensure_run_bold(run):
                                        changed = True
                                    rpr = run.find("w:rPr", W)
                                    if rpr is not None:
                                        has_italic = rpr.find("w:i", W) is not None or rpr.find("w:iCs", W) is not None
                                        if has_italic:
                                            cleaned = _strip_italic_nodes_from_rpr(rpr)
                                            run.remove(rpr)
                                            if len(list(cleaned)):
                                                run.insert(0, cleaned)
                                            italic_cleared_runs += 1
                                            changed = True
                                if changed:
                                    title_updated += 1
                                in_bibliography = True
                                continue

                            if in_bibliography and text_norm in _BIBLIOGRAPHY_STOP_MARKERS:
                                in_bibliography = False
                                continue

                            if not in_bibliography or not text:
                                continue

                            ppr = child.find("w:pPr", W)
                            if ppr is None:
                                ppr = ET.Element(_w_attr("pPr"))
                                child.insert(0, ppr)

                            ind = ppr.find("w:ind", W)
                            if ind is None:
                                ind = ET.Element(_w_attr("ind"))
                                ppr.append(ind)

                            desired = {
                                "left": "480",
                                "hanging": "480",
                                "hangingChars": "200",
                                "firstLine": "0",
                                "firstLineChars": "0",
                            }
                            changed = False
                            for key, value in desired.items():
                                if ind.get(_w_attr(key), "") != value:
                                    ind.set(_w_attr(key), value)
                                    changed = True

                            if _ensure_paragraph_spacing(child, before="0", after="0", line="300", line_rule="auto"):
                                changed = True

                            jc = ppr.find("w:jc", W)
                            if jc is None:
                                jc = ET.Element(_w_attr("jc"))
                                ppr.append(jc)
                            if jc.get(_w_attr("val"), "") != "both":
                                jc.set(_w_attr("val"), "both")
                                changed = True

                            for run in child.findall("w:r", W):
                                if _ensure_run_fonts(
                                    run,
                                    eastasia="宋体",
                                    ascii_font="Times New Roman",
                                    hansi="Times New Roman",
                                    cs="Times New Roman",
                                    hint="eastAsia",
                                ):
                                    changed = True
                                if _ensure_run_font_size(run, "24"):
                                    changed = True
                                rpr = run.find("w:rPr", W)
                                if rpr is None:
                                    continue
                                has_italic = rpr.find("w:i", W) is not None or rpr.find("w:iCs", W) is not None
                                if not has_italic:
                                    continue
                                cleaned = _strip_italic_nodes_from_rpr(rpr)
                                run.remove(rpr)
                                if len(list(cleaned)):
                                    run.insert(0, cleaned)
                                italic_cleared_runs += 1
                                changed = True

                            if changed:
                                updated += 1

                    data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

        _replace_docx_with_lock_hint(tmp_path, docx_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)

    return {
        "applied": updated > 0 or title_updated > 0,
        "title_updated": title_updated,
        "updated": updated,
        "italic_cleared_runs": italic_cleared_runs,
    }


def _normalize_body_list_paragraph_indent(docx_path: Path) -> Dict[str, object]:
    updated = 0

    heading_style_norms = {
        "1",
        "2",
        "3",
        "4",
        "5",
        "6",
        "heading1",
        "heading2",
        "heading3",
        "heading4",
        "heading5",
        "heading6",
    }

    with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
        tmp_path = Path(tmp.name)

    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
            tmp_path, "w", compression=zipfile.ZIP_DEFLATED
        ) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    root = ET.fromstring(data)
                    body = root.find("w:body", W)
                    if body is not None:
                        in_bibliography = False
                        for para in body.findall("w:p", W):
                            text = _extract_para_text(para).strip()
                            text_norm = _norm_plain_text(text)
                            if text_norm == _norm_plain_text("参考文献"):
                                in_bibliography = True
                                continue
                            if in_bibliography and text_norm in _BIBLIOGRAPHY_STOP_MARKERS:
                                in_bibliography = False
                                continue
                            if in_bibliography or not text:
                                continue

                            ppr = para.find("w:pPr", W)
                            if ppr is None:
                                continue
                            pstyle = ppr.find("w:pStyle", W)
                            sid = pstyle.get(_w_attr("val"), "") if pstyle is not None else ""
                            sid_norm = re.sub(r"[\s_-]+", "", sid.casefold())
                            if sid_norm in heading_style_norms or sid_norm.startswith("toc"):
                                continue

                            numpr = ppr.find("w:numPr", W)
                            if numpr is None:
                                continue
                            ilvl = numpr.find("w:ilvl", W)
                            try:
                                level = int(ilvl.get(_w_attr("val"), "0")) if ilvl is not None else 0
                            except ValueError:
                                level = 0

                            ind = ppr.find("w:ind", W)
                            if ind is None:
                                ind = ET.Element(_w_attr("ind"))
                                ppr.append(ind)

                            desired = {
                                "left": str(360 * (level + 1)),
                                "leftChars": "0",
                                "hanging": "360",
                                "hangingChars": "0",
                                "firstLine": "0",
                                "firstLineChars": "0",
                            }
                            changed = False
                            for key, value in desired.items():
                                if ind.get(_w_attr(key), "") != value:
                                    ind.set(_w_attr(key), value)
                                    changed = True
                            if changed:
                                updated += 1

                    data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

        _replace_docx_with_lock_hint(tmp_path, docx_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)

    return {
        "applied": updated > 0,
        "updated": updated,
    }


def _normalize_body_first_line_indent(docx_path: Path) -> Dict[str, object]:
    updated = 0
    style_updated = 0

    normal_style_ids: Set[str] = {"a", "normal"}
    skip_style_norms = {
        "1",
        "2",
        "3",
        "4",
        "5",
        "6",
        "heading1",
        "heading2",
        "heading3",
        "heading4",
        "heading5",
        "heading6",
        "caption",
    }

    def _style_norm(style_id: str) -> str:
        return re.sub(r"[\s_-]+", "", (style_id or "").casefold())

    def _ensure_two_char_first_line(para: ET.Element) -> bool:
        ppr = para.find("w:pPr", W)
        if ppr is None:
            ppr = ET.Element(_w_attr("pPr"))
            para.insert(0, ppr)
        ind = ppr.find("w:ind", W)
        if ind is None:
            ind = ET.Element(_w_attr("ind"))
            ppr.append(ind)

        changed = False
        if ind.get(_w_attr("firstLineChars"), "") != "200":
            ind.set(_w_attr("firstLineChars"), "200")
            changed = True
        for key in ["firstLine", "hanging", "hangingChars"]:
            if ind.get(_w_attr(key)) is not None:
                del ind.attrib[_w_attr(key)]
                changed = True
        return changed

    def _ensure_normal_style_uses_char_indent(styles_root: ET.Element) -> int:
        local_updated = 0
        for style in styles_root.findall("w:style", W):
            if style.get(_w_attr("type"), "") != "paragraph":
                continue
            style_id = style.get(_w_attr("styleId"), "")
            name_node = style.find("w:name", W)
            name = name_node.get(_w_attr("val"), "") if name_node is not None else ""
            if _style_norm(style_id) not in normal_style_ids and name.casefold() != "normal":
                continue

            ppr = style.find("w:pPr", W)
            if ppr is None:
                ppr = ET.Element(_w_attr("pPr"))
                style.append(ppr)
            ind = ppr.find("w:ind", W)
            if ind is None:
                ind = ET.Element(_w_attr("ind"))
                ppr.append(ind)

            changed = False
            if ind.get(_w_attr("firstLineChars"), "") != "200":
                ind.set(_w_attr("firstLineChars"), "200")
                changed = True
            if ind.get(_w_attr("firstLine")) is not None:
                del ind.attrib[_w_attr("firstLine")]
                changed = True
            if changed:
                local_updated += 1
        return local_updated

    def _is_candidate_body_paragraph(para: ET.Element, text: str, *, in_bibliography: bool) -> bool:
        if in_bibliography or not text:
            return False
        if _paragraph_contains_embedded_image(para):
            return False
        if _is_keyword_paragraph_text(text):
            return False
        if _norm_plain_text(text) in {_norm_plain_text("摘要"), _norm_plain_text("abstract")}:
            return False
        if _is_figure_caption_paragraph_text(text) or _is_table_caption_paragraph_text(text):
            return False
        if re.match(r"(?i)^(注[:：]|note[s]?[:：])", text.strip()):
            return False

        ppr = para.find("w:pPr", W)
        if ppr is None:
            return True
        if ppr.find("w:numPr", W) is not None:
            return False
        pstyle = ppr.find("w:pStyle", W)
        style_id = pstyle.get(_w_attr("val"), "") if pstyle is not None else ""
        style_norm = _style_norm(style_id)
        if style_norm in skip_style_norms or style_norm.startswith("toc"):
            return False
        return style_norm in normal_style_ids or style_norm == ""

    def _is_body_indent_scope_start(para: ET.Element, text_norm: str) -> bool:
        if text_norm in {
            _norm_plain_text("摘要"),
            _norm_plain_text("abstract"),
            _norm_plain_text("目录"),
            _norm_plain_text("图表目录"),
            _norm_plain_text("图目录"),
            _norm_plain_text("表目录"),
        }:
            return True
        if re.match(r"^第[一二三四五六七八九十]+章", text_norm):
            return True
        ppr = para.find("w:pPr", W)
        if ppr is None:
            return False
        pstyle = ppr.find("w:pStyle", W)
        style_id = pstyle.get(_w_attr("val"), "") if pstyle is not None else ""
        return _style_norm(style_id) in {"1", "heading1"}

    def _is_unindented_front_matter_title(text_norm: str) -> bool:
        return text_norm in {
            _norm_plain_text("摘要"),
            _norm_plain_text("abstract"),
            _norm_plain_text("目录"),
            _norm_plain_text("图表目录"),
            _norm_plain_text("图目录"),
            _norm_plain_text("表目录"),
        }

    with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
        tmp_path = Path(tmp.name)

    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
            tmp_path, "w", compression=zipfile.ZIP_DEFLATED
        ) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/styles.xml":
                    root = ET.fromstring(data)
                    style_updated += _ensure_normal_style_uses_char_indent(root)
                    data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                elif item.filename == "word/document.xml":
                    root = ET.fromstring(data)
                    body = root.find("w:body", W)
                    if body is not None:
                        in_bibliography = False
                        body_indent_scope = False
                        for para in body.findall("w:p", W):
                            text = _extract_para_text(para).strip()
                            text_norm = _norm_plain_text(text)
                            if not body_indent_scope:
                                if _is_body_indent_scope_start(para, text_norm):
                                    body_indent_scope = True
                                elif text:
                                    if _ensure_paragraph_indent_zero(para):
                                        updated += 1
                                    continue

                            if _is_unindented_front_matter_title(text_norm):
                                if _ensure_paragraph_indent_zero(para):
                                    updated += 1
                                continue

                            if text_norm == _norm_plain_text("参考文献"):
                                in_bibliography = True
                                continue
                            if in_bibliography and text_norm in _BIBLIOGRAPHY_STOP_MARKERS:
                                in_bibliography = False
                                continue
                            if _is_candidate_body_paragraph(para, text, in_bibliography=in_bibliography):
                                if _ensure_two_char_first_line(para):
                                    updated += 1
                    data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

        _replace_docx_with_lock_hint(tmp_path, docx_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)

    return {
        "applied": updated > 0 or style_updated > 0,
        "updated": updated,
        "style_updated": style_updated,
    }


def _normalize_visible_latex_tilde_spacing(docx_path: Path) -> Dict[str, object]:
    """Remove visible LaTeX nonbreaking-space tildes that survived into DOCX text."""
    updated = 0
    tmp_path: Optional[Path] = None

    def _normalize_text(value: str) -> str:
        normalized = re.sub(r"(?<=\S)~(?=\S)", " ", value or "")
        return _normalize_cjk_number_classifier_spacing(normalized)

    with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
        tmp_path = Path(tmp.name)

    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
            tmp_path, "w", compression=zipfile.ZIP_DEFLATED
        ) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    root = ET.fromstring(data)
                    for text_node in root.findall(".//w:t", W):
                        original = text_node.text or ""
                        normalized = _normalize_text(original)
                        if normalized != original:
                            text_node.text = normalized
                            updated += 1
                    data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

        _replace_docx_with_lock_hint(tmp_path, docx_path)
    finally:
        if tmp_path is not None and tmp_path.exists():
            tmp_path.unlink(missing_ok=True)

    return {"applied": updated > 0, "updated": updated}


def _rewrite_figure_caption_placeholders(docx_path: Path) -> Dict[str, object]:
    caption_style_id = _resolve_docx_caption_style_id(docx_path)
    updated = 0
    samples: List[str] = []

    with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
        tmp_path = Path(tmp.name)

    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
            tmp_path, "w", compression=zipfile.ZIP_DEFLATED
        ) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    root = ET.fromstring(data)
                    body = root.find("w:body", W)
                    if body is not None:
                        for para in body.findall("w:p", W):
                            text = _extract_para_text(para)
                            match = CAPTION_MARKER_RE.match(text)
                            if not match:
                                continue
                            kind, lang, sequence, title = match.groups()
                            formatted = _format_caption_text(kind, lang, sequence.strip(), title)
                            _replace_paragraph_text_keep_format(
                                para,
                                formatted,
                                species_italics=(
                                    lang.lower() == "en" and bool(_SPECIES_TOKEN_RE.search(formatted))
                                ),
                            )
                            _set_paragraph_style_id(para, caption_style_id)
                            _ensure_paragraph_spacing_zero(para)
                            _ensure_paragraph_keep_lines(para)
                            if lang.lower() == "zh":
                                for run in para.findall("w:r", W):
                                    run_text = "".join((node.text or "") for node in run.findall("w:t", W))
                                    if _run_contains_cjk_char(run_text):
                                        _ensure_run_eastasia_font(run, "宋体")
                            updated += 1
                            if len(samples) < 6:
                                samples.append(formatted)
                    data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

        _replace_docx_with_lock_hint(tmp_path, docx_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)

    return {
        "applied": updated > 0,
        "updated": updated,
        "caption_style_id": caption_style_id,
        "samples": samples,
    }


def _paragraph_contains_embedded_image(para: ET.Element) -> bool:
    return para.find(".//w:drawing", W) is not None or para.find(".//w:pict", W) is not None


def _fit_oversized_images_to_page(docx_path: Path) -> Dict[str, object]:
    updated = 0
    samples: List[Dict[str, int]] = []

    with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
        tmp_path = Path(tmp.name)

    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
            tmp_path, "w", compression=zipfile.ZIP_DEFLATED
        ) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    root = ET.fromstring(data)
                    body = root.find("w:body", W)
                    if body is not None:
                        body_sect = body.find("w:sectPr", W)
                        pg_sz = body_sect.find("w:pgSz", W) if body_sect is not None else None
                        pg_mar = body_sect.find("w:pgMar", W) if body_sect is not None else None
                        if pg_sz is not None and pg_mar is not None:
                            try:
                                page_h = int(pg_sz.get(_w_attr("h"), "0"))
                                page_w = int(pg_sz.get(_w_attr("w"), "0"))
                                mar_top = int(pg_mar.get(_w_attr("top"), "0"))
                                mar_bottom = int(pg_mar.get(_w_attr("bottom"), "0"))
                                mar_left = int(pg_mar.get(_w_attr("left"), "0"))
                                mar_right = int(pg_mar.get(_w_attr("right"), "0"))
                            except ValueError:
                                page_h = page_w = mar_top = mar_bottom = mar_left = mar_right = 0

                            usable_h_twips = max(0, page_h - mar_top - mar_bottom)
                            usable_w_twips = max(0, page_w - mar_left - mar_right)

                            if usable_h_twips > 0 and usable_w_twips > 0:
                                body_children = list(body)
                                for idx, para in enumerate(body_children):
                                    if para.tag != _w_attr("p") or not _paragraph_contains_embedded_image(para):
                                        continue

                                    drawing = para.find(".//w:drawing", W)
                                    if drawing is None:
                                        continue
                                    inline = drawing.find(".//wp:inline", OOXML_NAMESPACE_PREFIXES)
                                    extent = inline.find("wp:extent", OOXML_NAMESPACE_PREFIXES) if inline is not None else None
                                    xfrm_ext = drawing.find(".//a:xfrm/a:ext", OOXML_NAMESPACE_PREFIXES)
                                    if extent is None or xfrm_ext is None:
                                        continue

                                    try:
                                        cx = int(extent.get("cx", "0"))
                                        cy = int(extent.get("cy", "0"))
                                    except ValueError:
                                        continue
                                    if cx <= 0 or cy <= 0:
                                        continue

                                    caption_chain: List[ET.Element] = []
                                    for next_child in body_children[idx + 1 :]:
                                        if next_child.tag != _w_attr("p"):
                                            break
                                        next_text = _extract_para_text(next_child).strip()
                                        if not next_text:
                                            continue
                                        if _is_figure_caption_paragraph_text(next_text):
                                            caption_chain.append(next_child)
                                            continue
                                        break

                                    # Word 在分页时对图题、段前后距和行高的估计偏保守。
                                    # 对接近页高上限的图片，需要额外留出图题与页底安全余量，
                                    # 否则图本体虽未超高，仍容易把中英图题挤到下一页。
                                    caption_reserve_twips = 1080 + 480 * max(1, len(caption_chain))
                                    pagination_safety_twips = 720
                                    max_h_twips = max(0, usable_h_twips - caption_reserve_twips - pagination_safety_twips)
                                    max_w_emu = usable_w_twips * EMU_PER_TWIP
                                    max_h_emu = max_h_twips * EMU_PER_TWIP
                                    if max_h_emu <= 0:
                                        continue

                                    scale = min(1.0, max_w_emu / cx, max_h_emu / cy)
                                    if scale >= 0.999:
                                        continue

                                    new_cx = max(1, int(round(cx * scale)))
                                    new_cy = max(1, int(round(cy * scale)))
                                    extent.set("cx", str(new_cx))
                                    extent.set("cy", str(new_cy))
                                    xfrm_ext.set("cx", str(new_cx))
                                    xfrm_ext.set("cy", str(new_cy))
                                    updated += 1
                                    if len(samples) < 6:
                                        samples.append(
                                            {
                                                "index": idx,
                                                "old_cx": cx,
                                                "old_cy": cy,
                                                "new_cx": new_cx,
                                                "new_cy": new_cy,
                                            }
                                        )

                    data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

        _replace_docx_with_lock_hint(tmp_path, docx_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)

    return {
        "applied": updated > 0,
        "updated": updated,
        "samples": samples,
    }


def _normalize_image_paragraph_indent(docx_path: Path) -> Dict[str, object]:
    updated = 0
    indent_updated = 0
    spacing_updated = 0
    center_updated = 0

    with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
        tmp_path = Path(tmp.name)

    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
            tmp_path, "w", compression=zipfile.ZIP_DEFLATED
        ) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    root = ET.fromstring(data)
                    body = root.find("w:body", W)
                    if body is not None:
                        body_children = list(body)
                        for idx, para in enumerate(body_children):
                            if para.tag != _w_attr("p"):
                                continue
                            if not _paragraph_contains_embedded_image(para):
                                continue
                            changed = False
                            if _ensure_paragraph_indent_zero(para):
                                indent_updated += 1
                                changed = True
                            if _ensure_paragraph_spacing_zero(para):
                                spacing_updated += 1
                                changed = True
                            if _ensure_paragraph_center(para):
                                center_updated += 1
                                changed = True
                            caption_chain: List[ET.Element] = []
                            for next_child in body_children[idx + 1 :]:
                                if next_child.tag != _w_attr("p"):
                                    break
                                next_text = _extract_para_text(next_child).strip()
                                if not next_text:
                                    continue
                                if _is_figure_caption_paragraph_text(next_text):
                                    caption_chain.append(next_child)
                                    continue
                                break
                            if caption_chain:
                                if _ensure_paragraph_keep_next(para):
                                    changed = True
                                for chain_idx, caption_para in enumerate(caption_chain):
                                    if _ensure_paragraph_keep_lines(caption_para):
                                        changed = True
                                    if chain_idx < len(caption_chain) - 1 and _ensure_paragraph_keep_next(caption_para):
                                        changed = True
                            if changed:
                                updated += 1
                    data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

        _replace_docx_with_lock_hint(tmp_path, docx_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)

    return {
        "applied": updated > 0,
        "updated": updated,
        "indent_updated": indent_updated,
        "spacing_updated": spacing_updated,
        "center_updated": center_updated,
    }


def _apply_three_line_table_format(docx_path: Path) -> Dict[str, object]:
    updated = 0
    layout_updated = 0
    wide_layout_updated = 0
    landscape_section_updated = 0
    pagination_updated = 0

    with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
        tmp_path = Path(tmp.name)

    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
            tmp_path, "w", compression=zipfile.ZIP_DEFLATED
        ) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    root = ET.fromstring(data)
                    body = root.find("w:body", W)
                    if body is not None:
                        pending_table_caption_paras: List[ET.Element] = []
                        for child in list(body):
                            if child.tag == _w_attr("p"):
                                text = _extract_para_text(child).strip()
                                if not text:
                                    continue
                                if _is_table_caption_paragraph_text(text):
                                    pending_table_caption_paras.append(child)
                                else:
                                    pending_table_caption_paras = []
                            elif child.tag == _w_attr("tbl"):
                                if pending_table_caption_paras:
                                    for caption_para in pending_table_caption_paras:
                                        if _ensure_paragraph_keep_next(caption_para):
                                            pagination_updated += 1
                                        if _ensure_paragraph_keep_lines(caption_para):
                                            pagination_updated += 1
                                    if _apply_three_line_table_borders(child):
                                        updated += 1
                                    if _normalize_table_paragraph_layout(child):
                                        layout_updated += 1
                                    wide_width_override = _apply_landscape_section_around_table(
                                        body,
                                        child,
                                        pending_table_caption_paras[0],
                                    )
                                    if wide_width_override is not None:
                                        landscape_section_updated += 1
                                    if _rebalance_wide_table_columns(child, preferred_total_width=wide_width_override):
                                        wide_layout_updated += 1
                                    if _set_table_row_pagination(child):
                                        pagination_updated += 1
                                pending_table_caption_paras = []
                            else:
                                pending_table_caption_paras = []
                    data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

        _replace_docx_with_lock_hint(tmp_path, docx_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)

    return {
        "applied": updated > 0 or layout_updated > 0 or wide_layout_updated > 0 or pagination_updated > 0,
        "updated": updated,
        "layout_updated": layout_updated,
        "wide_layout_updated": wide_layout_updated,
        "landscape_section_updated": landscape_section_updated,
        "pagination_updated": pagination_updated,
    }


def _normalize_backmatter_heading_text(text: str) -> str:
    normalized = (text or "").strip()
    if re.fullmatch(r"致[\s\u3000]*谢", normalized):
        return "致  谢"
    if normalized.startswith("附"):
        normalized = re.sub(r"^附[\s\u3000]*录", "附录", normalized)
        normalized = re.sub(r"[\s\u3000]{2,}", " ", normalized)
    return normalized


def _is_keyword_paragraph_text(text: str) -> bool:
    return bool(re.match(r"(?i)^(关键词\s*[:：]|key\s*words?\s*[:：])", (text or "").strip()))


def _normalize_abstract_keyword_paragraphs(docx_path: Path) -> Dict[str, object]:
    updated = 0
    with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
        tmp_path = Path(tmp.name)

    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
            tmp_path, "w", compression=zipfile.ZIP_DEFLATED
        ) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    root = ET.fromstring(data)
                    body = root.find("w:body", W)
                    if body is not None:
                        idx = 0
                        while idx < len(body):
                            child = body[idx]
                            if child.tag != _w_attr("p"):
                                idx += 1
                                continue

                            text = _extract_para_text(child).strip()
                            if not text:
                                idx += 1
                                continue

                            compact = re.sub(r"[\s\u3000]+", "", text).casefold()
                            if compact == _norm_plain_text("摘要"):
                                if text != "摘  要":
                                    _replace_paragraph_text_keep_format(child, "摘  要")
                                    updated += 1
                            elif _is_keyword_paragraph_text(text):
                                prev = body[idx - 1] if idx > 0 else None
                                prev_text = (
                                    _extract_para_text(prev).strip()
                                    if prev is not None and prev.tag == _w_attr("p")
                                    else ""
                                )
                                if prev is None or prev.tag != _w_attr("p") or prev_text:
                                    body.insert(idx, ET.Element(_w_attr("p")))
                                    updated += 1
                                    idx += 1
                                if _ensure_paragraph_indent_zero(child):
                                    updated += 1
                            idx += 1
                    data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

        _replace_docx_with_lock_hint(tmp_path, docx_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)

    return {"applied": updated > 0, "updated": updated}


def _force_cv_subheading_bold(docx_path: Path) -> Dict[str, object]:
    updated = 0
    target_norms = {
        _norm_plain_text("作者简历："),
        _norm_plain_text("已发表（或正式接受）的学术论文："),
        _norm_plain_text("申请或已获得的专利："),
        _norm_plain_text("参加的研究项目及获奖情况："),
    }

    with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
        tmp_path = Path(tmp.name)

    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
            tmp_path, "w", compression=zipfile.ZIP_DEFLATED
        ) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    root = ET.fromstring(data)
                    body = root.find("w:body", W)
                    if body is not None:
                        for para in body.findall("w:p", W):
                            text_norm = _norm_plain_text(_extract_para_text(para))
                            if text_norm not in target_norms:
                                continue
                            para_changed = False
                            for run in para.findall("w:r", W):
                                if _ensure_run_bold(run):
                                    para_changed = True
                            if para_changed:
                                updated += 1
                    data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

        _replace_docx_with_lock_hint(tmp_path, docx_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)

    return {"applied": updated > 0, "updated": updated}


def _is_structural_backmatter_heading(text: str, text_norm: str) -> bool:
    if text_norm in {
        _norm_plain_text("参考文献"),
        _norm_plain_text("致谢"),
        _norm_plain_text("作者简历及攻读学位期间发表的学术论文与其他相关学术成果"),
        _norm_plain_text("附录"),
        _norm_plain_text("附录示例"),
        _norm_plain_text("附录一"),
        _norm_plain_text("附录二"),
    }:
        return True

    compact = re.sub(r"[\s\u3000]+", "", (text or ""))
    # 仅将短标题形式的“附录X”识别为结构标题，避免把“附录二测试内容...”正文误纳入目录。
    return bool(re.fullmatch(r"附录(?:[一二三四五六七八九十百零\dA-Za-z]{1,6})?", compact))


def _is_appendix_heading_text(text: str) -> bool:
    compact = re.sub(r"[\s\u3000]+", "", (text or ""))
    if not compact:
        return False
    if compact.startswith("附录"):
        return True
    if re.fullmatch(r"第[一二三四五六七八九十百零\d]+章补充[表图](?:（[^）]+）)?", compact):
        return True
    if "调查表" in compact and len(compact) <= 40:
        return True
    return False


def _set_paragraph_numbering_suppressed(ppr: ET.Element, level: str = "0") -> None:
    numpr = ppr.find("w:numPr", W)
    if numpr is None:
        numpr = ET.Element(_w_attr("numPr"))
        ppr.append(numpr)

    for node in list(numpr):
        numpr.remove(node)

    ilvl = ET.Element(_w_attr("ilvl"))
    ilvl.set(_w_attr("val"), level)
    numpr.append(ilvl)

    num_id = ET.Element(_w_attr("numId"))
    num_id.set(_w_attr("val"), "0")
    numpr.append(num_id)


def _strip_appendix_heading_numbering(docx_path: Path) -> Dict[str, object]:
    """Ensure appendix headings and subheadings stay unnumbered."""
    tmp_path: Optional[Path] = None
    with zipfile.ZipFile(docx_path, "r") as zin:
        out_doc_root = ET.fromstring(zin.read("word/document.xml"))
        out_body = out_doc_root.find("w:body", W)
        if out_body is None:
            raise RuntimeError("导出 docx 缺少 w:body，无法修正附录标题编号")

        fixed: List[str] = []
        appendix_scope = False
        appendix_heading_styles = {
            "1",
            "2",
            "3",
            "4",
            "5",
            "6",
            "heading1",
            "heading2",
            "heading3",
            "heading4",
            "heading5",
            "heading6",
        }
        appendix_heading_level_map = {
            "1": "0",
            "heading1": "0",
            "2": "1",
            "heading2": "1",
            "3": "2",
            "heading3": "2",
            "4": "3",
            "heading4": "3",
            "5": "4",
            "heading5": "4",
            "6": "5",
            "heading6": "5",
        }
        stop_norms = {
            _norm_plain_text("参考文献"),
            _norm_plain_text("致谢"),
            _norm_plain_text("作者简历及攻读学位期间发表的学术论文与其他相关学术成果"),
        }

        for para in out_body.findall("w:p", W):
            text = _extract_para_text(para).strip()
            if not text:
                continue

            text_norm = _norm_plain_text(text)
            if text_norm in stop_norms:
                appendix_scope = False
                continue

            ppr = para.find("w:pPr", W)
            if ppr is None:
                continue
            pstyle = ppr.find("w:pStyle", W)
            sid = pstyle.get(_w_attr("val"), "") if pstyle is not None else ""
            sid_norm = re.sub(r"[\s_-]+", "", sid.casefold())
            is_h1 = sid_norm in {"1", "heading1"}

            # Start appendix scope at appendix chapter heading.
            if is_h1 and _is_appendix_heading_text(text):
                appendix_scope = True

            if not appendix_scope:
                continue
            if sid_norm not in appendix_heading_styles:
                continue

            level = appendix_heading_level_map.get(sid_norm, "0")
            _set_paragraph_numbering_suppressed(ppr, level)
            changed = True

            if changed:
                fixed.append(text)

        _restore_markup_compatibility(out_doc_root, None)
        if not fixed:
            return {"applied": False, "fixed": []}

        with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
            tmp_path = Path(tmp.name)

        with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    data = ET.tostring(out_doc_root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

    if tmp_path is not None:
        try:
            _replace_docx_with_lock_hint(tmp_path, docx_path)
        finally:
            if tmp_path.exists():
                tmp_path.unlink(missing_ok=True)

    return {"applied": True, "fixed": fixed}


def _strip_backmatter_heading_numbering(docx_path: Path, reference_doc: Path) -> Dict[str, object]:
    """From '参考文献' to the end, keep TOC headings but suppress chapter numbering."""
    tmp_path: Optional[Path] = None
    with zipfile.ZipFile(docx_path, "r") as zin:
        out_doc_root = ET.fromstring(zin.read("word/document.xml"))
        out_body = out_doc_root.find("w:body", W)
        if out_body is None:
            raise RuntimeError("导出 docx 缺少 w:body，无法修正后置标题样式")

        style_root = ET.fromstring(zin.read("word/styles.xml"))
        heading1_id = "1"
        normal_id = "a"
        normal_candidate_ids: List[str] = []
        for style in style_root.findall("w:style", W):
            if style.get(_w_attr("type"), "") != "paragraph":
                continue
            sid = style.get(_w_attr("styleId"), "")
            if not sid:
                continue
            sid_norm = re.sub(r"[\s_-]+", "", sid.casefold())
            is_default = style.get(_w_attr("default"), "") == "1"
            name = style.find("w:name", W)
            nval = ""
            nval_norm = ""
            if name is not None:
                nval = name.get(_w_attr("val"), "")
                nval_norm = re.sub(r"[\s_-]+", "", nval.casefold())
            if is_default:
                normal_candidate_ids.insert(0, sid)
            elif sid_norm in {"a", "normal", "正文"} or nval_norm in {"normal", "正文"}:
                normal_candidate_ids.append(sid)

            if sid_norm in {"1", "heading1"}:
                heading1_id = sid
            elif nval_norm == "heading1":
                heading1_id = sid

        if normal_candidate_ids:
            normal_id = normal_candidate_ids[0]

        split_idx = _find_body_child_split_by_heading(out_doc_root, ["参考文献"])
        if split_idx is None:
            return {"applied": False, "fixed": [], "demoted": [], "normalized": []}

        fixed: List[str] = []
        demoted: List[str] = []
        normalized: List[str] = []
        body_children = list(out_body)
        heading1_style_keys = {"1", "heading1"}
        subheading_style_keys = {"2", "3", "4", "heading2", "heading3", "heading4"}
        for child in body_children[split_idx:]:
            if child.tag != _w_attr("p"):
                continue
            text = _extract_para_text(child).strip()
            if not text:
                continue

            ppr = child.find("w:pPr", W)
            if ppr is None:
                ppr = ET.Element(_w_attr("pPr"))
                child.insert(0, ppr)

            pstyle = ppr.find("w:pStyle", W)
            sid = ""
            if pstyle is not None:
                sid = pstyle.get(_w_attr("val"), "")
            sid_norm = re.sub(r"[\s_-]+", "", sid.casefold())

            text_norm = _norm_plain_text(text)
            is_heading1_style = sid_norm in heading1_style_keys
            is_subheading_style = sid_norm in subheading_style_keys
            is_structural_heading = _is_structural_backmatter_heading(text, text_norm) or is_heading1_style
            if not is_structural_heading and not is_subheading_style:
                continue

            normalized_text = _normalize_backmatter_heading_text(text)
            if normalized_text != text:
                _replace_paragraph_text_keep_format(child, normalized_text)
                normalized.append(normalized_text)

            if pstyle is None:
                pstyle = ET.Element(_w_attr("pStyle"))
                ppr.insert(0, pstyle)

            if is_structural_heading:
                pstyle.set(_w_attr("val"), heading1_id)

                _set_paragraph_numbering_suppressed(ppr, "0")
                fixed.append(normalized_text)
            else:
                # 作者简历下的条目保持为正文，不参与章级目录。
                pstyle.set(_w_attr("val"), normal_id)
                numpr = ppr.find("w:numPr", W)
                if numpr is not None:
                    ppr.remove(numpr)
                demoted.append(normalized_text)

        _restore_markup_compatibility(out_doc_root, None)
        if not (fixed or demoted or normalized):
            return {"applied": False, "fixed": [], "demoted": [], "normalized": []}

        with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
            tmp_path = Path(tmp.name)

        with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    data = ET.tostring(out_doc_root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

    if tmp_path is not None:
        try:
            _replace_docx_with_lock_hint(tmp_path, docx_path)
        finally:
            if tmp_path.exists():
                tmp_path.unlink(missing_ok=True)

    return {"applied": True, "fixed": fixed, "demoted": demoted, "normalized": normalized}


def _build_toc_field_paragraph(style_id: str, instr_code: str, hint_text: str) -> ET.Element:
    para = ET.Element(_w_attr("p"))

    ppr = ET.Element(_w_attr("pPr"))
    if style_id:
        pstyle = ET.Element(_w_attr("pStyle"))
        pstyle.set(_w_attr("val"), style_id)
        ppr.append(pstyle)
    para.append(ppr)

    run_begin = ET.Element(_w_attr("r"))
    fld_begin = ET.Element(_w_attr("fldChar"))
    fld_begin.set(_w_attr("fldCharType"), "begin")
    fld_begin.set(_w_attr("dirty"), "true")
    run_begin.append(fld_begin)
    para.append(run_begin)

    run_instr = ET.Element(_w_attr("r"))
    instr = ET.Element(_w_attr("instrText"))
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = instr_code
    run_instr.append(instr)
    para.append(run_instr)

    run_sep = ET.Element(_w_attr("r"))
    fld_sep = ET.Element(_w_attr("fldChar"))
    fld_sep.set(_w_attr("fldCharType"), "separate")
    run_sep.append(fld_sep)
    para.append(run_sep)

    run_hint = ET.Element(_w_attr("r"))
    hint = ET.Element(_w_attr("t"))
    hint.text = hint_text
    run_hint.append(hint)
    para.append(run_hint)

    run_end = ET.Element(_w_attr("r"))
    fld_end = ET.Element(_w_attr("fldChar"))
    fld_end.set(_w_attr("fldCharType"), "end")
    run_end.append(fld_end)
    para.append(run_end)
    return para


def _build_hidden_tc_field_runs(entry_text: str, list_flag: str) -> List[ET.Element]:
    text = re.sub(r"\s+", " ", (entry_text or "")).strip().replace('"', "'")
    if not text or not list_flag:
        return []
    flag = list_flag[:1].upper()
    instr_code = f' TC "{text}" \\f {flag} \\l 1 '

    def _new_hidden_run() -> ET.Element:
        run = ET.Element(_w_attr("r"))
        rpr = ET.Element(_w_attr("rPr"))
        vanish = ET.Element(_w_attr("vanish"))
        vanish.set(_w_attr("val"), "1")
        rpr.append(vanish)
        run.append(rpr)
        return run

    run_begin = _new_hidden_run()
    fld_begin = ET.Element(_w_attr("fldChar"))
    fld_begin.set(_w_attr("fldCharType"), "begin")
    run_begin.append(fld_begin)

    run_instr = _new_hidden_run()
    instr = ET.Element(_w_attr("instrText"))
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = instr_code
    run_instr.append(instr)

    run_sep = _new_hidden_run()
    fld_sep = ET.Element(_w_attr("fldChar"))
    fld_sep.set(_w_attr("fldCharType"), "separate")
    run_sep.append(fld_sep)

    run_end = _new_hidden_run()
    fld_end = ET.Element(_w_attr("fldChar"))
    fld_end.set(_w_attr("fldCharType"), "end")
    run_end.append(fld_end)
    return [run_begin, run_instr, run_sep, run_end]


def _field_instr_text_from_runs(runs: List[ET.Element]) -> str:
    parts: List[str] = []
    for run in runs:
        if run.tag != _w_attr("r"):
            continue
        parts.extend((node.text or "") for node in run.findall("w:instrText", W))
    return re.sub(r"\s+", " ", "".join(parts)).strip()


def _remove_tc_field_runs_from_paragraph(para: ET.Element) -> int:
    removed = 0

    while True:
        children = list(para)
        block_removed = False
        idx = 0
        while idx < len(children):
            child = children[idx]
            if child.tag == _w_attr("fldSimple"):
                instr_raw = child.get(_w_attr("instr"), "") or ""
                instr_norm = re.sub(r"\s+", " ", instr_raw).casefold()
                if " tc " in f" {instr_norm} " and ("\\f f" in instr_norm or "\\f t" in instr_norm):
                    para.remove(child)
                    removed += 1
                    block_removed = True
                    break

            if child.tag != _w_attr("r"):
                idx += 1
                continue

            fld_char = child.find("w:fldChar", W)
            if fld_char is None or fld_char.get(_w_attr("fldCharType"), "") != "begin":
                idx += 1
                continue

            end_idx: Optional[int] = None
            for probe_idx in range(idx + 1, len(children)):
                probe = children[probe_idx]
                if probe.tag != _w_attr("r"):
                    continue
                probe_fld = probe.find("w:fldChar", W)
                if probe_fld is not None and probe_fld.get(_w_attr("fldCharType"), "") == "end":
                    end_idx = probe_idx
                    break

            if end_idx is None:
                idx += 1
                continue

            field_runs = children[idx : end_idx + 1]
            instr_norm = _field_instr_text_from_runs(field_runs).casefold()
            if " tc " not in f" {instr_norm} " or ("\\f f" not in instr_norm and "\\f t" not in instr_norm):
                idx = end_idx + 1
                continue

            for node in field_runs:
                if node in list(para):
                    para.remove(node)
            removed += 1
            block_removed = True
            break

        if not block_removed:
            break

    return removed


def _build_plain_text_paragraph(text: str, style_id: str = "") -> ET.Element:
    para = ET.Element(_w_attr("p"))
    if style_id:
        ppr = ET.Element(_w_attr("pPr"))
        pstyle = ET.Element(_w_attr("pStyle"))
        pstyle.set(_w_attr("val"), style_id)
        ppr.append(pstyle)
        para.append(ppr)
    run = ET.Element(_w_attr("r"))
    t = ET.Element(_w_attr("t"))
    if text[:1].isspace() or text[-1:].isspace():
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    run.append(t)
    para.append(run)
    return para


def _build_toc_like_paragraph(entry_text: str, page_text: str, style_id: str = "") -> ET.Element:
    para = ET.Element(_w_attr("p"))
    if style_id:
        ppr = ET.Element(_w_attr("pPr"))
        pstyle = ET.Element(_w_attr("pStyle"))
        pstyle.set(_w_attr("val"), style_id)
        ppr.append(pstyle)
        para.append(ppr)

    run_entry = ET.Element(_w_attr("r"))
    t_entry = ET.Element(_w_attr("t"))
    if entry_text[:1].isspace() or entry_text[-1:].isspace():
        t_entry.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t_entry.text = entry_text
    run_entry.append(t_entry)
    para.append(run_entry)

    page = (page_text or "").strip()
    if page:
        run_tab = ET.Element(_w_attr("r"))
        tab = ET.Element(_w_attr("tab"))
        run_tab.append(tab)
        para.append(run_tab)

        run_page = ET.Element(_w_attr("r"))
        t_page = ET.Element(_w_attr("t"))
        t_page.text = page
        run_page.append(t_page)
        para.append(run_page)

    return para


def _ensure_main_toc_field(docx_path: Path) -> Dict[str, object]:
    tmp_path: Optional[Path] = None
    with zipfile.ZipFile(docx_path, "r") as zin:
        out_doc_root = ET.fromstring(zin.read("word/document.xml"))
        out_body = out_doc_root.find("w:body", W)
        if out_body is None:
            raise RuntimeError("导出 docx 缺少 w:body，无法插入目录域")

        for instr in out_doc_root.findall(".//w:instrText", W):
            text = "".join(instr.itertext())
            if "TOC" in text and "\\o" in text:
                return {"applied": False, "reason": "already_exists"}

        toc_heading_idx = _find_body_child_split_by_heading(out_doc_root, ["目录"])
        if toc_heading_idx is None:
            return {"applied": False, "reason": "heading_not_found"}

        body_children = list(out_body)
        insert_idx = toc_heading_idx + 1
        removed_placeholder = 0
        style_id = "TOC1"
        if insert_idx < len(body_children):
                next_node = body_children[insert_idx]
                if next_node.tag == _w_attr("p"):
                    next_text_norm = _norm_plain_text(_extract_para_text(next_node))
                    if "目录域请在word中右键更新" in next_text_norm:
                        out_body.remove(next_node)
                        removed_placeholder = 1

        toc_para = _build_toc_field_paragraph(
            style_id=style_id,
            instr_code=' TOC \\o "1-3" \\h \\z \\u ',
            hint_text="目录（右键更新域）",
        )
        out_body.insert(insert_idx, toc_para)

        _restore_markup_compatibility(out_doc_root, None)
        with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
            tmp_path = Path(tmp.name)

        with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    data = ET.tostring(out_doc_root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

    if tmp_path is not None:
        try:
            _replace_docx_with_lock_hint(tmp_path, docx_path)
        finally:
            if tmp_path.exists():
                tmp_path.unlink(missing_ok=True)

    return {
        "applied": True,
        "insert_index": insert_idx,
        "removed_placeholder": removed_placeholder,
    }


def _ensure_figure_table_toc_fields(docx_path: Path) -> Dict[str, object]:
    config = [
        ("图目录", "图目录请在word中右键更新", ' TOC \\h \\z \\f F ', "图目录（右键更新域）", "F"),
        ("表目录", "表目录请在word中右键更新", ' TOC \\h \\z \\f T ', "表目录（右键更新域）", "T"),
    ]
    tmp_path: Optional[Path] = None
    with zipfile.ZipFile(docx_path, "r") as zin:
        out_doc_root = ET.fromstring(zin.read("word/document.xml"))
        out_body = out_doc_root.find("w:body", W)
        if out_body is None:
            raise RuntimeError("导出 docx 缺少 w:body，无法插入图表目录域")

        body_children = list(out_body)
        inserted = 0
        removed_placeholder = 0
        touched_kinds: List[str] = []

        for heading, placeholder_norm, instr_code, hint_text, flag in config:
            has_field = False
            for instr in out_doc_root.findall(".//w:instrText", W):
                text = "".join(instr.itertext())
                if "TOC" in text and f"\\f {flag}" in text:
                    has_field = True
                    break
            if has_field:
                continue

            heading_idx = _find_body_child_split_by_heading(out_doc_root, [heading])
            if heading_idx is None:
                continue

            body_children = list(out_body)
            insert_idx = heading_idx + 1
            style_id = "TOC1"
            if insert_idx < len(body_children):
                next_node = body_children[insert_idx]
                if next_node.tag == _w_attr("p"):
                    next_text_norm = _norm_plain_text(_extract_para_text(next_node))
                    if placeholder_norm in next_text_norm:
                        out_body.remove(next_node)
                        removed_placeholder += 1

            toc_para = _build_toc_field_paragraph(style_id=style_id, instr_code=instr_code, hint_text=hint_text)
            out_body.insert(insert_idx, toc_para)
            inserted += 1
            touched_kinds.append(flag)

        _restore_markup_compatibility(out_doc_root, None)
        if inserted == 0 and removed_placeholder == 0:
            return {"applied": False, "inserted": 0, "removed_placeholder": 0, "kinds": []}

        with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
            tmp_path = Path(tmp.name)
        with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    data = ET.tostring(out_doc_root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

    if tmp_path is not None:
        try:
            _replace_docx_with_lock_hint(tmp_path, docx_path)
        finally:
            if tmp_path.exists():
                tmp_path.unlink(missing_ok=True)

    return {
        "applied": True,
        "inserted": inserted,
        "removed_placeholder": removed_placeholder,
        "kinds": touched_kinds,
    }


def _refresh_figure_table_catalog_fallback(
    docx_path: Path,
    *,
    force: bool = False,
    project_dir: Optional[Path] = None,
) -> Dict[str, object]:
    tmp_path: Optional[Path] = None
    with zipfile.ZipFile(docx_path, "r") as zin:
        out_doc_root = ET.fromstring(zin.read("word/document.xml"))
        out_body = out_doc_root.find("w:body", W)
        if out_body is None:
            return {"applied": False, "reason": "missing_body"}
        style_doc_root: Optional[ET.Element] = None
        try:
            style_doc_root = ET.fromstring(zin.read("word/styles.xml"))
        except Exception:
            style_doc_root = None

        figure_entries: List[Tuple[str, str]] = []
        table_entries: List[Tuple[str, str]] = []
        seen_fig: Set[str] = set()
        seen_tbl: Set[str] = set()

        def _append_unique(entries: List[Tuple[str, str]], seen: Set[str], text: str, page: str = "") -> None:
            normalized = re.sub(r"\s+", " ", (text or "")).strip()
            if not normalized or normalized in seen:
                return
            seen.add(normalized)
            entries.append((normalized, re.sub(r"\s+", " ", (page or "")).strip()))

        def _read_braced_group(source: str, start_idx: int) -> Tuple[str, int]:
            if start_idx < 0 or start_idx >= len(source) or source[start_idx] != "{":
                return "", -1
            depth = 0
            chars: List[str] = []
            for idx in range(start_idx, len(source)):
                ch = source[idx]
                if ch == "{":
                    depth += 1
                    if depth > 1:
                        chars.append(ch)
                    continue
                if ch == "}":
                    depth -= 1
                    if depth == 0:
                        return ("".join(chars), idx + 1)
                    if depth > 0:
                        chars.append(ch)
                    continue
                if depth > 0:
                    chars.append(ch)
            return "", -1

        def _latex_to_plain(text: str) -> str:
            plain = text or ""
            plain = plain.replace(r"\ignorespaces", " ")
            plain = plain.replace(r"\protect", " ")
            plain = plain.replace(r"\p@", "")
            plain = plain.replace(r"\%", "%")
            plain = plain.replace(r"\_", "_")
            plain = plain.replace(r"\&", "&")
            plain = plain.replace(r"\$", "")
            plain = re.sub(r"\\mitDelta", "Δ", plain)
            plain = re.sub(r"\\log", "log", plain)
            plain = re.sub(r"\\[a-zA-Z@]+", "", plain)
            plain = plain.replace("{", "").replace("}", "")
            plain = re.sub(r"\s+", " ", plain).strip()
            return plain

        def _alpha_ordinal(value: str) -> int:
            total = 0
            for ch in value.upper():
                if "A" <= ch <= "Z":
                    total = total * 26 + (ord(ch) - ord("A") + 1)
            return total or 1

        def _collect_from_material_file(path: Path, kind: str, prefix: str) -> List[Tuple[str, str]]:
            if not path.exists():
                return []
            results: List[Tuple[str, str]] = []
            appendix_table_ordinals: Dict[str, int] = {}
            for line in path.read_text(encoding="utf-8", errors="ignore").splitlines():
                if f"\\contentsline {{{kind}}}" not in line:
                    continue
                marker_idx = line.find("\\numberline")
                if marker_idx < 0:
                    continue
                number_start = line.find("{", marker_idx)
                number, next_idx = _read_braced_group(line, number_start)
                if not number or next_idx < 0:
                    continue
                title_start = line.find("{", next_idx)
                title, title_end_idx = _read_braced_group(line, title_start)
                if not title:
                    continue
                title_plain = _strip_caption_terminal_punctuation(_latex_to_plain(title))
                page_plain = ""
                if title_end_idx > 0:
                    page_start = line.find("{", title_end_idx)
                    page_raw, _ = _read_braced_group(line, page_start)
                    page_plain = _latex_to_plain(page_raw)
                display_prefix = prefix
                display_number = number
                if re.fullmatch(r"-\d+", number):
                    display_prefix = "附图" if kind == "figure" else "附表"
                    display_number = f"1-{number[1:]}"
                elif re.fullmatch(r"-\d+-\d+", number):
                    display_prefix = "附图" if kind == "figure" else "附表"
                    display_number = number[1:]
                else:
                    letter_match = re.fullmatch(r"([A-Z]+)-(\d+)", number)
                    if letter_match:
                        appendix_key, item_no = letter_match.groups()
                        if kind == "figure":
                            appendix_no = _alpha_ordinal(appendix_key)
                        else:
                            appendix_no = appendix_table_ordinals.setdefault(
                                appendix_key,
                                len(appendix_table_ordinals) + 1,
                            )
                        display_prefix = "附图" if kind == "figure" else "附表"
                        display_number = f"{appendix_no}-{item_no}"
                entry = f"{display_prefix}{display_number} {title_plain}".strip()
                if entry:
                    results.append((entry, page_plain))
            return results

        def _collect_from_tc_fields() -> Tuple[List[Tuple[str, str]], List[Tuple[str, str]]]:
            figure_results: List[Tuple[str, str]] = []
            table_results: List[Tuple[str, str]] = []
            tc_pattern = re.compile(r'\bTC\s+"([^"]+)"\s+\\f\s+([FT])\s+\\l\s+\d+\b', re.IGNORECASE)

            def _append_from_instr(instr_raw: str) -> None:
                instr = re.sub(r"\s+", " ", (instr_raw or "")).strip()
                match = tc_pattern.search(instr)
                if match is None:
                    return
                entry_text = re.sub(r"\s+", " ", match.group(1)).strip()
                list_flag = match.group(2).upper()
                if list_flag == "F":
                    _append_unique(figure_entries, seen_fig, entry_text, "")
                    figure_results.append((entry_text, ""))
                elif list_flag == "T":
                    _append_unique(table_entries, seen_tbl, entry_text, "")
                    table_results.append((entry_text, ""))

            for para in out_body.findall("w:p", W):
                for simple in para.findall("w:fldSimple", W):
                    _append_from_instr(simple.get(_w_attr("instr"), "") or "")

                children = list(para)
                for idx, child in enumerate(children):
                    if child.tag != _w_attr("r"):
                        continue
                    fld_char = child.find("w:fldChar", W)
                    if fld_char is None or fld_char.get(_w_attr("fldCharType"), "") != "begin":
                        continue
                    if idx + 1 >= len(children):
                        continue
                    instr_run = children[idx + 1]
                    if instr_run.tag != _w_attr("r"):
                        continue
                    instr_text = "".join((t.text or "") for t in instr_run.findall("w:instrText", W))
                    _append_from_instr(instr_text)

            return figure_results, table_results

        cache_candidates: List[Path] = []
        if project_dir is not None:
            cache_candidates.append(project_dir / ".latex-cache")
        cache_candidates.append(docx_path.parent / ".latex-cache")

        cache_dir: Optional[Path] = None
        for candidate in cache_candidates:
            if (candidate / "main.lof").exists() or (candidate / "main.lot").exists():
                cache_dir = candidate
                break
        entry_source = ""
        if cache_dir is not None:
            for entry_text, page_text in _collect_from_material_file(cache_dir / "main.lof", "figure", "图"):
                _append_unique(figure_entries, seen_fig, entry_text, page_text)
            for entry_text, page_text in _collect_from_material_file(cache_dir / "main.lot", "table", "表"):
                _append_unique(table_entries, seen_tbl, entry_text, page_text)
            if figure_entries or table_entries:
                entry_source = "material_cache"

        if not figure_entries and not table_entries:
            tc_figures, tc_tables = _collect_from_tc_fields()
            if tc_figures or tc_tables:
                entry_source = "tc_fields"

        if not figure_entries and not table_entries:
            if cache_dir is None:
                return {"applied": False, "reason": "missing_material_cache"}
            return {"applied": False, "reason": "empty_material_cache"}

        def _find_range(heading: str, end_markers: List[str]) -> Optional[Tuple[int, int]]:
            start_idx = _find_body_child_split_by_heading(out_doc_root, [heading])
            if start_idx is None:
                return None
            marker_set = {_norm_plain_text(x) for x in end_markers}
            end_idx = len(list(out_body))
            for idx, child in enumerate(list(out_body)[start_idx + 1 :], start=start_idx + 1):
                if child.tag != _w_attr("p"):
                    continue
                text_norm = _norm_plain_text(_extract_para_text(child))
                if text_norm in marker_set:
                    end_idx = idx
                    break
            return (start_idx, end_idx)

        def _has_unresolved_item(start_idx: int, end_idx: int) -> bool:
            body_children = list(out_body)
            for child in body_children[start_idx + 1 : end_idx]:
                if child.tag != _w_attr("p"):
                    continue
                if "未找到目录项" in _norm_plain_text(_extract_para_text(child)):
                    return True
            return False

        figure_range = _find_range("图目录", ["表目录", "绪论"])
        table_range = _find_range("表目录", ["绪论"])
        if figure_range is None or table_range is None:
            return {"applied": False, "reason": "heading_not_found"}

        entry_style_id = ""
        if style_doc_root is not None:
            for style in style_doc_root.findall("w:style", W):
                sid = style.get(_w_attr("styleId"), "") or ""
                if not sid:
                    continue
                name = style.find("w:name", W)
                name_val = (name.get(_w_attr("val"), "") if name is not None else "").strip().casefold()
                if sid == "TOC1" or name_val == "toc 1":
                    entry_style_id = sid
                    break

        needs_fix = force or _has_unresolved_item(*figure_range) or _has_unresolved_item(*table_range)
        if not needs_fix:
            return {"applied": False, "reason": "not_needed"}

        replacement_plan: List[Tuple[int, int, List[Tuple[str, str]], str]] = [
            (figure_range[0], figure_range[1], figure_entries, "figure"),
            (table_range[0], table_range[1], table_entries, "table"),
        ]
        replacement_plan.sort(key=lambda item: item[0], reverse=True)

        replaced_sections = 0
        replaced_paragraphs = 0
        for start_idx, end_idx, entries, _ in replacement_plan:
            body_children = list(out_body)
            section_children = body_children[start_idx + 1 : end_idx]
            style_id = ""
            for child in section_children:
                if child.tag != _w_attr("p"):
                    continue
                style_id = _extract_para_style_id(child)
                if style_id:
                    break

            for child in section_children:
                out_body.remove(child)
            replaced_paragraphs += len(section_children)

            insert_idx = start_idx + 1
            effective_style_id = entry_style_id or style_id
            if entries:
                for entry_text, page_text in entries:
                    out_body.insert(insert_idx, _build_toc_like_paragraph(entry_text, page_text, effective_style_id))
                    insert_idx += 1
            else:
                out_body.insert(insert_idx, _build_plain_text_paragraph("（未检索到目录条目）", effective_style_id))
            replaced_sections += 1

        _restore_markup_compatibility(out_doc_root, None)
        with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
            tmp_path = Path(tmp.name)
        with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    data = ET.tostring(out_doc_root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

    if tmp_path is not None:
        try:
            _replace_docx_with_lock_hint(tmp_path, docx_path)
        finally:
            if tmp_path.exists():
                tmp_path.unlink(missing_ok=True)

    return {
        "applied": True,
        "sections": replaced_sections,
        "replaced_paragraphs": replaced_paragraphs,
        "figure_entries": len(figure_entries),
        "table_entries": len(table_entries),
        "entry_source": entry_source or "unknown",
    }


def _materialize_figure_table_catalogs_before_word_update(
    docx_path: Path,
    *,
    project_dir: Optional[Path] = None,
) -> Dict[str, object]:
    return _refresh_figure_table_catalog_fallback(docx_path, force=True, project_dir=project_dir)


def _rebuild_figure_table_tc_fields(docx_path: Path) -> Dict[str, object]:
    tmp_path: Optional[Path] = None
    caption_style_id = _resolve_docx_caption_style_id(docx_path)
    caption_style_norms = {
        re.sub(r"[\s_-]+", "", item.casefold())
        for item in {caption_style_id, "ae", "caption", "ImageCaption", "Image Caption"}
        if item
    }

    def _is_caption_style(style_id: str) -> bool:
        return re.sub(r"[\s_-]+", "", (style_id or "").casefold()) in caption_style_norms

    def _matches_catalog_caption_text(text: str) -> bool:
        return bool(
            re.match(r"^图\s*\d", text)
            or re.match(r"^附图\s*\d+(?:-\d+)?", text)
            or re.match(r"^表\s*\d", text)
            or re.match(r"^附表\s*\d+(?:-\d+)?", text)
            or text.startswith("表：")
        )

    with zipfile.ZipFile(docx_path, "r") as zin:
        out_doc_root = ET.fromstring(zin.read("word/document.xml"))
        out_body = out_doc_root.find("w:body", W)
        if out_body is None:
            return {"applied": False, "reason": "missing_body", "added": 0, "removed": 0}

        body_children = list(out_body)
        skip_para_ids: Set[int] = set()

        def _mark_skip_range(start_markers: List[str], end_markers: List[str]) -> None:
            start_idx = _find_body_child_split_by_heading(out_doc_root, start_markers)
            if start_idx is None:
                return
            marker_set = {_norm_plain_text(x) for x in end_markers}
            end_idx = len(body_children)
            for idx, child in enumerate(body_children[start_idx + 1 :], start=start_idx + 1):
                if child.tag != _w_attr("p"):
                    continue
                if _norm_plain_text(_extract_para_text(child)) in marker_set:
                    end_idx = idx
                    break
            for child in body_children[start_idx + 1 : end_idx]:
                if child.tag == _w_attr("p"):
                    skip_para_ids.add(id(child))

        _mark_skip_range(["图目录"], ["表目录", "绪论"])
        _mark_skip_range(["表目录"], ["绪论"])

        removed = 0
        added = 0
        fig_count = 0
        tbl_count = 0
        has_styled_caption_candidates = any(
            child.tag == _w_attr("p")
            and _is_caption_style(_extract_para_style_id(child))
            and _matches_catalog_caption_text(re.sub(r"\s+", " ", _extract_para_text(child)).strip())
            for child in body_children
            if id(child) not in skip_para_ids
        )

        for para in out_body.findall("w:p", W):
            if id(para) in skip_para_ids:
                continue
            removed += _remove_tc_field_runs_from_paragraph(para)

            text = re.sub(r"\s+", " ", _extract_para_text(para)).strip()
            if not text:
                continue
            style_id = _extract_para_style_id(para)
            if has_styled_caption_candidates and not _is_caption_style(style_id):
                continue

            tc_runs: List[ET.Element] = []
            if re.match(r"^图\s*\d", text) or re.match(r"^附图\s*\d+(?:-\d+)?", text):
                tc_runs = _build_hidden_tc_field_runs(text, "F")
                fig_count += 1
            elif re.match(r"^表\s*\d", text) or re.match(r"^附表\s*\d+(?:-\d+)?", text):
                tc_runs = _build_hidden_tc_field_runs(text, "T")
                tbl_count += 1
            elif text.startswith("表："):
                cleaned = text
                cleaned = re.sub(r"（需在 Word 中人工整理）\s*$", "", cleaned)
                cleaned = re.sub(r"\s+/\s+.*$", "", cleaned)
                sequence = tbl_count + 1
                synthetic = f"表{sequence} {cleaned[2:].strip()}"
                tc_runs = _build_hidden_tc_field_runs(synthetic, "T")
                tbl_count += 1

            if tc_runs:
                for run in tc_runs:
                    para.append(run)
                added += 1

        _restore_markup_compatibility(out_doc_root, None)
        if added == 0 and removed == 0:
            return {"applied": False, "added": 0, "removed": 0, "figure_entries": 0, "table_entries": 0}

        with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
            tmp_path = Path(tmp.name)
        with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    data = ET.tostring(out_doc_root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

    if tmp_path is not None:
        try:
            _replace_docx_with_lock_hint(tmp_path, docx_path)
        finally:
            if tmp_path.exists():
                tmp_path.unlink(missing_ok=True)

    return {
        "applied": True,
        "added": added,
        "removed": removed,
        "figure_entries": fig_count,
        "table_entries": tbl_count,
    }


def _enforce_frontmatter_page_breaks(docx_path: Path) -> Dict[str, object]:
    """Ensure required front-matter headings start on a new page."""
    tmp_path: Optional[Path] = None
    with zipfile.ZipFile(docx_path, "r") as zin:
        out_doc_root = ET.fromstring(zin.read("word/document.xml"))
        out_body = out_doc_root.find("w:body", W)
        if out_body is None:
            raise RuntimeError("导出 docx 缺少 w:body，无法修正前置分页规则")

        targets = {_norm_plain_text("图表目录")}
        updated_headings: List[str] = []

        for para in out_body.findall("w:p", W):
            text = _extract_para_text(para).strip()
            if not text:
                continue
            if _norm_plain_text(text) not in targets:
                continue

            ppr = para.find("w:pPr", W)
            if ppr is None:
                ppr = ET.Element(_w_attr("pPr"))
                para.insert(0, ppr)

            page_break_before = ppr.find("w:pageBreakBefore", W)
            if page_break_before is None:
                page_break_before = ET.Element(_w_attr("pageBreakBefore"))
                ppr.append(page_break_before)

            current_val = (page_break_before.get(_w_attr("val"), "") or "").strip().casefold()
            if current_val not in {"1", "true", "on"}:
                page_break_before.set(_w_attr("val"), "1")
                updated_headings.append(text)

        _restore_markup_compatibility(out_doc_root, None)
        if not updated_headings:
            return {"applied": False, "updated": []}

        with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
            tmp_path = Path(tmp.name)

        with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    data = ET.tostring(out_doc_root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

    if tmp_path is not None:
        try:
            _replace_docx_with_lock_hint(tmp_path, docx_path)
        finally:
            if tmp_path.exists():
                tmp_path.unlink(missing_ok=True)

    return {"applied": True, "updated": updated_headings}


def _enforce_heading1_page_breaks(docx_path: Path) -> Dict[str, object]:
    """Ensure all chapter-level headings (Heading 1) start on a new page."""
    tmp_path: Optional[Path] = None
    with zipfile.ZipFile(docx_path, "r") as zin:
        out_doc_root = ET.fromstring(zin.read("word/document.xml"))
        out_body = out_doc_root.find("w:body", W)
        if out_body is None:
            raise RuntimeError("导出 docx 缺少 w:body，无法修正章节分页规则")

        updated_headings: List[str] = []
        for para in out_body.findall("w:p", W):
            text = _extract_para_text(para).strip()
            if not text:
                continue

            ppr = para.find("w:pPr", W)
            if ppr is None:
                continue
            pstyle = ppr.find("w:pStyle", W)
            sid = pstyle.get(_w_attr("val"), "") if pstyle is not None else ""
            sid_norm = re.sub(r"[\s_-]+", "", sid.casefold())
            if sid_norm not in {"1", "heading1"}:
                continue

            page_break_before = ppr.find("w:pageBreakBefore", W)
            if page_break_before is None:
                page_break_before = ET.Element(_w_attr("pageBreakBefore"))
                ppr.append(page_break_before)

            current_val = (page_break_before.get(_w_attr("val"), "") or "").strip().casefold()
            if current_val not in {"1", "true", "on"}:
                page_break_before.set(_w_attr("val"), "1")
                updated_headings.append(text)

        _restore_markup_compatibility(out_doc_root, None)
        if not updated_headings:
            return {"applied": False, "updated": []}

        with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
            tmp_path = Path(tmp.name)

        with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    data = ET.tostring(out_doc_root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

    if tmp_path is not None:
        try:
            _replace_docx_with_lock_hint(tmp_path, docx_path)
        finally:
            if tmp_path.exists():
                tmp_path.unlink(missing_ok=True)

    return {"applied": True, "updated": updated_headings}


def _merge_template_front_matter(docx_path: Path, reference_doc: Path, mode: str) -> Dict[str, object]:
    with zipfile.ZipFile(reference_doc, "r") as zref:
        ref_doc_root = ET.fromstring(zref.read("word/document.xml"))
        ref_rels_root = ET.fromstring(zref.read("word/_rels/document.xml.rels"))
        ref_rel_by_id = {
            rel.get("Id", ""): rel
            for rel in ref_rels_root.findall("{http://schemas.openxmlformats.org/package/2006/relationships}Relationship")
        }

        template_split_idx = _find_body_child_split_by_heading(ref_doc_root, ["摘 要", "摘要"])
        if template_split_idx is None:
            if mode == "strict":
                raise RuntimeError("strict 模式未在参考模板中找到“摘 要/摘要”锚点")
            return {"applied": False}

        tmp_path: Optional[Path] = None
        with zipfile.ZipFile(docx_path, "r") as zin:
            out_doc_root = ET.fromstring(zin.read("word/document.xml"))
            out_rels_root = ET.fromstring(zin.read("word/_rels/document.xml.rels"))
            output_split_idx = _find_body_child_split_by_heading(out_doc_root, ["摘 要", "摘要"])
            if output_split_idx is None:
                if mode == "strict":
                    raise RuntimeError("strict 模式未在导出文档中找到“摘 要/摘要”锚点")
                return {"applied": False}

            ref_body = ref_doc_root.find("w:body", W)
            out_body = out_doc_root.find("w:body", W)
            if ref_body is None or out_body is None:
                raise RuntimeError("DOCX 缺少 w:body，无法合并模板首页")

            ref_children = list(ref_body)
            out_children = list(out_body)

            prefix_children = [_clone_xml(x) for x in ref_children[:template_split_idx]]
            filtered_prefix_children: List[ET.Element] = []
            removed_spine_blocks = 0
            skip_next_table = False
            for child in prefix_children:
                text = ""
                if child.tag == _w_attr("p"):
                    text = _extract_para_text(child).strip()
                    normalized = _norm_plain_text(text)
                    if "书脊" in normalized:
                        removed_spine_blocks += 1
                        skip_next_table = True
                        continue
                elif child.tag == _w_attr("tbl"):
                    if skip_next_table:
                        removed_spine_blocks += 1
                        skip_next_table = False
                        continue
                    text = " ".join(
                        _extract_para_text(para).strip()
                        for para in child.findall(".//w:p", W)
                        if _extract_para_text(para).strip()
                    )
                    normalized = _norm_plain_text(text)
                    if "论文题目" in normalized and "中国科学院大学" in normalized and "作者姓名" in normalized:
                        removed_spine_blocks += 1
                        continue
                filtered_prefix_children.append(child)
            prefix_children = filtered_prefix_children
            tail_children = [_clone_xml(x) for x in out_children[output_split_idx:]]
            if not tail_children:
                raise RuntimeError("导出文档“摘要”后无正文内容，无法合并模板首页")

            for child in list(out_body):
                out_body.remove(child)
            for child in prefix_children + tail_children:
                out_body.append(child)

            required_rel_ids = _collect_relationship_ids_from_elements(prefix_children)
            custom_xml_rel_ids = [
                rid
                for rid, rel in ref_rel_by_id.items()
                if (rel.get("Type", "") or "").endswith("/customXml")
            ]
            required_rel_ids.extend(custom_xml_rel_ids)

            rid_map, required_targets = _merge_required_relationships(required_rel_ids, ref_rel_by_id, out_rels_root)
            _remap_relationship_ids(prefix_children, rid_map)
            _restore_markup_compatibility(out_doc_root, ref_doc_root)

            required_part_files = _collect_related_parts_from_targets(
                zref,
                source_part="word/document.xml",
                targets=required_targets,
            )

            with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
                tmp_path = Path(tmp.name)

            with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    if item.filename in required_part_files:
                        continue
                    if item.filename == "word/document.xml":
                        data = _serialize_xml(out_doc_root)
                        zout.writestr(item, data)
                    elif item.filename == "word/_rels/document.xml.rels":
                        data = _serialize_xml(out_rels_root, default_namespace=PKG_REL_NS)
                        zout.writestr(item, data)
                    else:
                        zout.writestr(item, zin.read(item.filename))

                for part in sorted(required_part_files):
                    if part in zref.namelist():
                        zout.writestr(part, zref.read(part))

        if tmp_path is not None:
            try:
                _replace_docx_with_lock_hint(tmp_path, docx_path)
            finally:
                if tmp_path.exists():
                    tmp_path.unlink(missing_ok=True)

    return {
        "applied": True,
        "template_split_idx": template_split_idx,
        "output_split_idx": output_split_idx,
        "removed_spine_blocks": removed_spine_blocks,
    }


def _set_run_text(run, text: str) -> bool:
    if run.text == text:
        return False
    run.text = text
    return True


def _set_run_sequence(paragraph, values: List[str]) -> int:
    runs = list(paragraph.runs)
    if not runs:
        return 0

    changed = 0
    for idx, run in enumerate(runs):
        target = values[idx] if idx < len(values) else ""
        if _set_run_text(run, target):
            changed += 1
    return changed


_SPECIES_TOKEN_RE = re.compile(
    r"(Artemisia\s+selengensis\*?|A\.\s*selengensis\*?)",
    re.IGNORECASE,
)
_SPECIES_MARKED_TOKEN_RE = re.compile(
    r"\*?(Artemisia\s+selengensis|A\.\s*selengensis)\*?",
    re.IGNORECASE,
)


def _set_cover_english_title_with_species_italics(paragraph, text: str) -> int:
    """Set cover English title while forcing species tokens to italic runs."""
    raw_text = text or ""
    # Remove markdown-style emphasis markers that may come from latex_inline_to_text.
    raw_text = re.sub(
        r"\*(Artemisia\s+selengensis|A\.\s*selengensis)\*",
        r"\1",
        raw_text,
        flags=re.IGNORECASE,
    )

    pieces: List[Tuple[str, bool]] = []
    last = 0
    for m in _SPECIES_TOKEN_RE.finditer(raw_text):
        if m.start() > last:
            pieces.append((raw_text[last : m.start()], False))
        pieces.append((m.group(0), True))
        last = m.end()
    if last < len(raw_text):
        pieces.append((raw_text[last:], False))
    if not pieces:
        pieces = [(raw_text, False)]

    runs = list(paragraph.runs)
    if not runs:
        paragraph.add_run("")
        runs = list(paragraph.runs)
    style_seed_run = next((r for r in runs if (r.text or "").strip()), runs[0])
    seed_bold = style_seed_run.bold
    seed_font_name = style_seed_run.font.name
    seed_font_size = style_seed_run.font.size
    while len(runs) < len(pieces):
        paragraph.add_run("")
        runs = list(paragraph.runs)

    changed = 0
    for idx, run in enumerate(runs):
        if idx < len(pieces):
            txt, italic = pieces[idx]
        else:
            txt, italic = "", False
        if _set_run_text(run, txt):
            changed += 1
        # 统一封面英文标题的基础字符样式，避免“前半句粗体、学名段落细体”不一致。
        if seed_bold is not None and run.bold != seed_bold:
            run.bold = seed_bold
            changed += 1
        if seed_font_name and run.font.name != seed_font_name:
            run.font.name = seed_font_name
            changed += 1
        if seed_font_size is not None and run.font.size != seed_font_size:
            run.font.size = seed_font_size
            changed += 1
        if run.italic != italic:
            run.italic = italic
            changed += 1
        # 防止封面模板旧 run 残留“下划线/上下标”污染英文标题显示。
        if run.underline:
            run.underline = False
            changed += 1
        if run.font.subscript:
            run.font.subscript = False
            changed += 1
        if run.font.superscript:
            run.font.superscript = False
            changed += 1
    return changed


def _set_selected_runs(paragraph, updates: Dict[int, str]) -> int:
    runs = list(paragraph.runs)
    changed = 0
    for idx, text in updates.items():
        if 0 <= idx < len(runs) and _set_run_text(runs[idx], text):
            changed += 1
    return changed


def _compose_supervisor_cn_line(supervisor_cn: str, supervisor_org_cn: str) -> str:
    parts = [part.strip() for part in (supervisor_cn, supervisor_org_cn) if part and part.strip()]
    return "  ".join(parts)


def _delete_doc_paragraph(paragraph) -> bool:
    element = getattr(paragraph, "_element", None)
    if element is None:
        return False
    parent = element.getparent()
    if parent is None:
        return False
    parent.remove(element)
    return True


def _find_doc_paragraph_index(doc, needle: str, *, start: int = 0) -> Optional[int]:
    if not needle:
        return None
    for idx in range(start, len(doc.paragraphs)):
        text = (doc.paragraphs[idx].text or "").strip()
        if needle in text:
            return idx
    return None


def _find_first_doc_paragraph_index_any(doc, needles: List[str]) -> Optional[int]:
    found: List[int] = []
    for needle in needles:
        idx = _find_doc_paragraph_index(doc, needle)
        if idx is not None:
            found.append(idx)
    return min(found) if found else None


def _measure_underlined_chars(paragraph) -> int:
    total = 0
    for run in paragraph.runs:
        if run.underline:
            total += len(run.text or "")
    return total


def _set_cover_underlined_field(
    paragraph,
    *,
    text_run_idx: int,
    text: str,
    target_total_chars: Optional[int] = None,
) -> int:
    runs = list(paragraph.runs)
    if not (0 <= text_run_idx < len(runs)):
        return 0

    underline_indices = [idx for idx, run in enumerate(runs) if run.underline]
    if text_run_idx not in underline_indices:
        return _set_run_text(runs[text_run_idx], text)

    changed = 0
    if _set_run_text(runs[text_run_idx], text):
        changed += 1

    if target_total_chars is None:
        target_total_chars = sum(len(runs[idx].text or "") for idx in underline_indices)

    current_total = sum(len(runs[idx].text or "") for idx in underline_indices)
    filler_indices = [idx for idx in underline_indices if idx != text_run_idx]
    if not filler_indices:
        return changed

    last_idx = filler_indices[-1]
    last_text = runs[last_idx].text or ""
    deficit = target_total_chars - current_total
    if deficit > 0:
        new_text = last_text + (" " * deficit)
    elif deficit < 0:
        trim = min(len(last_text), -deficit)
        new_text = last_text[: len(last_text) - trim]
    else:
        new_text = last_text

    if new_text != last_text and _set_run_text(runs[last_idx], new_text):
        changed += 1
    return changed


def _derive_supervisor_org_en(supervisor_org_cn: str, institute_en: str) -> str:
    mapping = {
        "中国科学院南京土壤研究所": "Institute of Soil Science, Chinese Academy of Sciences",
    }
    org_en = mapping.get(supervisor_org_cn, "").strip()
    institute_en = (institute_en or "").strip()
    institute_en = re.sub(r"\s*/\s*", " / ", institute_en)
    institute_en = re.sub(r"\s+", " ", institute_en).strip()

    if "fill in institute here" in institute_en.casefold():
        institute_en = ""

    if org_en and institute_en and org_en not in institute_en:
        return f"{org_en} / {institute_en}"
    return org_en or institute_en


def _compact_plain_text(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def _normalize_spine_title_text(zh_title: str) -> str:
    """书脊标题优先去空白/标点，避免窄版芯下被强制换行导致跨页。"""
    if not zh_title:
        return ""
    compact = re.sub(r"\s+", "", zh_title)
    compact = re.sub(r"[，,。；;：:！？!?、“”\"'（）()【】\\[\\]《》<>·•—-]", "", compact)
    return compact or zh_title


def _trim_spine_cell_paragraphs(cell, title_idx: int, author_idx: int, school_idx: int | None = None) -> None:
    """压缩书脊占位表格中的空白段，确保内容落在单页且位置稳定。"""
    paragraphs = list(cell.paragraphs)
    if not paragraphs:
        return

    keep: set[int] = {title_idx, author_idx}
    if school_idx is not None:
        keep.add(school_idx)

    anchors = sorted(idx for idx in (title_idx, author_idx, school_idx) if idx is not None)
    if len(anchors) < 2:
        return
    title_idx = anchors[0]
    author_idx = anchors[1]
    school_anchor = anchors[2] if len(anchors) > 2 else None

    lead_kept = 0
    for idx in range(title_idx - 1, -1, -1):
        if lead_kept >= 3:
            break
        if not (paragraphs[idx].text or "").strip():
            keep.add(idx)
            lead_kept += 1

    mid_kept = 0
    for idx in range(title_idx + 1, author_idx):
        if mid_kept >= 2:
            break
        if not (paragraphs[idx].text or "").strip():
            keep.add(idx)
            mid_kept += 1

    if school_anchor is not None:
        lower_mid_kept = 0
        for idx in range(author_idx + 1, school_anchor):
            if lower_mid_kept >= 2:
                break
            if not (paragraphs[idx].text or "").strip():
                keep.add(idx)
                lower_mid_kept += 1

        tail_start = school_anchor + 1
    else:
        tail_start = author_idx + 1

    tail_kept = 0
    for idx in range(tail_start, len(paragraphs)):
        if tail_kept >= 1:
            break
        if not (paragraphs[idx].text or "").strip():
            keep.add(idx)
            tail_kept += 1

    for idx in range(len(paragraphs) - 1, -1, -1):
        if idx in keep:
            continue
        _delete_doc_paragraph(paragraphs[idx])


def _sync_spine_metadata_from_info(doc, zh_title: str, author_cn: str, school_cn: str = "中国科学院大学") -> int:
    """同步书脊页的题目/作者/学校占位内容。"""
    if not zh_title and not author_cn and not school_cn:
        return 0

    spine_title = _normalize_spine_title_text(zh_title)
    changed = 0
    for table in doc.tables:
        rows = list(getattr(table, "rows", []))
        if len(rows) != 1:
            continue
        cells = list(rows[0].cells)
        if len(cells) != 1:
            continue

        paragraphs = list(cells[0].paragraphs)
        if len(paragraphs) < 12:
            continue

        non_empty = [(idx, (para.text or "").strip()) for idx, para in enumerate(paragraphs) if (para.text or "").strip()]
        if not non_empty:
            continue

        normalized_non_empty = [(idx, _compact_plain_text(text)) for idx, text in non_empty]
        has_spine_marker = any(
            marker in norm_text
            for _, norm_text in normalized_non_empty
            for marker in ("论文题目", "书脊", "中国科学院大学", "作者姓名", "张帅")
        )
        if not has_spine_marker:
            continue

        title_idx = next(
            (idx for idx, norm in normalized_non_empty if "论文题目" in norm or "书脊" in norm),
            None,
        )
        author_idx = next(
            (idx for idx, norm in normalized_non_empty if "作者姓名" in norm or "张帅" in norm),
            None,
        )
        school_idx = next(
            (idx for idx, norm in normalized_non_empty if "中国科学院大学" in norm),
            None,
        )

        if title_idx is None:
            split_idx = len(paragraphs) // 2
            upper_non_empty = [idx for idx, _ in non_empty if idx < split_idx]
            title_idx = upper_non_empty[0] if upper_non_empty else non_empty[0][0]
        if author_idx is None:
            split_idx = len(paragraphs) // 2
            lower_non_empty = [idx for idx, _ in non_empty if idx >= split_idx]
            author_idx = lower_non_empty[0] if lower_non_empty else non_empty[-1][0]
        if school_idx is None:
            lower_after_author = [idx for idx, _ in non_empty if idx > author_idx]
            school_idx = lower_after_author[-1] if lower_after_author else len(paragraphs) - 1

        if spine_title:
            changed += _set_run_sequence(paragraphs[title_idx], [spine_title])
        if author_cn:
            changed += _set_run_sequence(paragraphs[author_idx], [author_cn])
        if school_cn:
            changed += _set_run_sequence(paragraphs[school_idx], [school_cn])

        # 清理其余非关键占位文本，避免干扰书脊版式。
        for idx, norm in normalized_non_empty:
            if idx in (title_idx, author_idx, school_idx):
                continue
            if any(marker in norm for marker in ("3cm左右", "中国科学院大学", "论文题目", "作者姓名", "张帅", "书脊")):
                changed += _set_run_sequence(paragraphs[idx], [""])

        _trim_spine_cell_paragraphs(cells[0], title_idx, author_idx, school_idx)

        break

    return changed


def _sync_cover_metadata_from_info(docx_path: Path, project_dir: Path) -> Dict[str, object]:
    if DocxDocument is None:
        return {"applied": False, "updated": 0, "reason": "python-docx unavailable"}

    info_path = project_dir / "extraTex" / "info.tex"
    if not info_path.exists():
        return {"applied": False, "updated": 0, "reason": f"missing info.tex: {info_path}"}

    info = _parse_info_fields(info_path.read_text(encoding="utf-8"))
    zh_title = _strip_markdown_emphasis_markers(_clean_info_value(info["title"][0]))
    en_title = _strip_markdown_emphasis_markers(_clean_info_value(info["title"][1]))
    degree_level_cn = _clean_info_value(info["degree_level"][0])
    author_cn = _clean_info_value(info["author"][0])
    author_en = _clean_info_value(info["author"][1])
    supervisor_cn = _clean_info_value(info["supervisor"][0])
    supervisor_en = _clean_info_value(info["supervisor"][1])
    supervisor_org_cn = _clean_info_value(info["supervisor"][2])
    degree_type_cn = _clean_info_value(info["degree_type"][0])
    degree_type_en = _clean_info_value(info["degree_type"][1])
    subject_cn = _clean_info_value(info["subject"][0])
    subject_en = _clean_info_value(info["subject"][1])
    institute_cn = _clean_info_value(info["institute"][0])
    institute_en = _clean_info_value(info["institute"][1])
    grad_year = _clean_info_value(info["grad_year"][0])
    grad_month_cn = _clean_info_value(info["grad_month"][0])
    grad_month_en = _clean_info_value(info["grad_month"][1])

    supervisor_cn_inline = _compose_supervisor_cn_line(supervisor_cn, supervisor_org_cn)

    supervisor_affiliation_en = _derive_supervisor_org_en(supervisor_org_cn, institute_en)
    supervisor_en_line = f"Supervisor: {supervisor_en}" if supervisor_en else ""
    dissertation_line = "A dissertation submitted to" if degree_level_cn == "博士" else "A thesis submitted to"

    doc = DocxDocument(str(docx_path))
    changed = 0

    first_abstract_idx = _find_first_doc_paragraph_index_any(doc, ["摘 要", "摘要", "Abstract"])
    has_full_cover_window = first_abstract_idx is None or first_abstract_idx > 44

    if has_full_cover_window:
        if 4 < len(doc.paragraphs):
            changed += _set_run_sequence(doc.paragraphs[4], [degree_level_cn, "", "学位论文"])
        if 7 < len(doc.paragraphs):
            changed += _set_run_sequence(doc.paragraphs[7], [zh_title])
        if 12 < len(doc.paragraphs):
            changed += _set_selected_runs(doc.paragraphs[12], {4: author_cn, 5: ""})
        if 13 < len(doc.paragraphs):
            changed += _set_selected_runs(doc.paragraphs[13], {3: supervisor_cn_inline, 4: "        "})
        if 14 < len(doc.paragraphs):
            # Keep this line empty; we collapse it after label-based sync to avoid a gap before "学位类别".
            changed += _set_run_sequence(doc.paragraphs[14], [""] * len(doc.paragraphs[14].runs))
        if 15 < len(doc.paragraphs):
            changed += _set_selected_runs(doc.paragraphs[15], {5: degree_type_cn})
        if 16 < len(doc.paragraphs):
            changed += _set_selected_runs(doc.paragraphs[16], {5: subject_cn})
        if 17 < len(doc.paragraphs):
            changed += _set_selected_runs(doc.paragraphs[17], {5: institute_cn})
        if 21 < len(doc.paragraphs):
            changed += _set_run_sequence(doc.paragraphs[21], [grad_year, "年", grad_month_cn, "月", ""])
        if 26 < len(doc.paragraphs):
            changed += _set_cover_english_title_with_species_italics(doc.paragraphs[26], en_title)
        if 31 < len(doc.paragraphs):
            changed += _set_run_sequence(doc.paragraphs[31], [dissertation_line])
        if 35 < len(doc.paragraphs):
            changed += _set_run_sequence(doc.paragraphs[35], [degree_type_en, ""])
        if 36 < len(doc.paragraphs):
            changed += _set_run_sequence(doc.paragraphs[36], ["in " if subject_en else "", subject_en])
        if 38 < len(doc.paragraphs):
            changed += _set_run_sequence(doc.paragraphs[38], [author_en, "", "", "", ""])
        if 39 < len(doc.paragraphs):
            changed += _set_run_sequence(doc.paragraphs[39], [supervisor_en_line, "", "", "", "", "", ""])
        if 40 < len(doc.paragraphs):
            changed += _set_run_sequence(
                doc.paragraphs[40],
                [supervisor_affiliation_en] + [""] * max(0, len(doc.paragraphs[40].runs) - 1),
            )
        if 43 < len(doc.paragraphs):
            changed += _set_run_sequence(doc.paragraphs[43], [institute_en])
        if 44 < len(doc.paragraphs):
            changed += _set_run_sequence(doc.paragraphs[44], [grad_month_en, " " if grad_month_en and grad_year else "", grad_year, ""])

    zh_field_target = 0
    for marker in ["作者姓名", "学位类别", "学科专业", "培养单位"]:
        idx = _find_doc_paragraph_index(doc, marker)
        if idx is not None:
            zh_field_target = max(zh_field_target, _measure_underlined_chars(doc.paragraphs[idx]))

    supervisor_idx = _find_doc_paragraph_index(doc, "指导教师")
    if supervisor_idx is not None:
        target_chars = max(zh_field_target or 0, _measure_underlined_chars(doc.paragraphs[supervisor_idx]))
        changed += _set_cover_underlined_field(
            doc.paragraphs[supervisor_idx],
            text_run_idx=3,
            text=supervisor_cn_inline,
            target_total_chars=target_chars or None,
        )

        # Remove intermediate blank/placeholder lines so "学位类别" follows supervisor directly.
        degree_idx = _find_doc_paragraph_index(doc, "学位类别", start=supervisor_idx + 1)
        if degree_idx is not None and degree_idx > supervisor_idx + 1:
            remove_idx = degree_idx - 1
            while remove_idx > supervisor_idx:
                mid_para = doc.paragraphs[remove_idx]
                mid_text = (mid_para.text or "").strip()
                if any(marker in mid_text for marker in ("学位类别", "学科专业", "培养单位")):
                    break
                if _delete_doc_paragraph(mid_para):
                    changed += 1
                remove_idx -= 1

    supervisor_en_idx = _find_doc_paragraph_index(doc, "Supervisor:")
    if supervisor_en_idx is not None:
        changed += _set_run_sequence(
            doc.paragraphs[supervisor_en_idx],
            [supervisor_en_line] + [""] * max(0, len(doc.paragraphs[supervisor_en_idx].runs) - 1),
        )

    affiliation_en_idx = _find_doc_paragraph_index(doc, "Fill in institute here", start=supervisor_en_idx or 0)
    if affiliation_en_idx is None:
        affiliation_en_idx = _find_doc_paragraph_index(doc, "University of Chinese Academy of Sciences", start=supervisor_en_idx or 0)
    if affiliation_en_idx is not None:
        changed += _set_run_sequence(
            doc.paragraphs[affiliation_en_idx],
            [supervisor_affiliation_en] + [""] * max(0, len(doc.paragraphs[affiliation_en_idx].runs) - 1),
        )

    spine_changed = 0

    if changed:
        doc.save(str(docx_path))

    return {"applied": bool(changed), "updated": changed, "spine_updated": spine_changed}


def _remove_front_matter_annotation_shapes(docx_path: Path) -> Dict[str, object]:
    annotations = (
        "根据实际申请学位类别选择",
        "填写信息后不会自动居中",
        "dissertation用于博士论文",
        "黑体小四号，论文题目、作者、中国科学院大学英文和阿拉伯数字用Times New Roman",
        "3cm左右",
    )

    with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
        tmp_path = Path(tmp.name)

    removed = 0
    try:
        with zipfile.ZipFile(docx_path, "r") as zin:
            doc_root = ET.fromstring(zin.read("word/document.xml"))
            parent_map = {child: parent for parent in doc_root.iter() for child in parent}
            paragraphs_to_remove = []
            runs_to_remove = []

            for txbx in doc_root.findall(".//w:txbxContent", W):
                text = "".join((t.text or "") for t in txbx.findall(".//w:t", W)).strip()
                shape_id = ""
                shape_type = ""
                current = txbx
                while current is not None:
                    if current.tag == _v_attr("shape"):
                        shape_id = (current.get("id") or "").strip()
                        shape_type = (current.get("type") or "").strip()
                        break
                    current = parent_map.get(current)

                shape_hint = f"{shape_id} {shape_type}".strip()
                is_annotation_shape = any(token in shape_hint for token in ("对话气泡", "标注", "_x0000_s"))

                # 非空文本框：仅移除命中注释说明文案的框。
                if text:
                    if not any(marker in text for marker in annotations):
                        continue
                else:
                    # 空文本框：仅移除已识别为注释/气泡的残留框，避免误删其它有效对象。
                    if not is_annotation_shape:
                        continue

                current = txbx
                paragraph = None
                while current is not None:
                    if current.tag == _w_attr("p"):
                        paragraph = current
                        break
                    current = parent_map.get(current)
                if paragraph is not None and paragraph not in paragraphs_to_remove:
                    paragraphs_to_remove.append(paragraph)

            for line in doc_root.findall(".//v:line", OOXML_NAMESPACE_PREFIXES):
                line_id = (line.get("id") or "").strip()
                if "直接连接符" not in line_id:
                    continue

                current = line
                paragraph = None
                run = None
                while current is not None:
                    if run is None and current.tag == _w_attr("r"):
                        run = current
                    if current.tag == _w_attr("p"):
                        paragraph = current
                        break
                    current = parent_map.get(current)

                if paragraph is None:
                    continue

                paragraph_text = "".join((t.text or "") for t in paragraph.findall(".//w:t", W)).strip()
                if paragraph_text and paragraph_text != "中国科学院大学":
                    continue

                if not paragraph_text:
                    if paragraph not in paragraphs_to_remove:
                        paragraphs_to_remove.append(paragraph)
                    continue

                if run is not None and run not in runs_to_remove:
                    runs_to_remove.append(run)

            for paragraph in paragraphs_to_remove:
                parent = parent_map.get(paragraph)
                if parent is not None:
                    parent.remove(paragraph)
                    removed += 1

            for run in runs_to_remove:
                parent = parent_map.get(run)
                if parent is not None:
                    parent.remove(run)
                    removed += 1

            with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    if item.filename == "word/document.xml":
                        data = ET.tostring(doc_root, encoding="utf-8", xml_declaration=True)
                        zout.writestr(item, data)
                    else:
                        zout.writestr(item, zin.read(item.filename))

        _replace_docx_with_lock_hint(tmp_path, docx_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)

    return {"applied": bool(removed), "removed": removed}


def _strip_front_matter_hint_text(docx_path: Path) -> Dict[str, object]:
    literals = [
        "填写信息后不会自动居中，可以根据实际情况手动调整居中。",
    ]

    with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
        tmp_path = Path(tmp.name)

    replaced = 0
    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
            tmp_path, "w", compression=zipfile.ZIP_DEFLATED
        ) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    text = data.decode("utf-8")
                    for literal in literals:
                        if literal in text:
                            replaced += text.count(literal)
                            text = text.replace(literal, "")
                    data = text.encode("utf-8")
                zout.writestr(item, data)

        _replace_docx_with_lock_hint(tmp_path, docx_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)

    return {"applied": bool(replaced), "removed": replaced}


def _cleanup_known_word_artifacts(docx_path: Path) -> Dict[str, object]:
    macro_keywords = ("macrobutton", "acceptallchangesshown")
    pt_artifact_patterns = [
        re.compile(r"^\[\d+\]\s*\d+(?:\.\d+)?\s*pt\s*#\s*\d+(?:\.\d+)?\s*pt$", re.IGNORECASE),
        re.compile(r"^\[\d+\]\s*#\s*\d+(?:\.\d+)?\s*pt$", re.IGNORECASE),
        re.compile(r"^#\s*\d+(?:\.\d+)?\s*pt$", re.IGNORECASE),
        re.compile(r"^\d+(?:\.\d+)?\s*pt$", re.IGNORECASE),
    ]

    with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
        tmp_path = Path(tmp.name)

    removed_macro_runs = 0
    removed_macro_paragraphs = 0
    removed_pt_paragraphs = 0
    try:
        with zipfile.ZipFile(docx_path, "r") as zin:
            doc_root = ET.fromstring(zin.read("word/document.xml"))
            body = doc_root.find("w:body", W)
            if body is None:
                return {
                    "applied": False,
                    "removed_macro_runs": 0,
                    "removed_macro_paragraphs": 0,
                    "removed_pt_paragraphs": 0,
                    "reason": "missing_body",
                }

            for para in list(body):
                if para.tag != _w_attr("p"):
                    continue

                para_text_raw = "".join((t.text or "") for t in para.findall(".//w:t", W))
                para_instr_raw = "".join((t.text or "") for t in para.findall(".//w:instrText", W))
                para_payload = f"{para_text_raw} {para_instr_raw}".casefold()
                has_macro_artifact = all(k in para_payload for k in macro_keywords) or any(
                    k in para_payload for k in macro_keywords
                )

                if has_macro_artifact:
                    for child in list(para):
                        if child.tag == _w_attr("r"):
                            run_text = "".join((t.text or "") for t in child.findall(".//w:t", W)).casefold()
                            run_instr = "".join((t.text or "") for t in child.findall(".//w:instrText", W)).casefold()
                            run_payload = f"{run_text} {run_instr}"
                            run_has_macro = any(k in run_payload for k in macro_keywords)
                            run_has_field_structure = (
                                child.find("w:instrText", W) is not None or child.find("w:fldChar", W) is not None
                            )
                            if run_has_macro or run_has_field_structure:
                                para.remove(child)
                                removed_macro_runs += 1
                                continue
                        if child.tag == _w_attr("fldSimple"):
                            fld_instr = (child.get(_w_attr("instr"), "") or "").casefold()
                            fld_text = "".join((t.text or "") for t in child.findall(".//w:t", W)).casefold()
                            if any(k in fld_instr or k in fld_text for k in macro_keywords):
                                para.remove(child)
                                removed_macro_runs += 1

                para_text_clean = re.sub(r"\s+", " ", _extract_para_text(para)).strip()
                is_pt_artifact = bool(
                    para_text_clean and any(p.fullmatch(para_text_clean) for p in pt_artifact_patterns)
                )
                if is_pt_artifact:
                    body.remove(para)
                    removed_pt_paragraphs += 1
                    continue

                if has_macro_artifact and not para_text_clean:
                    body.remove(para)
                    removed_macro_paragraphs += 1

            with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    if item.filename == "word/document.xml":
                        data = ET.tostring(doc_root, encoding="utf-8", xml_declaration=True)
                        zout.writestr(item, data)
                    else:
                        zout.writestr(item, zin.read(item.filename))

        _replace_docx_with_lock_hint(tmp_path, docx_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)

    return {
        "applied": bool(removed_macro_runs or removed_macro_paragraphs or removed_pt_paragraphs),
        "removed_macro_runs": removed_macro_runs,
        "removed_macro_paragraphs": removed_macro_paragraphs,
        "removed_pt_paragraphs": removed_pt_paragraphs,
    }


def _twip_to_cm(value: str) -> str:
    try:
        return f"{(float(value) * 2.54 / 1440):.2f}"
    except Exception:
        return "NA"


def _analyze_docx_styles(docx_path: Path) -> Dict[str, object]:
    style_name_by_id: Dict[str, str] = {}
    style_name_to_id: Dict[str, str] = {}
    style_ids: set[str] = set()
    usage: Counter[str] = Counter()
    usage_body: Counter[str] = Counter()
    layout: Dict[str, str] = {}
    style_props: Dict[str, Dict[str, str]] = {}
    body_paragraphs: List[Tuple[str, str]] = []
    body_paragraph_details: List[Dict[str, str]] = []
    image_paragraph_details: List[Dict[str, str]] = []
    section_count = 0

    with zipfile.ZipFile(docx_path, "r") as zf:
        styles_xml = zf.read("word/styles.xml")
        styles_root = ET.fromstring(styles_xml)
        for st in styles_root.findall("w:style", W):
            if st.get(_w_attr("type"), "") != "paragraph":
                continue
            sid = st.get(_w_attr("styleId"), "")
            if not sid:
                continue
            style_ids.add(sid)
            name_node = st.find("w:name", W)
            sname = name_node.get(_w_attr("val"), sid) if name_node is not None else sid
            style_name_by_id[sid] = sname
            style_name_to_id[sname.casefold()] = sid

            ppr = st.find("w:pPr", W)
            rpr = st.find("w:rPr", W)
            spacing = ppr.find("w:spacing", W) if ppr is not None else None
            ind = ppr.find("w:ind", W) if ppr is not None else None
            jc = ppr.find("w:jc", W) if ppr is not None else None
            fonts = rpr.find("w:rFonts", W) if rpr is not None else None
            sz = rpr.find("w:sz", W) if rpr is not None else None
            style_props[sid] = {
                "line": spacing.get(_w_attr("line"), "") if spacing is not None else "",
                "lineRule": spacing.get(_w_attr("lineRule"), "") if spacing is not None else "",
                "before": spacing.get(_w_attr("before"), "") if spacing is not None else "",
                "after": spacing.get(_w_attr("after"), "") if spacing is not None else "",
                "firstLine": ind.get(_w_attr("firstLine"), "") if ind is not None else "",
                "firstLineChars": ind.get(_w_attr("firstLineChars"), "") if ind is not None else "",
                "left": ind.get(_w_attr("left"), "") if ind is not None else "",
                "jc": jc.get(_w_attr("val"), "") if jc is not None else "",
                "font_ascii": fonts.get(_w_attr("ascii"), "") if fonts is not None else "",
                "font_hAnsi": fonts.get(_w_attr("hAnsi"), "") if fonts is not None else "",
                "font_eastAsia": fonts.get(_w_attr("eastAsia"), "") if fonts is not None else "",
                "size_half_points": sz.get(_w_attr("val"), "") if sz is not None else "",
            }

        for name in zf.namelist():
            if not name.endswith(".xml"):
                continue
            if not (
                name == "word/document.xml"
                or name.startswith("word/header")
                or name.startswith("word/footer")
                or name in {"word/footnotes.xml", "word/endnotes.xml", "word/comments.xml"}
            ):
                continue
            data = zf.read(name)
            try:
                root = ET.fromstring(data)
            except Exception:
                continue
            for para in root.findall(".//w:p", W):
                ppr = para.find("w:pPr", W)
                if ppr is None:
                    continue
                pstyle = ppr.find("w:pStyle", W)
                if pstyle is None:
                    continue
                sid = pstyle.get(_w_attr("val"), "")
                if sid:
                    usage[sid] += 1

        try:
            doc_root = ET.fromstring(zf.read("word/document.xml"))
            for para in doc_root.findall(".//w:body//w:p", W):
                ppr = para.find("w:pPr", W)
                sid = ""
                ind = None
                jc = None
                has_image = _paragraph_contains_embedded_image(para)
                if ppr is not None:
                    pstyle = ppr.find("w:pStyle", W)
                    if pstyle is not None:
                        sid = pstyle.get(_w_attr("val"), "")
                    ind = ppr.find("w:ind", W)
                    jc = ppr.find("w:jc", W)
                if sid:
                    usage_body[sid] += 1
                text = "".join(t.text or "" for t in para.findall(".//w:t", W)).strip()
                if has_image:
                    image_paragraph_details.append(
                        {
                            "sid": sid,
                            "text": text,
                            "left": ind.get(_w_attr("left"), "") if ind is not None else "",
                            "firstLine": ind.get(_w_attr("firstLine"), "") if ind is not None else "",
                            "firstLineChars": ind.get(_w_attr("firstLineChars"), "") if ind is not None else "",
                            "hanging": ind.get(_w_attr("hanging"), "") if ind is not None else "",
                            "hangingChars": ind.get(_w_attr("hangingChars"), "") if ind is not None else "",
                            "jc": jc.get(_w_attr("val"), "") if jc is not None else "",
                        }
                    )
                if text:
                    body_paragraphs.append((sid, text))
                    body_paragraph_details.append(
                        {
                            "sid": sid,
                            "text": text,
                            "left": ind.get(_w_attr("left"), "") if ind is not None else "",
                            "firstLine": ind.get(_w_attr("firstLine"), "") if ind is not None else "",
                            "firstLineChars": ind.get(_w_attr("firstLineChars"), "") if ind is not None else "",
                            "hanging": ind.get(_w_attr("hanging"), "") if ind is not None else "",
                            "hangingChars": ind.get(_w_attr("hangingChars"), "") if ind is not None else "",
                            "jc": jc.get(_w_attr("val"), "") if jc is not None else "",
                        }
                    )
        except Exception:
            pass

        try:
            doc_root = ET.fromstring(zf.read("word/document.xml"))
            section_count = len(doc_root.findall(".//w:sectPr", W))
            sect = doc_root.find(".//w:sectPr", W)
            if sect is not None:
                pg_sz = sect.find("w:pgSz", W)
                pg_mar = sect.find("w:pgMar", W)
                if pg_sz is not None:
                    layout["page_w_twip"] = pg_sz.get(_w_attr("w"), "")
                    layout["page_h_twip"] = pg_sz.get(_w_attr("h"), "")
                if pg_mar is not None:
                    layout["margin_top_twip"] = pg_mar.get(_w_attr("top"), "")
                    layout["margin_bottom_twip"] = pg_mar.get(_w_attr("bottom"), "")
                    layout["margin_left_twip"] = pg_mar.get(_w_attr("left"), "")
                    layout["margin_right_twip"] = pg_mar.get(_w_attr("right"), "")
                    layout["margin_header_twip"] = pg_mar.get(_w_attr("header"), "")
                    layout["margin_footer_twip"] = pg_mar.get(_w_attr("footer"), "")
        except Exception:
            pass

    unknown = Counter({sid: cnt for sid, cnt in usage.items() if sid not in style_ids})
    return {
        "style_ids": style_ids,
        "style_name_by_id": style_name_by_id,
        "style_name_to_id": style_name_to_id,
        "usage": usage,
        "usage_body": usage_body,
        "unknown": unknown,
        "layout": layout,
        "style_props": style_props,
        "body_paragraphs": body_paragraphs,
        "body_paragraph_details": body_paragraph_details,
        "image_paragraph_details": image_paragraph_details,
        "section_count": section_count,
    }


def _resolve_style_id(
    style_ids: set[str],
    style_name_to_id: Dict[str, str],
    preferred_ids: List[str],
    preferred_names: List[str],
    fallback: str,
) -> str:
    for sid in preferred_ids:
        if sid in style_ids:
            return sid
    for name in preferred_names:
        sid = style_name_to_id.get(name.casefold())
        if sid and sid in style_ids:
            return sid
    if fallback in style_ids:
        return fallback
    return next(iter(style_ids))


def _build_style_remap(style_info: Dict[str, object]) -> Dict[str, str]:
    style_ids = style_info["style_ids"]  # type: ignore[assignment]
    style_name_to_id = style_info["style_name_to_id"]  # type: ignore[assignment]

    normal_id = _resolve_style_id(style_ids, style_name_to_id, ["a", "Normal"], ["normal"], "a")
    caption_id = _resolve_style_id(
        style_ids,
        style_name_to_id,
        ["ae", "Caption", "caption"],
        ["caption"],
        normal_id,
    )
    h1_id = _resolve_style_id(style_ids, style_name_to_id, ["1", "Heading1"], ["heading 1"], normal_id)
    h2_id = _resolve_style_id(style_ids, style_name_to_id, ["2", "Heading2"], ["heading 2"], normal_id)
    h3_id = _resolve_style_id(style_ids, style_name_to_id, ["3", "Heading3"], ["heading 3"], normal_id)
    h4_id = _resolve_style_id(style_ids, style_name_to_id, ["4", "Heading4"], ["heading 4"], normal_id)

    return {
        "FirstParagraph": normal_id,
        "First Paragraph": normal_id,
        "BodyText": normal_id,
        "Body Text": normal_id,
        "Compact": normal_id,
        "BlockText": normal_id,
        "Block Text": normal_id,
        "Bibliography": normal_id,
        "ImageCaption": caption_id,
        "Image Caption": caption_id,
        "Heading1": h1_id,
        "Heading2": h2_id,
        "Heading3": h3_id,
        "Heading4": h4_id,
    }


def _normalize_paragraph_styles_with_python_docx(docx_path: Path, remap: Dict[str, str]) -> Tuple[int, bool]:
    if DocxDocument is None:
        return 0, False

    doc = DocxDocument(str(docx_path))
    changed = 0
    for para in doc.paragraphs:
        ppr = para._p.pPr
        if ppr is None or ppr.pStyle is None:
            continue
        sid = ppr.pStyle.val
        if sid in remap:
            ppr.pStyle.val = remap[sid]
            changed += 1

    if changed:
        doc.save(str(docx_path))
    return changed, True


def _rewrite_docx_styles_xml(docx_path: Path, remap: Dict[str, str]) -> Counter[str]:
    replaced: Counter[str] = Counter()

    with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
        tmp_path = Path(tmp.name)

    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
            tmp_path, "w", compression=zipfile.ZIP_DEFLATED
        ) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                    try:
                        root = ET.fromstring(data)
                        changed = False
                        for pstyle in root.findall(".//w:pStyle", W):
                            sid = pstyle.get(_w_attr("val"), "")
                            if sid in remap:
                                pstyle.set(_w_attr("val"), remap[sid])
                                replaced[sid] += 1
                                changed = True
                        if changed:
                            data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                    except Exception:
                        pass
                zout.writestr(item, data)

        _replace_docx_with_lock_hint(tmp_path, docx_path)
    finally:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)

    return replaced


def _align_figure_table_style_to_toc1(docx_path: Path) -> Dict[str, object]:
    tmp_path: Optional[Path] = None
    with zipfile.ZipFile(docx_path, "r") as zin:
        try:
            styles_root = ET.fromstring(zin.read("word/styles.xml"))
        except Exception:
            return {"applied": False, "reason": "missing_styles_xml"}

        toc1_style: Optional[ET.Element] = None
        figure_style: Optional[ET.Element] = None
        for style in styles_root.findall("w:style", W):
            sid = (style.get(_w_attr("styleId"), "") or "").strip()
            name_node = style.find("w:name", W)
            name_val = (name_node.get(_w_attr("val"), "") if name_node is not None else "").strip().casefold()
            if sid == "TOC1" or name_val == "toc 1":
                toc1_style = style
            if sid == "af1" or name_val == "table of figures":
                figure_style = style

        if toc1_style is None or figure_style is None:
            return {"applied": False, "reason": "style_not_found"}

        changed = False
        toc1_ppr = toc1_style.find("w:pPr", W)
        toc1_rpr = toc1_style.find("w:rPr", W)

        if toc1_ppr is not None:
            old = figure_style.find("w:pPr", W)
            if old is not None:
                figure_style.remove(old)
            figure_style.append(_clone_xml(toc1_ppr))
            changed = True

        if toc1_rpr is not None:
            old = figure_style.find("w:rPr", W)
            if old is not None:
                figure_style.remove(old)
            figure_style.append(_clone_xml(toc1_rpr))
            changed = True

        based_on = figure_style.find("w:basedOn", W)
        if based_on is None:
            based_on = ET.Element(_w_attr("basedOn"))
            figure_style.insert(0, based_on)
        if based_on.get(_w_attr("val"), "") != "TOC1":
            based_on.set(_w_attr("val"), "TOC1")
            changed = True

        if not changed:
            return {"applied": False, "reason": "already_aligned"}
        _restore_markup_compatibility(styles_root, None)

        with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
            tmp_path = Path(tmp.name)
        with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/styles.xml":
                    data = ET.tostring(styles_root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

    if tmp_path is not None:
        try:
            _replace_docx_with_lock_hint(tmp_path, docx_path)
        finally:
            if tmp_path.exists():
                tmp_path.unlink(missing_ok=True)

    return {"applied": True}


def _extract_para_text_with_tabs(para: ET.Element) -> str:
    parts: List[str] = []
    for node in para.iter():
        if node.tag == _w_attr("tab"):
            parts.append("\t")
        elif node.tag == _w_attr("t"):
            parts.append(node.text or "")
    return "".join(parts)


def _rewrite_paragraph_text_with_tabs(para: ET.Element, text_with_tabs: str) -> None:
    ppr = para.find("w:pPr", W)
    ppr_clone = _clone_xml(ppr) if ppr is not None else None
    for child in list(para):
        para.remove(child)
    if ppr_clone is not None:
        para.append(ppr_clone)

    chunks = text_with_tabs.split("\t")
    for idx, chunk in enumerate(chunks):
        if chunk:
            run_text = ET.Element(_w_attr("r"))
            t = ET.Element(_w_attr("t"))
            if chunk[:1].isspace() or chunk[-1:].isspace():
                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t.text = chunk
            run_text.append(t)
            para.append(run_text)
        if idx < len(chunks) - 1:
            run_tab = ET.Element(_w_attr("r"))
            tab = ET.Element(_w_attr("tab"))
            run_tab.append(tab)
            para.append(run_tab)


def _normalize_toc_entry_spacing_text(text: str, *, in_catalog: bool) -> str:
    if not text:
        return text
    entry = text
    page = ""
    if "\t" in entry:
        entry, page = entry.split("\t", 1)

    if not in_catalog:
        entry = re.sub(r"^致[\s\u3000]+谢", "致谢", entry)
        entry = re.sub(r"^(附录(?:[一二三四五六七八九十百零A-Za-z0-9\-]*)?)[\s\u3000]{2,}", r"\1 ", entry)
    else:
        entry = re.sub(r"^([图表])\s*([0-9]+(?:\s*[-－—–]\s*[0-9]+)*)[\s\u3000]{2,}", r"\1\2 ", entry)
        entry = re.sub(r"^(附录(?:[一二三四五六七八九十百零A-Za-z0-9\-]*)?)[\s\u3000]{2,}", r"\1 ", entry)

    entry = re.sub(r"(?<=\S)~(?=\S)", " ", entry)

    if page:
        return f"{entry}\t{page.strip()}"
    return entry


def _normalize_toc_and_catalog_spacing(docx_path: Path) -> Dict[str, object]:
    tmp_path: Optional[Path] = None
    with zipfile.ZipFile(docx_path, "r") as zin:
        out_doc_root = ET.fromstring(zin.read("word/document.xml"))
        out_body = out_doc_root.find("w:body", W)
        if out_body is None:
            return {"applied": False, "reason": "missing_body"}

        body_children = list(out_body)

        def _find_range(start_markers: List[str], end_markers: List[str]) -> Optional[Tuple[int, int]]:
            start_idx = _find_body_child_split_by_heading(out_doc_root, start_markers)
            if start_idx is None:
                return None
            marker_set = {_norm_plain_text(x) for x in end_markers}
            end_idx = len(body_children)
            for idx, child in enumerate(body_children[start_idx + 1 :], start=start_idx + 1):
                if child.tag != _w_attr("p"):
                    continue
                if _norm_plain_text(_extract_para_text(child)) in marker_set:
                    end_idx = idx
                    break
            return (start_idx + 1, end_idx)

        ranges: List[Tuple[int, int, bool]] = []
        main_range = _find_range(["目录"], ["图表目录", "图目录", "绪论"])
        if main_range is not None:
            ranges.append((main_range[0], main_range[1], False))
        fig_range = _find_range(["图目录"], ["表目录", "绪论"])
        if fig_range is not None:
            ranges.append((fig_range[0], fig_range[1], True))
        tbl_range = _find_range(["表目录"], ["绪论"])
        if tbl_range is not None:
            ranges.append((tbl_range[0], tbl_range[1], True))

        if not ranges:
            return {"applied": False, "reason": "heading_not_found"}

        changed = 0
        style_changed = 0

        def _sanitize_toc_template_rpr(rpr: ET.Element) -> ET.Element:
            cloned = _clone_xml(rpr)
            for tag in ("w:spacing", "w:kern"):
                node = cloned.find(tag, W)
                if node is not None:
                    cloned.remove(node)
            return cloned

        def _first_visible_run_rpr(para: ET.Element) -> Optional[ET.Element]:
            for run in para.findall(".//w:r", W):
                if run.find("w:t", W) is None:
                    continue
                rpr = run.find("w:rPr", W)
                if rpr is not None:
                    return _sanitize_toc_template_rpr(rpr)
            return None

        def _apply_visible_run_rpr(para: ET.Element, template_rpr: Optional[ET.Element]) -> int:
            local_changed = 0
            for run in para.findall(".//w:r", W):
                if run.find("w:t", W) is None:
                    continue
                existing = run.find("w:rPr", W)
                if template_rpr is None:
                    if existing is not None:
                        run.remove(existing)
                        local_changed += 1
                    continue
                cloned = _clone_xml(template_rpr)
                if existing is not None and ET.tostring(existing, encoding="utf-8") == ET.tostring(
                    cloned,
                    encoding="utf-8",
                ):
                    continue
                if existing is not None:
                    run.remove(existing)
                run.insert(0, cloned)
                local_changed += 1
            return local_changed

        for start_idx, end_idx, in_catalog in ranges:
            style_templates: Dict[str, Optional[ET.Element]] = {}
            for child in body_children[start_idx:end_idx]:
                if child.tag != _w_attr("p"):
                    continue
                sid = _extract_para_style_id(child)
                if sid not in {"TOC1", "TOC2", "TOC3", "af1"}:
                    continue
                raw = _extract_para_text_with_tabs(child)
                entry = raw.split("\t", 1)[0].strip()
                if not entry or entry.startswith("附录"):
                    continue
                style_templates.setdefault(sid, _first_visible_run_rpr(child))

            for child in body_children[start_idx:end_idx]:
                if child.tag != _w_attr("p"):
                    continue
                sid = _extract_para_style_id(child)
                if sid not in {"TOC1", "TOC2", "TOC3", "af1"}:
                    continue
                raw = _extract_para_text_with_tabs(child)
                if not raw.strip():
                    continue
                normalized = _normalize_toc_entry_spacing_text(raw, in_catalog=in_catalog)
                if normalized != raw:
                    _rewrite_paragraph_text_with_tabs(child, normalized)
                    changed += 1
                entry = normalized.split("\t", 1)[0].strip()
                if not in_catalog and entry.startswith("附录") and sid in style_templates:
                    style_changed += _apply_visible_run_rpr(child, style_templates[sid])

        if changed <= 0 and style_changed <= 0:
            return {"applied": False, "reason": "no_change"}

        _restore_markup_compatibility(out_doc_root, None)
        with NamedTemporaryFile(delete=False, suffix=".docx", dir=str(docx_path.parent)) as tmp:
            tmp_path = Path(tmp.name)
        with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/document.xml":
                    data = ET.tostring(out_doc_root, encoding="utf-8", xml_declaration=True)
                zout.writestr(item, data)

    if tmp_path is not None:
        try:
            _replace_docx_with_lock_hint(tmp_path, docx_path)
        finally:
            if tmp_path.exists():
                tmp_path.unlink(missing_ok=True)

    return {"applied": True, "changed": changed, "style_changed": style_changed}


def _norm(s: str) -> str:
    return re.sub(r"\s+", "", s or "").replace("　", "")


def _count_keywords(line: str) -> int:
    if "：" in line:
        payload = line.split("：", 1)[1]
    elif ":" in line:
        payload = line.split(":", 1)[1]
    else:
        payload = line
    parts = [p.strip() for p in re.split(r"[，,、;；]", payload) if p.strip()]
    return len(parts)


def _extract_labeled_value(body_paragraphs: List[Tuple[str, str]], labels: List[str]) -> Optional[str]:
    labels_norm = [_norm(x) for x in labels]
    for _, text in body_paragraphs:
        raw = text.strip()
        if not raw:
            continue
        nt = _norm(raw)
        for label in labels_norm:
            if not nt.startswith(label):
                continue
            if "：" in raw:
                return raw.split("：", 1)[1].strip()
            if ":" in raw:
                return raw.split(":", 1)[1].strip()
            tail = raw[len(label) :].strip()
            return tail.lstrip("：: ").strip()
    return None


def _count_nonspace_chars(text: str) -> int:
    return len(re.sub(r"\s+", "", text or ""))


def _count_ascii_letters(text: str) -> int:
    return len(re.findall(r"[A-Za-z]", text or ""))


def _style_level_from_name(style_name: str, prefix: str) -> Optional[int]:
    norm = re.sub(r"[\s_-]+", "", (style_name or "").casefold())
    m = re.fullmatch(re.escape(prefix) + r"(\d+)", norm)
    if not m:
        return None
    try:
        return int(m.group(1))
    except Exception:
        return None


def _spec_checks(after: Dict[str, object], strict_scan: Optional[Dict[str, int]] = None) -> List[Tuple[str, str, str]]:
    layout = after["layout"]  # type: ignore[assignment]
    usage_body = after["usage_body"]  # type: ignore[assignment]
    style_name_by_id = after["style_name_by_id"]  # type: ignore[assignment]
    style_name_to_id = after["style_name_to_id"]  # type: ignore[assignment]
    style_props = after["style_props"]  # type: ignore[assignment]
    body_paragraphs = after["body_paragraphs"]  # type: ignore[assignment]
    body_paragraph_details = after["body_paragraph_details"]  # type: ignore[assignment]
    image_paragraph_details = after.get("image_paragraph_details", [])  # type: ignore[assignment]
    unknown = after["unknown"]  # type: ignore[assignment]

    checks: List[Tuple[str, str, str]] = []

    unknown_total = int(sum(unknown.values()))
    checks.append(("模板结构：未知样式引用(strict=0)", "PASS" if unknown_total == 0 else "WARN", f"当前 {unknown_total}"))

    a4_ok = layout.get("page_w_twip") == "11906" and layout.get("page_h_twip") == "16838"
    checks.append(
        (
            "纸张A4(210x297mm, 表3)",
            "PASS" if a4_ok else "WARN",
            f"{layout.get('page_w_twip')} x {layout.get('page_h_twip')} twip",
        )
    )

    margin_ok = (
        layout.get("margin_top_twip") == "1440"
        and layout.get("margin_bottom_twip") == "1440"
        and layout.get("margin_left_twip") == "1797"
        and layout.get("margin_right_twip") == "1797"
        and layout.get("margin_header_twip") == "851"
        and layout.get("margin_footer_twip") == "851"
    )
    checks.append(
        (
            "页边距与页眉页脚距离(表3)",
            "PASS" if margin_ok else "WARN",
            "top/bottom={}/{} left/right={}/{} header/footer={}/{}".format(
                layout.get("margin_top_twip"),
                layout.get("margin_bottom_twip"),
                layout.get("margin_left_twip"),
                layout.get("margin_right_twip"),
                layout.get("margin_header_twip"),
                layout.get("margin_footer_twip"),
            ),
        )
    )

    normal_id = style_name_to_id.get("normal", "a")
    normal = style_props.get(normal_id, {})
    normal_font_ok = "宋体" in normal.get("font_eastAsia", "") and (
        "Times New Roman" in normal.get("font_ascii", "") or "Times New Roman" in normal.get("font_hAnsi", "")
    )
    checks.append(
        (
            "正文样式字体(宋体+Times New Roman, 表7)",
            "PASS" if normal_font_ok else "WARN",
            "eastAsia={} ascii={} hAnsi={}".format(
                normal.get("font_eastAsia", ""),
                normal.get("font_ascii", ""),
                normal.get("font_hAnsi", ""),
            ),
        )
    )

    normal_para_ok = normal.get("line") == "300" and (
        normal.get("firstLineChars") == "200" or normal.get("firstLine") == "200"
    )
    checks.append(
        (
            "正文样式行距/首行缩进(1.25倍/两字符, 表7)",
            "PASS" if normal_para_ok else "WARN",
            (
                f"line={normal.get('line', '')}, firstLine={normal.get('firstLine', '')}, "
                f"firstLineChars={normal.get('firstLineChars', '')}"
            ),
        )
    )

    h4 = style_name_to_id.get("heading 4", "4")
    h4_count = int(usage_body.get(h4, 0))
    checks.append(("正文是否超过三级标题(0083)", "PASS" if h4_count == 0 else "WARN", f"Heading4 count={h4_count}"))

    toc_level_over_count = 0
    toc_entry_count = 0
    for sid, cnt in usage_body.items():
        if cnt <= 0:
            continue
        level = _style_level_from_name(style_name_by_id.get(sid, sid), "toc")
        if level is None:
            continue
        toc_entry_count += int(cnt)
        if level > 3:
            toc_level_over_count += int(cnt)
    if toc_entry_count == 0:
        checks.append(("目录是否超过三级标题(0036)", "WARN", "未检测到目录域，建议在 Word 中更新目录后复检"))
    else:
        checks.append(
            (
                "目录是否超过三级标题(0036)",
                "PASS" if toc_level_over_count == 0 else "WARN",
                f"目录项 {toc_entry_count} 条，超三级 {toc_level_over_count} 条",
            )
        )

    numbered_over3 = 0
    numbered_pat = re.compile(r"^\d+(?:\.\d+){3,}(?:\s|$)")
    for sid, text in body_paragraphs:
        level = _style_level_from_name(style_name_by_id.get(sid, sid), "heading")
        if level is None:
            continue
        if numbered_pat.match(text.strip()):
            numbered_over3 += 1
    checks.append(
        (
            "标题编号是否超过三级(0083)",
            "PASS" if numbered_over3 == 0 else "WARN",
            f"检测到四级及以上编号标题 {numbered_over3} 处",
        )
    )

    zh_title = _extract_labeled_value(body_paragraphs, ["中文题目", "论文题目"])
    if zh_title:
        zh_len = _count_nonspace_chars(zh_title)
        checks.append(("中文题目长度(≤25字/符, 0016)", "PASS" if zh_len <= 25 else "WARN", f"当前 {zh_len}"))
    else:
        checks.append(("中文题目长度(≤25字/符, 0016)", "WARN", "未自动识别“中文题目”字段"))

    en_title = _extract_labeled_value(body_paragraphs, ["英文题目"])
    if en_title:
        en_letters = _count_ascii_letters(en_title)
        checks.append(("英文题目长度(≤150字母, 0016)", "PASS" if en_letters <= 150 else "WARN", f"当前约 {en_letters}"))
    else:
        checks.append(("英文题目长度(≤150字母, 0016)", "WARN", "未自动识别“英文题目”字段"))

    norm_texts = [_norm(t) for _, t in body_paragraphs]
    required = ["摘要", "Abstract", "参考文献", "致谢"]
    for sec in required:
        has = any(t == _norm(sec) or sec in t for t in norm_texts)
        checks.append((f"章节存在：{sec}", "PASS" if has else "WARN", "已检测" if has else "未检测到"))

    has_toc = any(t == "目录" for t in norm_texts) or toc_entry_count > 0
    toc_detail = "已检测到目录标题/目录域" if has_toc else "建议在 Word 中更新目录域"
    checks.append(("章节存在：目录", "PASS" if has_toc else "WARN", toc_detail))

    has_lot = any(t in {"图目录", "表目录", "图表目录"} for t in norm_texts)
    checks.append(("章节存在：图表目录", "PASS" if has_lot else "WARN", "已检测" if has_lot else "建议在 Word 中补图表目录"))
    has_figure_toc = any(t == "图目录" for t in norm_texts)
    has_table_toc = any(t == "表目录" for t in norm_texts)
    toc_bundle_ok = has_toc and has_figure_toc and has_table_toc
    missing_toc_parts: List[str] = []
    if not has_toc:
        missing_toc_parts.append("目录")
    if not has_figure_toc:
        missing_toc_parts.append("图目录")
    if not has_table_toc:
        missing_toc_parts.append("表目录")
    checks.append(
        (
            "模板结构：目录/图目录/表目录标题存在",
            "PASS" if toc_bundle_ok else "WARN",
            "已检测" if toc_bundle_ok else f"缺失：{', '.join(missing_toc_parts)}",
        )
    )

    bib_layout_status = "WARN"
    bib_layout_detail = "未定位到参考文献正文段落"
    bib_heading_seen = False
    for item in body_paragraph_details:
        text_norm = _norm(item.get("text", ""))
        if text_norm == _norm("参考文献"):
            bib_heading_seen = True
            continue
        if not bib_heading_seen:
            continue
        if text_norm in {
            _norm("附录"),
            _norm("致谢"),
            _norm("作者简历及攻读学位期间发表的学术论文与其他相关学术成果"),
        }:
            break
        if not item.get("text", "").strip():
            continue

        left = item.get("left", "")
        hanging = item.get("hanging", "")
        hanging_chars = item.get("hangingChars", "")
        first_line = item.get("firstLine", "")
        first_line_chars = item.get("firstLineChars", "")
        jc = item.get("jc", "")
        ok = (
            left == "480"
            and hanging == "480"
            and hanging_chars == "200"
            and first_line in {"", "0"}
            and first_line_chars in {"", "0"}
            and jc == "both"
        )
        bib_layout_status = "PASS" if ok else "WARN"
        bib_layout_detail = (
            f"left={left or 'NA'} hanging={hanging or 'NA'} "
            f"hangingChars={hanging_chars or 'NA'} firstLine={first_line or 'NA'} "
            f"firstLineChars={first_line_chars or 'NA'} jc={jc or 'NA'}"
        )
        break
    checks.append(("参考文献段落格式(两端对齐/两汉字悬挂缩进)", bib_layout_status, bib_layout_detail))

    if image_paragraph_details:
        bad_image_indent = [
            item
            for item in image_paragraph_details
            if item.get("firstLine", "") not in {"", "0"} or item.get("firstLineChars", "") not in {"", "0"}
        ]
        if bad_image_indent:
            sample = bad_image_indent[0]
            image_indent_detail = (
                f"total={len(image_paragraph_details)} bad={len(bad_image_indent)} "
                f"sample:firstLine={sample.get('firstLine') or 'NA'} "
                f"firstLineChars={sample.get('firstLineChars') or 'NA'} "
                f"jc={sample.get('jc') or 'NA'}"
            )
        else:
            image_indent_detail = f"total={len(image_paragraph_details)} bad=0"
        checks.append(
            (
                "图片段落首行缩进(应为0)",
                "PASS" if not bad_image_indent else "WARN",
                image_indent_detail,
            )
        )
    else:
        checks.append(("图片段落首行缩进(应为0)", "WARN", "未检测到图片段落"))

    zh_kw_count = -1
    en_kw_count = -1
    zh_abstract_chars = -1
    try:
        idx_zh = next(i for i, t in enumerate(norm_texts) if t == "摘要")
        idx_en = next(i for i, t in enumerate(norm_texts[idx_zh + 1 :], start=idx_zh + 1) if t.lower() == "abstract")
        zh_lines: List[str] = []
        for _, t in body_paragraphs[idx_zh + 1 : idx_en]:
            nt = _norm(t)
            nt_cf = nt.casefold()
            if nt.startswith("关键词") or nt_cf.startswith("keywords") or nt_cf.startswith("keyword"):
                continue
            if nt:
                zh_lines.append(t.strip())
        zh_abstract_chars = len(re.sub(r"\s+", "", "".join(zh_lines)))
    except Exception:
        pass

    for i, (_, t) in enumerate(body_paragraphs):
        nt = _norm(t)
        if nt.startswith("关键词"):
            zh_kw_count = _count_keywords(t)
            if zh_kw_count <= 1 and i + 1 < len(body_paragraphs):
                ntext = body_paragraphs[i + 1][1].strip()
                if ntext:
                    zh_kw_count = _count_keywords(ntext)
        nt_cf = nt.casefold()
        if nt_cf.startswith("keywords") or nt_cf.startswith("keyword"):
            en_kw_count = _count_keywords(t)
            if en_kw_count <= 1 and i + 1 < len(body_paragraphs):
                ntext = body_paragraphs[i + 1][1].strip()
                if ntext:
                    en_kw_count = _count_keywords(ntext)

    if zh_abstract_chars >= 0:
        ok = 1500 <= zh_abstract_chars <= 3000
        checks.append(("中文摘要字数(1500-3000, 0032)", "PASS" if ok else "WARN", f"当前约 {zh_abstract_chars} 字"))
    else:
        checks.append(("中文摘要字数(1500-3000, 0032)", "WARN", "未能自动识别摘要边界"))

    if zh_kw_count >= 0:
        ok = 3 <= zh_kw_count <= 6
        checks.append(("中文关键词数量(3-6, 0033)", "PASS" if ok else "WARN", f"当前 {zh_kw_count} 个"))
    else:
        checks.append(("中文关键词数量(3-6, 0033)", "WARN", "未检测到“关键词”行"))

    if en_kw_count >= 0:
        ok = 3 <= en_kw_count <= 6
        checks.append(("英文关键词数量(3-6, 0033)", "PASS" if ok else "WARN", f"当前 {en_kw_count} 个"))
    else:
        checks.append(("英文关键词数量(3-6, 0033)", "WARN", "未检测到“Key Words/Keywords”行"))

    if strict_scan is not None:
        bare_env = int(strict_scan.get("bare_list_env_lines", 0))
        abnormal_dollar = int(strict_scan.get("abnormal_double_dollar_lines", 0))
        duplicate_placeholder = int(strict_scan.get("duplicate_bibliography_placeholder", 0))
        omml_nodes = int(strict_scan.get("omml_oMath_nodes", -1))
        omml_gate = strict_scan.get("omml_gate_max")
        checks.append(
            (
                "内容伪影：裸 itemize/enumerate(strict=0)",
                "PASS" if bare_env == 0 else "WARN",
                f"当前 {bare_env}",
            )
        )
        checks.append(
            (
                "内容伪影：异常 $$ 残留(strict=0)",
                "PASS" if abnormal_dollar == 0 else "WARN",
                f"当前 {abnormal_dollar}",
            )
        )
        checks.append(
            (
                "内容伪影：重复参考文献占位(strict=0)",
                "PASS" if duplicate_placeholder == 0 else "WARN",
                f"当前 {duplicate_placeholder}",
            )
        )
        if omml_nodes >= 0:
            status = "PASS"
            detail = f"m:oMath={omml_nodes}"
            if isinstance(omml_gate, int):
                status = "PASS" if omml_nodes <= omml_gate else "WARN"
                detail = f"m:oMath={omml_nodes} / gate={omml_gate}"
            checks.append(("OMML 数学节点门禁", status, detail))

    return checks


def _write_quality_report(
    report_path: Path,
    docx_path: Path,
    mode: str,
    remap: Dict[str, str],
    py_changed: int,
    python_docx_available: bool,
    xml_changed: Counter[str],
    before: Dict[str, object],
    after: Dict[str, object],
    strict_scan: Optional[Dict[str, int]] = None,
    section_sync: Optional[Dict[str, object]] = None,
    front_matter_sync: Optional[Dict[str, object]] = None,
) -> None:
    before_unknown = before["unknown"]  # type: ignore[assignment]
    after_unknown = after["unknown"]  # type: ignore[assignment]
    before_usage = before["usage"]  # type: ignore[assignment]
    after_usage = after["usage"]  # type: ignore[assignment]
    style_name_by_id = after["style_name_by_id"]  # type: ignore[assignment]
    layout = after["layout"]  # type: ignore[assignment]
    style_name_to_id = after["style_name_to_id"]  # type: ignore[assignment]
    before_section_count = int(before.get("section_count", 0))  # type: ignore[arg-type]
    after_section_count = int(after.get("section_count", 0))  # type: ignore[arg-type]

    h1 = style_name_to_id.get("heading 1", "1")
    h2 = style_name_to_id.get("heading 2", "2")
    h3 = style_name_to_id.get("heading 3", "3")
    h4 = style_name_to_id.get("heading 4", "4")

    lines = [
        "# DOCX 质量报告",
        "",
        f"- 时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"- 导出模式：`{mode}`",
        f"- 文件：`{docx_path}`",
        "",
        "## 样式修正摘要",
        "",
        f"- python-docx 可用：`{'yes' if python_docx_available else 'no'}`",
        f"- python-docx 直接修正段落数：`{py_changed}`",
        f"- XML 全局修正段落数：`{sum(xml_changed.values())}`",
        f"- 修正规则数：`{len(remap)}`",
        f"- 分节数量（修正前/后）：`{before_section_count} / {after_section_count}`",
        "",
        "## 未知样式对比",
        "",
        f"- 修正前未知样式引用：`{sum(before_unknown.values())}`",
    ]

    if section_sync is not None:
        applied = bool(section_sync.get("applied"))
        template_section_count = section_sync.get("template_section_count", "NA")
        lines.extend(
            [
                f"- 模板分节骨架同步：`{'yes' if applied else 'no'}`",
                f"- 参考模板分节数量：`{template_section_count}`",
            ]
        )
    if front_matter_sync is not None:
        fm_applied = bool(front_matter_sync.get("applied"))
        lines.append(f"- 模板首页前置内容同步：`{'yes' if fm_applied else 'no'}`")

    for sid, cnt in before_unknown.most_common(20):
        lines.append(f"  - `{sid}`: {cnt}")

    lines.append(f"- 修正后未知样式引用：`{sum(after_unknown.values())}`")
    for sid, cnt in after_unknown.most_common(20):
        lines.append(f"  - `{sid}`: {cnt}")

    lines.extend(
        [
            "",
            "## 标题层级统计",
            "",
            f"- Heading 1（`{h1}`）：`{after_usage.get(h1, 0)}`",
            f"- Heading 2（`{h2}`）：`{after_usage.get(h2, 0)}`",
            f"- Heading 3（`{h3}`）：`{after_usage.get(h3, 0)}`",
            f"- Heading 4（`{h4}`）：`{after_usage.get(h4, 0)}`",
            "",
            "## 版面参数（首个分节）",
            "",
            f"- 纸张（twip）：`{layout.get('page_w_twip', 'NA')} x {layout.get('page_h_twip', 'NA')}`",
            f"- 纸张（cm）：`{_twip_to_cm(layout.get('page_w_twip', ''))} x {_twip_to_cm(layout.get('page_h_twip', ''))}`",
            f"- 页边距上/下（twip）：`{layout.get('margin_top_twip', 'NA')} / {layout.get('margin_bottom_twip', 'NA')}`",
            f"- 页边距左/右（twip）：`{layout.get('margin_left_twip', 'NA')} / {layout.get('margin_right_twip', 'NA')}`",
            "",
            "## 样式分布（Top 12）",
            "",
            "| 样式ID | 样式名 | 修正前 | 修正后 |",
            "|---|---|---:|---:|",
        ]
    )

    top_ids = [sid for sid, _ in after_usage.most_common(12)]
    for sid, _ in before_usage.most_common(12):
        if sid not in top_ids:
            top_ids.append(sid)
        if len(top_ids) >= 12:
            break

    for sid in top_ids[:12]:
        lines.append(
            "| `{sid}` | {name} | {b} | {a} |".format(
                sid=sid,
                name=style_name_by_id.get(sid, sid),
                b=before_usage.get(sid, 0),
                a=after_usage.get(sid, 0),
            )
        )

    checks = _spec_checks(after, strict_scan=strict_scan)
    lines.extend(
        [
            "",
            "## 资环规范一致性检查（自动）",
            "",
            "- 依据：`中国科学院大学资源与环境学位评定分委员会研究生学位论文撰写具体要`（2023）",
            "",
        ]
    )
    lines.append("| 检查项 | 结果 | 说明 |")
    lines.append("|---|---|---|")
    for item, status, detail in checks:
        lines.append(f"| {item} | {status} | {detail} |")

    lines.extend(
        [
            "",
            "## 后处理说明",
            "",
            (
                "- 当前环境已启用 `python-docx`，会先做段落级样式修正，"
                "再做 XML 全局样式归一化。"
                if python_docx_available
                else "- 当前环境未启用 `python-docx`，本次仅执行 XML 全局样式归一化。"
            ),
        ]
    )

    lines.extend(["", "## 样式映射规则", ""])
    for src, dst in remap.items():
        lines.append(f"- `{src}` -> `{dst}`")

    report_path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def normalize_docx_styles(
    docx_path: Path,
    report_path: Path,
    mode: str = "portable",
    strict_scan: Optional[Dict[str, int]] = None,
    section_sync: Optional[Dict[str, object]] = None,
    front_matter_sync: Optional[Dict[str, object]] = None,
) -> None:
    before = _analyze_docx_styles(docx_path)
    remap = _build_style_remap(before)
    py_changed, python_docx_available = _normalize_paragraph_styles_with_python_docx(docx_path, remap)
    xml_changed = _rewrite_docx_styles_xml(docx_path, remap)
    after = _analyze_docx_styles(docx_path)
    _write_quality_report(
        report_path,
        docx_path,
        mode,
        remap,
        py_changed,
        python_docx_available,
        xml_changed,
        before,
        after,
        strict_scan=strict_scan,
        section_sync=section_sync,
        front_matter_sync=front_matter_sync,
    )


def main() -> int:
    args = parse_args()
    project_dir = args.project_dir.resolve()
    reference_doc = discover_reference_doc(project_dir, args.reference_doc)
    if args.word_visible and not args.word_update_fields:
        print("[WARN] --word-visible 仅在 --word-update-fields 启用时生效，当前将忽略")

    if args.postprocess_only_docx is not None:
        if args.prepare_tex:
            print("[WARN] --prepare-tex 与 --postprocess-only-docx 同时出现时将忽略前处理。")
        if args.ensure_page_cache:
            print("[WARN] --ensure-page-cache 与 --postprocess-only-docx 同时出现时将忽略页码缓存构建。")
        docx_path = args.postprocess_only_docx.resolve()
        report_path = (
            args.quality_report.resolve()
            if args.quality_report
            else docx_path.with_name(f"{docx_path.stem}_质量报告.md")
        )
        _run_postprocess_preflight(
            args.mode,
            docx_path,
            report_path,
            word_update_fields=args.word_update_fields,
        )
        bibliography = list((project_dir / "bibs").glob("*.bib"))
        citation_locale_fix = _localize_chinese_intext_citations(docx_path, bibliography)
        if citation_locale_fix.get("applied"):
            print(
                "[OK] chinese citation locale fix: replacements={count} paragraphs={paras}".format(
                    count=citation_locale_fix.get("replacements", 0),
                    paras=citation_locale_fix.get("updated_paragraphs", 0),
                )
            )
        english_citation_fix = _localize_english_intext_citations(docx_path)
        if english_citation_fix.get("applied"):
            print(
                "[OK] english citation locale fix: replacements={count} paragraphs={paras}".format(
                    count=english_citation_fix.get("replacements", 0),
                    paras=english_citation_fix.get("updated_paragraphs", 0),
                )
            )
        bibliography_locale_fix = _localize_chinese_bibliography_author_terms(docx_path)
        if bibliography_locale_fix.get("applied"):
            print(
                "[OK] chinese bibliography locale fix: replacements={count} paragraphs={paras}".format(
                    count=bibliography_locale_fix.get("replacements", 0),
                    paras=bibliography_locale_fix.get("updated_paragraphs", 0),
                )
            )
        bibliography_entry_fix = _normalize_chinese_bibliography_entries(docx_path)
        if bibliography_entry_fix.get("applied"):
            print(
                "[OK] chinese bibliography entry fix: replacements={count} paragraphs={paras}".format(
                    count=bibliography_entry_fix.get("replacements", 0),
                    paras=bibliography_entry_fix.get("updated_paragraphs", 0),
                )
            )
        english_bibliography_fix = _localize_english_bibliography_terms(docx_path)
        if english_bibliography_fix.get("applied"):
            print(
                "[OK] english bibliography locale fix: replacements={count} paragraphs={paras}".format(
                    count=english_bibliography_fix.get("replacements", 0),
                    paras=english_bibliography_fix.get("updated_paragraphs", 0),
                )
            )
        bookmark_cleanup = _cleanup_unmatched_bookmarks(docx_path)
        if bookmark_cleanup["applied"]:
            print(
                "[OK] bookmark cleanup: removed_starts={starts} removed_ends={ends}".format(
                    starts=bookmark_cleanup["removed_starts"],
                    ends=bookmark_cleanup["removed_ends"],
                )
            )
        figure_caption_fix = _rewrite_figure_caption_placeholders(docx_path)
        if figure_caption_fix.get("applied"):
            print(
                "[OK] figure caption fix: updated={count} style={style}".format(
                    count=figure_caption_fix.get("updated", 0),
                    style=figure_caption_fix.get("caption_style_id", "NA"),
                )
            )
        image_fit_fix = _fit_oversized_images_to_page(docx_path)
        if image_fit_fix.get("applied"):
            print(
                "[OK] oversized image fit fix: updated={count}".format(
                    count=image_fit_fix.get("updated", 0),
                )
            )
        technical_route_landscape_fix = _apply_landscape_section_to_technical_route_figure(docx_path)
        if technical_route_landscape_fix.get("applied"):
            print(
                "[OK] technical route landscape fix: updated={count}".format(
                    count=technical_route_landscape_fix.get("updated", 0),
                )
            )
        image_paragraph_fix = _normalize_image_paragraph_indent(docx_path)
        if image_paragraph_fix.get("applied"):
            print(
                "[OK] image paragraph layout fix: updated={count} indent={indent} spacing={spacing} center={center}".format(
                    count=image_paragraph_fix.get("updated", 0),
                    indent=image_paragraph_fix.get("indent_updated", 0),
                    spacing=image_paragraph_fix.get("spacing_updated", 0),
                    center=image_paragraph_fix.get("center_updated", 0),
                )
            )
        three_line_table_fix = _apply_three_line_table_format(docx_path)
        if three_line_table_fix.get("applied"):
            print(
                "[OK] three-line table fix: updated={count} layout={layout} wide_layout={wide_layout} landscape_sections={landscape_sections} pagination={pagination}".format(
                    count=three_line_table_fix.get("updated", 0),
                    layout=three_line_table_fix.get("layout_updated", 0),
                    wide_layout=three_line_table_fix.get("wide_layout_updated", 0),
                    landscape_sections=three_line_table_fix.get("landscape_section_updated", 0),
                    pagination=three_line_table_fix.get("pagination_updated", 0),
                )
            )
        abstract_keyword_fix = _normalize_abstract_keyword_paragraphs(docx_path)
        if abstract_keyword_fix.get("applied"):
            print(
                "[OK] abstract keyword fix: updated={count}".format(
                    count=abstract_keyword_fix.get("updated", 0),
                )
            )
        cv_subheading_bold_fix = _force_cv_subheading_bold(docx_path)
        if cv_subheading_bold_fix.get("applied"):
            print(
                "[OK] cv subheading bold fix: updated={count}".format(
                    count=cv_subheading_bold_fix.get("updated", 0),
                )
            )
        cjk_quote_font_fix = _force_cjk_quote_runs_to_eastasia_font(docx_path)
        if cjk_quote_font_fix.get("applied"):
            print(
                "[OK] cjk quote font fix: updated_runs={count}".format(
                    count=cjk_quote_font_fix.get("updated_runs", 0),
                )
            )
        bibliography_layout_fix = _normalize_bibliography_paragraph_layout(docx_path)
        if bibliography_layout_fix.get("applied"):
            print(
                "[OK] bibliography layout fix: updated={count} italic_cleared_runs={italic}".format(
                    count=bibliography_layout_fix.get("updated", 0),
                    italic=bibliography_layout_fix.get("italic_cleared_runs", 0),
                )
            )
        body_list_indent_fix = _normalize_body_list_paragraph_indent(docx_path)
        if body_list_indent_fix.get("applied"):
            print(
                "[OK] body list indent fix: updated={count}".format(
                    count=body_list_indent_fix.get("updated", 0),
                )
            )
        tc_rebuild = _rebuild_figure_table_tc_fields(docx_path)
        if tc_rebuild.get("applied"):
            print(
                "[OK] figure/table tc rebuild: entries={entries} removed_old={removed} figs={figs} tbls={tbls}".format(
                    entries=tc_rebuild.get("added", 0),
                    removed=tc_rebuild.get("removed", 0),
                    figs=tc_rebuild.get("figure_entries", 0),
                    tbls=tc_rebuild.get("table_entries", 0),
                )
            )
        figure_table_toc_fix = _ensure_figure_table_toc_fields(docx_path)
        if figure_table_toc_fix.get("applied"):
            print(
                "[OK] figure/table toc field fix: inserted={inserted} removed_placeholder={removed}".format(
                    inserted=figure_table_toc_fix.get("inserted", 0),
                    removed=figure_table_toc_fix.get("removed_placeholder", 0),
                )
            )
        toc_style_align = _align_figure_table_style_to_toc1(docx_path)
        if toc_style_align.get("applied"):
            print("[OK] figure/table toc style fix: aligned_to=TOC1")
        cover_metadata_sync = _sync_cover_metadata_from_info(docx_path, project_dir)
        print(
            "[OK] cover metadata sync: updated={count} spine_updated={spine}".format(
                count=cover_metadata_sync.get("updated", 0),
                spine=cover_metadata_sync.get("spine_updated", 0),
            )
        )
        if _should_skip_header_rule_sync_after_word_update(docx_path, word_update_fields=True):
            print("[WARN] header rule sync skipped: existing Word TOC field results detected")
        else:
            header_rule_fix = _normalize_docx_headers_by_rules(docx_path, project_dir)
            if header_rule_fix.get("applied"):
                print(
                    "[OK] header rule sync: updated_sections={count} page_numbering_updates={page_updates}".format(
                        count=header_rule_fix.get("updated_sections", 0),
                        page_updates=header_rule_fix.get("page_numbering_updates", 0),
                    )
                )
            else:
                print(
                    "[WARN] header rule sync skipped: {reason}".format(
                        reason=header_rule_fix.get("reason", "no eligible section found"),
                    )
                )
        hint_cleanup = _strip_front_matter_hint_text(docx_path)
        print(
            "[OK] front matter hint cleanup: removed={count}".format(
                count=hint_cleanup.get("removed", 0),
            )
        )
        annotation_cleanup = _remove_front_matter_annotation_shapes(docx_path)
        print(
            "[OK] front matter annotation cleanup: removed={count}".format(
                count=annotation_cleanup.get("removed", 0),
            )
        )
        artifact_cleanup = _cleanup_known_word_artifacts(docx_path)
        if artifact_cleanup.get("applied"):
            print(
                "[OK] word artifact cleanup: macro_runs={runs} macro_paragraphs={paras} pt_paragraphs={pt}".format(
                    runs=artifact_cleanup.get("removed_macro_runs", 0),
                    paras=artifact_cleanup.get("removed_macro_paragraphs", 0),
                    pt=artifact_cleanup.get("removed_pt_paragraphs", 0),
                )
            )
        normalize_docx_styles(docx_path, report_path, mode=args.mode, strict_scan=None)
        backmatter_fix_post = _strip_backmatter_heading_numbering(docx_path, reference_doc)
        if backmatter_fix_post.get("applied"):
            print(
                "[OK] backmatter heading fix (post-normalize): removed_chapter_markers={count} demoted_subheadings={demoted} normalized_titles={normalized}".format(
                    count=len(backmatter_fix_post.get("fixed", [])),
                    demoted=len(backmatter_fix_post.get("demoted", [])),
                    normalized=len(backmatter_fix_post.get("normalized", [])),
                )
            )
        appendix_fix_post = _strip_appendix_heading_numbering(docx_path)
        if appendix_fix_post.get("applied"):
            print(
                "[OK] appendix heading fix: normalized_appendix_heading_numbering={count}".format(
                    count=len(appendix_fix_post.get("fixed", [])),
                )
            )
        typography_sync = _sync_project_typography_overrides(docx_path)
        if typography_sync.get("applied"):
            print(
                "[OK] typography sync: captions={captions} notes={notes} keywords={keywords} appendix_note={appendix}".format(
                    captions=typography_sync.get("caption_updated", 0),
                    notes=typography_sync.get("note_updated", 0),
                    keywords=typography_sync.get("keyword_updated", 0),
                    appendix=typography_sync.get("appendix_note_updated", 0),
                )
            )
        body_first_line_fix = _normalize_body_first_line_indent(docx_path)
        if body_first_line_fix.get("applied"):
            print(
                "[OK] body first-line indent fix: updated={count} styles={styles}".format(
                    count=body_first_line_fix.get("updated", 0),
                    styles=body_first_line_fix.get("style_updated", 0),
                )
            )
        latex_tilde_fix = _normalize_visible_latex_tilde_spacing(docx_path)
        if latex_tilde_fix.get("applied"):
            print(
                "[OK] latex tilde spacing fix: updated={count}".format(
                    count=latex_tilde_fix.get("updated", 0),
                )
            )
        math_font_fix = _ensure_docx_math_font(docx_path)
        if math_font_fix.get("applied"):
            print(
                "[OK] math font fix: font={font}".format(
                    font=math_font_fix.get("font", "Times New Roman"),
                )
            )
        plain_script_fix = _repair_plain_script_runs_in_docx(docx_path)
        if plain_script_fix.get("applied"):
            print(
                "[OK] plain script visual fix: updated_paragraphs={count}".format(
                    count=plain_script_fix.get("updated", 0),
                )
            )
        if args.word_update_fields:
            word_update = update_docx_fields_with_word(
                docx_path,
                visible=args.word_visible,
                timeout_seconds=args.word_update_timeout,
            )
            if not word_update.get("applied"):
                raise RuntimeError(f"Word 自动更新域失败：{word_update.get('reason', 'unknown error')}")
            print(
                "[OK] word field update: updated_all_fields=1 powershell={ps}".format(
                    ps=word_update.get("powershell", "NA"),
                )
            )
        toc_fallback = _refresh_figure_table_catalog_fallback(
            docx_path,
            force=not args.word_update_fields,
            project_dir=project_dir,
        )
        if toc_fallback.get("applied"):
            print(
                "[OK] figure/table toc materialize: sections={sections} figs={figs} tbls={tbls} replaced_paragraphs={paras} mode={mode}".format(
                    sections=toc_fallback.get("sections", 0),
                    figs=toc_fallback.get("figure_entries", 0),
                    tbls=toc_fallback.get("table_entries", 0),
                    paras=toc_fallback.get("replaced_paragraphs", 0),
                    mode="force-static" if not args.word_update_fields else "fallback",
                )
            )
            bookmark_cleanup_after = _cleanup_unmatched_bookmarks(docx_path)
            if bookmark_cleanup_after["applied"]:
                print(
                    "[OK] bookmark cleanup (post toc materialize): removed_starts={starts} removed_ends={ends}".format(
                        starts=bookmark_cleanup_after["removed_starts"],
                        ends=bookmark_cleanup_after["removed_ends"],
                    )
                )
        toc_spacing_fix = _normalize_toc_and_catalog_spacing(docx_path)
        if toc_spacing_fix.get("applied"):
            print(
                "[OK] toc spacing normalize: changed_paragraphs={count}".format(
                    count=toc_spacing_fix.get("changed", 0),
                )
            )
        if _should_skip_header_rule_sync_after_word_update(docx_path, word_update_fields=True):
            print("[WARN] header rule sync (final) skipped: existing Word TOC field results detected")
        else:
            header_rule_fix_final = _normalize_docx_headers_by_rules(docx_path, project_dir)
            if header_rule_fix_final.get("applied"):
                print(
                    "[OK] header rule sync (final): updated_sections={count} page_numbering_updates={page_updates}".format(
                        count=header_rule_fix_final.get("updated_sections", 0),
                        page_updates=header_rule_fix_final.get("page_numbering_updates", 0),
                    )
                )
        header_paragraph_layout_fix = _normalize_header_paragraph_layout(docx_path)
        if header_paragraph_layout_fix.get("applied"):
            print(
                "[OK] header paragraph layout fix: updated_parts={parts} updated_paragraphs={paras} cleared_indent_attrs={attrs}".format(
                    parts=header_paragraph_layout_fix.get("updated_parts", 0),
                    paras=header_paragraph_layout_fix.get("updated_paragraphs", 0),
                    attrs=header_paragraph_layout_fix.get("cleared_indent_attrs", 0),
                )
            )
        non_cjk_font_fix = _force_non_cjk_runs_to_times_new_roman(docx_path)
        if non_cjk_font_fix.get("applied"):
            print(
                "[OK] non-cjk font fix: font={font} updated_runs={runs} updated_parts={parts}".format(
                    font=non_cjk_font_fix.get("font", "Times New Roman"),
                    runs=non_cjk_font_fix.get("updated_runs", 0),
                    parts=non_cjk_font_fix.get("updated_parts", 0),
                )
            )
        if args.word_update_fields:
            word_repair = repair_docx_with_word_open_and_repair(
                docx_path,
                visible=args.word_visible,
                timeout_seconds=args.word_update_timeout,
            )
            if not word_repair.get("applied"):
                raise RuntimeError(f"Word OpenAndRepair 归一化失败：{word_repair.get('reason', 'unknown error')}")
            print(
                "[OK] word open-and-repair normalize: powershell={ps}".format(
                    ps=word_repair.get("powershell", "NA"),
                )
            )
        package_check = _validate_docx_package_integrity(docx_path)
        if not package_check["ok"]:
            details: List[str] = []
            for key in [
                "missing_entries",
                "xml_parse_failures",
                "missing_rel_targets",
                "missing_override_parts",
                "missing_rid_refs",
                "unmatched_bookmarks",
                "invalid_mc_ignorable",
            ]:
                items = package_check.get(key, [])
                if items:
                    details.append(f"{key}={len(items)}")
            detail_text = ", ".join(details) if details else "unknown package issue"
            if args.mode == "strict":
                raise RuntimeError(f"DOCX 结构完整性校验失败：{detail_text}")
            print(f"[WARN] package integrity: {detail_text}")
        else:
            print("[OK] package integrity: rel/content-types/xml checks passed")
        omml_gate_result = _enforce_omml_gate(docx_path, args.omml_gate_max)
        print(
            "[OK] omml scan: m:oMath={omath} m:oMathPara={opara} gate={gate}".format(
                omath=omml_gate_result.get("oMath", 0),
                opara=omml_gate_result.get("oMathPara", 0),
                gate=omml_gate_result.get("gate_max", "disabled"),
            )
        )
        print(f"[OK] postprocess docx: {docx_path}")
        print(f"[OK] quality report: {report_path}")
        return 0

    tex_file = (project_dir / args.tex_file).resolve()

    if not tex_file.exists():
        raise FileNotFoundError(f"主文件不存在: {tex_file}")
    tex_stem = tex_file.stem
    output_docx = args.output.resolve() if args.output else (project_dir / f"{tex_stem}_from_tex_资环模板.docx")
    markdown_path = (
        args.markdown_output.resolve()
        if args.markdown_output
        else (project_dir / f"{tex_stem}_from_tex_word_source.md")
    )

    markdown_path.parent.mkdir(parents=True, exist_ok=True)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    report_path = (
        args.quality_report.resolve()
        if args.quality_report
        else output_docx.with_name(f"{output_docx.stem}_质量报告.md")
    )

    _run_export_preflight(
        mode=args.mode,
        reference_doc=reference_doc,
        markdown_path=markdown_path,
        output_docx=output_docx,
        report_path=None if args.skip_style_normalization else report_path,
        word_update_fields=args.word_update_fields,
    )

    if args.prepare_tex:
        run_prepare_tex_for_export(project_dir)
    if args.ensure_page_cache:
        run_page_cache_build_for_word_export(project_dir)

    md_text = render_markdown(project_dir, tex_file)
    strict_scan = _scan_markdown_artifacts(md_text)
    markdown_path.write_text(md_text, encoding="utf-8")

    bibliography = list((project_dir / "bibs").glob("*.bib"))
    run_pandoc(markdown_path, output_docx, reference_doc, project_dir, bibliography)
    citation_locale_fix = _localize_chinese_intext_citations(output_docx, bibliography)
    if citation_locale_fix.get("applied"):
        print(
            "[OK] chinese citation locale fix: replacements={count} paragraphs={paras}".format(
                count=citation_locale_fix.get("replacements", 0),
                paras=citation_locale_fix.get("updated_paragraphs", 0),
            )
        )
    english_citation_fix = _localize_english_intext_citations(output_docx)
    if english_citation_fix.get("applied"):
        print(
            "[OK] english citation locale fix: replacements={count} paragraphs={paras}".format(
                count=english_citation_fix.get("replacements", 0),
                paras=english_citation_fix.get("updated_paragraphs", 0),
            )
        )
    bibliography_locale_fix = _localize_chinese_bibliography_author_terms(output_docx)
    if bibliography_locale_fix.get("applied"):
        print(
            "[OK] chinese bibliography locale fix: replacements={count} paragraphs={paras}".format(
                count=bibliography_locale_fix.get("replacements", 0),
                paras=bibliography_locale_fix.get("updated_paragraphs", 0),
            )
        )
    bibliography_entry_fix = _normalize_chinese_bibliography_entries(output_docx)
    if bibliography_entry_fix.get("applied"):
        print(
            "[OK] chinese bibliography entry fix: replacements={count} paragraphs={paras}".format(
                count=bibliography_entry_fix.get("replacements", 0),
                paras=bibliography_entry_fix.get("updated_paragraphs", 0),
            )
        )
    english_bibliography_fix = _localize_english_bibliography_terms(output_docx)
    if english_bibliography_fix.get("applied"):
        print(
            "[OK] english bibliography locale fix: replacements={count} paragraphs={paras}".format(
                count=english_bibliography_fix.get("replacements", 0),
                paras=english_bibliography_fix.get("updated_paragraphs", 0),
            )
        )
    section_sync = _sync_docx_section_layout(output_docx, reference_doc, mode=args.mode)
    front_matter_sync: Optional[Dict[str, object]] = None
    if args.mode == "strict":
        front_matter_sync = _merge_template_front_matter(output_docx, reference_doc, mode=args.mode)
        if front_matter_sync.get("applied"):
            print("[OK] front matter sync: preserve template cover/declaration pages")
    if section_sync.get("applied"):
        print(
            "[OK] section layout sync: sections={sections} template_sections={template}".format(
                sections=section_sync.get("section_count", "NA"),
                template=section_sync.get("template_section_count", "NA"),
            )
        )
    cover_metadata_sync = _sync_cover_metadata_from_info(output_docx, project_dir)
    print(
        "[OK] cover metadata sync: updated={count} spine_updated={spine}".format(
            count=cover_metadata_sync.get("updated", 0),
            spine=cover_metadata_sync.get("spine_updated", 0),
        )
    )
    header_rule_fix = _normalize_docx_headers_by_rules(output_docx, project_dir)
    if header_rule_fix.get("applied"):
        print(
            "[OK] header rule sync: updated_sections={count} page_numbering_updates={page_updates}".format(
                count=header_rule_fix.get("updated_sections", 0),
                page_updates=header_rule_fix.get("page_numbering_updates", 0),
            )
        )
    else:
        print(
            "[WARN] header rule sync skipped: {reason}".format(
                reason=header_rule_fix.get("reason", "no eligible section found"),
            )
        )
    hint_cleanup = _strip_front_matter_hint_text(output_docx)
    print(
        "[OK] front matter hint cleanup: removed={count}".format(
            count=hint_cleanup.get("removed", 0),
        )
    )
    annotation_cleanup = _remove_front_matter_annotation_shapes(output_docx)
    print(
        "[OK] front matter annotation cleanup: removed={count}".format(
            count=annotation_cleanup.get("removed", 0),
        )
    )
    artifact_cleanup = _cleanup_known_word_artifacts(output_docx)
    if artifact_cleanup.get("applied"):
        print(
            "[OK] word artifact cleanup: macro_runs={runs} macro_paragraphs={paras} pt_paragraphs={pt}".format(
                runs=artifact_cleanup.get("removed_macro_runs", 0),
                paras=artifact_cleanup.get("removed_macro_paragraphs", 0),
                pt=artifact_cleanup.get("removed_pt_paragraphs", 0),
            )
        )

    preface_fix = _fix_preface_heading_numbering(output_docx, reference_doc)
    if preface_fix.get("applied"):
        print(
            "[OK] preface heading fix: {count} heading(s) switched to unnumbered template layout".format(
                count=len(preface_fix.get("fixed", [])),
            )
        )

    preface_numbering_cleanup = _strip_preface_numbering_controls(output_docx)
    if preface_numbering_cleanup.get("applied"):
        print(
            "[OK] preface numbering cleanup: removed_numPr={numpr} removed_heading_styles={styles}".format(
                numpr=preface_numbering_cleanup.get("removed_numpr", 0),
                styles=preface_numbering_cleanup.get("removed_heading_styles", 0),
            )
        )

    heading_numbering_fix = _normalize_heading_multilevel_numbering(output_docx)
    if heading_numbering_fix.get("applied"):
        print(
            "[OK] heading numbering fix: normalized_starts={count} removed_restart_rules={restart}".format(
                count=heading_numbering_fix.get("updated_levels", 0),
                restart=heading_numbering_fix.get("removed_restart_rules", 0),
            )
        )

    heading_spacing_fix = _enforce_heading_number_spacing(output_docx)
    if heading_spacing_fix.get("applied"):
        print(
            "[OK] heading spacing fix: updated_formats={count}".format(
                count=heading_spacing_fix.get("updated_formats", 0),
            )
        )

    backmatter_fix = _strip_backmatter_heading_numbering(output_docx, reference_doc)
    if backmatter_fix.get("applied"):
        print(
            "[OK] backmatter heading fix: removed_chapter_markers={count} demoted_subheadings={demoted} normalized_titles={normalized}".format(
                count=len(backmatter_fix.get("fixed", [])),
                demoted=len(backmatter_fix.get("demoted", [])),
                normalized=len(backmatter_fix.get("normalized", [])),
            )
        )

    toc_field_fix = _ensure_main_toc_field(output_docx)
    if toc_field_fix.get("applied"):
        print(
            "[OK] toc field fix: inserted_main_toc=1 removed_placeholder={removed}".format(
                removed=toc_field_fix.get("removed_placeholder", 0),
            )
        )

    frontmatter_page_break_fix = _enforce_frontmatter_page_breaks(output_docx)
    if frontmatter_page_break_fix.get("applied"):
        print(
            "[OK] frontmatter page break fix: updated_headings={count}".format(
                count=len(frontmatter_page_break_fix.get("updated", [])),
            )
        )

    chapter_page_break_fix = _enforce_heading1_page_breaks(output_docx)
    if chapter_page_break_fix.get("applied"):
        print(
            "[OK] chapter page break fix: updated_headings={count}".format(
                count=len(chapter_page_break_fix.get("updated", [])),
            )
        )

    bookmark_cleanup = _cleanup_unmatched_bookmarks(output_docx)
    if bookmark_cleanup["applied"]:
        print(
            "[OK] bookmark cleanup: removed_starts={starts} removed_ends={ends}".format(
                starts=bookmark_cleanup["removed_starts"],
                ends=bookmark_cleanup["removed_ends"],
            )
        )

    figure_caption_fix = _rewrite_figure_caption_placeholders(output_docx)
    if figure_caption_fix.get("applied"):
        print(
            "[OK] figure caption fix: updated={count} style={style}".format(
                count=figure_caption_fix.get("updated", 0),
                style=figure_caption_fix.get("caption_style_id", "NA"),
            )
        )
    image_fit_fix = _fit_oversized_images_to_page(output_docx)
    if image_fit_fix.get("applied"):
        print(
            "[OK] oversized image fit fix: updated={count}".format(
                count=image_fit_fix.get("updated", 0),
            )
        )
    technical_route_landscape_fix = _apply_landscape_section_to_technical_route_figure(output_docx)
    if technical_route_landscape_fix.get("applied"):
        print(
            "[OK] technical route landscape fix: updated={count}".format(
                count=technical_route_landscape_fix.get("updated", 0),
            )
        )
    image_paragraph_fix = _normalize_image_paragraph_indent(output_docx)
    if image_paragraph_fix.get("applied"):
        print(
            "[OK] image paragraph layout fix: updated={count} indent={indent} spacing={spacing} center={center}".format(
                count=image_paragraph_fix.get("updated", 0),
                indent=image_paragraph_fix.get("indent_updated", 0),
                spacing=image_paragraph_fix.get("spacing_updated", 0),
                center=image_paragraph_fix.get("center_updated", 0),
            )
        )
    three_line_table_fix = _apply_three_line_table_format(output_docx)
    if three_line_table_fix.get("applied"):
        print(
            "[OK] three-line table fix: updated={count} layout={layout} wide_layout={wide_layout} landscape_sections={landscape_sections} pagination={pagination}".format(
                count=three_line_table_fix.get("updated", 0),
                layout=three_line_table_fix.get("layout_updated", 0),
                wide_layout=three_line_table_fix.get("wide_layout_updated", 0),
                landscape_sections=three_line_table_fix.get("landscape_section_updated", 0),
                pagination=three_line_table_fix.get("pagination_updated", 0),
            )
        )
    abstract_keyword_fix = _normalize_abstract_keyword_paragraphs(output_docx)
    if abstract_keyword_fix.get("applied"):
        print(
            "[OK] abstract keyword fix: updated={count}".format(
                count=abstract_keyword_fix.get("updated", 0),
            )
        )
    cv_subheading_bold_fix = _force_cv_subheading_bold(output_docx)
    if cv_subheading_bold_fix.get("applied"):
        print(
            "[OK] cv subheading bold fix: updated={count}".format(
                count=cv_subheading_bold_fix.get("updated", 0),
            )
        )
    cjk_quote_font_fix = _force_cjk_quote_runs_to_eastasia_font(output_docx)
    if cjk_quote_font_fix.get("applied"):
        print(
            "[OK] cjk quote font fix: updated_runs={count}".format(
                count=cjk_quote_font_fix.get("updated_runs", 0),
            )
        )
    bibliography_layout_fix = _normalize_bibliography_paragraph_layout(output_docx)
    if bibliography_layout_fix.get("applied"):
        print(
            "[OK] bibliography layout fix: updated={count} italic_cleared_runs={italic}".format(
                count=bibliography_layout_fix.get("updated", 0),
                italic=bibliography_layout_fix.get("italic_cleared_runs", 0),
            )
        )
    body_list_indent_fix = _normalize_body_list_paragraph_indent(output_docx)
    if body_list_indent_fix.get("applied"):
        print(
            "[OK] body list indent fix: updated={count}".format(
                count=body_list_indent_fix.get("updated", 0),
            )
        )
    tc_rebuild = _rebuild_figure_table_tc_fields(output_docx)
    if tc_rebuild.get("applied"):
        print(
            "[OK] figure/table tc rebuild: entries={entries} removed_old={removed} figs={figs} tbls={tbls}".format(
                entries=tc_rebuild.get("added", 0),
                removed=tc_rebuild.get("removed", 0),
                figs=tc_rebuild.get("figure_entries", 0),
                tbls=tc_rebuild.get("table_entries", 0),
            )
        )
    figure_table_toc_fix = _ensure_figure_table_toc_fields(output_docx)
    if figure_table_toc_fix.get("applied"):
        print(
            "[OK] figure/table toc field fix: inserted={inserted} removed_placeholder={removed}".format(
                inserted=figure_table_toc_fix.get("inserted", 0),
                removed=figure_table_toc_fix.get("removed_placeholder", 0),
            )
        )
    toc_style_align = _align_figure_table_style_to_toc1(output_docx)
    if toc_style_align.get("applied"):
        print("[OK] figure/table toc style fix: aligned_to=TOC1")
    omml_scan_for_report = _scan_docx_omml_nodes(output_docx)
    strict_scan["omml_oMath_nodes"] = int(omml_scan_for_report.get("oMath", 0))
    if args.omml_gate_max is not None:
        strict_scan["omml_gate_max"] = int(args.omml_gate_max)

    if not args.skip_style_normalization:
        normalize_docx_styles(
            output_docx,
            report_path,
            mode=args.mode,
            strict_scan=strict_scan,
            section_sync=section_sync,
            front_matter_sync=front_matter_sync,
        )
        print(f"[OK] quality report: {report_path}")
        backmatter_fix_post = _strip_backmatter_heading_numbering(output_docx, reference_doc)
        if backmatter_fix_post.get("applied"):
            print(
                "[OK] backmatter heading fix (post-normalize): removed_chapter_markers={count} demoted_subheadings={demoted} normalized_titles={normalized}".format(
                    count=len(backmatter_fix_post.get("fixed", [])),
                    demoted=len(backmatter_fix_post.get("demoted", [])),
                    normalized=len(backmatter_fix_post.get("normalized", [])),
                )
            )
        appendix_fix_post = _strip_appendix_heading_numbering(output_docx)
        if appendix_fix_post.get("applied"):
            print(
                "[OK] appendix heading fix: normalized_appendix_heading_numbering={count}".format(
                    count=len(appendix_fix_post.get("fixed", [])),
                )
            )
    typography_sync = _sync_project_typography_overrides(output_docx)
    if typography_sync.get("applied"):
        print(
            "[OK] typography sync: captions={captions} notes={notes} keywords={keywords} appendix_note={appendix}".format(
                captions=typography_sync.get("caption_updated", 0),
                notes=typography_sync.get("note_updated", 0),
                keywords=typography_sync.get("keyword_updated", 0),
                appendix=typography_sync.get("appendix_note_updated", 0),
            )
        )

    body_first_line_fix = _normalize_body_first_line_indent(output_docx)
    if body_first_line_fix.get("applied"):
        print(
            "[OK] body first-line indent fix: updated={count} styles={styles}".format(
                count=body_first_line_fix.get("updated", 0),
                styles=body_first_line_fix.get("style_updated", 0),
            )
        )
    latex_tilde_fix = _normalize_visible_latex_tilde_spacing(output_docx)
    if latex_tilde_fix.get("applied"):
        print(
            "[OK] latex tilde spacing fix: updated={count}".format(
                count=latex_tilde_fix.get("updated", 0),
            )
        )

    species_run_fix = _repair_species_token_runs_in_docx(output_docx)
    if species_run_fix.get("applied"):
        print(
            "[OK] species italics fix: updated_paragraphs={count}".format(
                count=species_run_fix.get("updated", 0),
            )
        )

    plain_script_fix = _repair_plain_script_runs_in_docx(output_docx)
    if plain_script_fix.get("applied"):
        print(
            "[OK] plain script visual fix: updated_paragraphs={count}".format(
                count=plain_script_fix.get("updated", 0),
            )
        )

    math_font_fix = _ensure_docx_math_font(output_docx)
    if math_font_fix.get("applied"):
        print(
            "[OK] math font fix: font={font}".format(
                font=math_font_fix.get("font", "Times New Roman"),
            )
        )

    first_main_odd_page_fix = _enforce_first_main_section_odd_page(output_docx)
    if first_main_odd_page_fix.get("applied"):
        print(
            "[OK] first main odd-page fix: previous_type={prev} new_type={new}".format(
                prev=first_main_odd_page_fix.get("previous_type", ""),
                new=first_main_odd_page_fix.get("new_type", ""),
            )
        )

    if args.word_update_fields:
        toc_fallback_pre_word = _materialize_figure_table_catalogs_before_word_update(
            output_docx,
            project_dir=project_dir,
        )
        if toc_fallback_pre_word.get("applied"):
            print(
                "[OK] figure/table toc materialize before Word update: sections={sections} figs={figs} tbls={tbls} replaced_paragraphs={paras} source={source}".format(
                    sections=toc_fallback_pre_word.get("sections", 0),
                    figs=toc_fallback_pre_word.get("figure_entries", 0),
                    tbls=toc_fallback_pre_word.get("table_entries", 0),
                    paras=toc_fallback_pre_word.get("replaced_paragraphs", 0),
                    source=toc_fallback_pre_word.get("entry_source", "unknown"),
                )
            )
        word_update = update_docx_fields_with_word(
            output_docx,
            visible=args.word_visible,
            timeout_seconds=args.word_update_timeout,
        )
        if not word_update.get("applied"):
            raise RuntimeError(f"Word 自动更新域失败：{word_update.get('reason', 'unknown error')}")
        print(
            "[OK] word field update: updated_all_fields=1 powershell={ps}".format(
                ps=word_update.get("powershell", "NA"),
            )
        )
    toc_fallback = _refresh_figure_table_catalog_fallback(
        output_docx,
        force=not args.word_update_fields,
        project_dir=project_dir,
    )
    if toc_fallback.get("applied"):
        print(
            "[OK] figure/table toc materialize: sections={sections} figs={figs} tbls={tbls} replaced_paragraphs={paras} mode={mode}".format(
                sections=toc_fallback.get("sections", 0),
                figs=toc_fallback.get("figure_entries", 0),
                tbls=toc_fallback.get("table_entries", 0),
                paras=toc_fallback.get("replaced_paragraphs", 0),
                mode="force-static" if not args.word_update_fields else "fallback",
            )
        )
        bookmark_cleanup_after = _cleanup_unmatched_bookmarks(output_docx)
        if bookmark_cleanup_after["applied"]:
            print(
                "[OK] bookmark cleanup (post toc materialize): removed_starts={starts} removed_ends={ends}".format(
                    starts=bookmark_cleanup_after["removed_starts"],
                    ends=bookmark_cleanup_after["removed_ends"],
                )
            )
    toc_spacing_fix = _normalize_toc_and_catalog_spacing(output_docx)
    if toc_spacing_fix.get("applied"):
        print(
            "[OK] toc spacing normalize: changed_paragraphs={count} style_runs={style}".format(
                count=toc_spacing_fix.get("changed", 0),
                style=toc_spacing_fix.get("style_changed", 0),
            )
        )
    if _should_skip_header_rule_sync_after_word_update(output_docx, word_update_fields=args.word_update_fields):
        print("[WARN] header rule sync (final) skipped: Word-updated TOC field results detected")
    else:
        header_rule_fix_final = _normalize_docx_headers_by_rules(output_docx, project_dir)
        if header_rule_fix_final.get("applied"):
            print(
                "[OK] header rule sync (final): updated_sections={count} page_numbering_updates={page_updates}".format(
                    count=header_rule_fix_final.get("updated_sections", 0),
                    page_updates=header_rule_fix_final.get("page_numbering_updates", 0),
                )
            )
    header_paragraph_layout_fix = _normalize_header_paragraph_layout(output_docx)
    if header_paragraph_layout_fix.get("applied"):
        print(
            "[OK] header paragraph layout fix: updated_parts={parts} updated_paragraphs={paras} cleared_indent_attrs={attrs}".format(
                parts=header_paragraph_layout_fix.get("updated_parts", 0),
                paras=header_paragraph_layout_fix.get("updated_paragraphs", 0),
                attrs=header_paragraph_layout_fix.get("cleared_indent_attrs", 0),
            )
        )
    non_cjk_font_fix = _force_non_cjk_runs_to_times_new_roman(output_docx)
    if non_cjk_font_fix.get("applied"):
        print(
            "[OK] non-cjk font fix: font={font} updated_runs={runs} updated_parts={parts}".format(
                font=non_cjk_font_fix.get("font", "Times New Roman"),
                runs=non_cjk_font_fix.get("updated_runs", 0),
                parts=non_cjk_font_fix.get("updated_parts", 0),
            )
        )
    if args.word_update_fields:
        word_repair = repair_docx_with_word_open_and_repair(
            output_docx,
            visible=args.word_visible,
            timeout_seconds=args.word_update_timeout,
        )
        if not word_repair.get("applied"):
            raise RuntimeError(f"Word OpenAndRepair 归一化失败：{word_repair.get('reason', 'unknown error')}")
        print(
            "[OK] word open-and-repair normalize: powershell={ps}".format(
                ps=word_repair.get("powershell", "NA"),
            )
        )

    package_check = _validate_docx_package_integrity(output_docx)
    if not package_check["ok"]:
        details: List[str] = []
        for key in [
            "missing_entries",
            "xml_parse_failures",
            "missing_rel_targets",
            "missing_override_parts",
            "missing_rid_refs",
            "unmatched_bookmarks",
            "invalid_mc_ignorable",
        ]:
            items = package_check.get(key, [])
            if items:
                details.append(f"{key}={len(items)}")
        detail_text = ", ".join(details) if details else "unknown package issue"
        if args.mode == "strict":
            raise RuntimeError(f"DOCX 结构完整性校验失败：{detail_text}")
        print(f"[WARN] package integrity: {detail_text}")
    else:
        print("[OK] package integrity: rel/content-types/xml checks passed")
    omml_gate_result = _enforce_omml_gate(output_docx, args.omml_gate_max)
    print(
        "[OK] omml scan: m:oMath={omath} m:oMathPara={opara} gate={gate}".format(
            omath=omml_gate_result.get("oMath", 0),
            opara=omml_gate_result.get("oMathPara", 0),
            gate=omml_gate_result.get("gate_max", "disabled"),
        )
    )

    print(f"[OK] markdown: {markdown_path}")
    print(f"[OK] docx: {output_docx}")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(f"[ERROR] {exc}", file=sys.stderr)
        raise
